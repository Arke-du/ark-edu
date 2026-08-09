import base64
import csv
import json
import os
import sqlite3
import uuid
import re
import unicodedata
import tempfile
import time
import requests

from dotenv import load_dotenv

# Carrega automaticamente as variáveis do arquivo .env
load_dotenv(override=False)

from datetime import datetime, timezone
from zoneinfo import ZoneInfo
from io import BytesIO
from pathlib import Path

import cv2
import numpy as np
import qrcode
from PIL import Image
from flask import Flask, flash, redirect, render_template, request, session, jsonify, url_for, send_file, send_from_directory, abort
from flask_mail import Mail, Message
from itsdangerous import URLSafeTimedSerializer
from werkzeug.security import check_password_hash, generate_password_hash
from werkzeug.utils import secure_filename
try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None
from lgpd_security import mascarar_cpf, preparar_cpf, somente_digitos, validar_cpf
from backup_manager import BackupManager, BackupError

from bncc_catalogo import garantir_tabela_bncc, consultar_bncc
from matrizes_catalogo import garantir_tabelas_matrizes, consultar_matrizes


BASE_DIR = os.path.dirname(os.path.abspath(__file__))


def _resolver_armazenamento_persistente():
    """Resolve banco e uploads sem depender do disco efêmero do Render.

    Prioridade:
    1. variáveis DATABASE_PATH / UPLOAD_FOLDER;
    2. RENDER_DISK_PATH (recomendado: /var/data);
    3. /var/data quando o diretório existir;
    4. caminhos locais do projeto apenas para desenvolvimento.
    """
    disk_path = os.environ.get("RENDER_DISK_PATH", "").strip()
    if not disk_path and os.path.isdir("/var/data"):
        disk_path = "/var/data"

    db_path = os.environ.get("DATABASE_PATH", "").strip()
    upload_folder = os.environ.get("UPLOAD_FOLDER", "").strip()

    if not db_path:
        db_path = (
            os.path.join(disk_path, "plataforma.db")
            if disk_path
            else os.path.join(BASE_DIR, "plataforma.db")
        )

    if not upload_folder:
        upload_folder = (
            os.path.join(disk_path, "uploads")
            if disk_path
            else os.path.join(BASE_DIR, "static", "uploads")
        )

    return os.path.abspath(db_path), os.path.abspath(upload_folder), disk_path


DB_PATH, UPLOAD_FOLDER, PERSISTENT_DISK_PATH = _resolver_armazenamento_persistente()

# Garante que o diretório do banco exista antes de qualquer CREATE TABLE.
os.makedirs(os.path.dirname(DB_PATH), exist_ok=True)

# Quando houver disco persistente, mantém também os backups no mesmo volume.
if PERSISTENT_DISK_PATH and not os.environ.get("BACKUP_DIR"):
    os.environ["BACKUP_DIR"] = os.path.join(PERSISTENT_DISK_PATH, "backups")

# Diagnóstico visível no log do Render. Não imprime segredos.
print(f"[ARK EDUS] Banco de dados: {DB_PATH}")
print(f"[ARK EDUS] Uploads: {UPLOAD_FOLDER}")
if os.environ.get("RENDER") and not PERSISTENT_DISK_PATH and not os.environ.get("DATABASE_PATH"):
    print(
        "[ARK EDUS][ATENCAO] Render sem armazenamento persistente configurado. "
        "Configure um Persistent Disk em /var/data e DATABASE_PATH=/var/data/plataforma.db."
    )

app = Flask(__name__)

TERMOS_VERSAO = os.environ.get("TERMOS_VERSAO", "1.0")
PRIVACIDADE_VERSAO = os.environ.get("PRIVACIDADE_VERSAO", "1.0")
CANAL_PRIVACIDADE = os.environ.get("CANAL_PRIVACIDADE", "privacidade@arkedus.com.br")

# Fuso horário oficial da plataforma (Guaraí/Tocantins).
FUSO_HORARIO_SISTEMA = ZoneInfo("America/Araguaina")


def agora_local():
    """Retorna a data e hora atuais de Guaraí/Tocantins."""
    return datetime.now(FUSO_HORARIO_SISTEMA)


@app.template_filter("formatar_data_hora_local")
def formatar_data_hora_local(valor):
    """
    Converte datas salvas pelo SQLite em UTC para o horário de Tocantins
    e exibe no formato DD/MM/AAAA às HH:mm.
    """
    if valor in (None, ""):
        return "—"

    if isinstance(valor, datetime):
        data = valor
    else:
        texto = str(valor).strip()
        try:
            data = datetime.fromisoformat(texto.replace("Z", "+00:00"))
        except ValueError:
            formatos = (
                "%Y-%m-%d %H:%M:%S.%f",
                "%Y-%m-%d %H:%M:%S",
                "%Y-%m-%dT%H:%M:%S.%f",
                "%Y-%m-%dT%H:%M:%S",
            )
            data = None
            for formato in formatos:
                try:
                    data = datetime.strptime(texto, formato)
                    break
                except ValueError:
                    continue
            if data is None:
                return texto

    # CURRENT_TIMESTAMP do SQLite é UTC. Quando o valor não traz fuso,
    # tratamos como UTC antes de converter para America/Araguaina.
    if data.tzinfo is None:
        data = data.replace(tzinfo=timezone.utc)

    data_local = data.astimezone(FUSO_HORARIO_SISTEMA)
    return data_local.strftime("%d/%m/%Y às %H:%M")


@app.template_filter("formatar_data_br")
def formatar_data_br(valor):
    """Exibe uma data no formato DD/MM/AAAA."""
    if valor in (None, ""):
        return "—"

    if isinstance(valor, datetime):
        data = valor
    else:
        texto = str(valor).strip()
        for formato in ("%Y-%m-%d", "%Y-%m-%d %H:%M:%S", "%d/%m/%Y"):
            try:
                data = datetime.strptime(texto, formato)
                break
            except ValueError:
                data = None
        if data is None:
            return texto

    return data.strftime("%d/%m/%Y")
app.secret_key = os.environ.get(
    "SECRET_KEY",
    "chave-temporaria-local-altere-no-render"
)

app.config.update(
    MAIL_SERVER=os.environ.get("MAIL_SERVER", "smtp.gmail.com"),
    MAIL_PORT=int(os.environ.get("MAIL_PORT", "587")),
    MAIL_USE_TLS=os.environ.get("MAIL_USE_TLS", "true").lower() == "true",
    MAIL_USE_SSL=os.environ.get("MAIL_USE_SSL", "false").lower() == "true",
    MAIL_USERNAME=os.environ.get("MAIL_USERNAME"),
    MAIL_PASSWORD=os.environ.get("MAIL_PASSWORD"),
    MAIL_DEFAULT_SENDER=os.environ.get("MAIL_DEFAULT_SENDER")
    or os.environ.get("MAIL_USERNAME"),
    UPLOAD_FOLDER=UPLOAD_FOLDER,
    MAX_CONTENT_LENGTH=10 * 1024 * 1024,
    SESSION_COOKIE_PATH="/",
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE="Lax",
    SESSION_COOKIE_SECURE=os.environ.get("SESSION_COOKIE_SECURE", "true").lower() == "true",
    PERMANENT_SESSION_LIFETIME=int(os.environ.get("SESSION_TIMEOUT_MINUTES", "30")) * 60,
)

os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)

def _nome_upload(valor):
    """Normaliza referências antigas e novas sem alterar o valor salvo no banco."""
    valor = (valor or "").replace("\\", "/").strip()
    if not valor:
        return ""
    if valor.startswith("/static/uploads/"):
        return valor[len("/static/uploads/"):]
    if valor.startswith("static/uploads/"):
        return valor[len("static/uploads/"):]
    if valor.startswith("/uploads/"):
        return valor[len("/uploads/"):]
    if valor.startswith("uploads/"):
        return valor[len("uploads/"):]
    return valor.lstrip("/")

@app.route("/uploads/<path:nome_arquivo>")
def servir_upload(nome_arquivo):
    return send_from_directory(app.config["UPLOAD_FOLDER"], nome_arquivo)

@app.context_processor
def disponibilizar_url_upload():
    def upload_url(valor):
        nome = _nome_upload(valor)
        return url_for("servir_upload", nome_arquivo=nome) if nome else ""
    return {"upload_url": upload_url}

@app.template_filter("normalizar_html_uploads")
def normalizar_html_uploads(valor):
    """Faz HTML antigo e novo de enunciados apontar para a rota persistente de uploads."""
    html = valor or ""
    html = html.replace('src="/static/uploads/', 'src="/uploads/')
    html = html.replace("src='/static/uploads/", "src='/uploads/")
    html = html.replace('src="static/uploads/', 'src="/uploads/')
    html = html.replace("src='static/uploads/", "src='/uploads/")
    return html


def _arquivo_upload_data_uri(valor):
    """Retorna um upload persistente como data URI, ideal para impressão/PDF."""
    nome = _nome_upload(valor)
    if not nome:
        return ""
    nome = os.path.basename(nome.split("?", 1)[0].split("#", 1)[0])
    caminho = os.path.join(app.config["UPLOAD_FOLDER"], nome)
    if not os.path.isfile(caminho):
        return ""
    ext = os.path.splitext(nome)[1].lower()
    mime = {
        ".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
        ".webp": "image/webp", ".gif": "image/gif"
    }.get(ext)
    if not mime:
        return ""
    try:
        with open(caminho, "rb") as arquivo:
            dados = base64.b64encode(arquivo.read()).decode("ascii")
        return f"data:{mime};base64,{dados}"
    except OSError:
        return ""


@app.template_filter("imagem_data_uri")
def imagem_data_uri(valor):
    """Converte o nome/caminho de uma imagem de upload em data URI quando possível."""
    nome = _nome_upload(valor)
    if not nome:
        return ""
    return _arquivo_upload_data_uri(nome) or url_for("servir_upload", nome_arquivo=nome)


@app.template_filter("embutir_imagens_html")
def embutir_imagens_html(valor):
    """Embutir imagens do editor em data URI para visualização e impressão/PDF."""
    html = normalizar_html_uploads(valor or "")
    padrao = re.compile(
        r'''src=(?P<q>["'])(?P<src>(?!data:)(?:https?://[^"']+)?(?:/)?(?:static/)?uploads/[^"']+)(?P=q)''',
        re.I,
    )

    def substituir(match):
        src = match.group("src")
        caminho_src = re.sub(r"^https?://[^/]+/", "/", src, flags=re.I)
        nome = _nome_upload(caminho_src.split("?", 1)[0].split("#", 1)[0])
        data_uri = _arquivo_upload_data_uri(nome)
        if not data_uri:
            return match.group(0)
        q = match.group("q")
        return f"src={q}{data_uri}{q}"

    return padrao.sub(substituir, html)


mail = Mail(app)

print("=" * 60)
print("MAIL_USERNAME:", app.config.get("MAIL_USERNAME"))
print("MAIL_PASSWORD:", app.config.get("MAIL_PASSWORD"))
print("MAIL_DEFAULT_SENDER:", app.config.get("MAIL_DEFAULT_SENDER"))
print("=" * 60)

serializer = URLSafeTimedSerializer(app.secret_key)
serializer = URLSafeTimedSerializer(app.secret_key)

garantir_tabela_bncc(DB_PATH)
garantir_tabelas_matrizes(DB_PATH)

backup_manager = BackupManager(app, DB_PATH, UPLOAD_FOLDER)
backup_manager.garantir_tabela()


def conectar_banco():
    banco = sqlite3.connect(DB_PATH)

    # Funções usadas pelos gatilhos de proteção do CPF. O banco armazena
    # somente a versão mascarada na coluna legada e mantém a versão
    # criptografada em coluna separada.
    def _cpf_encrypt(valor):
        if not valor or "*" in str(valor):
            return None
        try:
            return preparar_cpf(valor)[0]
        except (ValueError, RuntimeError):
            return None

    def _cpf_hash(valor):
        if not valor or "*" in str(valor):
            return None
        try:
            return preparar_cpf(valor)[1]
        except (ValueError, RuntimeError):
            return None

    def _cpf_final(valor):
        if not valor or "*" in str(valor):
            return None
        try:
            return preparar_cpf(valor)[2]
        except (ValueError, RuntimeError):
            return None

    def _cpf_mask(valor):
        return mascarar_cpf(valor) if valor and "*" not in str(valor) else valor

    banco.create_function("cpf_encrypt", 1, _cpf_encrypt)
    banco.create_function("cpf_hash", 1, _cpf_hash)
    banco.create_function("cpf_final", 1, _cpf_final)
    banco.create_function("cpf_mask", 1, _cpf_mask)
    banco.execute("PRAGMA foreign_keys = ON")
    banco.execute("PRAGMA busy_timeout = 5000")
    return banco


@app.template_filter("mascarar_cpf")
def filtro_mascarar_cpf(valor):
    return mascarar_cpf(valor)


def sincronizar_ano_letivo_instituicao(cursor, escola_id, ano, tornar_ativo=True):
    """Cria ou atualiza o ano letivo oficial de uma instituição.

    A tabela anos_letivos é a fonte oficial. O campo escolas.ano_letivo
    permanece sincronizado apenas por compatibilidade com telas antigas.
    """
    if not escola_id or ano in (None, ""):
        return None

    try:
        ano = int(ano)
    except (TypeError, ValueError):
        raise ValueError("O ano letivo informado é inválido.")

    cursor.execute("""
        SELECT id, ativo, encerrado
        FROM anos_letivos
        WHERE escola_id = ?
          AND ano = ?
        LIMIT 1
    """, (escola_id, ano))

    registro = cursor.fetchone()

    if tornar_ativo:
        cursor.execute("""
            UPDATE anos_letivos
            SET ativo = 0
            WHERE escola_id = ?
              AND ativo = 1
        """, (escola_id,))

    if registro:
        cursor.execute("""
            UPDATE anos_letivos
            SET ativo = ?,
                encerrado = 0
            WHERE id = ?
        """, (1 if tornar_ativo else registro["ativo"], registro["id"]))
        ano_letivo_id = registro["id"]
    else:
        cursor.execute("""
            INSERT INTO anos_letivos (
                escola_id,
                ano,
                ativo,
                encerrado
            )
            VALUES (?, ?, ?, 0)
        """, (escola_id, ano, 1 if tornar_ativo else 0))
        ano_letivo_id = cursor.lastrowid

    cursor.execute("""
        UPDATE escolas
        SET ano_letivo = ?
        WHERE id = ?
    """, (str(ano), escola_id))

    return ano_letivo_id


def sincronizar_anos_letivos_legados():
    """Migra automaticamente instituições antigas que só têm escolas.ano_letivo."""
    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        cursor.execute("""
            SELECT id, ano_letivo
            FROM escolas
            WHERE TRIM(COALESCE(ano_letivo, '')) <> ''
        """)

        for escola in cursor.fetchall():
            cursor.execute("""
                SELECT id
                FROM anos_letivos
                WHERE escola_id = ?
                  AND ativo = 1
                  AND encerrado = 0
                LIMIT 1
            """, (escola["id"],))

            if cursor.fetchone() is None:
                sincronizar_ano_letivo_instituicao(
                    cursor,
                    escola["id"],
                    escola["ano_letivo"],
                    tornar_ativo=True
                )

        banco.commit()

    except (sqlite3.Error, ValueError) as erro:
        banco.rollback()
        print("ERRO AO SINCRONIZAR ANOS LETIVOS LEGADOS:", erro)

    finally:
        banco.close()

def obter_escola_usuario(usuario_id=None):
    """
    Retorna o ID da instituição vinculada ao usuário.

    Primeiro tenta utilizar a escola registrada na sessão.
    Se não existir, consulta a tabela usuarios e atualiza
    a sessão automaticamente.
    """

    escola_id = session.get("escola_id")

    if escola_id:
        try:
            return int(escola_id)
        except (TypeError, ValueError):
            session.pop("escola_id", None)

    if not usuario_id:
        usuario_id = session.get("usuario_id")

    if not usuario_id:
        return None

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:

        cursor.execute("""
            SELECT escola_id
            FROM usuarios
            WHERE id = ?
            LIMIT 1
        """, (
            usuario_id,
        ))

        usuario = cursor.fetchone()

        if not usuario or not usuario["escola_id"]:
            return None

        escola_id = int(usuario["escola_id"])

        session["escola_id"] = escola_id

        return escola_id

    except sqlite3.Error as erro:

        print(
            "ERRO AO RECUPERAR A INSTITUIÇÃO DO USUÁRIO:",
            erro
        )

        return None

    finally:
        banco.close()


def obter_ano_letivo_ativo(escola_id):
    """
    Retorna o ano letivo oficialmente ativo da instituição.

    Retorno:
        sqlite3.Row com:
        - id
        - escola_id
        - ano
        - data_inicio
        - data_fim
        - ativo
        - encerrado

    Retorna None quando a instituição não possui ano ativo.
    """

    if not escola_id:
        return None

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:

        cursor.execute("""
            SELECT
                id,
                escola_id,
                ano,
                data_inicio,
                data_fim,
                ativo,
                encerrado
            FROM anos_letivos
            WHERE escola_id = ?
              AND ativo = 1
              AND encerrado = 0
            ORDER BY ano DESC
            LIMIT 1
        """, (
            escola_id,
        ))

        return cursor.fetchone()

    except sqlite3.Error as erro:

        print(
            "ERRO AO BUSCAR O ANO LETIVO ATIVO:",
            erro
        )

        return None

    finally:
        banco.close()


def obter_ano_letivo_selecionado(escola_id):
    """
    Retorna o ano letivo que deve ser utilizado pela plataforma.

    Futuramente, o administrador poderá selecionar um ano antigo
    para consulta. Quando não houver seleção manual, será utilizado
    automaticamente o ano letivo ativo.
    """

    if not escola_id:
        return None

    ano_selecionado_id = session.get("ano_letivo_selecionado_id")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:

        # =====================================================
        # ANO ESCOLHIDO MANUALMENTE PARA CONSULTA
        # =====================================================

        if ano_selecionado_id:

            cursor.execute("""
                SELECT
                    id,
                    escola_id,
                    ano,
                    data_inicio,
                    data_fim,
                    ativo,
                    encerrado
                FROM anos_letivos
                WHERE id = ?
                  AND escola_id = ?
                LIMIT 1
            """, (
                ano_selecionado_id,
                escola_id
            ))

            ano_selecionado = cursor.fetchone()

            if ano_selecionado:
                return ano_selecionado

            # Remove da sessão caso o ano não pertença à escola.
            session.pop(
                "ano_letivo_selecionado_id",
                None
            )

        # =====================================================
        # ANO ATIVO PADRÃO
        # =====================================================

        cursor.execute("""
            SELECT
                id,
                escola_id,
                ano,
                data_inicio,
                data_fim,
                ativo,
                encerrado
            FROM anos_letivos
            WHERE escola_id = ?
              AND ativo = 1
              AND encerrado = 0
            ORDER BY ano DESC
            LIMIT 1
        """, (
            escola_id,
        ))

        return cursor.fetchone()

    except sqlite3.Error as erro:

        print(
            "ERRO AO BUSCAR O ANO LETIVO SELECIONADO:",
            erro
        )

        return None

    finally:
        banco.close()


def atualizar_ano_letivo_na_sessao(escola_id):
    """
    Atualiza a sessão com o ano letivo utilizado pela plataforma.

    Retorna o registro do ano letivo ou None.
    """

    ano_letivo = obter_ano_letivo_selecionado(
        escola_id
    )

    if not ano_letivo:

        session.pop("ano_letivo_id", None)
        session.pop("ano_letivo", None)

        return None

    session["ano_letivo_id"] = ano_letivo["id"]
    session["ano_letivo"] = ano_letivo["ano"]

    return ano_letivo


# =========================================================
# GERENCIADOR GLOBAL DO ANO LETIVO
# =========================================================

def garantir_ano_atual_para_escola(escola_id):
    """
    Garante que o ano civil atual exista na tabela anos_letivos.

    A criação automática NÃO encerra nem troca o ano ativo existente.
    Quando a instituição ainda não possui nenhum ano letivo, o ano
    atual é criado como ativo. Caso já exista outro ano ativo, o novo
    ano é criado apenas como preparado, aguardando migração/ativação.
    """

    if not escola_id:
        return None

    ano_atual = datetime.now().year

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        cursor.execute("""
            SELECT id, escola_id, ano, ativo, encerrado
            FROM anos_letivos
            WHERE escola_id = ? AND ano = ?
            LIMIT 1
        """, (escola_id, ano_atual))

        existente = cursor.fetchone()
        if existente:
            # Correção de consistência: se o ano civil atual existe, mas a
            # instituição ficou sem nenhum ano letivo ativo (por exemplo, após
            # migração/backup), reabre automaticamente o ano atual. Isso evita
            # que o ano corrente apareça como "Encerrado" por engano.
            cursor.execute("""
                SELECT id
                FROM anos_letivos
                WHERE escola_id = ?
                  AND ativo = 1
                  AND encerrado = 0
                LIMIT 1
            """, (escola_id,))
            ativo_existente = cursor.fetchone()

            if ativo_existente is None and (existente["ativo"] != 1 or existente["encerrado"] != 0):
                cursor.execute("""
                    UPDATE anos_letivos
                    SET ativo = 1,
                        encerrado = 0
                    WHERE id = ?
                """, (existente["id"],))
                cursor.execute("""
                    UPDATE escolas
                    SET ano_letivo = ?
                    WHERE id = ?
                """, (str(ano_atual), escola_id))
                banco.commit()
                cursor.execute("""
                    SELECT id, escola_id, ano, ativo, encerrado
                    FROM anos_letivos
                    WHERE id = ?
                """, (existente["id"],))
                return cursor.fetchone()

            return existente

        cursor.execute("""
            SELECT id
            FROM anos_letivos
            WHERE escola_id = ?
              AND ativo = 1
              AND encerrado = 0
            LIMIT 1
        """, (escola_id,))

        possui_ativo = cursor.fetchone() is not None

        cursor.execute("""
            INSERT INTO anos_letivos (
                escola_id,
                ano,
                ativo,
                encerrado
            )
            VALUES (?, ?, ?, 0)
        """, (
            escola_id,
            ano_atual,
            0 if possui_ativo else 1
        ))

        novo_id = cursor.lastrowid
        banco.commit()

        cursor.execute("""
            SELECT id, escola_id, ano, ativo, encerrado
            FROM anos_letivos
            WHERE id = ?
        """, (novo_id,))

        return cursor.fetchone()

    except sqlite3.Error as erro:
        banco.rollback()
        print("ERRO AO GARANTIR ANO ATUAL:", erro)
        return None

    finally:
        banco.close()


def corrigir_ano_corrente_inconsistente():
    """Corrige estados claramente inconsistentes do ano letivo corrente.

    Se o ano civil atual estiver marcado como encerrado enquanto não existe
    ano ativo, ou enquanto um ano futuro foi ativado por engano, o ano atual
    volta a ser o ativo. O ano futuro permanece cadastrado como planejado.
    """
    ano_atual = datetime.now().year
    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()
    try:
        cursor.execute("SELECT id FROM escolas WHERE COALESCE(status, 1) = 1")
        escolas_ids = [r["id"] for r in cursor.fetchall()]

        for escola_id in escolas_ids:
            cursor.execute("""
                SELECT id, ano, ativo, encerrado
                FROM anos_letivos
                WHERE escola_id = ? AND ano = ?
                LIMIT 1
            """, (escola_id, ano_atual))
            atual = cursor.fetchone()
            if not atual:
                continue

            cursor.execute("""
                SELECT id, ano
                FROM anos_letivos
                WHERE escola_id = ? AND ativo = 1 AND encerrado = 0
                ORDER BY ano DESC
                LIMIT 1
            """, (escola_id,))
            ativo = cursor.fetchone()

            precisa_corrigir = (
                atual["ativo"] != 1 or atual["encerrado"] != 0
            ) and (
                ativo is None or int(ativo["ano"]) > ano_atual
            )

            if not precisa_corrigir:
                continue

            # Um ano futuro ativado antes da hora volta para "Planejado".
            cursor.execute("""
                UPDATE anos_letivos
                SET ativo = 0, encerrado = 0
                WHERE escola_id = ? AND ano > ? AND ativo = 1
            """, (escola_id, ano_atual))
            cursor.execute("""
                UPDATE anos_letivos
                SET ativo = 1, encerrado = 0
                WHERE id = ?
            """, (atual["id"],))
            cursor.execute("UPDATE escolas SET ano_letivo = ? WHERE id = ?", (str(ano_atual), escola_id))

        banco.commit()
    except sqlite3.Error as erro:
        banco.rollback()
        print("ERRO AO CORRIGIR ANO LETIVO CORRENTE:", erro)
    finally:
        banco.close()


def obter_ano_global_administrador():
    """Retorna o número do ano visualizado pelo Administrador Geral."""

    ano_sessao = session.get("ano_letivo_visualizado")

    if ano_sessao:
        try:
            return int(ano_sessao)
        except (TypeError, ValueError):
            session.pop("ano_letivo_visualizado", None)

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        ano_atual = datetime.now().year

        cursor.execute("""
            SELECT ano
            FROM anos_letivos
            WHERE ano = ?
            LIMIT 1
        """, (ano_atual,))

        registro = cursor.fetchone()

        if not registro:
            cursor.execute("""
                SELECT ano
                FROM anos_letivos
                WHERE ativo = 1
                  AND encerrado = 0
                ORDER BY ano DESC
                LIMIT 1
            """)
            registro = cursor.fetchone()

        if not registro:
            cursor.execute("""
                SELECT ano
                FROM anos_letivos
                ORDER BY ano DESC
                LIMIT 1
            """)
            registro = cursor.fetchone()

        if not registro:
            return None

        ano = int(registro["ano"])
        session["ano_letivo_visualizado"] = ano
        session["ano_letivo"] = ano
        return ano

    except sqlite3.Error as erro:
        print("ERRO AO OBTER ANO GLOBAL DO ADMINISTRADOR:", erro)
        return None

    finally:
        banco.close()


def obter_contexto_plataforma():
    """
    Retorna o contexto único usado pelas páginas da plataforma.

    Para usuários de uma instituição, ano_letivo_id identifica o
    registro exato da escola. Para o Administrador Geral, o filtro
    global usa o número do ano, porque cada escola possui seu próprio
    ID para o mesmo período.
    """

    usuario_id = session.get("usuario_id")
    cargo = session.get("usuario_cargo", "").strip()
    escola_id = obter_escola_usuario() if cargo != "Administrador Geral" else None

    contexto = {
        "usuario_id": usuario_id,
        "cargo": cargo,
        "escola_id": escola_id,
        "ano_letivo_id": None,
        "ano": None,
        "ano_ativo": None,
        "ano_ativo_id": None,
        "consultando_historico": False
    }

    if not usuario_id:
        return contexto

    if cargo == "Administrador Geral":
        contexto["ano"] = obter_ano_global_administrador()
        return contexto

    if not escola_id:
        return contexto

    garantir_ano_atual_para_escola(escola_id)

    ano_ativo = obter_ano_letivo_ativo(escola_id)
    ano_visualizado = atualizar_ano_letivo_na_sessao(escola_id)

    if ano_ativo:
        contexto["ano_ativo"] = ano_ativo["ano"]
        contexto["ano_ativo_id"] = ano_ativo["id"]

    if ano_visualizado:
        contexto["ano"] = ano_visualizado["ano"]
        contexto["ano_letivo_id"] = ano_visualizado["id"]
        contexto["consultando_historico"] = not (
            ano_visualizado["ativo"] == 1
            and ano_visualizado["encerrado"] == 0
        )

    return contexto


@app.context_processor
def contexto_global_ano_letivo():
    """Disponibiliza o seletor de ano letivo em todos os templates."""

    if "usuario_id" not in session:
        return {
            "ano_contexto": None,
            "ano_ativo_contexto": None,
            "anos_disponiveis": [],
            "consultando_ano_antigo": False
        }

    contexto = obter_contexto_plataforma()
    cargo = contexto["cargo"]

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        if cargo == "Administrador Geral":
            cursor.execute("""
                SELECT DISTINCT ano
                FROM anos_letivos
                ORDER BY ano DESC
            """)
        elif contexto["escola_id"]:
            cursor.execute("""
                SELECT id, ano, ativo, encerrado
                FROM anos_letivos
                WHERE escola_id = ?
                ORDER BY ano DESC
            """, (contexto["escola_id"],))
        else:
            return {
                "ano_contexto": None,
                "ano_ativo_contexto": None,
                "anos_disponiveis": [],
                "consultando_ano_antigo": False
            }

        return {
            "ano_contexto": contexto["ano"],
            "ano_ativo_contexto": contexto["ano_ativo"],
            "anos_disponiveis": cursor.fetchall(),
            "consultando_ano_antigo": contexto["consultando_historico"]
        }

    except sqlite3.Error as erro:
        print("ERRO NO CONTEXTO GLOBAL DO ANO LETIVO:", erro)
        return {
            "ano_contexto": contexto.get("ano"),
            "ano_ativo_contexto": contexto.get("ano_ativo"),
            "anos_disponiveis": [],
            "consultando_ano_antigo": False
        }

    finally:
        banco.close()


@app.route("/trocar-ano-letivo/<int:ano>", methods=["POST"])
def trocar_ano_letivo(ano):
    """Troca apenas o ano visualizado, sem alterar o ano ativo."""

    if "usuario_id" not in session:
        return redirect("/login")

    cargo = session.get("usuario_cargo", "").strip()
    escola_id = obter_escola_usuario()

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        if cargo == "Administrador Geral":
            cursor.execute("""
                SELECT ano
                FROM anos_letivos
                WHERE ano = ?
                LIMIT 1
            """, (ano,))

            registro = cursor.fetchone()

            if not registro:
                flash("Ano letivo não encontrado.", "erro")
                return redirect(request.referrer or "/")

            session["ano_letivo_visualizado"] = ano
            session["ano_letivo"] = ano
            session.pop("ano_letivo_id", None)
            session.pop("ano_letivo_selecionado_id", None)

        else:
            if not escola_id:
                flash("Não foi possível identificar sua instituição.", "erro")
                return redirect(request.referrer or "/")

            cursor.execute("""
                SELECT id, ano, ativo, encerrado
                FROM anos_letivos
                WHERE escola_id = ?
                  AND ano = ?
                LIMIT 1
            """, (escola_id, ano))

            registro = cursor.fetchone()

            if not registro:
                flash(
                    "O ano letivo selecionado não pertence à sua instituição.",
                    "erro"
                )
                return redirect(request.referrer or "/")

            session["ano_letivo_selecionado_id"] = registro["id"]
            session["ano_letivo_id"] = registro["id"]
            session["ano_letivo"] = registro["ano"]
            session["ano_letivo_visualizado"] = registro["ano"]

        flash(f"Visualizando o ano letivo {ano}.", "success")
        return redirect(request.referrer or "/")

    except sqlite3.Error as erro:
        print("ERRO AO TROCAR ANO LETIVO:", erro)
        flash(f"Erro ao trocar o ano letivo: {erro}", "erro")
        return redirect(request.referrer or "/")

    finally:
        banco.close()


def cargo_permitido(cargos_permitidos):

    if "usuario_id" not in session:
        return False

    cargo_atual = session.get("usuario_cargo")
    if cargo_atual in cargos_permitidos:
        return True
    if cargo_atual == "Professor Autônomo":
        # Nas rotas acadêmicas o autônomo administra o próprio workspace.
        # Áreas institucionais sensíveis permanecem bloqueadas no módulo SaaS.
        equivalentes = {"Administrador da Instituição", "Coordenador", "Secretaria", "Professor"}
        return bool(equivalentes.intersection(set(cargos_permitidos)))
    return False

def criar_tabelas():
    banco = conectar_banco()
    cursor = banco.cursor()

    cursor.execute("PRAGMA foreign_keys = ON")

    def garantir_coluna(tabela, coluna, definicao):
        cursor.execute(f"PRAGMA table_info({tabela})")
        colunas = {linha[1] for linha in cursor.fetchall()}

        if coluna not in colunas:
            cursor.execute(
                f"ALTER TABLE {tabela} ADD COLUMN {coluna} {definicao}"
            )

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS escolas (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nome_instituicao TEXT NOT NULL,
            codigo_inep TEXT,
            cnpj TEXT,
            cep TEXT,
            endereco TEXT,
            cidade TEXT,
            estado TEXT,
            telefone TEXT,
            whatsapp TEXT,
            email TEXT,
            site TEXT,
            diretor TEXT,
            coordenador1 TEXT,
            coordenador2 TEXT,
            coordenador3 TEXT,
            secretario TEXT,
            tipo_instituicao TEXT,
            ano_letivo TEXT,
            modalidade_ensino TEXT,
            etapas_ensino TEXT,
            logo TEXT,
            status INTEGER DEFAULT 1,
            criado_em TEXT
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS anos_letivos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            escola_id INTEGER NOT NULL,
            ano INTEGER NOT NULL,
            data_inicio TEXT,
            data_fim TEXT,
            ativo INTEGER NOT NULL DEFAULT 0,
            encerrado INTEGER NOT NULL DEFAULT 0,
            criado_em TEXT,
            FOREIGN KEY (escola_id) REFERENCES escolas(id) ON DELETE CASCADE,
            UNIQUE (escola_id, ano)
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS cargos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nome TEXT UNIQUE NOT NULL
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS usuarios (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nome TEXT NOT NULL,
            email TEXT UNIQUE NOT NULL,
            senha TEXT NOT NULL,
            cargo_id INTEGER,
            ativo INTEGER DEFAULT 1,
            escola_id INTEGER,
            cpf TEXT,
            telefone TEXT,
            data_nascimento TEXT,
            FOREIGN KEY (cargo_id) REFERENCES cargos(id),
            FOREIGN KEY (escola_id) REFERENCES escolas(id) ON DELETE SET NULL
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS turmas (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nome TEXT NOT NULL,
            etapa TEXT,
            ano TEXT NOT NULL,
            turno TEXT NOT NULL,
            escola_id INTEGER,
            FOREIGN KEY (escola_id) REFERENCES escolas(id) ON DELETE CASCADE
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS alunos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nome TEXT NOT NULL,
            matricula TEXT,
            turma_id INTEGER NOT NULL,
            escola_id INTEGER,
            FOREIGN KEY (turma_id) REFERENCES turmas(id) ON DELETE CASCADE,
            FOREIGN KEY (escola_id) REFERENCES escolas(id) ON DELETE CASCADE
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS aluno_matriculas (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            aluno_id INTEGER NOT NULL,
            escola_id INTEGER NOT NULL,
            ano_letivo_id INTEGER NOT NULL,
            turma_id INTEGER NOT NULL,
            situacao TEXT NOT NULL DEFAULT 'Cursando',
            data_matricula TEXT DEFAULT CURRENT_TIMESTAMP,
            data_encerramento TEXT,
            observacao TEXT,
            criado_em TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (aluno_id) REFERENCES alunos(id) ON DELETE CASCADE,
            FOREIGN KEY (escola_id) REFERENCES escolas(id) ON DELETE CASCADE,
            FOREIGN KEY (ano_letivo_id) REFERENCES anos_letivos(id) ON DELETE CASCADE,
            FOREIGN KEY (turma_id) REFERENCES turmas(id) ON DELETE CASCADE,
            UNIQUE (aluno_id, ano_letivo_id)
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS professores (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nome TEXT NOT NULL,
            email TEXT,
            disciplina TEXT,
            escola_id INTEGER,
            FOREIGN KEY (escola_id) REFERENCES escolas(id) ON DELETE CASCADE
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS professor_disciplinas (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            professor_id INTEGER NOT NULL,
            disciplina TEXT NOT NULL,
            FOREIGN KEY (professor_id) REFERENCES professores(id) ON DELETE CASCADE
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS professor_turmas (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            professor_id INTEGER NOT NULL,
            turma_id INTEGER NOT NULL,
            FOREIGN KEY (professor_id) REFERENCES professores(id) ON DELETE CASCADE,
            FOREIGN KEY (turma_id) REFERENCES turmas(id) ON DELETE CASCADE,
            UNIQUE (professor_id, turma_id)
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS componentes_curriculares (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            escola_id INTEGER NOT NULL,
            etapa_ensino TEXT NOT NULL,
            nome TEXT NOT NULL,
            tipo TEXT NOT NULL DEFAULT 'padrao',
            ativo INTEGER NOT NULL DEFAULT 1,
            FOREIGN KEY (escola_id) REFERENCES escolas(id) ON DELETE CASCADE
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS professor_componentes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            professor_id INTEGER NOT NULL,
            turma_id INTEGER NOT NULL,
            componente_id INTEGER NOT NULL,
            escola_id INTEGER NOT NULL,
            FOREIGN KEY (professor_id) REFERENCES usuarios(id) ON DELETE CASCADE,
            FOREIGN KEY (turma_id) REFERENCES turmas(id) ON DELETE CASCADE,
            FOREIGN KEY (componente_id) REFERENCES componentes_curriculares(id) ON DELETE CASCADE,
            FOREIGN KEY (escola_id) REFERENCES escolas(id) ON DELETE CASCADE,
            UNIQUE (professor_id, turma_id, componente_id)
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS assuntos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            escola_id INTEGER,
            disciplina TEXT NOT NULL,
            etapa_ensino TEXT NOT NULL,
            ano_serie TEXT NOT NULL,
            nome TEXT NOT NULL,
            ativo INTEGER NOT NULL DEFAULT 1,
            criado_em TEXT DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (escola_id) REFERENCES escolas(id) ON DELETE CASCADE
        )
    """)

    cursor.execute("""
        CREATE INDEX IF NOT EXISTS idx_assuntos_filtro
        ON assuntos (
            escola_id,
            disciplina,
            etapa_ensino,
            ano_serie,
            ativo
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS questoes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            disciplina TEXT NOT NULL,
            assunto TEXT,
            tipo_questao TEXT NOT NULL DEFAULT 'multipla_escolha',
            enunciado TEXT NOT NULL,
            imagem TEXT,
            alternativa_a TEXT NOT NULL DEFAULT '',
            alternativa_b TEXT NOT NULL DEFAULT '',
            alternativa_c TEXT NOT NULL DEFAULT '',
            alternativa_d TEXT NOT NULL DEFAULT '',
            correta TEXT NOT NULL DEFAULT '',
            alternativas_json TEXT,
            respostas_corretas TEXT,
            resposta_esperada TEXT,
            criterios_correcao TEXT,
            habilidade TEXT,
            dificuldade TEXT NOT NULL,
            observacoes TEXT,
            escola_id INTEGER,
            criado_por INTEGER,
            criado_em TEXT DEFAULT CURRENT_TIMESTAMP,
            atualizado_em TEXT,
            FOREIGN KEY (escola_id) REFERENCES escolas(id) ON DELETE CASCADE,
            FOREIGN KEY (criado_por) REFERENCES usuarios(id) ON DELETE SET NULL
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS provas (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nome TEXT NOT NULL,
            turma_id INTEGER NOT NULL,
            professor_id INTEGER,
            disciplina TEXT NOT NULL,
            quantidade INTEGER NOT NULL,
            data_geracao TEXT,
            data_aplicacao TEXT,
            escola_id INTEGER,
            FOREIGN KEY (turma_id) REFERENCES turmas(id),
            FOREIGN KEY (professor_id) REFERENCES professores(id),
            FOREIGN KEY (escola_id) REFERENCES escolas(id) ON DELETE CASCADE
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS prova_questoes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            prova_id INTEGER NOT NULL,
            questao_id INTEGER NOT NULL,
            FOREIGN KEY (prova_id) REFERENCES provas(id) ON DELETE CASCADE,
            FOREIGN KEY (questao_id) REFERENCES questoes(id) ON DELETE CASCADE
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS respostas_alunos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            prova_id INTEGER,
            aluno_id INTEGER,
            numero_questao INTEGER,
            resposta_aluno TEXT,
            resposta_correta TEXT,
            acertou INTEGER,
            FOREIGN KEY (prova_id) REFERENCES provas(id) ON DELETE CASCADE,
            FOREIGN KEY (aluno_id) REFERENCES alunos(id) ON DELETE CASCADE
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS resultados (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            prova_id INTEGER,
            aluno_id INTEGER,
            acertos INTEGER,
            erros INTEGER,
            nota REAL,
            FOREIGN KEY (prova_id) REFERENCES provas(id) ON DELETE CASCADE,
            FOREIGN KEY (aluno_id) REFERENCES alunos(id) ON DELETE CASCADE
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS instituicao (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nome TEXT,
            cidade TEXT,
            estado TEXT,
            diretor TEXT,
            coordenador TEXT,
            ano_letivo TEXT,
            logo TEXT
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS permissoes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            cargo_id INTEGER,
            modulo TEXT,
            visualizar INTEGER DEFAULT 0,
            cadastrar INTEGER DEFAULT 0,
            editar INTEGER DEFAULT 0,
            excluir INTEGER DEFAULT 0,
            pode_acessar INTEGER DEFAULT 0,
            FOREIGN KEY (cargo_id) REFERENCES cargos(id) ON DELETE CASCADE
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS codigos_recuperacao (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            usuario_id INTEGER NOT NULL,
            codigo TEXT NOT NULL,
            usado INTEGER DEFAULT 0,
            criado_em TEXT,
            FOREIGN KEY (usuario_id) REFERENCES usuarios(id) ON DELETE CASCADE
        )
    """)

    # Catálogo pedagógico inicial. Os registros com escola_id NULL
    # ficam disponíveis para todas as instituições. Assuntos temporários
    # digitados durante a criação de uma questão não entram nesta tabela.
    assuntos_padrao = {
        ("Língua Portuguesa", "Ensino Fundamental - Anos Finais", "6º ano"): [
            "Leitura e interpretação de texto", "Gêneros textuais",
            "Substantivos", "Adjetivos", "Artigos e numerais",
            "Pronomes", "Verbos", "Ortografia", "Pontuação",
            "Produção textual"
        ],
        ("Língua Portuguesa", "Ensino Fundamental - Anos Finais", "7º ano"): [
            "Leitura e interpretação de texto", "Gêneros textuais",
            "Verbos e locuções verbais", "Advérbios", "Preposições",
            "Conjunções", "Sujeito e predicado", "Ortografia",
            "Pontuação", "Produção textual"
        ],
        ("Língua Portuguesa", "Ensino Fundamental - Anos Finais", "8º ano"): [
            "Leitura e interpretação de texto", "Gêneros textuais",
            "Termos essenciais da oração", "Complementos verbais",
            "Adjunto adnominal e adjunto adverbial", "Vozes verbais",
            "Período simples e composto", "Figuras de linguagem",
            "Pontuação", "Produção textual"
        ],
        ("Língua Portuguesa", "Ensino Fundamental - Anos Finais", "9º ano"): [
            "Leitura e interpretação de texto", "Gêneros textuais",
            "Orações coordenadas", "Orações subordinadas",
            "Concordância verbal e nominal", "Regência verbal e nominal",
            "Crase", "Figuras de linguagem", "Variação linguística",
            "Produção textual"
        ],
        ("Matemática", "Ensino Fundamental - Anos Finais", "6º ano"): [
            "Números naturais", "Operações fundamentais", "Múltiplos e divisores",
            "Frações", "Números decimais", "Porcentagem", "Razão e proporção",
            "Regra de três", "Geometria plana", "Grandezas e medidas"
        ],
        ("Matemática", "Ensino Fundamental - Anos Finais", "7º ano"): [
            "Números inteiros", "Números racionais", "Expressões algébricas",
            "Equações do primeiro grau", "Razão e proporção", "Regra de três",
            "Porcentagem", "Ângulos", "Triângulos", "Probabilidade"
        ],
        ("Matemática", "Ensino Fundamental - Anos Finais", "8º ano"): [
            "Produtos notáveis", "Fatoração", "Sistemas de equações",
            "Potenciação e radiciação", "Porcentagem e juros",
            "Teorema de Pitágoras", "Polígonos", "Área e volume",
            "Estatística", "Probabilidade"
        ],
        ("Matemática", "Ensino Fundamental - Anos Finais", "9º ano"): [
            "Números reais", "Equações do segundo grau", "Funções",
            "Semelhança de triângulos", "Teorema de Tales",
            "Trigonometria no triângulo retângulo", "Circunferência",
            "Área e volume", "Estatística", "Probabilidade"
        ],
        ("História", "Ensino Fundamental - Anos Finais", "6º ano"): [
            "Introdução aos estudos históricos", "Pré-História",
            "Mesopotâmia", "Egito Antigo", "Povos Hebreus",
            "Fenícios e Persas", "Grécia Antiga", "Roma Antiga",
            "África Antiga", "Povos originários da América"
        ],
        ("História", "Ensino Fundamental - Anos Finais", "7º ano"): [
            "Feudalismo", "Mundo islâmico", "Renascimento cultural",
            "Reformas religiosas", "Expansão marítima europeia",
            "Povos pré-colombianos", "Colonização espanhola",
            "Colonização portuguesa", "Brasil Colonial",
            "Escravidão e resistência"
        ],
        ("História", "Ensino Fundamental - Anos Finais", "8º ano"): [
            "Iluminismo", "Revolução Industrial", "Independência dos Estados Unidos",
            "Revolução Francesa", "Era Napoleônica", "Independências na América",
            "Primeiro Reinado", "Período Regencial", "Segundo Reinado",
            "Abolição da escravidão"
        ],
        ("História", "Ensino Fundamental - Anos Finais", "9º ano"): [
            "Primeira República", "Primeira Guerra Mundial",
            "Revolução Russa", "Crise de 1929", "Nazifascismo",
            "Era Vargas", "Segunda Guerra Mundial", "Guerra Fria",
            "Ditadura Militar no Brasil", "Nova República"
        ],
        ("Geografia", "Ensino Fundamental - Anos Finais", "6º ano"): [
            "Paisagem e espaço geográfico", "Cartografia",
            "Orientação e coordenadas geográficas", "Relevo",
            "Clima", "Hidrografia", "Vegetação", "População",
            "Espaço urbano e rural", "Questões ambientais"
        ],
        ("Ciências", "Ensino Fundamental - Anos Finais", "6º ano"): [
            "Matéria e energia", "Misturas e separação de materiais",
            "Transformações químicas", "Célula", "Tecidos e órgãos",
            "Sistema locomotor", "Sistema nervoso", "Visão e audição",
            "Terra e Universo", "Camadas da Terra"
        ],
        ("Língua Portuguesa", "Ensino Médio", "1ª série"): [
            "Leitura e interpretação", "Gêneros discursivos",
            "Literatura medieval", "Renascimento e Classicismo",
            "Barroco", "Arcadismo", "Funções da linguagem",
            "Variação linguística", "Morfologia", "Produção textual"
        ],
        ("Língua Portuguesa", "Ensino Médio", "2ª série"): [
            "Leitura e interpretação", "Romantismo", "Realismo",
            "Naturalismo", "Parnasianismo", "Simbolismo",
            "Sintaxe do período simples", "Sintaxe do período composto",
            "Semântica", "Produção textual"
        ],
        ("Língua Portuguesa", "Ensino Médio", "3ª série"): [
            "Leitura e interpretação", "Pré-Modernismo", "Modernismo",
            "Literatura contemporânea", "Concordância", "Regência",
            "Crase", "Coesão e coerência", "Redação do ENEM",
            "Revisão gramatical"
        ]
    }

    for (disciplina_assunto, etapa_assunto, serie_assunto), nomes_assuntos in assuntos_padrao.items():
        for nome_assunto in nomes_assuntos:
            cursor.execute("""
                INSERT INTO assuntos (
                    escola_id,
                    disciplina,
                    etapa_ensino,
                    ano_serie,
                    nome,
                    ativo
                )
                SELECT NULL, ?, ?, ?, ?, 1
                WHERE NOT EXISTS (
                    SELECT 1
                    FROM assuntos
                    WHERE escola_id IS NULL
                      AND disciplina = ?
                      AND etapa_ensino = ?
                      AND ano_serie = ?
                      AND nome = ?
                )
            """, (
                disciplina_assunto, etapa_assunto, serie_assunto, nome_assunto,
                disciplina_assunto, etapa_assunto, serie_assunto, nome_assunto
            ))

    # Catálogo ampliado para o Ensino Fundamental - Anos Finais e Ensino Médio.
    # Apenas assuntos ainda inexistentes são incluídos.
    assuntos_ampliados = {
        # =====================================================
        # ENSINO FUNDAMENTAL - ANOS FINAIS
        # =====================================================
        ("Língua Portuguesa", "Ensino Fundamental - Anos Finais", "6º ano"): [
            "Tipos de linguagem", "Elementos da comunicação", "Sentido literal e figurado",
            "Sinônimos e antônimos", "Parônimos e homônimos", "Estrutura das palavras",
            "Formação de palavras", "Sílaba tônica", "Acentuação gráfica",
            "Frase, oração e período", "Tipos de frase", "Discurso direto e indireto",
            "Coesão textual", "Coerência textual", "Texto narrativo", "Texto descritivo",
            "Texto injuntivo", "Notícia", "Conto", "Fábula", "Poema",
            "História em quadrinhos", "Bilhete", "Carta pessoal"
        ],
        ("Língua Portuguesa", "Ensino Fundamental - Anos Finais", "7º ano"): [
            "Derivação e composição", "Modo indicativo", "Modo subjuntivo",
            "Modo imperativo", "Verbos regulares e irregulares", "Transitividade verbal",
            "Predicação verbal", "Tipos de sujeito", "Tipos de predicado",
            "Complemento nominal", "Aposto e vocativo", "Adjunto adnominal",
            "Adjunto adverbial", "Concordância verbal", "Concordância nominal",
            "Crônica", "Entrevista", "Reportagem", "Carta do leitor",
            "Resumo", "Relato pessoal", "Biografia", "Autobiografia"
        ],
        ("Língua Portuguesa", "Ensino Fundamental - Anos Finais", "8º ano"): [
            "Predicativo do sujeito", "Predicativo do objeto", "Objeto direto",
            "Objeto indireto", "Agente da passiva", "Voz ativa", "Voz passiva",
            "Voz reflexiva", "Orações coordenadas", "Orações subordinadas",
            "Conjunções coordenativas", "Conjunções subordinativas",
            "Regência verbal", "Regência nominal", "Crase",
            "Artigo de opinião", "Editorial", "Resenha crítica",
            "Texto publicitário", "Charge", "Cartum", "Infográfico"
        ],
        ("Língua Portuguesa", "Ensino Fundamental - Anos Finais", "9º ano"): [
            "Orações subordinadas substantivas", "Orações subordinadas adjetivas",
            "Orações subordinadas adverbiais", "Colocação pronominal",
            "Próclise", "Ênclise", "Mesóclise", "Pontuação no período composto",
            "Denotação e conotação", "Ambiguidade", "Intertextualidade",
            "Ironia", "Humor", "Tese e argumentação", "Estratégias argumentativas",
            "Dissertação argumentativa", "Manifesto", "Debate", "Seminário",
            "Podcast", "Editorial", "Carta aberta"
        ],

        ("Matemática", "Ensino Fundamental - Anos Finais", "6º ano"): [
            "Sistema de numeração decimal", "Expressões numéricas",
            "Critérios de divisibilidade", "Números primos", "MMC", "MDC",
            "Frações equivalentes", "Comparação de frações",
            "Adição e subtração de frações", "Multiplicação e divisão de frações",
            "Operações com números decimais", "Plano cartesiano", "Ângulos",
            "Retas e segmentos", "Polígonos", "Perímetro", "Área de figuras planas",
            "Sistema métrico decimal", "Leitura de tabelas", "Leitura de gráficos"
        ],
        ("Matemática", "Ensino Fundamental - Anos Finais", "7º ano"): [
            "Operações com números inteiros", "Expressões com números inteiros",
            "Operações com números racionais", "Proporcionalidade direta",
            "Proporcionalidade inversa", "Escala", "Equações com uma incógnita",
            "Inequações", "Sequências numéricas", "Linguagem algébrica",
            "Retas paralelas e transversais", "Congruência de triângulos",
            "Quadriláteros", "Circunferência", "Área de polígonos",
            "Volume de blocos retangulares", "Média aritmética",
            "Gráficos estatísticos", "Princípio multiplicativo"
        ],
        ("Matemática", "Ensino Fundamental - Anos Finais", "8º ano"): [
            "Monômios", "Polinômios", "Operações com polinômios",
            "Equações fracionárias", "Sistemas lineares", "Função afim",
            "Sequências", "Notação científica", "Radicais", "Razões notáveis",
            "Semelhança de figuras", "Construções geométricas",
            "Mediatriz e bissetriz", "Prismas", "Cilindros",
            "Área total", "Volume", "Medidas de tendência central",
            "Gráficos e tabelas", "Probabilidade experimental"
        ],
        ("Matemática", "Ensino Fundamental - Anos Finais", "9º ano"): [
            "Potências e raízes", "Notação científica", "Produtos notáveis",
            "Fatoração algébrica", "Equações biquadradas", "Sistemas quadráticos",
            "Função afim", "Função quadrática", "Gráficos de funções",
            "Razões trigonométricas", "Relações métricas no triângulo retângulo",
            "Polígonos regulares", "Comprimento da circunferência",
            "Área do círculo", "Prismas e cilindros", "Volume de sólidos",
            "Análise combinatória", "Probabilidade composta",
            "Média, moda e mediana", "Gráficos estatísticos"
        ],

        ("História", "Ensino Fundamental - Anos Finais", "6º ano"): [
            "Fontes históricas", "Tempo histórico", "Periodização da História",
            "Evolução humana", "Paleolítico", "Neolítico", "Idade dos Metais",
            "Revolução Agrícola", "Civilizações hidráulicas", "Código de Hamurábi",
            "Religião no Egito", "Sociedade egípcia", "Democracia ateniense",
            "Esparta", "Guerras Médicas", "Guerra do Peloponeso",
            "República Romana", "Império Romano", "Cristianismo",
            "Queda do Império Romano", "Reinos africanos antigos",
            "Povos indígenas do Brasil"
        ],
        ("História", "Ensino Fundamental - Anos Finais", "7º ano"): [
            "Império Bizantino", "Império Carolíngio", "Sociedade feudal",
            "Cruzadas", "Renascimento comercial e urbano", "Peste Negra",
            "Formação das monarquias nacionais", "Mercantilismo", "Humanismo",
            "Contrarreforma", "Grandes Navegações", "Conquista da América",
            "Astecas", "Maias", "Incas", "Capitanias hereditárias",
            "Governo-geral", "Economia açucareira", "União Ibérica",
            "Invasões holandesas", "Quilombos"
        ],
        ("História", "Ensino Fundamental - Anos Finais", "8º ano"): [
            "Liberalismo", "Despotismo esclarecido", "Independência do Haiti",
            "Congresso de Viena", "Revoluções liberais", "Nacionalismo",
            "Unificação italiana", "Unificação alemã", "Imperialismo",
            "Partilha da África", "Vinda da família real", "Independência do Brasil",
            "Constituição de 1824", "Confederação do Equador",
            "Revoltas Regenciais", "Golpe da Maioridade", "Guerra do Paraguai",
            "Café e industrialização", "Movimento abolicionista",
            "Proclamação da República"
        ],
        ("História", "Ensino Fundamental - Anos Finais", "9º ano"): [
            "República da Espada", "República Oligárquica", "Coronelismo",
            "Revolta da Vacina", "Guerra de Canudos", "Guerra do Contestado",
            "Revolução de 1930", "Estado Novo", "Holocausto",
            "Descolonização da África", "Descolonização da Ásia",
            "Revolução Chinesa", "Revolução Cubana", "Populismo no Brasil",
            "Golpe de 1964", "Regime Militar", "Redemocratização",
            "Constituição de 1988", "Globalização", "Conflitos contemporâneos"
        ],

        ("Geografia", "Ensino Fundamental - Anos Finais", "6º ano"): [
            "Lugar, território e região", "Escalas cartográficas", "Fusos horários",
            "Movimentos da Terra", "Estações do ano", "Estrutura interna da Terra",
            "Tectonismo", "Vulcanismo", "Abalos sísmicos", "Tipos de rocha",
            "Formação do solo", "Agentes do relevo", "Elementos do clima",
            "Fatores climáticos", "Bacias hidrográficas", "Biomas brasileiros",
            "Recursos naturais", "Impactos ambientais", "Urbanização",
            "Atividades econômicas"
        ],
        ("Geografia", "Ensino Fundamental - Anos Finais", "7º ano"): [
            "Formação territorial do Brasil", "Regiões brasileiras",
            "Regionalização do IBGE", "População brasileira", "Migrações internas",
            "Urbanização brasileira", "Industrialização brasileira",
            "Agropecuária", "Extrativismo", "Fontes de energia",
            "Transportes e comunicação", "Região Norte", "Região Nordeste",
            "Região Centro-Oeste", "Região Sudeste", "Região Sul",
            "Amazônia", "Cerrado", "Semiárido", "Desigualdades regionais"
        ],
        ("Geografia", "Ensino Fundamental - Anos Finais", "8º ano"): [
            "População mundial", "Indicadores demográficos", "Fluxos migratórios",
            "Globalização", "Blocos econômicos", "América Anglo-Saxônica",
            "América Latina", "Estados Unidos", "Canadá", "México",
            "América Central", "América do Sul", "África",
            "Colonização da África", "Economia africana", "Conflitos africanos",
            "Ásia Ocidental", "Oriente Médio", "Geopolítica mundial",
            "Redes e fluxos"
        ],
        ("Geografia", "Ensino Fundamental - Anos Finais", "9º ano"): [
            "Europa", "União Europeia", "Rússia", "Ásia", "China", "Japão",
            "Índia", "Tigres Asiáticos", "Oceania", "Austrália",
            "Nova Zelândia", "Antártida", "Nova ordem mundial",
            "Organizações internacionais", "Conflitos contemporâneos",
            "Questões ambientais globais", "Mudanças climáticas",
            "Economia global", "Revolução técnico-científica",
            "Desigualdades socioespaciais"
        ],

        ("Ciências", "Ensino Fundamental - Anos Finais", "6º ano"): [
            "Propriedades da matéria", "Estados físicos da matéria",
            "Mudanças de estado físico", "Substâncias e misturas",
            "Métodos de separação", "Transformações físicas",
            "Transformações químicas", "Organização celular", "Microscopia",
            "Sistema ósseo", "Sistema muscular", "Órgãos dos sentidos",
            "Sistema nervoso central", "Drogas e sistema nervoso",
            "Forma e estrutura da Terra", "Fósseis", "Placas tectônicas",
            "Sistema Solar", "Movimentos da Terra", "Lua"
        ],
        ("Ciências", "Ensino Fundamental - Anos Finais", "7º ano"): [
            "Máquinas simples", "Calor e temperatura", "Propagação do calor",
            "Equilíbrio térmico", "Combustíveis", "Efeito estufa",
            "Camada de ozônio", "Vírus", "Bactérias", "Protozoários",
            "Fungos", "Vacinas", "Ecossistemas", "Cadeias alimentares",
            "Teias alimentares", "Biomas brasileiros", "Impactos ambientais",
            "Atmosfera", "Fenômenos naturais", "Placas tectônicas"
        ],
        ("Ciências", "Ensino Fundamental - Anos Finais", "8º ano"): [
            "Fontes e tipos de energia", "Energia elétrica", "Circuitos elétricos",
            "Consumo de energia", "Reprodução humana", "Puberdade",
            "Sistema reprodutor masculino", "Sistema reprodutor feminino",
            "Métodos contraceptivos", "Infecções sexualmente transmissíveis",
            "Hereditariedade", "Genética básica", "Clima e tempo",
            "Previsão do tempo", "Sistema Sol, Terra e Lua",
            "Fases da Lua", "Eclipses", "Estações do ano",
            "Movimentos da Terra", "Saúde e sexualidade"
        ],
        ("Ciências", "Ensino Fundamental - Anos Finais", "9º ano"): [
            "Átomos e moléculas", "Elementos químicos", "Tabela periódica",
            "Ligações químicas", "Reações químicas", "Leis ponderais",
            "Estrutura da matéria", "Ondas", "Som", "Luz",
            "Espectro eletromagnético", "Genética", "Leis de Mendel",
            "Evolução", "Seleção natural", "Diversidade biológica",
            "Sistema Solar", "Estrelas", "Galáxias", "Universo"
        ],

        ("Língua Inglesa", "Ensino Fundamental - Anos Finais", "6º ano"): [
            "Greetings", "Personal information", "Alphabet", "Numbers",
            "Colors", "School objects", "Family", "Verb to be",
            "Personal pronouns", "Possessive adjectives", "Simple present",
            "Daily routine", "Days and months", "Reading comprehension"
        ],
        ("Língua Inglesa", "Ensino Fundamental - Anos Finais", "7º ano"): [
            "Simple present", "Present continuous", "There is and there are",
            "Prepositions of place", "Countable and uncountable nouns",
            "Some and any", "Can and cannot", "Adverbs of frequency",
            "Sports and leisure", "Food and drinks", "Reading comprehension",
            "Text genres"
        ],
        ("Língua Inglesa", "Ensino Fundamental - Anos Finais", "8º ano"): [
            "Simple past", "Regular and irregular verbs", "Past continuous",
            "Comparatives", "Superlatives", "Future with will",
            "Going to", "Modal verbs", "Technology vocabulary",
            "Environment vocabulary", "Reading comprehension", "Text genres"
        ],
        ("Língua Inglesa", "Ensino Fundamental - Anos Finais", "9º ano"): [
            "Present perfect", "Past perfect", "Conditional sentences",
            "Passive voice", "Reported speech", "Relative pronouns",
            "Modal verbs", "Phrasal verbs", "Global issues",
            "Media and communication", "Reading comprehension", "Text genres"
        ],

        # =====================================================
        # ENSINO MÉDIO
        # =====================================================
        ("Língua Portuguesa", "Ensino Médio", "1ª série"): [
            "Gêneros discursivos", "Funções da linguagem", "Variação linguística",
            "Fonologia", "Morfologia", "Estrutura e formação de palavras",
            "Substantivos", "Adjetivos", "Pronomes", "Verbos",
            "Trovadorismo", "Humanismo", "Classicismo", "Quinhentismo",
            "Barroco", "Arcadismo", "Interpretação textual",
            "Coesão e coerência", "Introdução à redação do ENEM"
        ],
        ("Língua Portuguesa", "Ensino Médio", "2ª série"): [
            "Sintaxe do período simples", "Termos da oração",
            "Coordenação", "Subordinação", "Concordância verbal",
            "Concordância nominal", "Regência verbal", "Regência nominal",
            "Crase", "Romantismo", "Realismo", "Naturalismo",
            "Parnasianismo", "Simbolismo", "Pré-Modernismo",
            "Interpretação textual", "Argumentação", "Redação do ENEM"
        ],
        ("Língua Portuguesa", "Ensino Médio", "3ª série"): [
            "Modernismo - primeira fase", "Modernismo - segunda fase",
            "Modernismo - terceira fase", "Literatura contemporânea",
            "Orações subordinadas", "Pontuação", "Colocação pronominal",
            "Semântica", "Figuras de linguagem", "Intertextualidade",
            "Gêneros digitais", "Tese e repertório sociocultural",
            "Competências da redação do ENEM", "Coesão argumentativa",
            "Proposta de intervenção", "Revisão gramatical"
        ],

        ("Matemática", "Ensino Médio", "1ª série"): [
            "Conjuntos", "Conjuntos numéricos", "Intervalos reais",
            "Função afim", "Função quadrática", "Função modular",
            "Função exponencial", "Função logarítmica", "Progressão aritmética",
            "Progressão geométrica", "Razão e proporção", "Porcentagem",
            "Trigonometria no triângulo retângulo", "Geometria plana",
            "Semelhança de triângulos", "Estatística básica"
        ],
        ("Matemática", "Ensino Médio", "2ª série"): [
            "Trigonometria no ciclo", "Funções trigonométricas",
            "Matrizes", "Determinantes", "Sistemas lineares",
            "Análise combinatória", "Probabilidade", "Binômio de Newton",
            "Geometria espacial", "Prismas", "Pirâmides", "Cilindros",
            "Cones", "Esferas", "Áreas e volumes", "Estatística"
        ],
        ("Matemática", "Ensino Médio", "3ª série"): [
            "Geometria analítica", "Distância entre pontos", "Ponto médio",
            "Equação da reta", "Circunferência", "Números complexos",
            "Polinômios", "Equações polinomiais", "Matemática financeira",
            "Juros simples", "Juros compostos", "Probabilidade",
            "Estatística", "Funções", "Revisão ENEM"
        ],

        ("História", "Ensino Médio", "1ª série"): [
            "Pré-História", "Antiguidade Oriental", "Grécia Antiga",
            "Roma Antiga", "África Antiga", "Idade Média",
            "Feudalismo", "Império Bizantino", "Islamismo",
            "Cruzadas", "Renascimento comercial e urbano",
            "Renascimento cultural", "Reformas religiosas",
            "Formação dos Estados Nacionais", "Absolutismo",
            "Mercantilismo", "Expansão marítima"
        ],
        ("História", "Ensino Médio", "2ª série"): [
            "Colonização da América", "Brasil Colonial", "Escravidão",
            "Iluminismo", "Revolução Industrial", "Independência dos Estados Unidos",
            "Revolução Francesa", "Era Napoleônica", "Independências da América",
            "Primeiro Reinado", "Período Regencial", "Segundo Reinado",
            "Imperialismo", "Unificações italiana e alemã",
            "Abolição e Proclamação da República"
        ],
        ("História", "Ensino Médio", "3ª série"): [
            "Primeira República", "Primeira Guerra Mundial",
            "Revolução Russa", "Crise de 1929", "Nazifascismo",
            "Era Vargas", "Segunda Guerra Mundial", "Guerra Fria",
            "Descolonização afro-asiática", "Ditadura Militar no Brasil",
            "Redemocratização", "Nova República", "Globalização",
            "Conflitos contemporâneos", "História do Tocantins"
        ],

        ("Geografia", "Ensino Médio", "1ª série"): [
            "Cartografia", "Escalas", "Projeções cartográficas",
            "Geologia", "Relevo", "Solos", "Climatologia",
            "Hidrografia", "Biomas", "Questões ambientais",
            "Demografia", "Urbanização", "Espaço rural",
            "Fontes de energia", "Geografia do Brasil"
        ],
        ("Geografia", "Ensino Médio", "2ª série"): [
            "Industrialização", "Globalização", "Redes e fluxos",
            "Comércio internacional", "Blocos econômicos", "Geopolítica",
            "Estados Unidos", "Europa", "Rússia", "África",
            "América Latina", "Oriente Médio", "Ásia",
            "China", "Índia", "Japão"
        ],
        ("Geografia", "Ensino Médio", "3ª série"): [
            "Nova ordem mundial", "Conflitos internacionais",
            "Organizações internacionais", "Migrações internacionais",
            "Desigualdades socioespaciais", "Mudanças climáticas",
            "Questões energéticas", "Agronegócio", "Urbanização brasileira",
            "Industrialização brasileira", "Regiões brasileiras",
            "Amazônia", "Cerrado", "Geografia do Tocantins",
            "Revisão ENEM"
        ],

        ("Biologia", "Ensino Médio", "1ª série"): [
            "Origem da vida", "Bioquímica celular", "Citologia",
            "Membrana plasmática", "Organelas celulares", "Metabolismo energético",
            "Fotossíntese", "Respiração celular", "Divisão celular",
            "Histologia animal", "Vírus", "Bactérias", "Protozoários",
            "Fungos", "Botânica"
        ],
        ("Biologia", "Ensino Médio", "2ª série"): [
            "Zoologia", "Fisiologia humana", "Sistema digestório",
            "Sistema respiratório", "Sistema circulatório", "Sistema excretor",
            "Sistema nervoso", "Sistema endócrino", "Sistema reprodutor",
            "Embriologia", "Imunologia", "Doenças infecciosas",
            "Parasitologia", "Saúde pública"
        ],
        ("Biologia", "Ensino Médio", "3ª série"): [
            "Genética", "Leis de Mendel", "Herança ligada ao sexo",
            "Biotecnologia", "Evolução", "Seleção natural",
            "Especiação", "Ecologia", "Cadeias e teias alimentares",
            "Ciclos biogeoquímicos", "Relações ecológicas",
            "Sucessão ecológica", "Biomas brasileiros",
            "Impactos ambientais", "Revisão ENEM"
        ],

        ("Física", "Ensino Médio", "1ª série"): [
            "Grandezas físicas", "Vetores", "Cinemática",
            "Movimento uniforme", "Movimento uniformemente variado",
            "Queda livre", "Lançamentos", "Dinâmica",
            "Leis de Newton", "Força de atrito", "Trabalho",
            "Energia", "Potência", "Impulso", "Quantidade de movimento"
        ],
        ("Física", "Ensino Médio", "2ª série"): [
            "Hidrostática", "Pressão", "Princípio de Pascal",
            "Princípio de Arquimedes", "Termologia", "Calorimetria",
            "Dilatação térmica", "Termodinâmica", "Ondulatória",
            "Som", "Óptica geométrica", "Espelhos", "Lentes",
            "Refração", "Instrumentos ópticos"
        ],
        ("Física", "Ensino Médio", "3ª série"): [
            "Eletrostática", "Campo elétrico", "Potencial elétrico",
            "Corrente elétrica", "Resistores", "Circuitos elétricos",
            "Potência elétrica", "Magnetismo", "Eletromagnetismo",
            "Indução eletromagnética", "Física moderna", "Relatividade",
            "Física quântica", "Radioatividade", "Revisão ENEM"
        ],

        ("Química", "Ensino Médio", "1ª série"): [
            "Matéria e energia", "Estados físicos", "Misturas",
            "Separação de misturas", "Modelos atômicos", "Estrutura atômica",
            "Tabela periódica", "Propriedades periódicas", "Ligações químicas",
            "Geometria molecular", "Forças intermoleculares",
            "Funções inorgânicas", "Reações químicas", "Balanceamento"
        ],
        ("Química", "Ensino Médio", "2ª série"): [
            "Mol", "Massa molar", "Estequiometria", "Soluções",
            "Concentração", "Diluição", "Mistura de soluções",
            "Termoquímica", "Cinética química", "Equilíbrio químico",
            "Equilíbrio iônico", "pH e pOH", "Hidrólise salina",
            "Eletroquímica"
        ],
        ("Química", "Ensino Médio", "3ª série"): [
            "Química orgânica", "Cadeias carbônicas", "Funções orgânicas",
            "Isomeria", "Reações orgânicas", "Polímeros",
            "Petróleo", "Bioquímica", "Radioatividade",
            "Química ambiental", "Química dos alimentos",
            "Combustíveis", "Pilhas e eletrólise", "Revisão ENEM"
        ],

        ("Filosofia", "Ensino Médio", "1ª série"): [
            "Origem da Filosofia", "Mito e razão", "Pré-socráticos",
            "Sócrates", "Platão", "Aristóteles", "Filosofia helenística",
            "Ética antiga", "Política na Antiguidade", "Conhecimento e verdade"
        ],
        ("Filosofia", "Ensino Médio", "2ª série"): [
            "Filosofia medieval", "Patrística", "Escolástica",
            "Racionalismo", "Empirismo", "Iluminismo", "Kant",
            "Filosofia política moderna", "Contrato social", "Ética moderna"
        ],
        ("Filosofia", "Ensino Médio", "3ª série"): [
            "Idealismo", "Marxismo", "Nietzsche", "Fenomenologia",
            "Existencialismo", "Escola de Frankfurt", "Filosofia da ciência",
            "Bioética", "Filosofia contemporânea", "Cidadania e democracia"
        ],

        ("Sociologia", "Ensino Médio", "1ª série"): [
            "Introdução à Sociologia", "Socialização", "Cultura",
            "Identidade", "Etnocentrismo", "Relativismo cultural",
            "Instituições sociais", "Grupos sociais", "Estratificação social",
            "Desigualdade social"
        ],
        ("Sociologia", "Ensino Médio", "2ª série"): [
            "Trabalho e sociedade", "Capitalismo", "Classes sociais",
            "Karl Marx", "Émile Durkheim", "Max Weber",
            "Poder e política", "Estado", "Democracia",
            "Movimentos sociais", "Cidadania"
        ],
        ("Sociologia", "Ensino Médio", "3ª série"): [
            "Globalização", "Indústria cultural", "Mídia e sociedade",
            "Consumo", "Gênero e sociedade", "Relações étnico-raciais",
            "Violência", "Urbanização", "Meio ambiente",
            "Direitos humanos", "Juventude e participação social"
        ],

        ("Língua Inglesa", "Ensino Médio", "1ª série"): [
            "Reading strategies", "Cognates and false cognates",
            "Simple present", "Present continuous", "Simple past",
            "Pronouns", "Articles", "Prepositions", "Text genres",
            "Vocabulary in context", "Interpretação de textos"
        ],
        ("Língua Inglesa", "Ensino Médio", "2ª série"): [
            "Present perfect", "Past continuous", "Future forms",
            "Modal verbs", "Comparatives and superlatives",
            "Passive voice", "Relative pronouns", "Conditionals",
            "Text genres", "Interpretação de textos"
        ],
        ("Língua Inglesa", "Ensino Médio", "3ª série"): [
            "Reported speech", "Passive voice", "Conditionals",
            "Phrasal verbs", "Linking words", "Argumentative texts",
            "Scientific texts", "Media texts", "ENEM reading strategies",
            "Interpretação de textos"
        ]
    }

    for (disciplina_assunto, etapa_assunto, serie_assunto), nomes_assuntos in assuntos_ampliados.items():
        for nome_assunto in nomes_assuntos:
            cursor.execute("""
                INSERT INTO assuntos (
                    escola_id,
                    disciplina,
                    etapa_ensino,
                    ano_serie,
                    nome,
                    ativo
                )
                SELECT NULL, ?, ?, ?, ?, 1
                WHERE NOT EXISTS (
                    SELECT 1
                    FROM assuntos
                    WHERE escola_id IS NULL
                      AND disciplina = ?
                      AND etapa_ensino = ?
                      AND ano_serie = ?
                      AND nome = ?
                )
            """, (
                disciplina_assunto, etapa_assunto, serie_assunto, nome_assunto,
                disciplina_assunto, etapa_assunto, serie_assunto, nome_assunto
            ))

    # Catálogo completo do Ensino Fundamental - Anos Iniciais.
    # Abrange do 1º ao 5º ano e os componentes curriculares mais comuns.
    assuntos_anos_iniciais = {
        # =====================================================
        # LÍNGUA PORTUGUESA
        # =====================================================
        ("Língua Portuguesa", "Ensino Fundamental - Anos Iniciais", "1º ano"): [
            "Alfabeto", "Ordem alfabética", "Vogais e consoantes",
            "Consciência fonológica", "Sílabas", "Formação de palavras",
            "Palavras e imagens", "Leitura de palavras", "Leitura de frases",
            "Escrita do nome", "Letra maiúscula e minúscula",
            "Segmentação de palavras", "Rimas", "Parlendas", "Cantigas",
            "Trava-línguas", "Adivinhas", "Listas", "Bilhetes",
            "Contos infantis", "Interpretação de texto", "Produção de frases"
        ],
        ("Língua Portuguesa", "Ensino Fundamental - Anos Iniciais", "2º ano"): [
            "Ordem alfabética", "Sílabas simples e complexas",
            "Separação silábica", "Classificação quanto ao número de sílabas",
            "Substantivos", "Nomes próprios e comuns", "Gênero do substantivo",
            "Número do substantivo", "Adjetivos", "Artigos",
            "Sinônimos e antônimos", "Ortografia", "Pontuação",
            "Frase e parágrafo", "Leitura e interpretação",
            "Fábulas", "Contos", "Poemas", "Bilhetes", "Convites",
            "Receitas", "Histórias em quadrinhos", "Produção textual"
        ],
        ("Língua Portuguesa", "Ensino Fundamental - Anos Iniciais", "3º ano"): [
            "Substantivos próprios e comuns", "Substantivos coletivos",
            "Adjetivos", "Artigos", "Pronomes pessoais", "Verbos",
            "Tempos verbais", "Sílaba tônica", "Acentuação gráfica",
            "Uso de M e N", "Uso de R e RR", "Uso de S, SS, C e Ç",
            "Pontuação", "Frase, oração e parágrafo",
            "Leitura e interpretação", "Contos", "Fábulas", "Lendas",
            "Notícias", "Cartas", "Receitas", "Poemas",
            "Histórias em quadrinhos", "Produção textual"
        ],
        ("Língua Portuguesa", "Ensino Fundamental - Anos Iniciais", "4º ano"): [
            "Substantivos", "Adjetivos", "Artigos", "Pronomes",
            "Verbos e tempos verbais", "Advérbios", "Preposições",
            "Concordância nominal", "Concordância verbal",
            "Sílaba tônica", "Acentuação", "Encontro vocálico",
            "Encontro consonantal", "Dígrafos", "Ortografia",
            "Pontuação", "Discurso direto e indireto",
            "Leitura e interpretação", "Conto", "Crônica", "Notícia",
            "Reportagem", "Poema", "Carta", "Resumo", "Produção textual"
        ],
        ("Língua Portuguesa", "Ensino Fundamental - Anos Iniciais", "5º ano"): [
            "Classes gramaticais", "Substantivos e classificações",
            "Adjetivos e locuções adjetivas", "Pronomes", "Numerais",
            "Verbos e conjugações", "Tempos e modos verbais",
            "Advérbios", "Preposições", "Conjunções",
            "Sujeito e predicado", "Tipos de sujeito", "Concordância verbal",
            "Concordância nominal", "Acentuação gráfica", "Crase",
            "Pontuação", "Figuras de linguagem", "Denotação e conotação",
            "Leitura e interpretação", "Gêneros textuais",
            "Artigo de opinião", "Notícia", "Reportagem", "Crônica",
            "Poema", "Biografia", "Produção textual"
        ],

        # =====================================================
        # MATEMÁTICA
        # =====================================================
        ("Matemática", "Ensino Fundamental - Anos Iniciais", "1º ano"): [
            "Números até 10", "Números até 20", "Números até 100",
            "Contagem", "Sequência numérica", "Antecessor e sucessor",
            "Comparação de quantidades", "Maior, menor e igual",
            "Adição", "Subtração", "Situações-problema",
            "Dezenas e unidades", "Formas geométricas planas",
            "Localização e posição", "Noções de comprimento",
            "Noções de massa", "Noções de capacidade", "Calendário",
            "Dias da semana", "Horas", "Sistema monetário",
            "Tabelas e gráficos simples"
        ],
        ("Matemática", "Ensino Fundamental - Anos Iniciais", "2º ano"): [
            "Números até 100", "Números até 1000", "Valor posicional",
            "Centenas, dezenas e unidades", "Composição e decomposição",
            "Adição com e sem reagrupamento", "Subtração com e sem reagrupamento",
            "Multiplicação como adição de parcelas iguais",
            "Divisão como repartição", "Dobro", "Triplo", "Metade",
            "Situações-problema", "Sequências", "Figuras geométricas",
            "Sólidos geométricos", "Comprimento", "Massa", "Capacidade",
            "Tempo", "Calendário", "Sistema monetário", "Tabelas e gráficos"
        ],
        ("Matemática", "Ensino Fundamental - Anos Iniciais", "3º ano"): [
            "Números naturais", "Sistema de numeração decimal",
            "Valor posicional", "Composição e decomposição",
            "Adição", "Subtração", "Multiplicação", "Divisão",
            "Tabuada", "Problemas com as quatro operações",
            "Frações simples", "Metade, terça parte e quarta parte",
            "Sequências numéricas", "Par e ímpar", "Figuras planas",
            "Sólidos geométricos", "Perímetro", "Comprimento",
            "Massa", "Capacidade", "Tempo", "Sistema monetário",
            "Leitura de tabelas", "Leitura de gráficos"
        ],
        ("Matemática", "Ensino Fundamental - Anos Iniciais", "4º ano"): [
            "Números naturais", "Ordens e classes", "Valor posicional",
            "Composição e decomposição", "Adição", "Subtração",
            "Multiplicação", "Divisão", "Expressões numéricas",
            "Múltiplos e divisores", "Frações", "Frações equivalentes",
            "Números decimais", "Problemas com as quatro operações",
            "Ângulos", "Retas", "Polígonos", "Perímetro", "Área",
            "Medidas de comprimento", "Massa", "Capacidade", "Tempo",
            "Sistema monetário", "Tabelas e gráficos"
        ],
        ("Matemática", "Ensino Fundamental - Anos Iniciais", "5º ano"): [
            "Sistema de numeração decimal", "Ordens e classes",
            "Números naturais", "Números decimais", "Adição e subtração",
            "Multiplicação e divisão", "Expressões numéricas",
            "Múltiplos e divisores", "Números primos", "MMC e MDC",
            "Frações", "Operações com frações", "Porcentagem",
            "Razão e proporção", "Regra de três simples",
            "Plano cartesiano", "Ângulos", "Polígonos", "Triângulos",
            "Quadriláteros", "Perímetro", "Área", "Volume",
            "Grandezas e medidas", "Média aritmética",
            "Tabelas, gráficos e probabilidade"
        ],

        # =====================================================
        # CIÊNCIAS
        # =====================================================
        ("Ciências", "Ensino Fundamental - Anos Iniciais", "1º ano"): [
            "Corpo humano", "Partes do corpo", "Órgãos dos sentidos",
            "Hábitos de higiene", "Alimentação saudável", "Saúde",
            "Seres vivos e não vivos", "Animais", "Plantas",
            "Ambientes", "Dia e noite", "Sol", "Lua",
            "Água", "Ar", "Solo", "Materiais do cotidiano",
            "Cuidados com o ambiente", "Lixo e reciclagem"
        ],
        ("Ciências", "Ensino Fundamental - Anos Iniciais", "2º ano"): [
            "Seres vivos", "Ciclo de vida", "Animais vertebrados e invertebrados",
            "Habitat dos animais", "Plantas e suas partes",
            "Germinação", "Necessidades das plantas", "Corpo humano",
            "Higiene e saúde", "Alimentação", "Água", "Ar", "Solo",
            "Luz e sombra", "Dia e noite", "Materiais e objetos",
            "Reutilização e reciclagem", "Preservação ambiental"
        ],
        ("Ciências", "Ensino Fundamental - Anos Iniciais", "3º ano"): [
            "Características dos animais", "Classificação dos animais",
            "Alimentação dos animais", "Reprodução dos animais",
            "Plantas", "Fotossíntese", "Cadeia alimentar",
            "Corpo humano", "Sistema locomotor", "Órgãos dos sentidos",
            "Saúde e prevenção", "Estados físicos da água",
            "Ciclo da água", "Ar e atmosfera", "Solo",
            "Som", "Luz", "Terra", "Lua", "Sistema Solar",
            "Impactos ambientais"
        ],
        ("Ciências", "Ensino Fundamental - Anos Iniciais", "4º ano"): [
            "Célula", "Seres unicelulares e pluricelulares",
            "Microrganismos", "Cadeias alimentares", "Relações ecológicas",
            "Ecossistemas", "Corpo humano", "Sistema digestório",
            "Sistema respiratório", "Sistema circulatório",
            "Hábitos saudáveis", "Misturas", "Transformações da matéria",
            "Água e saneamento", "Solo", "Rochas e minerais",
            "Pontos cardeais", "Movimentos da Terra", "Fases da Lua",
            "Preservação ambiental"
        ],
        ("Ciências", "Ensino Fundamental - Anos Iniciais", "5º ano"): [
            "Matéria e energia", "Propriedades da matéria",
            "Misturas e separação", "Transformações físicas e químicas",
            "Ciclo da água", "Uso sustentável da água",
            "Nutrição", "Sistema digestório", "Sistema respiratório",
            "Sistema circulatório", "Sistema excretor",
            "Alimentação saudável", "Reprodução humana",
            "Puberdade", "Sistema Solar", "Movimentos da Terra",
            "Fases da Lua", "Constelações", "Fontes de energia",
            "Consumo consciente", "Reciclagem e sustentabilidade"
        ],

        # =====================================================
        # HISTÓRIA
        # =====================================================
        ("História", "Ensino Fundamental - Anos Iniciais", "1º ano"): [
            "Identidade", "História pessoal", "Nome e sobrenome",
            "Família", "Diferentes tipos de família", "Linha do tempo",
            "Memórias", "Brinquedos e brincadeiras", "Escola",
            "Regras de convivência", "Casa e moradia", "Bairro",
            "Datas comemorativas", "Direitos das crianças"
        ],
        ("História", "Ensino Fundamental - Anos Iniciais", "2º ano"): [
            "História pessoal e familiar", "Documentos pessoais",
            "Fontes históricas", "Memórias familiares", "Comunidade",
            "Bairro", "Profissões", "Trabalho no passado e no presente",
            "Meios de transporte", "Meios de comunicação",
            "Brincadeiras antigas e atuais", "Mudanças e permanências",
            "Patrimônio cultural", "Direitos e deveres"
        ],
        ("História", "Ensino Fundamental - Anos Iniciais", "3º ano"): [
            "Município", "História da cidade", "Formação da comunidade",
            "Grupos sociais", "Povos indígenas", "Povos africanos",
            "Imigração", "Trabalho e profissões", "Espaços públicos",
            "Patrimônio histórico", "Patrimônio cultural",
            "Festas e tradições", "Mudanças na cidade",
            "Poder público municipal", "Cidadania"
        ],
        ("História", "Ensino Fundamental - Anos Iniciais", "4º ano"): [
            "Nomadismo e sedentarização", "Primeiros grupos humanos",
            "Povos indígenas do Brasil", "Grandes navegações",
            "Chegada dos portugueses", "Colonização do Brasil",
            "Escravidão indígena", "Escravidão africana",
            "Economia açucareira", "Bandeirantes", "Mineração",
            "Formação do território brasileiro", "Migrações",
            "Patrimônio histórico", "Diversidade cultural"
        ],
        ("História", "Ensino Fundamental - Anos Iniciais", "5º ano"): [
            "Povos e culturas antigas", "Formação das primeiras cidades",
            "Cidadania na Antiguidade", "Grécia Antiga", "Roma Antiga",
            "Democracia", "Direitos humanos", "Constituição",
            "Formação do povo brasileiro", "Povos indígenas",
            "Povos africanos", "Imigração no Brasil",
            "Abolição da escravidão", "Proclamação da República",
            "Patrimônio material e imaterial", "Diversidade religiosa",
            "Cidadania e participação social"
        ],

        # =====================================================
        # GEOGRAFIA
        # =====================================================
        ("Geografia", "Ensino Fundamental - Anos Iniciais", "1º ano"): [
            "Lugar de vivência", "Casa", "Escola", "Bairro",
            "Paisagem", "Elementos naturais e culturais",
            "Localização", "Direita e esquerda", "Frente e atrás",
            "Perto e longe", "Trajetos", "Meios de transporte",
            "Meios de comunicação", "Campo e cidade",
            "Tempo atmosférico", "Cuidados com o ambiente"
        ],
        ("Geografia", "Ensino Fundamental - Anos Iniciais", "2º ano"): [
            "Lugar e paisagem", "Bairro", "Zona urbana e rural",
            "Tipos de moradia", "Trabalho no campo e na cidade",
            "Meios de transporte", "Meios de comunicação",
            "Representação dos lugares", "Mapas simples",
            "Pontos de referência", "Orientação", "Recursos naturais",
            "Água", "Solo", "Vegetação", "Impactos ambientais"
        ],
        ("Geografia", "Ensino Fundamental - Anos Iniciais", "3º ano"): [
            "Município", "Cidade e campo", "Paisagens urbanas e rurais",
            "Atividades econômicas", "Agricultura", "Pecuária",
            "Indústria", "Comércio e serviços", "Trabalho",
            "Representação cartográfica", "Mapas", "Legendas",
            "Pontos cardeais", "Relevo", "Hidrografia",
            "Vegetação", "Clima", "Problemas ambientais"
        ],
        ("Geografia", "Ensino Fundamental - Anos Iniciais", "4º ano"): [
            "Território brasileiro", "Divisão política do Brasil",
            "Estados e capitais", "Regiões brasileiras",
            "Município e estado", "População brasileira",
            "Migrações", "Diversidade cultural", "Campo e cidade",
            "Urbanização", "Atividades econômicas", "Agropecuária",
            "Indústria", "Comércio", "Relevo brasileiro",
            "Clima", "Hidrografia", "Biomas brasileiros",
            "Questões ambientais"
        ],
        ("Geografia", "Ensino Fundamental - Anos Iniciais", "5º ano"): [
            "Território e fronteiras", "Regiões brasileiras",
            "População brasileira", "Distribuição da população",
            "Migrações internas", "Urbanização", "Rede urbana",
            "Industrialização", "Agropecuária", "Fontes de energia",
            "Transportes e comunicação", "Cartografia",
            "Escala", "Coordenadas geográficas", "Relevo",
            "Clima", "Hidrografia", "Biomas",
            "Desigualdades regionais", "Sustentabilidade"
        ],

        # =====================================================
        # ARTE
        # =====================================================
        ("Arte", "Ensino Fundamental - Anos Iniciais", "1º ano"): [
            "Cores primárias", "Formas", "Linhas", "Texturas",
            "Desenho", "Pintura", "Colagem", "Modelagem",
            "Música e sons", "Expressão corporal", "Teatro",
            "Brincadeiras cantadas", "Arte e identidade"
        ],
        ("Arte", "Ensino Fundamental - Anos Iniciais", "2º ano"): [
            "Cores primárias e secundárias", "Mistura de cores",
            "Formas geométricas na arte", "Texturas", "Desenho",
            "Pintura", "Colagem", "Escultura", "Música",
            "Ritmo", "Dança", "Teatro", "Arte popular"
        ],
        ("Arte", "Ensino Fundamental - Anos Iniciais", "3º ano"): [
            "Elementos visuais", "Linha, forma e cor", "Luz e sombra",
            "Desenho de observação", "Pintura", "Gravura",
            "Escultura", "Fotografia", "Música", "Ritmo e melodia",
            "Dança", "Teatro", "Cultura popular", "Arte indígena"
        ],
        ("Arte", "Ensino Fundamental - Anos Iniciais", "4º ano"): [
            "Artes visuais", "Composição visual", "Perspectiva",
            "Pintura", "Escultura", "Gravura", "Fotografia",
            "Música brasileira", "Instrumentos musicais", "Dança",
            "Teatro", "Arte indígena", "Arte africana",
            "Patrimônio cultural"
        ],
        ("Arte", "Ensino Fundamental - Anos Iniciais", "5º ano"): [
            "Elementos da linguagem visual", "História da arte",
            "Arte brasileira", "Arte indígena", "Arte afro-brasileira",
            "Arte popular", "Pintura", "Escultura", "Fotografia",
            "Cinema", "Música", "Dança", "Teatro",
            "Cultura e patrimônio", "Produção artística"
        ],

        # =====================================================
        # EDUCAÇÃO FÍSICA
        # =====================================================
        ("Educação Física", "Ensino Fundamental - Anos Iniciais", "1º ano"): [
            "Brincadeiras e jogos", "Coordenação motora", "Equilíbrio",
            "Lateralidade", "Esquema corporal", "Movimentos básicos",
            "Jogos cooperativos", "Ritmo", "Expressão corporal"
        ],
        ("Educação Física", "Ensino Fundamental - Anos Iniciais", "2º ano"): [
            "Brincadeiras populares", "Jogos de regras simples",
            "Coordenação motora", "Agilidade", "Equilíbrio",
            "Lateralidade", "Ginástica geral", "Dança",
            "Jogos cooperativos", "Hábitos saudáveis"
        ],
        ("Educação Física", "Ensino Fundamental - Anos Iniciais", "3º ano"): [
            "Jogos populares", "Jogos cooperativos", "Esportes de marca",
            "Esportes de precisão", "Ginástica", "Dança",
            "Lutas de contexto comunitário", "Coordenação motora",
            "Capacidades físicas", "Saúde e movimento"
        ],
        ("Educação Física", "Ensino Fundamental - Anos Iniciais", "4º ano"): [
            "Jogos e brincadeiras tradicionais", "Esportes de invasão",
            "Esportes de rede e parede", "Esportes de campo e taco",
            "Ginástica geral", "Danças populares", "Lutas",
            "Capacidades físicas", "Regras esportivas",
            "Saúde e qualidade de vida"
        ],
        ("Educação Física", "Ensino Fundamental - Anos Iniciais", "5º ano"): [
            "Jogos eletrônicos", "Esportes de invasão", "Esportes de rede",
            "Esportes de marca", "Esportes de precisão",
            "Ginástica geral", "Danças do Brasil", "Lutas do Brasil",
            "Capacidades físicas", "Fair play", "Regras e arbitragem",
            "Saúde, exercício e qualidade de vida"
        ],

        # =====================================================
        # ENSINO RELIGIOSO
        # =====================================================
        ("Ensino Religioso", "Ensino Fundamental - Anos Iniciais", "1º ano"): [
            "Identidade", "Respeito", "Convivência", "Família",
            "Amizade", "Diferenças", "Valores", "Solidariedade",
            "Símbolos religiosos", "Festas e celebrações"
        ],
        ("Ensino Religioso", "Ensino Fundamental - Anos Iniciais", "2º ano"): [
            "Identidade e alteridade", "Respeito às diferenças",
            "Família e comunidade", "Valores humanos", "Solidariedade",
            "Símbolos religiosos", "Espaços sagrados",
            "Festas religiosas", "Tradições culturais"
        ],
        ("Ensino Religioso", "Ensino Fundamental - Anos Iniciais", "3º ano"): [
            "Diversidade religiosa", "Tradições religiosas",
            "Espaços sagrados", "Práticas religiosas", "Orações",
            "Festas e celebrações", "Valores éticos",
            "Respeito e tolerância", "Cultura de paz"
        ],
        ("Ensino Religioso", "Ensino Fundamental - Anos Iniciais", "4º ano"): [
            "Religiões do mundo", "Textos sagrados", "Símbolos religiosos",
            "Ritos e celebrações", "Lideranças religiosas",
            "Tradições indígenas", "Tradições afro-brasileiras",
            "Ética", "Direitos humanos", "Cultura de paz"
        ],
        ("Ensino Religioso", "Ensino Fundamental - Anos Iniciais", "5º ano"): [
            "Diversidade de crenças", "Religiões monoteístas",
            "Religiões orientais", "Matrizes indígenas",
            "Matrizes africanas", "Textos sagrados",
            "Mitos e narrativas religiosas", "Ética e cidadania",
            "Liberdade religiosa", "Diálogo inter-religioso"
        ],

        # =====================================================
        # LÍNGUA INGLESA
        # =====================================================
        ("Língua Inglesa", "Ensino Fundamental - Anos Iniciais", "1º ano"): [
            "Greetings", "Colors", "Numbers 1 to 10", "School objects",
            "Family", "Body parts", "Animals", "Songs and games"
        ],
        ("Língua Inglesa", "Ensino Fundamental - Anos Iniciais", "2º ano"): [
            "Greetings", "Alphabet", "Numbers", "Colors",
            "Family members", "School objects", "Toys",
            "Animals", "Food", "Days of the week"
        ],
        ("Língua Inglesa", "Ensino Fundamental - Anos Iniciais", "3º ano"): [
            "Personal information", "Numbers", "Colors",
            "Family", "Parts of the house", "School subjects",
            "Food and drinks", "Animals", "Clothes",
            "Days and months", "Verb to be"
        ],
        ("Língua Inglesa", "Ensino Fundamental - Anos Iniciais", "4º ano"): [
            "Personal information", "Daily routine", "Simple present",
            "Family", "House", "School", "Food",
            "Clothes", "Weather", "Places in town",
            "Prepositions of place", "Can and cannot"
        ],
        ("Língua Inglesa", "Ensino Fundamental - Anos Iniciais", "5º ano"): [
            "Simple present", "Daily routine", "Adverbs of frequency",
            "Personal pronouns", "Possessive adjectives",
            "There is and there are", "Places in town",
            "Food and drinks", "Sports", "Weather",
            "Dates and time", "Reading comprehension"
        ]
    }

    for (disciplina_assunto, etapa_assunto, serie_assunto), nomes_assuntos in assuntos_anos_iniciais.items():
        for nome_assunto in nomes_assuntos:
            cursor.execute("""
                INSERT INTO assuntos (
                    escola_id,
                    disciplina,
                    etapa_ensino,
                    ano_serie,
                    nome,
                    ativo
                )
                SELECT NULL, ?, ?, ?, ?, 1
                WHERE NOT EXISTS (
                    SELECT 1
                    FROM assuntos
                    WHERE escola_id IS NULL
                      AND disciplina = ?
                      AND etapa_ensino = ?
                      AND ano_serie = ?
                      AND nome = ?
                )
            """, (
                disciplina_assunto, etapa_assunto, serie_assunto, nome_assunto,
                disciplina_assunto, etapa_assunto, serie_assunto, nome_assunto
            ))

    garantir_coluna("escolas", "tipo_instituicao", "TEXT")
    garantir_coluna("usuarios", "escola_id", "INTEGER")
    garantir_coluna("usuarios", "cpf", "TEXT")
    garantir_coluna("usuarios", "foto", "TEXT")
    garantir_coluna("turmas", "escola_id", "INTEGER")
    garantir_coluna("turmas", "etapa", "TEXT")
    garantir_coluna("turmas", "ano_letivo", "TEXT")
    garantir_coluna("turmas", "ano_letivo_id", "INTEGER")
    garantir_coluna("alunos", "escola_id", "INTEGER")
    garantir_coluna("alunos", "ano_letivo_id", "INTEGER")
    garantir_coluna("professores", "escola_id", "INTEGER")
    garantir_coluna("questoes", "escola_id", "INTEGER")
    garantir_coluna("questoes", "assunto", "TEXT")
    garantir_coluna("questoes", "assunto_temporario", "INTEGER DEFAULT 0")
    garantir_coluna("questoes", "tipo_questao", "TEXT DEFAULT 'multipla_escolha'")
    garantir_coluna("questoes", "enunciado_html", "TEXT")
    garantir_coluna("questoes", "alternativas_json", "TEXT")
    garantir_coluna("questoes", "respostas_corretas", "TEXT")
    garantir_coluna("questoes", "resposta_esperada", "TEXT")
    garantir_coluna("questoes", "criterios_correcao", "TEXT")
    garantir_coluna("questoes", "observacoes", "TEXT")
    garantir_coluna("questoes", "criado_por", "INTEGER")
    garantir_coluna("questoes", "criado_em", "TEXT")
    garantir_coluna("questoes", "atualizado_em", "TEXT")
    garantir_coluna("questoes", "etapa_ensino", "TEXT")
    garantir_coluna("questoes", "ano_serie", "TEXT")
    garantir_coluna("questoes", "subassunto", "TEXT")
    garantir_coluna("questoes", "unidade_tematica", "TEXT")
    garantir_coluna("questoes", "objeto_conhecimento", "TEXT")
    garantir_coluna("questoes", "habilidade_bncc", "TEXT")
    garantir_coluna("questoes", "matriz_referencia", "TEXT")
    garantir_coluna("questoes", "descritor_saeb", "TEXT")
    garantir_coluna("questoes", "sistema_matriz", "TEXT")
    garantir_coluna("questoes", "taxonomia_bloom", "TEXT")
    garantir_coluna("questoes", "fonte", "TEXT")
    garantir_coluna("questoes", "ano_fonte", "INTEGER")
    garantir_coluna("questoes", "tags", "TEXT")
    garantir_coluna("questoes", "tempo_estimado", "INTEGER")
    garantir_coluna("questoes", "linhas_resposta", "INTEGER DEFAULT 5")
    garantir_coluna("provas", "professor_id", "INTEGER")
    garantir_coluna("provas", "data_geracao", "TEXT")
    garantir_coluna("provas", "data_aplicacao", "TEXT")
    garantir_coluna("provas", "escola_id", "INTEGER")
    garantir_coluna("provas", "ano_letivo_id", "INTEGER")
    garantir_coluna("provas", "status", "TEXT DEFAULT 'rascunho'")
    garantir_coluna("provas", "atualizado_em", "TEXT")
    garantir_coluna("provas", "media_ativa", "INTEGER DEFAULT 0")
    garantir_coluna("provas", "media_aprovacao", "REAL")
    garantir_coluna("provas", "peso_total", "REAL DEFAULT 10")
    garantir_coluna("provas", "tipo_peso", "TEXT DEFAULT 'automatico'")
    garantir_coluna("provas", "tem_nota", "INTEGER NOT NULL DEFAULT 0")
    garantir_coluna("prova_questoes", "peso", "REAL DEFAULT 0")
    garantir_coluna("prova_questoes", "ordem", "INTEGER DEFAULT 0")
    garantir_coluna("prova_questoes", "anulada", "INTEGER NOT NULL DEFAULT 0")
    garantir_coluna("instituicao", "logo", "TEXT")
    garantir_coluna("permissoes", "pode_acessar", "INTEGER DEFAULT 0")

    garantir_coluna("componentes_curriculares", "escola_id", "INTEGER")
    garantir_coluna("componentes_curriculares", "etapa_ensino", "TEXT")
    garantir_coluna("componentes_curriculares", "nome", "TEXT")
    garantir_coluna(
        "componentes_curriculares",
        "tipo",
        "TEXT DEFAULT 'padrao'"
    )
    garantir_coluna(
        "componentes_curriculares",
        "ativo",
        "INTEGER DEFAULT 1"
    )

    cursor.execute("PRAGMA table_info(turmas)")
    colunas_turmas = {linha[1] for linha in cursor.fetchall()}

    if "etapa_ensino" in colunas_turmas and "etapa" in colunas_turmas:
        cursor.execute("""
            UPDATE turmas
            SET etapa = etapa_ensino
            WHERE (etapa IS NULL OR TRIM(etapa) = '')
              AND etapa_ensino IS NOT NULL
        """)

    cursor.execute("""
        UPDATE componentes_curriculares
        SET tipo = 'padrao'
        WHERE tipo IS NULL OR TRIM(tipo) = ''
    """)

    cursor.execute("""
        UPDATE componentes_curriculares
        SET ativo = 1
        WHERE ativo IS NULL
    """)

    cursor.execute("""
        CREATE INDEX IF NOT EXISTS idx_componentes_escola_etapa
        ON componentes_curriculares (escola_id, etapa_ensino)
    """)

    cursor.execute("""
        CREATE INDEX IF NOT EXISTS idx_usuarios_escola
        ON usuarios (escola_id)
    """)

    cursor.execute("""
        CREATE INDEX IF NOT EXISTS idx_turmas_escola
        ON turmas (escola_id)
    """)

    cursor.execute("""
        CREATE INDEX IF NOT EXISTS idx_aluno_matriculas_aluno
        ON aluno_matriculas (aluno_id)
    """)

    cursor.execute("""
        CREATE INDEX IF NOT EXISTS idx_aluno_matriculas_ano
        ON aluno_matriculas (ano_letivo_id)
    """)

    cursor.execute("""
        CREATE INDEX IF NOT EXISTS idx_aluno_matriculas_turma
        ON aluno_matriculas (turma_id)
    """)

    cargos = [
        "Administrador Geral",
        "Administrador da Instituição",
        "Coordenador",
        "Professor",
        "Secretaria"
    ]

    for nome_cargo in cargos:
        cursor.execute(
            "INSERT OR IGNORE INTO cargos (nome) VALUES (?)",
            (nome_cargo,)
        )

    cursor.execute("""
        SELECT id
        FROM cargos
        WHERE nome = ?
        LIMIT 1
    """, ("Administrador Geral",))

    cargo_admin = cursor.fetchone()
    cargo_admin_id = cargo_admin[0] if cargo_admin else None

    if cargo_admin_id:
        cursor.execute("""
            INSERT OR IGNORE INTO usuarios (
                nome,
                email,
                senha,
                cargo_id,
                ativo
            )
            VALUES (?, ?, ?, ?, 1)
        """, (
            "Administrador",
            "admin",
            generate_password_hash("admin123"),
            cargo_admin_id
        ))

    # =====================================================
    # CONFIGURAÇÕES E AUDITORIA DO ANO LETIVO
    # =====================================================

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS configuracao_ano_letivo_global (
            id INTEGER PRIMARY KEY CHECK (id = 1),
            ativo INTEGER NOT NULL DEFAULT 0,
            ano INTEGER,
            data_execucao TEXT,
            data_inicio TEXT,
            data_fim TEXT,
            copiar_turmas INTEGER NOT NULL DEFAULT 1,
            copiar_vinculos INTEGER NOT NULL DEFAULT 1,
            encerrar_anterior INTEGER NOT NULL DEFAULT 1,
            executado INTEGER NOT NULL DEFAULT 0,
            atualizado_em TEXT DEFAULT CURRENT_TIMESTAMP
        )
    """)

    cursor.execute("""
        INSERT OR IGNORE INTO configuracao_ano_letivo_global (
            id,
            ativo,
            copiar_turmas,
            copiar_vinculos,
            encerrar_anterior,
            executado
        )
        VALUES (1, 0, 1, 1, 1, 0)
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS configuracao_ano_letivo_instituicao (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            escola_id INTEGER NOT NULL UNIQUE,
            modo TEXT NOT NULL DEFAULT 'global',
            ativo INTEGER NOT NULL DEFAULT 0,
            ano INTEGER,
            data_execucao TEXT,
            data_inicio TEXT,
            data_fim TEXT,
            copiar_turmas INTEGER NOT NULL DEFAULT 1,
            copiar_vinculos INTEGER NOT NULL DEFAULT 1,
            encerrar_anterior INTEGER NOT NULL DEFAULT 1,
            executado INTEGER NOT NULL DEFAULT 0,
            atualizado_em TEXT DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (escola_id) REFERENCES escolas(id) ON DELETE CASCADE
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS ano_letivo_auditoria (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            escola_id INTEGER,
            ano_letivo_id INTEGER,
            usuario_id INTEGER,
            acao TEXT NOT NULL,
            detalhes TEXT,
            criado_em TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (escola_id) REFERENCES escolas(id) ON DELETE SET NULL,
            FOREIGN KEY (ano_letivo_id) REFERENCES anos_letivos(id) ON DELETE SET NULL,
            FOREIGN KEY (usuario_id) REFERENCES usuarios(id) ON DELETE SET NULL
        )
    """)

    banco.commit()
    banco.close()

# =========================================================
# SITE COMERCIAL PÚBLICO
# =========================================================

@app.route("/inicio")
def site_inicio():
    """Página inicial pública do site comercial, inclusive para usuários logados."""
    return render_template("site/index.html")


@app.route("/recursos")
def site_recursos():
    return render_template("site/recursos.html")


@app.route("/para-professores")
def site_professores():
    return render_template("site/professores.html")


@app.route("/para-instituicoes")
def site_instituicoes():
    return render_template("site/instituicoes.html")


@app.route("/planos")
def site_planos():
    return render_template("site/planos.html")


@app.route("/sobre")
def site_sobre():
    return render_template("site/sobre.html")


@app.route("/termos-de-uso")
def site_termos():
    return render_template("site/termos.html")


@app.route("/politica-de-privacidade")
def site_privacidade():
    return render_template("site/privacidade.html")


@app.route("/contato", methods=["GET", "POST"])
def site_contato():
    if request.method == "POST":
        nome = request.form.get("nome", "").strip()
        email = request.form.get("email", "").strip()
        telefone = request.form.get("telefone", "").strip()
        assunto = request.form.get("assunto", "Contato pelo site").strip()
        mensagem = request.form.get("mensagem", "").strip()

        if not nome or not email or not mensagem:
            flash("Preencha nome, e-mail e mensagem.", "erro")
            return render_template("site/contato.html")

        destino = os.environ.get("EMAIL_COMERCIAL") or app.config.get("MAIL_DEFAULT_SENDER")
        if destino and app.config.get("MAIL_USERNAME") and app.config.get("MAIL_PASSWORD"):
            try:
                corpo = (
                    f"Novo contato pelo site ARK EDUS\n\n"
                    f"Nome: {nome}\nE-mail: {email}\nTelefone: {telefone or 'Não informado'}\n"
                    f"Assunto: {assunto}\n\nMensagem:\n{mensagem}"
                )
                mail.send(Message(subject=f"ARK EDUS — {assunto}", recipients=[destino], body=corpo, reply_to=email))
                flash("Mensagem enviada com sucesso. Em breve entraremos em contato.", "success")
            except Exception:
                app.logger.exception("Falha ao enviar contato comercial")
                flash("Recebemos seus dados, mas o envio automático está temporariamente indisponível.", "aviso")
        else:
            flash("Contato registrado. Configure EMAIL_COMERCIAL e as variáveis de e-mail no Render para receber as mensagens.", "aviso")
        return redirect(url_for("site_contato"))

    return render_template("site/contato.html")


# =========================================================
# DASHBOARD
# =========================================================

@app.route("/")
@app.route("/dashboard")
def index():
    """
    Dashboard global da plataforma.

    Os totais acadêmicos (turmas, alunos e provas) obedecem ao ano
    selecionado na barra superior. Cadastros permanentes, como usuários,
    instituições, professores e questões, não são zerados ao trocar o ano.
    """

    if "usuario_id" not in session:
        if request.path == "/":
            return render_template("site/index.html")
        return redirect(url_for("login"))

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    contexto = obter_contexto_plataforma()

    usuario_id = contexto["usuario_id"]
    usuario_cargo = contexto["cargo"]
    escola_id = contexto["escola_id"]
    ano_letivo_id = contexto["ano_letivo_id"]
    ano_visualizado = contexto["ano"]
    ano_letivo_id_ativo = contexto["ano_ativo_id"]
    ano_letivo_ativo = contexto["ano_ativo"]

    total_instituicoes = 0
    total_usuarios = 0
    total_professores = 0
    total_alunos = 0
    total_turmas = 0
    total_questoes = 0
    total_provas = 0

    nome_instituicao = None
    permissoes_usuario = []

    try:
        # =====================================================
        # PERMISSÕES DO USUÁRIO
        # =====================================================

        if usuario_cargo == "Administrador Geral":
            permissoes_usuario = [
                "Dashboard",
                "Instituições",
                "Usuários",
                "Anos Letivos",
                "Turmas",
                "Professores",
                "Alunos",
                "Questões",
                "Provas",
                "Relatórios"
            ]
        else:
            cursor.execute("""
                SELECT modulo
                FROM usuario_permissoes
                WHERE usuario_id = ?
                  AND pode_acessar = 1
            """, (usuario_id,))

            permissoes_usuario = [
                linha["modulo"]
                for linha in cursor.fetchall()
            ]

            if (
                usuario_cargo == "Administrador da Instituição"
                and "Anos Letivos" not in permissoes_usuario
            ):
                permissoes_usuario.append("Anos Letivos")

        # =====================================================
        # ADMINISTRADOR GERAL
        # Usa o número do ano, porque cada instituição possui
        # um ID próprio para o mesmo ano letivo.
        # =====================================================

        if usuario_cargo == "Administrador Geral":
            cursor.execute("""
                SELECT COUNT(*) AS total
                FROM escolas
                WHERE COALESCE(status, 1) = 1
            """)
            total_instituicoes = cursor.fetchone()["total"]

            cursor.execute("""
                SELECT COUNT(*) AS total
                FROM usuarios
                WHERE ativo = 1
            """)
            total_usuarios = cursor.fetchone()["total"]

            cursor.execute("""
                SELECT COUNT(*) AS total
                FROM professores
            """)
            total_professores = cursor.fetchone()["total"]

            cursor.execute("""
                SELECT COUNT(*) AS total
                FROM questoes
            """)
            total_questoes = cursor.fetchone()["total"]

            if ano_visualizado is not None:
                # Turmas do ano selecionado em todas as instituições.
                cursor.execute("""
                    SELECT COUNT(*) AS total
                    FROM turmas AS t
                    INNER JOIN anos_letivos AS al
                        ON al.id = t.ano_letivo_id
                       AND al.escola_id = t.escola_id
                    WHERE al.ano = ?
                """, (ano_visualizado,))
                total_turmas = cursor.fetchone()["total"]

                # Alunos do ano selecionado. O UNION mantém compatibilidade
                # com os registros antigos da tabela alunos e com a nova
                # tabela de histórico aluno_matriculas.
                cursor.execute("""
                    SELECT COUNT(*) AS total
                    FROM (
                        SELECT
                            am.aluno_id AS aluno_id,
                            am.escola_id AS escola_id
                        FROM aluno_matriculas AS am
                        INNER JOIN anos_letivos AS al
                            ON al.id = am.ano_letivo_id
                           AND al.escola_id = am.escola_id
                        WHERE al.ano = ?

                        UNION

                        SELECT
                            a.id AS aluno_id,
                            a.escola_id AS escola_id
                        FROM alunos AS a
                        INNER JOIN anos_letivos AS al
                            ON al.id = a.ano_letivo_id
                           AND al.escola_id = a.escola_id
                        WHERE al.ano = ?
                    ) AS alunos_do_ano
                """, (ano_visualizado, ano_visualizado))
                total_alunos = cursor.fetchone()["total"]

                cursor.execute("""
                    SELECT COUNT(*) AS total
                    FROM provas AS p
                    INNER JOIN anos_letivos AS al
                        ON al.id = p.ano_letivo_id
                       AND al.escola_id = p.escola_id
                    WHERE al.ano = ?
                """, (ano_visualizado,))
                total_provas = cursor.fetchone()["total"]

        # =====================================================
        # USUÁRIOS VINCULADOS A UMA INSTITUIÇÃO
        # Usa o ID exato do ano selecionado para aquela escola.
        # =====================================================

        elif escola_id:
            cursor.execute("""
                SELECT nome_instituicao
                FROM escolas
                WHERE id = ?
                LIMIT 1
            """, (escola_id,))

            escola = cursor.fetchone()
            if escola:
                nome_instituicao = escola["nome_instituicao"]

            total_instituicoes = 1

            cursor.execute("""
                SELECT COUNT(*) AS total
                FROM usuarios
                WHERE escola_id = ?
                  AND ativo = 1
            """, (escola_id,))
            total_usuarios = cursor.fetchone()["total"]

            cursor.execute("""
                SELECT COUNT(*) AS total
                FROM professores
                WHERE escola_id = ?
            """, (escola_id,))
            total_professores = cursor.fetchone()["total"]

            cursor.execute("""
                SELECT COUNT(*) AS total
                FROM questoes
                WHERE escola_id = ?
            """, (escola_id,))
            total_questoes = cursor.fetchone()["total"]

            if ano_letivo_id:
                cursor.execute("""
                    SELECT COUNT(*) AS total
                    FROM turmas
                    WHERE escola_id = ?
                      AND ano_letivo_id = ?
                """, (escola_id, ano_letivo_id))
                total_turmas = cursor.fetchone()["total"]

                cursor.execute("""
                    SELECT COUNT(*) AS total
                    FROM (
                        SELECT aluno_id
                        FROM aluno_matriculas
                        WHERE escola_id = ?
                          AND ano_letivo_id = ?

                        UNION

                        SELECT id AS aluno_id
                        FROM alunos
                        WHERE escola_id = ?
                          AND ano_letivo_id = ?
                    ) AS alunos_do_ano
                """, (
                    escola_id,
                    ano_letivo_id,
                    escola_id,
                    ano_letivo_id
                ))
                total_alunos = cursor.fetchone()["total"]

                cursor.execute("""
                    SELECT COUNT(*) AS total
                    FROM provas
                    WHERE escola_id = ?
                      AND ano_letivo_id = ?
                """, (escola_id, ano_letivo_id))
                total_provas = cursor.fetchone()["total"]

        else:
            nome_instituicao = "Usuário sem instituição vinculada"

        return render_template(
            "dashboard/index.html",
            total_instituicoes=total_instituicoes,
            total_usuarios=total_usuarios,
            total_professores=total_professores,
            total_alunos=total_alunos,
            total_turmas=total_turmas,
            total_questoes=total_questoes,
            total_provas=total_provas,
            nome_instituicao=nome_instituicao,
            ano_letivo_id_ativo=ano_letivo_id_ativo,
            ano_letivo_ativo=ano_letivo_ativo,
            ano_letivo_id_visualizado=ano_letivo_id,
            ano_letivo_visualizado=ano_visualizado,
            consultando_historico=contexto["consultando_historico"],
            permissoes_usuario=permissoes_usuario
        )

    except sqlite3.Error as erro:
        import traceback
        traceback.print_exc()

        print("ERRO AO CARREGAR O DASHBOARD:", erro)
        flash(
            f"Erro ao carregar os dados do dashboard: {erro}",
            "erro"
        )

        return render_template(
            "dashboard/index.html",
            total_instituicoes=0,
            total_usuarios=0,
            total_professores=0,
            total_alunos=0,
            total_turmas=0,
            total_questoes=0,
            total_provas=0,
            nome_instituicao=nome_instituicao,
            ano_letivo_id_ativo=ano_letivo_id_ativo,
            ano_letivo_ativo=ano_letivo_ativo,
            ano_letivo_id_visualizado=ano_letivo_id,
            ano_letivo_visualizado=ano_visualizado,
            consultando_historico=contexto["consultando_historico"],
            permissoes_usuario=permissoes_usuario
        )

    finally:
        banco.close()


@app.route("/esqueci_senha")
def esqueci_senha():
    return render_template("esqueci_senha.html")

# =========================================================
# LISTAR TURMAS PELO ANO LETIVO SELECIONADO
# =========================================================

@app.route("/turmas")
def turmas():

    if not permissao_modulo("Turmas"):
        return redirect("/acesso_negado")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    cargo = session.get("usuario_cargo", "").strip()
    usuario_id = session.get("usuario_id")
    escola_id = obter_escola_usuario()

    pode_gerenciar = cargo in [
        "Administrador Geral",
        "Administrador da Instituição",
        "Coordenador",
        "Secretaria",
        "Professor Autônomo"
    ]

    escolas = []
    lista_turmas = []

    ano_letivo_visualizado = session.get("ano_letivo_visualizado")
    ano_letivo_id = session.get("ano_letivo_id")
    consultando_ano_antigo = False

    try:

        # =====================================================
        # ADMINISTRADOR GERAL
        #
        # Para ele, o mesmo número de ano é aplicado a todas
        # as instituições. Cada escola possui um ID diferente
        # para o seu próprio registro em anos_letivos.
        # =====================================================

        if cargo == "Administrador Geral":

            if not ano_letivo_visualizado:
                ano_letivo_visualizado = datetime.now().year
                session["ano_letivo_visualizado"] = ano_letivo_visualizado
                session["ano_letivo"] = ano_letivo_visualizado

            try:
                ano_letivo_visualizado = int(ano_letivo_visualizado)
            except (TypeError, ValueError):
                ano_letivo_visualizado = datetime.now().year
                session["ano_letivo_visualizado"] = ano_letivo_visualizado
                session["ano_letivo"] = ano_letivo_visualizado

            cursor.execute("""
                SELECT COUNT(*) AS total
                FROM anos_letivos
                WHERE ano = ?
                  AND ativo = 1
                  AND encerrado = 0
            """, (
                ano_letivo_visualizado,
            ))

            resultado_ativo = cursor.fetchone()

            consultando_ano_antigo = (
                not resultado_ativo
                or resultado_ativo["total"] == 0
            )

            cursor.execute("""
                SELECT
                    t.*,
                    e.nome_instituicao,

                    al.id AS ano_letivo_id_atual,
                    al.ano AS ano_letivo_atual,
                    al.ativo AS ano_letivo_ativo,
                    al.encerrado AS ano_letivo_encerrado,

                    CASE
                        WHEN (
                            SELECT COUNT(*)
                            FROM aluno_matriculas AS am
                            WHERE am.turma_id = t.id
                              AND am.ano_letivo_id = t.ano_letivo_id
                        ) > 0
                        THEN (
                            SELECT COUNT(*)
                            FROM aluno_matriculas AS am
                            WHERE am.turma_id = t.id
                              AND am.ano_letivo_id = t.ano_letivo_id
                        )
                        ELSE (
                            SELECT COUNT(*)
                            FROM alunos AS a
                            WHERE a.turma_id = t.id
                              AND a.ano_letivo_id = t.ano_letivo_id
                        )
                    END AS total_alunos,

                    (
                        SELECT COUNT(DISTINCT pv.professor_id)
                        FROM professor_vinculos AS pv
                        WHERE pv.turma_id = t.id
                    ) AS total_professores

                FROM turmas AS t

                INNER JOIN anos_letivos AS al
                    ON al.id = t.ano_letivo_id
                   AND al.escola_id = t.escola_id

                INNER JOIN escolas AS e
                    ON e.id = t.escola_id

                WHERE al.ano = ?

                ORDER BY
                    e.nome_instituicao COLLATE NOCASE ASC,
                    t.etapa COLLATE NOCASE ASC,
                    t.ano COLLATE NOCASE ASC,
                    t.nome COLLATE NOCASE ASC,
                    t.turno COLLATE NOCASE ASC
            """, (
                ano_letivo_visualizado,
            ))

            lista_turmas = cursor.fetchall()

            # Todas as instituições ativas devem aparecer no cadastro.
            # O LEFT JOIN mantém também as instituições que ainda não
            # possuem ano letivo ativo configurado.
            cursor.execute("""
                SELECT
                    e.id,
                    e.nome_instituicao,
                    al.id AS ano_letivo_id,
                    al.ano AS ano_letivo_ativo

                FROM escolas AS e

                LEFT JOIN anos_letivos AS al
                    ON al.escola_id = e.id
                   AND al.ativo = 1
                   AND al.encerrado = 0

                WHERE COALESCE(e.status, 1) = 1

                ORDER BY
                    e.nome_instituicao COLLATE NOCASE ASC
            """)

            escolas = cursor.fetchall()

        # =====================================================
        # USUÁRIOS VINCULADOS A UMA INSTITUIÇÃO
        # =====================================================

        else:

            if not escola_id:
                flash(
                    "Não foi possível identificar sua instituição.",
                    "erro"
                )

                return render_template(
                    "gestao/turmas.html",
                    turmas=[],
                    escolas=[],
                    cargo=cargo,
                    pode_gerenciar=pode_gerenciar,
                    ano_letivo_ativo=None,
                    ano_letivo_visualizado=None,
                    consultando_ano_antigo=False
                )

            # Recupera o ano selecionado ou, por padrão, o ativo.
            ano_selecionado = atualizar_ano_letivo_na_sessao(escola_id)

            if not ano_selecionado:
                flash(
                    "A instituição não possui um ano letivo disponível.",
                    "erro"
                )

                return render_template(
                    "gestao/turmas.html",
                    turmas=[],
                    escolas=[],
                    cargo=cargo,
                    pode_gerenciar=pode_gerenciar,
                    ano_letivo_ativo=None,
                    ano_letivo_visualizado=None,
                    consultando_ano_antigo=False
                )

            ano_letivo_id = ano_selecionado["id"]
            ano_letivo_visualizado = ano_selecionado["ano"]

            session["ano_letivo_id"] = ano_letivo_id
            session["ano_letivo"] = ano_letivo_visualizado
            session["ano_letivo_visualizado"] = ano_letivo_visualizado

            consultando_ano_antigo = not (
                ano_selecionado["ativo"] == 1
                and ano_selecionado["encerrado"] == 0
            )

            # Professor vê apenas turmas em que possui vínculo.
            if cargo == "Professor":

                cursor.execute("""
                    SELECT DISTINCT
                        t.*,
                        e.nome_instituicao,

                        al.id AS ano_letivo_id_atual,
                        al.ano AS ano_letivo_atual,
                        al.ativo AS ano_letivo_ativo,
                        al.encerrado AS ano_letivo_encerrado,

                        CASE
                            WHEN (
                                SELECT COUNT(*)
                                FROM aluno_matriculas AS am
                                WHERE am.turma_id = t.id
                                  AND am.ano_letivo_id = t.ano_letivo_id
                            ) > 0
                            THEN (
                                SELECT COUNT(*)
                                FROM aluno_matriculas AS am
                                WHERE am.turma_id = t.id
                                  AND am.ano_letivo_id = t.ano_letivo_id
                            )
                            ELSE (
                                SELECT COUNT(*)
                                FROM alunos AS a
                                WHERE a.turma_id = t.id
                                  AND a.ano_letivo_id = t.ano_letivo_id
                            )
                        END AS total_alunos,

                        (
                            SELECT COUNT(DISTINCT pv_total.professor_id)
                            FROM professor_vinculos AS pv_total
                            WHERE pv_total.turma_id = t.id
                        ) AS total_professores

                    FROM turmas AS t

                    INNER JOIN professor_vinculos AS pv
                        ON pv.turma_id = t.id

                    INNER JOIN anos_letivos AS al
                        ON al.id = t.ano_letivo_id
                       AND al.escola_id = t.escola_id

                    INNER JOIN escolas AS e
                        ON e.id = t.escola_id

                    WHERE pv.professor_id = ?
                      AND t.escola_id = ?
                      AND t.ano_letivo_id = ?

                    ORDER BY
                        t.etapa COLLATE NOCASE ASC,
                        t.ano COLLATE NOCASE ASC,
                        t.nome COLLATE NOCASE ASC,
                        t.turno COLLATE NOCASE ASC
                """, (
                    usuario_id,
                    escola_id,
                    ano_letivo_id
                ))

            else:

                cursor.execute("""
                    SELECT
                        t.*,
                        e.nome_instituicao,

                        al.id AS ano_letivo_id_atual,
                        al.ano AS ano_letivo_atual,
                        al.ativo AS ano_letivo_ativo,
                        al.encerrado AS ano_letivo_encerrado,

                        CASE
                            WHEN (
                                SELECT COUNT(*)
                                FROM aluno_matriculas AS am
                                WHERE am.turma_id = t.id
                                  AND am.ano_letivo_id = t.ano_letivo_id
                            ) > 0
                            THEN (
                                SELECT COUNT(*)
                                FROM aluno_matriculas AS am
                                WHERE am.turma_id = t.id
                                  AND am.ano_letivo_id = t.ano_letivo_id
                            )
                            ELSE (
                                SELECT COUNT(*)
                                FROM alunos AS a
                                WHERE a.turma_id = t.id
                                  AND a.ano_letivo_id = t.ano_letivo_id
                            )
                        END AS total_alunos,

                        (
                            SELECT COUNT(DISTINCT pv.professor_id)
                            FROM professor_vinculos AS pv
                            WHERE pv.turma_id = t.id
                        ) AS total_professores

                    FROM turmas AS t

                    INNER JOIN anos_letivos AS al
                        ON al.id = t.ano_letivo_id
                       AND al.escola_id = t.escola_id

                    INNER JOIN escolas AS e
                        ON e.id = t.escola_id

                    WHERE t.escola_id = ?
                      AND t.ano_letivo_id = ?

                    ORDER BY
                        t.etapa COLLATE NOCASE ASC,
                        t.ano COLLATE NOCASE ASC,
                        t.nome COLLATE NOCASE ASC,
                        t.turno COLLATE NOCASE ASC
                """, (
                    escola_id,
                    ano_letivo_id
                ))

            lista_turmas = cursor.fetchall()

        return render_template(
            "gestao/turmas.html",
            turmas=lista_turmas,
            escolas=escolas,
            cargo=cargo,
            pode_gerenciar=pode_gerenciar,

            # Mantido para compatibilidade com seu HTML atual.
            ano_letivo_ativo=ano_letivo_visualizado,

            # Novos nomes.
            ano_letivo_visualizado=ano_letivo_visualizado,
            consultando_ano_antigo=consultando_ano_antigo
        )

    except sqlite3.Error as erro:

        import traceback
        traceback.print_exc()

        print("ERRO AO LISTAR TURMAS:", erro)

        flash(
            f"Erro ao carregar as turmas: {erro}",
            "erro"
        )

        return render_template(
            "gestao/turmas.html",
            turmas=[],
            escolas=[],
            cargo=cargo,
            pode_gerenciar=pode_gerenciar,
            ano_letivo_ativo=None,
            ano_letivo_visualizado=None,
            consultando_ano_antigo=False
        )

    finally:
        banco.close()


# =========================================================
# CADASTRAR TURMA NO ANO LETIVO VISUALIZADO
# =========================================================

@app.route("/cadastrar_turma", methods=["POST"])
def cadastrar_turma():

    if not cargo_permitido([
        "Administrador Geral",
        "Administrador da Instituição",
        "Coordenador",
        "Secretaria",
        "Professor Autônomo"
    ]):

        flash(
            "Você não possui permissão para cadastrar turmas.",
            "erro"
        )

        return redirect("/acesso_negado")

    etapa = request.form.get("etapa", "").strip()
    ano_serie = request.form.get("ano", "").strip()
    identificacao = request.form.get("nome", "").strip()
    turno = request.form.get("turno", "").strip()

    cargo = session.get("usuario_cargo", "").strip()
    escola_id = obter_escola_usuario()

    if not etapa:
        flash("Selecione a etapa de ensino.", "erro")
        return redirect("/turmas")

    if not ano_serie:
        flash("Selecione o ano ou série.", "erro")
        return redirect("/turmas")

    if not identificacao:
        flash("Informe a identificação da turma.", "erro")
        return redirect("/turmas")

    if not turno:
        flash("Selecione o turno.", "erro")
        return redirect("/turmas")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:

        # =====================================================
        # DESCOBRIR A INSTITUIÇÃO
        # =====================================================

        if cargo == "Administrador Geral":

            escola_formulario = request.form.get(
                "escola_id",
                ""
            ).strip()

            if not escola_formulario:
                flash("Selecione uma instituição.", "erro")
                return redirect("/turmas")

            try:
                escola_id = int(escola_formulario)
            except (TypeError, ValueError):
                flash(
                    "A instituição selecionada é inválida.",
                    "erro"
                )
                return redirect("/turmas")

        elif not escola_id:

            flash(
                "Não foi possível identificar a instituição da turma.",
                "erro"
            )

            return redirect("/turmas")

        # =====================================================
        # VALIDAR INSTITUIÇÃO
        # =====================================================

        cursor.execute("""
            SELECT
                id,
                nome_instituicao
            FROM escolas
            WHERE id = ?
              AND COALESCE(status, 1) = 1
            LIMIT 1
        """, (
            escola_id,
        ))

        escola = cursor.fetchone()

        if escola is None:

            flash(
                "A instituição selecionada não existe ou está inativa.",
                "erro"
            )

            return redirect("/turmas")

        # =====================================================
        # DEFINIR O ANO EM QUE A TURMA SERÁ CRIADA
        # =====================================================

        if cargo == "Administrador Geral":

            numero_ano_letivo = session.get(
                "ano_letivo_visualizado"
            ) or session.get("ano_letivo")

            if not numero_ano_letivo:
                flash(
                    "Selecione um ano letivo no topo da plataforma.",
                    "erro"
                )
                return redirect("/turmas")

            try:
                numero_ano_letivo = int(numero_ano_letivo)
            except (TypeError, ValueError):
                flash("O ano letivo selecionado é inválido.", "erro")
                return redirect("/turmas")

            cursor.execute("""
                SELECT
                    id,
                    ano,
                    ativo,
                    encerrado
                FROM anos_letivos
                WHERE escola_id = ?
                  AND ano = ?
                LIMIT 1
            """, (
                escola_id,
                numero_ano_letivo
            ))

            ano_letivo = cursor.fetchone()

        else:

            ano_letivo = atualizar_ano_letivo_na_sessao(
                escola_id
            )

        if ano_letivo is None:

            flash(
                "A instituição não possui o ano letivo "
                "selecionado. Cadastre ou prepare esse ano antes "
                "de criar uma turma.",
                "erro"
            )

            return redirect("/turmas")

        ano_letivo_id = ano_letivo["id"]
        numero_ano_letivo = ano_letivo["ano"]

        # Impede alterações acidentais em anos encerrados.
        if (
            ano_letivo["encerrado"] == 1
            or ano_letivo["ativo"] != 1
        ):

            flash(
                f"O ano letivo {numero_ano_letivo} está em modo "
                "de consulta. Volte ao ano ativo para cadastrar turmas.",
                "erro"
            )

            return redirect("/turmas")

        # =====================================================
        # VERIFICAR TURMA DUPLICADA NO MESMO ANO
        # =====================================================

        cursor.execute("""
            SELECT id
            FROM turmas
            WHERE LOWER(TRIM(nome)) = LOWER(TRIM(?))
              AND LOWER(TRIM(etapa)) = LOWER(TRIM(?))
              AND LOWER(TRIM(ano)) = LOWER(TRIM(?))
              AND LOWER(TRIM(turno)) = LOWER(TRIM(?))
              AND escola_id = ?
              AND ano_letivo_id = ?
            LIMIT 1
        """, (
            identificacao,
            etapa,
            ano_serie,
            turno,
            escola_id,
            ano_letivo_id
        ))

        if cursor.fetchone():

            flash(
                f"Já existe uma turma igual cadastrada em "
                f"{numero_ano_letivo}.",
                "erro"
            )

            return redirect("/turmas")

        # =====================================================
        # CADASTRAR TURMA
        # =====================================================

        cursor.execute("""
            INSERT INTO turmas (
                nome,
                etapa,
                ano,
                turno,
                escola_id,
                ano_letivo,
                ano_letivo_id
            )
            VALUES (?, ?, ?, ?, ?, ?, ?)
        """, (
            identificacao,
            etapa,
            ano_serie,
            turno,
            escola_id,
            numero_ano_letivo,
            ano_letivo_id
        ))

        banco.commit()

        flash(
            f"Turma cadastrada com sucesso no ano letivo "
            f"{numero_ano_letivo}.",
            "success"
        )

        return redirect("/turmas")

    except sqlite3.Error as erro:

        banco.rollback()

        import traceback
        traceback.print_exc()

        print(
            "ERRO DO BANCO AO CADASTRAR TURMA:",
            erro
        )

        flash(
            f"Erro ao cadastrar turma: {erro}",
            "erro"
        )

        return redirect("/turmas")

    finally:
        banco.close()

# =========================================================
# VISUALIZAR TURMA
# =========================================================

@app.route("/turmas/<int:turma_id>")
def visualizar_turma(turma_id):

    if not permissao_modulo("Turmas"):
        return redirect("/acesso_negado")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    cargo = session.get("usuario_cargo", "").strip()
    usuario_id = session.get("usuario_id")
    escola_id = session.get("escola_id")

    pode_editar = cargo in [
        "Administrador Geral",
        "Administrador da Instituição",
        "Coordenador",
        "Secretaria"
    ]

    try:

        cursor.execute("""
            SELECT
                turmas.*,
                escolas.nome_instituicao,
                escolas.logo AS escola_logo,
                escolas.coordenador1,
                escolas.coordenador2,
                escolas.coordenador3,
                escolas.secretario
            FROM turmas
            LEFT JOIN escolas
                ON escolas.id = turmas.escola_id
            WHERE turmas.id = ?
            LIMIT 1
        """, (
            turma_id,
        ))

        turma = cursor.fetchone()

        if turma is None:

            flash(
                "Turma não encontrada.",
                "erro"
            )

            return redirect("/turmas")

        # Professor só acessa turma em que possui vínculo.
        if cargo == "Professor":

            cursor.execute("""
                SELECT 1
                FROM professor_vinculos
                WHERE professor_id = ?
                  AND turma_id = ?
                LIMIT 1
            """, (
                usuario_id,
                turma_id
            ))

            if cursor.fetchone() is None:

                flash(
                    "Você não está vinculado a esta turma.",
                    "erro"
                )

                return redirect("/turmas")

        elif cargo != "Administrador Geral":

            if turma["escola_id"] != escola_id:

                flash(
                    "Você não possui acesso a esta turma.",
                    "erro"
                )

                return redirect("/turmas")

        # Lista os estudantes matriculados nesta turma no respectivo ano letivo.
        # O UNION mantém compatibilidade com registros antigos existentes
        # somente na tabela alunos.
        cursor.execute("""
            SELECT
                a.id,
                a.nome,
                a.matricula,
                am.situacao,
                am.id AS matricula_id
            FROM aluno_matriculas AS am
            INNER JOIN alunos AS a
                ON a.id = am.aluno_id
            WHERE am.turma_id = ?
              AND am.ano_letivo_id = ?
              AND COALESCE(am.situacao, 'Cursando') NOT IN (
                  'Transferido',
                  'Cancelado'
              )

            UNION ALL

            SELECT
                a.id,
                a.nome,
                a.matricula,
                'Cursando' AS situacao,
                NULL AS matricula_id
            FROM alunos AS a
            WHERE a.turma_id = ?
              AND a.ano_letivo_id = ?
              AND NOT EXISTS (
                  SELECT 1
                  FROM aluno_matriculas AS am
                  WHERE am.aluno_id = a.id
                    AND am.ano_letivo_id = a.ano_letivo_id
              )

            ORDER BY nome COLLATE NOCASE ASC
        """, (
            turma_id,
            turma["ano_letivo_id"],
            turma_id,
            turma["ano_letivo_id"]
        ))

        alunos = cursor.fetchall()

        # Cada linha representa um vínculo real entre:
        # professor + turma + componente curricular.
        cursor.execute("""
            SELECT
                pv.id AS vinculo_id,
                pv.professor_id,
                pv.componente_id,
                u.nome,
                u.email,
                c.nome AS cargo,
                cc.nome AS componente_nome,
                cc.etapa_ensino AS componente_etapa
            FROM professor_vinculos AS pv
            INNER JOIN usuarios AS u
                ON u.id = pv.professor_id
            LEFT JOIN cargos AS c
                ON c.id = u.cargo_id
            INNER JOIN componentes_curriculares AS cc
                ON cc.id = pv.componente_id
            WHERE pv.turma_id = ?
            ORDER BY
                u.nome COLLATE NOCASE ASC,
                cc.nome COLLATE NOCASE ASC
        """, (turma_id,))

        professores = cursor.fetchall()

        # Dados prontos para a impressão da lista da turma. Mantemos somente
        # nomes únicos de professores e usamos os responsáveis institucionais
        # cadastrados na própria escola (coordenadores e secretário/a).
        nomes_professores_impressao = []
        vistos_professores = set()
        for professor in professores:
            nome_professor = (professor["nome"] or "").strip()
            chave = nome_professor.casefold()
            if nome_professor and chave not in vistos_professores:
                vistos_professores.add(chave)
                nomes_professores_impressao.append(nome_professor)

        coordenadores_impressao = [
            (turma["coordenador1"] or "").strip(),
            (turma["coordenador2"] or "").strip(),
            (turma["coordenador3"] or "").strip(),
        ]
        coordenadores_impressao = [nome for nome in coordenadores_impressao if nome]

        secretario_impressao = (turma["secretario"] or "").strip()
        alunos_impressao = [
            {
                "nome": (aluno["nome"] or "").strip(),
                "matricula": str(aluno["matricula"] or "").strip(),
            }
            for aluno in alunos
        ]

        pode_gerenciar_professores = cargo in [
            "Administrador Geral",
            "Administrador da Instituição",
            "Coordenador"
        ]

        professores_disponiveis = []
        componentes_disponiveis = []

        if pode_gerenciar_professores:
            # Professores ativos pertencentes à mesma instituição da turma.
            cursor.execute("""
                SELECT
                    u.id,
                    u.nome,
                    u.email
                FROM usuarios AS u
                INNER JOIN cargos AS c
                    ON c.id = u.cargo_id
                WHERE c.nome = 'Professor'
                  AND COALESCE(u.ativo, 1) = 1
                  AND u.escola_id = ?
                ORDER BY u.nome COLLATE NOCASE ASC
            """, (turma["escola_id"],))

            professores_disponiveis = cursor.fetchall()

            # Busca todos os componentes ativos da instituição e compara
            # a etapa em Python. Isso evita falhas causadas por hífen,
            # travessão, abreviações ou pequenas diferenças de escrita.
            cursor.execute("""
                SELECT
                    MIN(id) AS id,
                    TRIM(nome) AS nome,
                    TRIM(etapa_ensino) AS etapa_ensino
                FROM componentes_curriculares
                WHERE escola_id = ?
                  AND ativo = 1
                  AND TRIM(COALESCE(nome, '')) <> ''
                GROUP BY
                    LOWER(TRIM(nome)),
                    LOWER(TRIM(COALESCE(etapa_ensino, '')))
                ORDER BY TRIM(nome) COLLATE NOCASE ASC
            """, (turma["escola_id"],))

            todos_componentes = cursor.fetchall()
            etapa_turma_normalizada = _identificar_etapa_ensino(
                turma["etapa"]
            )

            componentes_disponiveis = [
                componente
                for componente in todos_componentes
                if (
                    not etapa_turma_normalizada
                    or _identificar_etapa_ensino(
                        componente["etapa_ensino"]
                    ) == etapa_turma_normalizada
                )
            ]

            # Compatibilidade com cadastros antigos:
            # se a etapa estiver escrita de forma diferente ou vazia,
            # não deixa o campo de componentes sem opções.
            if not componentes_disponiveis:
                componentes_disponiveis = todos_componentes

        cursor.execute("""
            SELECT *
            FROM provas
            WHERE turma_id = ?
            ORDER BY id DESC
        """, (turma_id,))

        avaliacoes = cursor.fetchall()

        return render_template(
            "gestao/visualizar_turma.html",
            turma=turma,
            alunos=alunos,
            professores=professores,
            professores_disponiveis=professores_disponiveis,
            componentes_disponiveis=componentes_disponiveis,
            avaliacoes=avaliacoes,
            cargo=cargo,
            pode_editar=pode_editar,
            pode_gerenciar_professores=pode_gerenciar_professores,
            nomes_professores_impressao=nomes_professores_impressao,
            coordenadores_impressao=coordenadores_impressao,
            secretario_impressao=secretario_impressao,
            alunos_impressao=alunos_impressao
        )

    except sqlite3.Error as erro:

        import traceback
        traceback.print_exc()

        print("ERRO AO VISUALIZAR TURMA:", erro)

        flash(
            f"Erro ao carregar os dados da turma: {erro}",
            "erro"
        )

        return redirect("/turmas")

    finally:
        banco.close()



# =========================================================
# NORMALIZAÇÃO DAS ETAPAS DE ENSINO
# =========================================================

def _normalizar_texto_etapa(valor):
    """Remove acentos, pontuação e diferenças de espaçamento."""
    texto = unicodedata.normalize(
        "NFKD",
        str(valor or "")
    ).encode(
        "ascii",
        "ignore"
    ).decode(
        "ascii"
    ).lower()

    return re.sub(r"[^a-z0-9]+", " ", texto).strip()


def _identificar_etapa_ensino(valor):
    """
    Converte diferentes formas de escrita da etapa para uma chave única.
    """
    texto = _normalizar_texto_etapa(valor)

    if not texto:
        return ""

    if "educacao infantil" in texto or texto in {"infantil", "ei"}:
        return "educacao_infantil"

    if (
        "anos iniciais" in texto
        or "fundamental i" in texto
        or "fundamental 1" in texto
        or texto in {"ef i", "ef 1"}
    ):
        return "fundamental_anos_iniciais"

    if (
        "anos finais" in texto
        or "fundamental ii" in texto
        or "fundamental 2" in texto
        or texto in {"ef ii", "ef 2"}
    ):
        return "fundamental_anos_finais"

    if "ensino medio" in texto or texto in {"medio", "em"}:
        return "ensino_medio"

    if "ensino superior" in texto or texto in {"superior", "es"}:
        return "ensino_superior"

    return texto.replace(" ", "_")


# =========================================================
# VÍNCULOS DE PROFESSORES DIRETAMENTE NA PÁGINA DA TURMA
# =========================================================

def _pode_gerenciar_professores_turma():
    return cargo_permitido([
        "Administrador Geral",
        "Administrador da Instituição",
        "Coordenador"
    ])


def _buscar_turma_para_gerenciar_vinculos(cursor, turma_id):
    cursor.execute("""
        SELECT id, nome, etapa, ano, turno, escola_id
        FROM turmas
        WHERE id = ?
        LIMIT 1
    """, (turma_id,))

    turma = cursor.fetchone()

    if turma is None:
        return None, "Turma não encontrada."

    cargo = session.get("usuario_cargo", "").strip()
    escola_logada_id = session.get("escola_id")

    if (
        cargo != "Administrador Geral"
        and turma["escola_id"] != escola_logada_id
    ):
        return None, "Você não possui acesso a esta turma."

    return turma, None


def _validar_professor_e_componente_turma(
    cursor,
    turma,
    professor_id,
    componente_id
):
    cursor.execute("""
        SELECT
            u.id,
            u.nome,
            u.escola_id,
            COALESCE(u.ativo, 1) AS ativo,
            c.nome AS cargo
        FROM usuarios AS u
        INNER JOIN cargos AS c
            ON c.id = u.cargo_id
        WHERE u.id = ?
        LIMIT 1
    """, (professor_id,))

    professor = cursor.fetchone()

    if professor is None:
        return None, None, "Professor não encontrado."

    if professor["cargo"] != "Professor":
        return None, None, "O usuário selecionado não possui o cargo Professor."

    if professor["ativo"] != 1:
        return None, None, "O professor selecionado está inativo."

    if professor["escola_id"] != turma["escola_id"]:
        return None, None, (
            "O professor selecionado não pertence à instituição desta turma."
        )

    cursor.execute("""
        SELECT id, nome, etapa_ensino, escola_id, ativo
        FROM componentes_curriculares
        WHERE id = ?
        LIMIT 1
    """, (componente_id,))

    componente = cursor.fetchone()

    if componente is None:
        return None, None, "Componente curricular não encontrado."

    if componente["ativo"] != 1:
        return None, None, "O componente curricular selecionado está inativo."

    if componente["escola_id"] != turma["escola_id"]:
        return None, None, (
            "O componente curricular não pertence à instituição desta turma."
        )

    # Não bloqueia o vínculo por diferença no texto da etapa.
    #
    # Em cadastros antigos, alguns componentes curriculares podem ter sido
    # registrados com outra etapa, mesmo pertencendo à mesma instituição.
    # A instituição continua sendo validada acima, portanto o vínculo permanece
    # restrito à escola correta.
    return professor, componente, None


@app.route(
    "/turmas/<int:turma_id>/professores/adicionar",
    methods=["POST"]
)
def adicionar_professor_turma(turma_id):
    if not _pode_gerenciar_professores_turma():
        flash(
            "Você não possui permissão para gerenciar professores da turma.",
            "erro"
        )
        return redirect("/acesso_negado")

    professor_id = request.form.get("professor_id", type=int)
    componente_id = request.form.get("componente_id", type=int)

    if not professor_id or not componente_id:
        flash("Selecione o professor e o componente curricular.", "erro")
        return redirect(f"/turmas/{turma_id}#professores")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        cursor.execute("PRAGMA foreign_keys = ON")

        turma, erro = _buscar_turma_para_gerenciar_vinculos(
            cursor,
            turma_id
        )

        if erro:
            flash(erro, "erro")
            return redirect("/turmas")

        professor, componente, erro = (
            _validar_professor_e_componente_turma(
                cursor,
                turma,
                professor_id,
                componente_id
            )
        )

        if erro:
            flash(erro, "erro")
            return redirect(f"/turmas/{turma_id}#professores")

        cursor.execute("""
            SELECT id
            FROM professor_vinculos
            WHERE professor_id = ?
              AND turma_id = ?
              AND componente_id = ?
            LIMIT 1
        """, (
            professor_id,
            turma_id,
            componente_id
        ))

        if cursor.fetchone():
            flash(
                "Este professor já está vinculado a esse componente nesta turma.",
                "erro"
            )
            return redirect(f"/turmas/{turma_id}#professores")

        cursor.execute("""
            INSERT INTO professor_vinculos (
                professor_id,
                turma_id,
                componente_id
            )
            VALUES (?, ?, ?)
        """, (
            professor_id,
            turma_id,
            componente_id
        ))

        banco.commit()

        flash(
            f"{professor['nome']} foi vinculado(a) ao componente "
            f"{componente['nome']}.",
            "success"
        )

        return redirect(f"/turmas/{turma_id}#professores")

    except sqlite3.IntegrityError:
        banco.rollback()
        flash(
            "Este professor já está vinculado a esse componente nesta turma.",
            "erro"
        )
        return redirect(f"/turmas/{turma_id}#professores")

    except sqlite3.Error as erro:
        banco.rollback()
        print("ERRO AO ADICIONAR PROFESSOR À TURMA:", erro)
        flash(f"Não foi possível salvar o vínculo: {erro}", "erro")
        return redirect(f"/turmas/{turma_id}#professores")

    finally:
        banco.close()


@app.route(
    "/turmas/<int:turma_id>/professores/<int:vinculo_id>/editar",
    methods=["POST"]
)
def editar_professor_turma(turma_id, vinculo_id):
    if not _pode_gerenciar_professores_turma():
        flash(
            "Você não possui permissão para gerenciar professores da turma.",
            "erro"
        )
        return redirect("/acesso_negado")

    professor_id = request.form.get("professor_id", type=int)
    componente_id = request.form.get("componente_id", type=int)

    if not professor_id or not componente_id:
        flash("Selecione o professor e o componente curricular.", "erro")
        return redirect(f"/turmas/{turma_id}#professores")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        cursor.execute("PRAGMA foreign_keys = ON")

        turma, erro = _buscar_turma_para_gerenciar_vinculos(
            cursor,
            turma_id
        )

        if erro:
            flash(erro, "erro")
            return redirect("/turmas")

        cursor.execute("""
            SELECT id
            FROM professor_vinculos
            WHERE id = ?
              AND turma_id = ?
            LIMIT 1
        """, (
            vinculo_id,
            turma_id
        ))

        if cursor.fetchone() is None:
            flash("Vínculo não encontrado nesta turma.", "erro")
            return redirect(f"/turmas/{turma_id}#professores")

        professor, componente, erro = (
            _validar_professor_e_componente_turma(
                cursor,
                turma,
                professor_id,
                componente_id
            )
        )

        if erro:
            flash(erro, "erro")
            return redirect(f"/turmas/{turma_id}#professores")

        cursor.execute("""
            SELECT id
            FROM professor_vinculos
            WHERE professor_id = ?
              AND turma_id = ?
              AND componente_id = ?
              AND id != ?
            LIMIT 1
        """, (
            professor_id,
            turma_id,
            componente_id,
            vinculo_id
        ))

        if cursor.fetchone():
            flash(
                "Este professor já está vinculado a esse componente nesta turma.",
                "erro"
            )
            return redirect(f"/turmas/{turma_id}#professores")

        cursor.execute("""
            UPDATE professor_vinculos
            SET
                professor_id = ?,
                componente_id = ?
            WHERE id = ?
              AND turma_id = ?
        """, (
            professor_id,
            componente_id,
            vinculo_id,
            turma_id
        ))

        banco.commit()

        flash(
            f"Vínculo de {professor['nome']} com "
            f"{componente['nome']} atualizado com sucesso.",
            "success"
        )

        return redirect(f"/turmas/{turma_id}#professores")

    except sqlite3.IntegrityError:
        banco.rollback()
        flash(
            "Este professor já está vinculado a esse componente nesta turma.",
            "erro"
        )
        return redirect(f"/turmas/{turma_id}#professores")

    except sqlite3.Error as erro:
        banco.rollback()
        print("ERRO AO EDITAR VÍNCULO DA TURMA:", erro)
        flash(f"Não foi possível atualizar o vínculo: {erro}", "erro")
        return redirect(f"/turmas/{turma_id}#professores")

    finally:
        banco.close()


@app.route(
    "/turmas/<int:turma_id>/professores/<int:vinculo_id>/excluir",
    methods=["POST"]
)
def excluir_professor_turma(turma_id, vinculo_id):
    if not _pode_gerenciar_professores_turma():
        flash(
            "Você não possui permissão para gerenciar professores da turma.",
            "erro"
        )
        return redirect("/acesso_negado")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        cursor.execute("PRAGMA foreign_keys = ON")

        turma, erro = _buscar_turma_para_gerenciar_vinculos(
            cursor,
            turma_id
        )

        if erro:
            flash(erro, "erro")
            return redirect("/turmas")

        cursor.execute("""
            SELECT
                pv.id,
                u.nome AS professor_nome,
                cc.nome AS componente_nome
            FROM professor_vinculos AS pv
            INNER JOIN usuarios AS u
                ON u.id = pv.professor_id
            INNER JOIN componentes_curriculares AS cc
                ON cc.id = pv.componente_id
            WHERE pv.id = ?
              AND pv.turma_id = ?
            LIMIT 1
        """, (
            vinculo_id,
            turma_id
        ))

        vinculo = cursor.fetchone()

        if vinculo is None:
            flash("Vínculo não encontrado nesta turma.", "erro")
            return redirect(f"/turmas/{turma_id}#professores")

        cursor.execute("""
            DELETE FROM professor_vinculos
            WHERE id = ?
              AND turma_id = ?
        """, (
            vinculo_id,
            turma_id
        ))

        banco.commit()

        flash(
            f"O vínculo de {vinculo['professor_nome']} com "
            f"{vinculo['componente_nome']} foi removido.",
            "success"
        )

        return redirect(f"/turmas/{turma_id}#professores")

    except sqlite3.Error as erro:
        banco.rollback()
        print("ERRO AO EXCLUIR VÍNCULO DA TURMA:", erro)
        flash(f"Não foi possível excluir o vínculo: {erro}", "erro")
        return redirect(f"/turmas/{turma_id}#professores")

    finally:
        banco.close()


# =========================================================
# EDITAR TURMA
# =========================================================

@app.route(
    "/turmas/<int:turma_id>/editar",
    methods=["GET", "POST"]
)
def editar_turma(turma_id):

    if not cargo_permitido([
        "Administrador Geral",
        "Administrador da Instituição",
        "Coordenador",
        "Secretaria"
    ]):

        flash(
            "Você não possui permissão para editar turmas.",
            "erro"
        )

        return redirect("/acesso_negado")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    cargo = session.get("usuario_cargo", "").strip()
    escola_id = session.get("escola_id")

    try:

        cursor.execute("""
            SELECT *
            FROM turmas
            WHERE id = ?
            LIMIT 1
        """, (
            turma_id,
        ))

        turma = cursor.fetchone()

        if turma is None:

            flash(
                "Turma não encontrada.",
                "erro"
            )

            return redirect("/turmas")

        if (
            cargo != "Administrador Geral"
            and turma["escola_id"] != escola_id
        ):

            flash(
                "Você não possui acesso a esta turma.",
                "erro"
            )

            return redirect("/turmas")

        if request.method == "POST":

            etapa = request.form.get("etapa", "").strip()
            ano = request.form.get("ano", "").strip()
            nome = request.form.get("nome", "").strip()
            turno = request.form.get("turno", "").strip()

            if not etapa or not ano or not nome or not turno:

                flash(
                    "Preencha todos os dados da turma.",
                    "erro"
                )

                return redirect(
                    f"/turmas/{turma_id}/editar"
                )

            cursor.execute("""
                SELECT id
                FROM turmas
                WHERE LOWER(TRIM(nome)) = LOWER(TRIM(?))
                  AND LOWER(TRIM(etapa)) = LOWER(TRIM(?))
                  AND LOWER(TRIM(ano)) = LOWER(TRIM(?))
                  AND LOWER(TRIM(turno)) = LOWER(TRIM(?))
                  AND escola_id = ?
                  AND id != ?
                LIMIT 1
            """, (
                nome,
                etapa,
                ano,
                turno,
                turma["escola_id"],
                turma_id
            ))

            if cursor.fetchone():

                flash(
                    "Já existe outra turma com esses mesmos dados.",
                    "erro"
                )

                return redirect(
                    f"/turmas/{turma_id}/editar"
                )

            cursor.execute("""
                UPDATE turmas
                SET
                    nome = ?,
                    etapa = ?,
                    ano = ?,
                    turno = ?
                WHERE id = ?
            """, (
                nome,
                etapa,
                ano,
                turno,
                turma_id
            ))

            banco.commit()

            flash(
                "Turma atualizada com sucesso.",
                "success"
            )

            return redirect(
                f"/turmas/{turma_id}"
            )

        return render_template(
            "gestao/editar_turma.html",
            turma=turma
        )

    except sqlite3.Error as erro:

        banco.rollback()

        import traceback
        traceback.print_exc()

        print("ERRO AO EDITAR TURMA:", erro)

        flash(
            f"Erro ao editar turma: {erro}",
            "erro"
        )

        return redirect("/turmas")

    finally:
        banco.close()


# =========================================================
# EXCLUIR TURMA
# =========================================================

@app.route(
    "/turmas/<int:turma_id>/excluir",
    methods=["POST"]
)
def excluir_turma(turma_id):

    if not cargo_permitido([
        "Administrador Geral",
        "Administrador da Instituição",
        "Coordenador",
        "Secretaria"
    ]):

        flash(
            "Você não possui permissão para excluir turmas.",
            "erro"
        )

        return redirect("/acesso_negado")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    cargo = session.get("usuario_cargo", "").strip()
    escola_id = session.get("escola_id")

    try:

        cursor.execute("""
            SELECT *
            FROM turmas
            WHERE id = ?
            LIMIT 1
        """, (
            turma_id,
        ))

        turma = cursor.fetchone()

        if turma is None:

            flash(
                "Turma não encontrada.",
                "erro"
            )

            return redirect("/turmas")

        if (
            cargo != "Administrador Geral"
            and turma["escola_id"] != escola_id
        ):

            flash(
                "Você não possui permissão para excluir esta turma.",
                "erro"
            )

            return redirect("/turmas")

        cursor.execute("""
            SELECT COUNT(*) AS total
            FROM alunos
            WHERE turma_id = ?
        """, (
            turma_id,
        ))

        total_alunos = cursor.fetchone()["total"]

        cursor.execute("""
            SELECT COUNT(*) AS total
            FROM provas
            WHERE turma_id = ?
        """, (
            turma_id,
        ))

        total_provas = cursor.fetchone()["total"]

        if total_alunos > 0 or total_provas > 0:

            flash(
                "Não é possível excluir a turma porque existem alunos ou avaliações vinculados.",
                "erro"
            )

            return redirect("/turmas")

        cursor.execute("""
            DELETE FROM professor_vinculos
            WHERE turma_id = ?
        """, (
            turma_id,
        ))

        cursor.execute("""
            DELETE FROM professor_turmas
            WHERE turma_id = ?
        """, (
            turma_id,
        ))

        cursor.execute("""
            DELETE FROM coordenador_turmas
            WHERE turma_id = ?
        """, (
            turma_id,
        ))

        cursor.execute("""
            DELETE FROM turmas
            WHERE id = ?
        """, (
            turma_id,
        ))

        banco.commit()

        flash(
            "Turma excluída com sucesso.",
            "success"
        )

        return redirect("/turmas")

    except sqlite3.Error as erro:

        banco.rollback()

        import traceback
        traceback.print_exc()

        print("ERRO AO EXCLUIR TURMA:", erro)

        flash(
            f"Erro ao excluir turma: {erro}",
            "erro"
        )

        return redirect("/turmas")

    finally:
        banco.close()

# =========================================================
# LISTAR ALUNOS E MATRÍCULAS PELO ANO SELECIONADO
# =========================================================

@app.route("/alunos")
def alunos():
    """Lista alunos e prepara o formulário de matrícula por instituição."""

    if not permissao_modulo("Alunos"):
        return redirect("/acesso_negado")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    contexto = obter_contexto_plataforma()
    cargo = contexto["cargo"]
    escola_id = contexto["escola_id"]
    ano_letivo_id = contexto["ano_letivo_id"]
    ano_visualizado = contexto["ano"]
    consultando_historico = contexto["consultando_historico"]

    lista_escolas = []
    lista_turmas = []
    lista_alunos = []

    def renderizar(**extras):
        dados = {
            "alunos": lista_alunos,
            "turmas": lista_turmas,
            "escolas": lista_escolas,
            "cargo": cargo,
            "escola_id_usuario": escola_id,
            "ano_letivo_ativo": ano_visualizado,
            "ano_letivo_visualizado": ano_visualizado,
            "consultando_ano_antigo": consultando_historico,
        }
        dados.update(extras)
        return render_template("alunos.html", **dados)

    try:
        if cargo == "Administrador Geral":
            if ano_visualizado is None:
                ano_visualizado = obter_ano_global_administrador()

            if ano_visualizado is None:
                return renderizar(
                    ano_letivo_ativo=None,
                    ano_letivo_visualizado=None,
                    consultando_ano_antigo=False,
                )

            cursor.execute("""
                SELECT
                    COUNT(*) AS total_registros,
                    SUM(CASE WHEN ativo = 1 AND encerrado = 0 THEN 1 ELSE 0 END) AS total_ativos
                FROM anos_letivos
                WHERE ano = ?
            """, (ano_visualizado,))
            situacao_ano = cursor.fetchone()
            total_registros = int(situacao_ano["total_registros"] or 0) if situacao_ano else 0
            total_ativos = int(situacao_ano["total_ativos"] or 0) if situacao_ano else 0
            # Um ano sem nenhuma instituição cadastrada ainda não é um ano
            # "encerrado". Isso evita bloquear/confundir o Administrador Geral
            # em instalações novas, antes do primeiro cadastro de instituição.
            consultando_historico = total_registros > 0 and total_ativos == 0

            # Instituições disponíveis no ano selecionado.
            cursor.execute("""
                SELECT DISTINCT e.id, e.nome_instituicao
                FROM escolas AS e
                INNER JOIN anos_letivos AS al
                    ON al.escola_id = e.id
                   AND al.ano = ?
                WHERE COALESCE(e.status, 1) = 1
                ORDER BY e.nome_instituicao COLLATE NOCASE ASC
            """, (ano_visualizado,))
            lista_escolas = cursor.fetchall()

            # Todas as turmas são enviadas ao template com escola_id.
            # O JavaScript da página exibe somente as da instituição escolhida.
            cursor.execute("""
                SELECT
                    t.id, t.nome, t.etapa, t.ano, t.turno,
                    t.escola_id, t.ano_letivo_id,
                    e.nome_instituicao,
                    al.ano AS ano_letivo
                FROM turmas AS t
                INNER JOIN anos_letivos AS al
                    ON al.id = t.ano_letivo_id
                   AND al.escola_id = t.escola_id
                INNER JOIN escolas AS e ON e.id = t.escola_id
                WHERE al.ano = ?
                ORDER BY
                    e.nome_instituicao COLLATE NOCASE ASC,
                    t.etapa COLLATE NOCASE ASC,
                    t.ano COLLATE NOCASE ASC,
                    t.nome COLLATE NOCASE ASC,
                    t.turno COLLATE NOCASE ASC
            """, (ano_visualizado,))
            lista_turmas = cursor.fetchall()

            cursor.execute("""
                SELECT * FROM (
                    SELECT
                        a.id, a.nome, a.matricula,
                        am.turma_id, am.escola_id, am.ano_letivo_id,
                        am.id AS matricula_id, am.situacao,
                        t.nome AS nome_turma, t.ano AS ano_turma,
                        t.etapa, t.turno,
                        e.nome_instituicao, al.ano AS ano_letivo
                    FROM aluno_matriculas AS am
                    INNER JOIN alunos AS a ON a.id = am.aluno_id
                    INNER JOIN turmas AS t
                        ON t.id = am.turma_id
                       AND t.escola_id = am.escola_id
                       AND t.ano_letivo_id = am.ano_letivo_id
                    INNER JOIN anos_letivos AS al
                        ON al.id = am.ano_letivo_id
                       AND al.escola_id = am.escola_id
                    INNER JOIN escolas AS e ON e.id = am.escola_id
                    WHERE al.ano = ?

                    UNION ALL

                    SELECT
                        a.id, a.nome, a.matricula,
                        a.turma_id, a.escola_id, a.ano_letivo_id,
                        NULL AS matricula_id, 'Cursando' AS situacao,
                        t.nome AS nome_turma, t.ano AS ano_turma,
                        t.etapa, t.turno,
                        e.nome_instituicao, al.ano AS ano_letivo
                    FROM alunos AS a
                    INNER JOIN turmas AS t
                        ON t.id = a.turma_id
                       AND t.escola_id = a.escola_id
                       AND t.ano_letivo_id = a.ano_letivo_id
                    INNER JOIN anos_letivos AS al
                        ON al.id = a.ano_letivo_id
                       AND al.escola_id = a.escola_id
                    INNER JOIN escolas AS e ON e.id = a.escola_id
                    WHERE al.ano = ?
                      AND NOT EXISTS (
                          SELECT 1 FROM aluno_matriculas AS am
                          WHERE am.aluno_id = a.id
                            AND am.ano_letivo_id = a.ano_letivo_id
                      )
                ) AS alunos_ano
                ORDER BY nome_instituicao COLLATE NOCASE ASC, nome COLLATE NOCASE ASC
            """, (ano_visualizado, ano_visualizado))
            lista_alunos = cursor.fetchall()

        else:
            if not escola_id:
                flash("Não foi possível identificar sua instituição.", "erro")
                return renderizar(
                    ano_letivo_ativo=None,
                    ano_letivo_visualizado=None,
                    consultando_ano_antigo=False,
                )

            ano_selecionado = atualizar_ano_letivo_na_sessao(escola_id)
            if not ano_selecionado:
                flash("A instituição não possui um ano letivo disponível.", "erro")
                return renderizar(
                    ano_letivo_ativo=None,
                    ano_letivo_visualizado=None,
                    consultando_ano_antigo=False,
                )

            ano_letivo_id = ano_selecionado["id"]
            ano_visualizado = ano_selecionado["ano"]
            consultando_historico = not (
                ano_selecionado["ativo"] == 1 and ano_selecionado["encerrado"] == 0
            )

            session["ano_letivo_id"] = ano_letivo_id
            session["ano_letivo"] = ano_visualizado
            session["ano_letivo_visualizado"] = ano_visualizado

            # Para usuários institucionais, somente a própria instituição.
            cursor.execute("""
                SELECT id, nome_instituicao
                FROM escolas
                WHERE id = ? AND COALESCE(status, 1) = 1
                LIMIT 1
            """, (escola_id,))
            escola_usuario = cursor.fetchone()
            lista_escolas = [escola_usuario] if escola_usuario else []

            cursor.execute("""
                SELECT
                    t.id, t.nome, t.etapa, t.ano, t.turno,
                    t.escola_id, t.ano_letivo_id,
                    e.nome_instituicao,
                    al.ano AS ano_letivo
                FROM turmas AS t
                INNER JOIN anos_letivos AS al
                    ON al.id = t.ano_letivo_id
                   AND al.escola_id = t.escola_id
                INNER JOIN escolas AS e ON e.id = t.escola_id
                WHERE t.escola_id = ? AND t.ano_letivo_id = ?
                ORDER BY
                    t.etapa COLLATE NOCASE ASC,
                    t.ano COLLATE NOCASE ASC,
                    t.nome COLLATE NOCASE ASC,
                    t.turno COLLATE NOCASE ASC
            """, (escola_id, ano_letivo_id))
            lista_turmas = cursor.fetchall()

            cursor.execute("""
                SELECT * FROM (
                    SELECT
                        a.id, a.nome, a.matricula,
                        am.turma_id, am.escola_id, am.ano_letivo_id,
                        am.id AS matricula_id, am.situacao,
                        t.nome AS nome_turma, t.ano AS ano_turma,
                        t.etapa, t.turno,
                        e.nome_instituicao, al.ano AS ano_letivo
                    FROM aluno_matriculas AS am
                    INNER JOIN alunos AS a ON a.id = am.aluno_id
                    INNER JOIN turmas AS t
                        ON t.id = am.turma_id
                       AND t.escola_id = am.escola_id
                       AND t.ano_letivo_id = am.ano_letivo_id
                    INNER JOIN anos_letivos AS al
                        ON al.id = am.ano_letivo_id
                       AND al.escola_id = am.escola_id
                    INNER JOIN escolas AS e ON e.id = am.escola_id
                    WHERE am.escola_id = ? AND am.ano_letivo_id = ?

                    UNION ALL

                    SELECT
                        a.id, a.nome, a.matricula,
                        a.turma_id, a.escola_id, a.ano_letivo_id,
                        NULL AS matricula_id, 'Cursando' AS situacao,
                        t.nome AS nome_turma, t.ano AS ano_turma,
                        t.etapa, t.turno,
                        e.nome_instituicao, al.ano AS ano_letivo
                    FROM alunos AS a
                    INNER JOIN turmas AS t
                        ON t.id = a.turma_id
                       AND t.escola_id = a.escola_id
                       AND t.ano_letivo_id = a.ano_letivo_id
                    INNER JOIN anos_letivos AS al
                        ON al.id = a.ano_letivo_id
                       AND al.escola_id = a.escola_id
                    INNER JOIN escolas AS e ON e.id = a.escola_id
                    WHERE a.escola_id = ? AND a.ano_letivo_id = ?
                      AND NOT EXISTS (
                          SELECT 1 FROM aluno_matriculas AS am
                          WHERE am.aluno_id = a.id
                            AND am.ano_letivo_id = a.ano_letivo_id
                      )
                ) AS alunos_ano
                ORDER BY nome COLLATE NOCASE ASC
            """, (escola_id, ano_letivo_id, escola_id, ano_letivo_id))
            lista_alunos = cursor.fetchall()

        return renderizar(
            ano_letivo_ativo=ano_visualizado,
            ano_letivo_visualizado=ano_visualizado,
            consultando_ano_antigo=consultando_historico,
        )

    except sqlite3.Error as erro:
        import traceback
        traceback.print_exc()
        print("ERRO AO LISTAR ALUNOS:", erro)
        flash(f"Erro ao carregar os alunos: {erro}", "erro")
        lista_escolas = []
        lista_turmas = []
        lista_alunos = []
        return renderizar()

    finally:
        banco.close()


# =========================================================
# GERAR NÚMERO DE MATRÍCULA AUTOMATICAMENTE
# =========================================================

def gerar_numero_matricula(cursor, escola_id, numero_ano):
    """
    Gera uma matrícula no formato AAAA + sequência de 4 dígitos.

    Exemplo:
        20260001
        20260002

    A sequência é independente para cada instituição e ano letivo.
    """

    prefixo = str(numero_ano)

    cursor.execute("""
        SELECT matricula
        FROM alunos
        WHERE escola_id = ?
          AND matricula LIKE ?
        ORDER BY id ASC
    """, (
        escola_id,
        f"{prefixo}%"
    ))

    maior_sequencia = 0

    for registro in cursor.fetchall():
        matricula_existente = str(registro["matricula"] or "").strip()

        if not matricula_existente.startswith(prefixo):
            continue

        sufixo = matricula_existente[len(prefixo):]

        if sufixo.isdigit():
            maior_sequencia = max(
                maior_sequencia,
                int(sufixo)
            )

    proxima_sequencia = maior_sequencia + 1

    while True:
        matricula = f"{prefixo}{proxima_sequencia:04d}"

        cursor.execute("""
            SELECT id
            FROM alunos
            WHERE escola_id = ?
              AND matricula = ?
            LIMIT 1
        """, (
            escola_id,
            matricula
        ))

        if cursor.fetchone() is None:
            return matricula

        proxima_sequencia += 1


# =========================================================
# IMPORTAÇÃO DE ALUNOS POR PDF
# =========================================================

IMPORTACAO_ALUNOS_DIR = os.path.join(tempfile.gettempdir(), "arkedu_importacoes_alunos")
os.makedirs(IMPORTACAO_ALUNOS_DIR, exist_ok=True)


def _normalizar_nome_importacao(valor):
    """Limpa espaços sem alterar acentos ou a grafia original do estudante."""
    return re.sub(r"\s+", " ", str(valor or "")).strip()


def _extrair_alunos_texto_pdf(texto):
    """
    Extrai pares matrícula/nome de relações alfabéticas escolares.

    O formato principal esperado é semelhante a:
        2712   1   ANA JÚLIA SANTOS FRAGOSO

    Há também suporte para PDFs cujo extrator quebra as três colunas
    em linhas consecutivas (matrícula, número de ordem e nome).
    """
    if not texto:
        return []

    linhas = [
        re.sub(r"\s+", " ", linha.replace("\xa0", " ")).strip()
        for linha in texto.splitlines()
    ]
    linhas = [linha for linha in linhas if linha]

    encontrados = []
    matriculas_vistas = set()

    # 1) Caso mais comum: todas as colunas permanecem na mesma linha.
    padrao_linha = re.compile(
        r"^(?P<matricula>[A-Za-z0-9./_-]{2,30})\s+"
        r"(?P<ordem>\d{1,4})\s+"
        r"(?P<nome>[A-Za-zÀ-ÖØ-öø-ÿ][A-Za-zÀ-ÖØ-öø-ÿ'’ .-]{2,})$"
    )

    cabecalhos = {
        "matricula", "matrícula", "nº", "no", "n°", "nome do aluno",
        "nome", "aluno", "turma", "turno", "curso", "serie", "série"
    }

    for linha in linhas:
        m = padrao_linha.match(linha)
        if not m:
            continue
        matricula = m.group("matricula").strip()
        nome = _normalizar_nome_importacao(m.group("nome"))
        if nome.casefold() in cabecalhos or len(nome) < 3:
            continue
        chave = matricula.casefold()
        if chave not in matriculas_vistas:
            encontrados.append({"matricula": matricula, "nome": nome})
            matriculas_vistas.add(chave)

    # 2) Alguns PDFs devolvem "matrícula", "ordem" e "nome" em linhas separadas.
    if not encontrados:
        for i in range(len(linhas) - 2):
            matricula = linhas[i]
            ordem = linhas[i + 1]
            nome = _normalizar_nome_importacao(linhas[i + 2])

            if not re.fullmatch(r"[A-Za-z0-9./_-]{2,30}", matricula):
                continue
            if not re.fullmatch(r"\d{1,4}", ordem):
                continue
            if not re.fullmatch(r"[A-Za-zÀ-ÖØ-öø-ÿ][A-Za-zÀ-ÖØ-öø-ÿ'’ .-]{2,}", nome):
                continue
            if nome.casefold() in cabecalhos:
                continue

            chave = matricula.casefold()
            if chave not in matriculas_vistas:
                encontrados.append({"matricula": matricula, "nome": nome})
                matriculas_vistas.add(chave)

    return encontrados


def _salvar_previa_importacao(usuario_id, turma_id, alunos):
    token = uuid.uuid4().hex
    caminho = os.path.join(IMPORTACAO_ALUNOS_DIR, f"{token}.json")
    dados = {
        "usuario_id": usuario_id,
        "turma_id": turma_id,
        "criado_em": time.time(),
        "alunos": alunos,
    }
    with open(caminho, "w", encoding="utf-8") as arquivo:
        json.dump(dados, arquivo, ensure_ascii=False)
    return token


def _carregar_previa_importacao(token):
    if not token or not re.fullmatch(r"[a-f0-9]{32}", token):
        return None, None
    caminho = os.path.join(IMPORTACAO_ALUNOS_DIR, f"{token}.json")
    if not os.path.exists(caminho):
        return None, None
    try:
        with open(caminho, "r", encoding="utf-8") as arquivo:
            dados = json.load(arquivo)
    except (OSError, json.JSONDecodeError):
        return None, None

    # Prévia válida por 30 minutos.
    if time.time() - float(dados.get("criado_em", 0)) > 1800:
        try:
            os.remove(caminho)
        except OSError:
            pass
        return None, None
    return dados, caminho


@app.route("/alunos/importar", methods=["GET"])
def importar_alunos_pagina():
    if not cargo_permitido([
        "Administrador Geral",
        "Administrador da Instituição",
        "Coordenador",
        "Secretaria",
        "Professor Autônomo"
    ]):
        return redirect("/acesso_negado")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()
    cargo = session.get("usuario_cargo", "").strip()
    escola_id_usuario = obter_escola_usuario()
    contexto = obter_contexto_plataforma()
    ano_visualizado = contexto.get("ano")

    try:
        parametros = []
        filtros = ["al.ativo = 1", "al.encerrado = 0"]

        if cargo != "Administrador Geral":
            if not escola_id_usuario:
                flash("Não foi possível identificar sua instituição.", "erro")
                return redirect("/alunos")
            filtros.append("t.escola_id = ?")
            parametros.append(escola_id_usuario)

        if ano_visualizado:
            filtros.append("al.ano = ?")
            parametros.append(ano_visualizado)

        cursor.execute(f"""
            SELECT
                t.id, t.nome, t.ano, t.turno, t.etapa,
                t.escola_id, t.ano_letivo_id,
                e.nome_instituicao,
                al.ano AS ano_letivo
            FROM turmas AS t
            INNER JOIN anos_letivos AS al
                ON al.id = t.ano_letivo_id
               AND al.escola_id = t.escola_id
            INNER JOIN escolas AS e ON e.id = t.escola_id
            WHERE {' AND '.join(filtros)}
            ORDER BY e.nome_instituicao COLLATE NOCASE,
                     t.ano COLLATE NOCASE,
                     t.nome COLLATE NOCASE
        """, parametros)
        turmas = cursor.fetchall()
    finally:
        banco.close()

    return render_template(
        "importar_alunos.html",
        turmas=turmas,
        turma_preselecionada=request.args.get("turma_id", ""),
        ano_letivo_visualizado=ano_visualizado,
        etapa="upload"
    )


@app.route("/alunos/importar/pdf", methods=["POST"])
def importar_alunos_pdf():
    if not cargo_permitido([
        "Administrador Geral",
        "Administrador da Instituição",
        "Coordenador",
        "Secretaria",
        "Professor Autônomo"
    ]):
        return redirect("/acesso_negado")

    arquivo = request.files.get("arquivo_pdf")
    turma_id = request.form.get("turma_id", type=int)

    if not turma_id:
        flash("Selecione a turma que receberá os alunos.", "erro")
        return redirect("/alunos/importar")
    if not arquivo or not arquivo.filename:
        flash("Selecione um arquivo PDF.", "erro")
        return redirect("/alunos/importar")
    if not arquivo.filename.lower().endswith(".pdf"):
        flash("O arquivo precisa estar no formato PDF.", "erro")
        return redirect("/alunos/importar")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()
    cargo = session.get("usuario_cargo", "").strip()
    escola_id_usuario = obter_escola_usuario()

    try:
        cursor.execute("""
            SELECT t.*, e.nome_instituicao, al.ano AS ano_letivo,
                   al.ativo, al.encerrado
            FROM turmas t
            INNER JOIN escolas e ON e.id = t.escola_id
            INNER JOIN anos_letivos al
                ON al.id = t.ano_letivo_id
               AND al.escola_id = t.escola_id
            WHERE t.id = ?
            LIMIT 1
        """, (turma_id,))
        turma = cursor.fetchone()

        if not turma:
            flash("A turma selecionada não existe.", "erro")
            return redirect("/alunos/importar")
        if cargo != "Administrador Geral" and turma["escola_id"] != escola_id_usuario:
            return redirect("/acesso_negado")
        if turma["ativo"] != 1 or turma["encerrado"] == 1:
            flash("Só é possível importar alunos para o ano letivo ativo.", "erro")
            return redirect("/alunos/importar")

        if PdfReader is None:
            flash("A biblioteca de leitura de PDF ainda não está instalada. Execute: pip install -r requirements.txt", "erro")
            return redirect("/alunos/importar")

        try:
            leitor = PdfReader(arquivo.stream)
            textos = []
            for pagina in leitor.pages:
                try:
                    textos.append(pagina.extract_text() or "")
                except Exception:
                    textos.append("")
            texto = "\n".join(textos)
        except Exception:
            flash("Não foi possível ler este PDF. Verifique se o arquivo não está corrompido ou protegido por senha.", "erro")
            return redirect("/alunos/importar")

        alunos_extraidos = _extrair_alunos_texto_pdf(texto)

        if not alunos_extraidos:
            flash(
                "Não encontrei matrícula e nome dos alunos neste PDF. "
                "Ele pode ser uma digitalização/imagem ou ter um formato diferente.",
                "erro"
            )
            return redirect("/alunos/importar")

        token = _salvar_previa_importacao(
            session.get("usuario_id"),
            turma_id,
            alunos_extraidos
        )

        return render_template(
            "importar_alunos.html",
            etapa="revisao",
            turma=turma,
            alunos_extraidos=alunos_extraidos,
            token=token,
            ano_letivo_visualizado=turma["ano_letivo"]
        )
    finally:
        banco.close()


@app.route("/alunos/importar/confirmar", methods=["POST"])
def confirmar_importacao_alunos():
    if not cargo_permitido([
        "Administrador Geral",
        "Administrador da Instituição",
        "Coordenador",
        "Secretaria",
        "Professor Autônomo"
    ]):
        return redirect("/acesso_negado")

    token = request.form.get("token", "").strip()
    previa, caminho_previa = _carregar_previa_importacao(token)
    if not previa:
        flash("A prévia da importação expirou. Envie o PDF novamente.", "erro")
        return redirect("/alunos/importar")
    if previa.get("usuario_id") != session.get("usuario_id"):
        return redirect("/acesso_negado")

    turma_id = int(previa.get("turma_id"))
    matriculas = request.form.getlist("matricula")
    nomes = request.form.getlist("nome")
    incluir = set(request.form.getlist("incluir"))

    # Limita a confirmação à quantidade originalmente extraída.
    limite = len(previa.get("alunos", []))
    matriculas = matriculas[:limite]
    nomes = nomes[:limite]

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()
    cargo = session.get("usuario_cargo", "").strip()
    escola_id_usuario = obter_escola_usuario()

    importados = 0
    ja_matriculados = 0
    ignorados = 0
    erros = 0
    vistos = set()

    try:
        cursor.execute("""
            SELECT t.id, t.escola_id, t.ano_letivo_id,
                   al.ano, al.ativo, al.encerrado
            FROM turmas t
            INNER JOIN anos_letivos al
                ON al.id = t.ano_letivo_id
               AND al.escola_id = t.escola_id
            WHERE t.id = ?
            LIMIT 1
        """, (turma_id,))
        turma = cursor.fetchone()
        if not turma:
            flash("A turma selecionada não existe mais.", "erro")
            return redirect("/alunos/importar")
        if cargo != "Administrador Geral" and turma["escola_id"] != escola_id_usuario:
            return redirect("/acesso_negado")
        if turma["ativo"] != 1 or turma["encerrado"] == 1:
            flash("O ano letivo foi encerrado e a importação não pode ser concluída.", "erro")
            return redirect("/alunos")

        escola_id = turma["escola_id"]
        ano_letivo_id = turma["ano_letivo_id"]

        for indice, (matricula_bruta, nome_bruto) in enumerate(zip(matriculas, nomes)):
            if str(indice) not in incluir:
                ignorados += 1
                continue

            matricula = str(matricula_bruta or "").strip()
            nome = _normalizar_nome_importacao(nome_bruto)
            chave = matricula.casefold()

            if not matricula or not nome or len(nome) < 3:
                erros += 1
                continue
            if chave in vistos:
                ignorados += 1
                continue
            vistos.add(chave)

            cursor.execute("""
                SELECT id
                FROM alunos
                WHERE escola_id = ?
                  AND LOWER(TRIM(matricula)) = LOWER(TRIM(?))
                ORDER BY id ASC
                LIMIT 1
            """, (escola_id, matricula))
            existente = cursor.fetchone()

            if existente:
                aluno_id = existente["id"]
                cursor.execute("""
                    SELECT id
                    FROM aluno_matriculas
                    WHERE aluno_id = ? AND ano_letivo_id = ?
                    LIMIT 1
                """, (aluno_id, ano_letivo_id))
                if cursor.fetchone():
                    ja_matriculados += 1
                    continue

                cursor.execute("""
                    UPDATE alunos
                    SET nome = ?, turma_id = ?, ano_letivo_id = ?
                    WHERE id = ?
                """, (nome, turma_id, ano_letivo_id, aluno_id))
            else:
                cursor.execute("""
                    INSERT INTO alunos (nome, matricula, turma_id, escola_id, ano_letivo_id)
                    VALUES (?, ?, ?, ?, ?)
                """, (nome, matricula, turma_id, escola_id, ano_letivo_id))
                aluno_id = cursor.lastrowid

            cursor.execute("""
                INSERT INTO aluno_matriculas
                    (aluno_id, escola_id, ano_letivo_id, turma_id, situacao)
                VALUES (?, ?, ?, ?, 'Cursando')
            """, (aluno_id, escola_id, ano_letivo_id, turma_id))
            importados += 1

        banco.commit()
    except Exception as erro:
        banco.rollback()
        app.logger.exception("Erro ao importar alunos por PDF")
        flash(f"A importação não pôde ser concluída: {erro}", "erro")
        return redirect("/alunos/importar")
    finally:
        banco.close()
        if caminho_previa:
            try:
                os.remove(caminho_previa)
            except OSError:
                pass

    partes = [f"{importados} aluno(s) importado(s)"]
    if ja_matriculados:
        partes.append(f"{ja_matriculados} já matriculado(s) no ano")
    if ignorados:
        partes.append(f"{ignorados} ignorado(s)")
    if erros:
        partes.append(f"{erros} com dados inválidos")
    flash("Importação concluída: " + ", ".join(partes) + ".", "sucesso")
    return redirect("/alunos")


# =========================================================
# CADASTRAR ALUNO E CRIAR MATRÍCULA ANUAL
# =========================================================

@app.route("/cadastrar_aluno", methods=["POST"])
def cadastrar_aluno():

    if not cargo_permitido([
        "Administrador Geral",
        "Administrador da Instituição",
        "Coordenador",
        "Secretaria",
        "Professor Autônomo"
    ]):
        return redirect("/acesso_negado")

    nome = request.form.get("nome", "").strip()
    modo_matricula = request.form.get(
        "modo_matricula",
        "automatica"
    ).strip().lower()
    matricula = request.form.get("matricula", "").strip()
    turma_id = request.form.get("turma_id", "").strip()
    escola_id_form = request.form.get("escola_id", "").strip()
    origem_turma_id = request.form.get("origem_turma_id", type=int)

    destino_retorno = (
        f"/turmas/{origem_turma_id}#estudantes"
        if origem_turma_id
        else "/alunos"
    )

    cargo = session.get("usuario_cargo", "").strip()
    escola_id_usuario = obter_escola_usuario()

    if not nome:
        flash("Informe o nome do aluno.", "erro")
        return redirect(destino_retorno)

    if modo_matricula not in ["automatica", "manual"]:
        modo_matricula = "automatica"

    if modo_matricula == "manual" and not matricula:
        flash("Informe o número da matrícula.", "erro")
        return redirect(destino_retorno)

    if not turma_id:
        flash("Selecione uma turma.", "erro")
        return redirect(destino_retorno)

    try:
        turma_id = int(turma_id)
    except (TypeError, ValueError):
        flash("A turma selecionada é inválida.", "erro")
        return redirect(destino_retorno)

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:

        cursor.execute("""
            SELECT
                t.id,
                t.escola_id,
                t.ano_letivo_id,
                al.ano,
                al.ativo,
                al.encerrado
            FROM turmas AS t
            INNER JOIN anos_letivos AS al
                ON al.id = t.ano_letivo_id
               AND al.escola_id = t.escola_id
            WHERE t.id = ?
            LIMIT 1
        """, (turma_id,))

        turma = cursor.fetchone()

        if turma is None:
            flash("A turma selecionada não existe.", "erro")
            return redirect(destino_retorno)

        escola_id = turma["escola_id"]
        ano_letivo_id = turma["ano_letivo_id"]
        numero_ano = turma["ano"]

        if cargo == "Administrador Geral":
            if not escola_id_form:
                flash("Selecione uma instituição.", "erro")
                return redirect(destino_retorno)
            try:
                escola_id_selecionada = int(escola_id_form)
            except (TypeError, ValueError):
                flash("A instituição selecionada é inválida.", "erro")
                return redirect(destino_retorno)
            if escola_id_selecionada != escola_id:
                flash(
                    "A turma selecionada não pertence à instituição informada.",
                    "erro"
                )
                return redirect(destino_retorno)

        if (
            cargo != "Administrador Geral"
            and escola_id != escola_id_usuario
        ):
            flash(
                "A turma selecionada não pertence à sua instituição.",
                "erro"
            )
            return redirect(destino_retorno)

        # Matrículas e novos cadastros só podem ser feitos no ano ativo.
        if turma["ativo"] != 1 or turma["encerrado"] == 1:
            flash(
                f"O ano letivo {numero_ano} está em modo de consulta. "
                "Volte ao ano ativo para matricular alunos.",
                "erro"
            )
            return redirect(destino_retorno)

        # Confirma que a turma pertence ao ano exibido no topo.
        if cargo == "Administrador Geral":
            ano_topo = session.get("ano_letivo_visualizado")
            if ano_topo and int(ano_topo) != int(numero_ano):
                flash(
                    "A turma não pertence ao ano letivo selecionado.",
                    "erro"
                )
                return redirect(destino_retorno)
        else:
            ano_selecionado = atualizar_ano_letivo_na_sessao(escola_id)
            if (
                not ano_selecionado
                or ano_selecionado["id"] != ano_letivo_id
            ):
                flash(
                    "A turma não pertence ao ano letivo selecionado.",
                    "erro"
                )
                return redirect(destino_retorno)

        # Gera o número somente depois de identificar corretamente
        # a instituição e o ano letivo da turma selecionada.
        if modo_matricula == "automatica":
            matricula = gerar_numero_matricula(
                cursor,
                escola_id,
                numero_ano
            )

        # A matrícula identifica o cadastro permanente do estudante
        # dentro da instituição. Se já existir, apenas criamos o novo
        # vínculo anual; não duplicamos o aluno.
        cursor.execute("""
            SELECT id, nome
            FROM alunos
            WHERE escola_id = ?
              AND LOWER(TRIM(matricula)) = LOWER(TRIM(?))
            ORDER BY id ASC
            LIMIT 1
        """, (escola_id, matricula))

        aluno_existente = cursor.fetchone()

        if aluno_existente:
            aluno_id = aluno_existente["id"]

            cursor.execute("""
                SELECT id
                FROM aluno_matriculas
                WHERE aluno_id = ?
                  AND ano_letivo_id = ?
                LIMIT 1
            """, (aluno_id, ano_letivo_id))

            if cursor.fetchone():
                flash(
                    "Este aluno já possui matrícula neste ano letivo.",
                    "erro"
                )
                return redirect(destino_retorno)

            # Mantém o nome atualizado quando o cadastro foi localizado
            # pela matrícula permanente.
            cursor.execute("""
                UPDATE alunos
                SET
                    nome = ?,
                    turma_id = ?,
                    ano_letivo_id = ?
                WHERE id = ?
            """, (
                nome,
                turma_id,
                ano_letivo_id,
                aluno_id
            ))

        else:
            cursor.execute("""
                INSERT INTO alunos (
                    nome,
                    matricula,
                    turma_id,
                    escola_id,
                    ano_letivo_id
                )
                VALUES (?, ?, ?, ?, ?)
            """, (
                nome,
                matricula,
                turma_id,
                escola_id,
                ano_letivo_id
            ))

            aluno_id = cursor.lastrowid

        cursor.execute("""
            INSERT INTO aluno_matriculas (
                aluno_id,
                escola_id,
                ano_letivo_id,
                turma_id,
                situacao
            )
            VALUES (?, ?, ?, ?, 'Cursando')
        """, (
            aluno_id,
            escola_id,
            ano_letivo_id,
            turma_id
        ))

        banco.commit()

        flash(
            f"Aluno matriculado com sucesso. Número da matrícula: "
            f"{matricula}.",
            "success"
        )

        return redirect(destino_retorno)

    except sqlite3.IntegrityError as erro:

        banco.rollback()
        print("ERRO DE INTEGRIDADE AO MATRICULAR ALUNO:", erro)

        flash(
            "Não foi possível concluir a matrícula. Verifique se o "
            "aluno já está matriculado neste ano letivo.",
            "erro"
        )

        return redirect(destino_retorno)

    except sqlite3.Error as erro:

        banco.rollback()

        import traceback
        traceback.print_exc()

        print("ERRO AO MATRICULAR ALUNO:", erro)

        flash(
            f"Erro ao matricular aluno: {erro}",
            "erro"
        )

        return redirect(destino_retorno)

    finally:
        banco.close()


# =========================================================
# HISTÓRICO DE MATRÍCULAS DO ALUNO
# =========================================================

@app.route("/alunos/<int:aluno_id>")
def historico_aluno(aluno_id):

    if not permissao_modulo("Alunos"):
        return redirect("/acesso_negado")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    cargo = session.get("usuario_cargo", "").strip()
    escola_id_usuario = obter_escola_usuario()

    try:
        cursor.execute("""
            SELECT
                a.id,
                a.nome,
                a.matricula,
                a.escola_id,
                e.nome_instituicao
            FROM alunos AS a
            LEFT JOIN escolas AS e
                ON e.id = a.escola_id
            WHERE a.id = ?
            LIMIT 1
        """, (aluno_id,))

        aluno = cursor.fetchone()

        if aluno is None:
            flash("Aluno não encontrado.", "erro")
            return redirect("/alunos")

        if (
            cargo != "Administrador Geral"
            and aluno["escola_id"] != escola_id_usuario
        ):
            flash("Você não possui acesso a este aluno.", "erro")
            return redirect("/alunos")

        cursor.execute("""
            SELECT
                am.id,
                am.ano_letivo_id,
                am.turma_id,
                am.situacao,
                am.data_matricula,
                am.data_encerramento,
                am.observacao,
                al.ano AS ano_letivo,
                al.ativo,
                al.encerrado,
                t.nome AS nome_turma,
                t.ano AS ano_turma,
                t.etapa,
                t.turno
            FROM aluno_matriculas AS am
            INNER JOIN anos_letivos AS al
                ON al.id = am.ano_letivo_id
            INNER JOIN turmas AS t
                ON t.id = am.turma_id
            WHERE am.aluno_id = ?
            ORDER BY al.ano DESC
        """, (aluno_id,))

        historico = cursor.fetchall()

        # Compatibilidade para um registro antigo que ainda não tenha
        # sido inserido em aluno_matriculas.
        if not historico:
            cursor.execute("""
                SELECT
                    NULL AS id,
                    a.ano_letivo_id,
                    a.turma_id,
                    'Cursando' AS situacao,
                    NULL AS data_matricula,
                    NULL AS data_encerramento,
                    NULL AS observacao,
                    al.ano AS ano_letivo,
                    al.ativo,
                    al.encerrado,
                    t.nome AS nome_turma,
                    t.ano AS ano_turma,
                    t.etapa,
                    t.turno
                FROM alunos AS a
                INNER JOIN anos_letivos AS al
                    ON al.id = a.ano_letivo_id
                INNER JOIN turmas AS t
                    ON t.id = a.turma_id
                WHERE a.id = ?
                LIMIT 1
            """, (aluno_id,))

            registro_antigo = cursor.fetchone()
            historico = [registro_antigo] if registro_antigo else []

        return render_template(
            "aluno_historico.html",
            aluno=aluno,
            historico=historico
        )

    except sqlite3.Error as erro:
        import traceback
        traceback.print_exc()

        print("ERRO AO CARREGAR HISTÓRICO DO ALUNO:", erro)
        flash(f"Erro ao carregar o histórico do aluno: {erro}", "erro")
        return redirect("/alunos")

    finally:
        banco.close()


@app.route("/professores")
def professores():

    if not permissao_modulo("Professores"):
        return redirect("/acesso_negado")

    banco = conectar_banco()
    cursor = banco.cursor()

    cursor.execute("""
        SELECT 
            professores.id,
            professores.nome,
            professores.email,
            COALESCE(
                GROUP_CONCAT(DISTINCT professor_disciplinas.disciplina),
                professores.disciplina,
                ''
            ) AS disciplinas
        FROM professores
        LEFT JOIN professor_disciplinas 
            ON professores.id = professor_disciplinas.professor_id
        GROUP BY professores.id
        ORDER BY professores.nome
    """)
    lista_professores = cursor.fetchall()

    cursor.execute("""
        SELECT *
        FROM turmas
        ORDER BY nome
    """)
    lista_turmas = cursor.fetchall()

    banco.close()

    return render_template(
        "professores.html",
        professores=lista_professores,
        turmas=lista_turmas
    )

@app.route("/cadastrar_professor", methods=["POST"])
def cadastrar_professor():

    if not cargo_permitido([
        "Administrador Geral",
        "Administrador da Instituição",
        "Coordenador"
    ]):
        return redirect("/login")

    nome = request.form["nome"]
    email = request.form["email"]

    disciplinas = request.form.getlist("disciplinas")
    turmas = request.form.getlist("turmas")

    disciplina_principal = ", ".join(disciplinas)

    banco = conectar_banco()
    cursor = banco.cursor()

    cursor.execute("""
        INSERT INTO professores (nome, email, disciplina)
        VALUES (?, ?, ?)
    """, (nome, email, disciplina_principal))

    professor_id = cursor.lastrowid

    for disciplina in disciplinas:
        cursor.execute("""
            INSERT INTO professor_disciplinas (professor_id, disciplina)
            VALUES (?, ?)
        """, (professor_id, disciplina))

    for turma_id in turmas:
        cursor.execute("""
            INSERT INTO professor_turmas (professor_id, turma_id)
            VALUES (?, ?)
        """, (professor_id, turma_id))

    banco.commit()
    banco.close()

    return redirect("/professores")

@app.route("/questoes")
def questoes():

    if not permissao_modulo("Questões"):
        return redirect("/acesso_negado")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    cargo = (session.get("usuario_cargo") or "").strip()
    usuario_id = session.get("usuario_id")
    escola_id = obter_escola_usuario()

    busca = (request.args.get("busca") or "").strip()
    disciplina = (request.args.get("disciplina") or "").strip()
    etapa = (request.args.get("etapa") or "").strip()
    ano_serie = (request.args.get("ano_serie") or "").strip()
    assunto = (request.args.get("assunto") or "").strip()
    tipo = (request.args.get("tipo") or "").strip()
    dificuldade = (request.args.get("dificuldade") or "").strip()
    bloom = (request.args.get("bloom") or "").strip()
    origem = (request.args.get("origem") or "todas").strip()

    pagina = max(request.args.get("pagina", default=1, type=int) or 1, 1)
    por_pagina = 12
    deslocamento = (pagina - 1) * por_pagina

    filtros = []
    parametros = []

    # Banco de questões compartilhado: todos os usuários podem consultar
    # questões criadas por qualquer professor ou instituição.

    if origem == "minhas":
        filtros.append("q.criado_por = ?")
        parametros.append(usuario_id)

    if busca:
        termo = f"%{busca}%"
        filtros.append("(" + " OR ".join([
            "q.enunciado LIKE ?", "q.assunto LIKE ?", "q.subassunto LIKE ?",
            "q.habilidade_bncc LIKE ?", "q.descritor_saeb LIKE ?", "q.tags LIKE ?"
        ]) + ")")
        parametros.extend([termo] * 6)

    campos = [
        (disciplina, "q.disciplina = ?"),
        (etapa, "q.etapa_ensino = ?"),
        (ano_serie, "q.ano_serie = ?"),
        (assunto, "q.assunto = ?"),
        (tipo, "q.tipo_questao = ?"),
        (dificuldade, "q.dificuldade = ?"),
        (bloom, "q.taxonomia_bloom = ?")
    ]
    for valor, sql in campos:
        if valor:
            filtros.append(sql)
            parametros.append(valor)

    where = "WHERE " + " AND ".join(filtros) if filtros else ""

    try:
        cursor.execute(f"SELECT COUNT(*) AS total FROM questoes q {where}", parametros)
        total = cursor.fetchone()["total"]
        total_paginas = max((total + por_pagina - 1) // por_pagina, 1)
        if pagina > total_paginas:
            pagina = total_paginas
            deslocamento = (pagina - 1) * por_pagina

        cursor.execute(f"""
            SELECT
                q.*,
                COALESCE(u.nome, 'ARK EDUS') AS autor_nome,
                COALESCE(e.nome_instituicao, 'Banco compartilhado') AS instituicao_nome,
                (SELECT COUNT(*) FROM prova_questoes pq WHERE pq.questao_id = q.id) AS total_usos
            FROM questoes q
            LEFT JOIN usuarios u ON u.id = q.criado_por
            LEFT JOIN escolas e ON e.id = q.escola_id
            {where}
            ORDER BY COALESCE(q.atualizado_em, q.criado_em) DESC, q.id DESC
            LIMIT ? OFFSET ?
        """, parametros + [por_pagina, deslocamento])
        lista_questoes = []
        cargos_administrativos = {
            "Administrador Geral",
            "Administrador da Instituição"
        }

        for registro in cursor.fetchall():
            item = dict(registro)
            item["pode_excluir"] = (
                cargo in cargos_administrativos
                or item.get("criado_por") == usuario_id
            )

            # Um usuário de instituição nunca pode excluir questão de outra escola.
            if (
                cargo != "Administrador Geral"
                and item.get("escola_id") not in (None, escola_id)
            ):
                item["pode_excluir"] = False

            lista_questoes.append(item)

        def distintos(campo):
            cursor.execute(f"""
                SELECT DISTINCT {campo} AS valor
                FROM questoes
                WHERE {campo} IS NOT NULL
                  AND TRIM({campo}) <> ''
                ORDER BY {campo} COLLATE NOCASE
            """)
            return [linha["valor"] for linha in cursor.fetchall()]

        opcoes = {
            "disciplinas": distintos("disciplina"),
            "etapas": distintos("etapa_ensino"),
            "anos": distintos("ano_serie"),
            "assuntos": distintos("assunto"),
            "blooms": distintos("taxonomia_bloom")
        }

        cursor.execute("SELECT COUNT(*) AS total FROM questoes")
        total_geral = cursor.fetchone()["total"]

        cursor.execute("SELECT COUNT(*) AS total FROM questoes WHERE criado_por = ?", (usuario_id,))
        total_minhas = cursor.fetchone()["total"]

        cursor.execute("SELECT COUNT(*) AS total FROM questoes WHERE tipo_questao IN ('multipla_escolha','multiplas_respostas','verdadeiro_falso')")
        total_objetivas = cursor.fetchone()["total"]

        cursor.execute("SELECT COUNT(*) AS total FROM questoes WHERE tipo_questao IN ('discursiva','resposta_curta','numerica')")
        total_discursivas = cursor.fetchone()["total"]

        return render_template(
            "questoes/index.html",
            questoes=lista_questoes,
            opcoes=opcoes,
            total=total,
            total_geral=total_geral,
            total_minhas=total_minhas,
            total_objetivas=total_objetivas,
            total_discursivas=total_discursivas,
            pagina=pagina,
            total_paginas=total_paginas,
            filtros_atuais={
                "busca": busca, "disciplina": disciplina, "etapa": etapa,
                "ano_serie": ano_serie, "assunto": assunto, "tipo": tipo,
                "dificuldade": dificuldade, "bloom": bloom, "origem": origem
            }
        )

    except sqlite3.Error as erro:
        print("ERRO AO CARREGAR BANCO DE QUESTÕES:", erro)
        flash("Não foi possível carregar o banco de questões.", "erro")
        return render_template(
            "questoes/index.html", questoes=[], opcoes={}, total=0,
            total_geral=0, total_minhas=0, total_objetivas=0,
            total_discursivas=0, pagina=1, total_paginas=1,
            filtros_atuais={}
        )
    finally:
        banco.close()


@app.route("/questoes/<int:questao_id>")
def visualizar_questao(questao_id):
    if not permissao_modulo("Questões"):
        return redirect("/acesso_negado")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        cursor.execute("""
            SELECT
                q.*,
                COALESCE(u.nome, 'ARK EDUS') AS autor_nome,
                COALESCE(e.nome_instituicao, 'Banco compartilhado') AS instituicao_nome,
                (
                    SELECT COUNT(*)
                    FROM prova_questoes AS pq
                    WHERE pq.questao_id = q.id
                ) AS total_usos
            FROM questoes AS q
            LEFT JOIN usuarios AS u
                ON u.id = q.criado_por
            LEFT JOIN escolas AS e
                ON e.id = q.escola_id
            WHERE q.id = ?
            LIMIT 1
        """, (questao_id,))

        questao = cursor.fetchone()

        if not questao:
            flash("Questão não encontrada.", "erro")
            return redirect("/questoes")

        cargo = (session.get("usuario_cargo") or "").strip()
        escola_usuario = obter_escola_usuario()

        if (
            cargo != "Administrador Geral"
            and questao["escola_id"] is not None
            and escola_usuario is not None
            and int(questao["escola_id"]) != int(escola_usuario)
        ):
            return redirect("/acesso_negado")

        try:
            alternativas = json.loads(questao["alternativas_json"] or "[]")
        except (TypeError, ValueError, json.JSONDecodeError):
            alternativas = []

        if not alternativas:
            alternativas = []
            for indice, campo in enumerate(
                ["alternativa_a", "alternativa_b", "alternativa_c", "alternativa_d"]
            ):
                texto = questao[campo] or ""
                if texto:
                    alternativas.append({
                        "letra": chr(65 + indice),
                        "texto": texto,
                        "imagem": ""
                    })

        try:
            respostas_corretas = json.loads(
                questao["respostas_corretas"] or "[]"
            )
        except (TypeError, ValueError, json.JSONDecodeError):
            respostas_corretas = []

        if not respostas_corretas and questao["correta"]:
            respostas_corretas = [questao["correta"]]

        tags = [
            item.strip()
            for item in (questao["tags"] or "").split(",")
            if item.strip()
        ]

        pode_editar = (
            cargo in [
                "Administrador Geral",
                "Administrador da Instituição",
                "Coordenador"
            ]
            or questao["criado_por"] == session.get("usuario_id")
        )

        return render_template(
            "questoes/visualizar.html",
            questao=questao,
            alternativas=alternativas,
            respostas_corretas=respostas_corretas,
            tags=tags,
            pode_editar=pode_editar
        )

    except sqlite3.Error as erro:
        print("ERRO AO VISUALIZAR QUESTÃO:", erro)
        flash("Não foi possível abrir a questão.", "erro")
        return redirect("/questoes")

    finally:
        banco.close()


@app.route("/api/bncc/opcoes")
def api_bncc_opcoes():
    if not permissao_modulo("Questões"):
        return jsonify({"erro": "Acesso negado."}), 403

    etapa = (request.args.get("etapa") or "").strip()
    componente = (request.args.get("componente") or "").strip()
    ano_serie = (request.args.get("ano_serie") or "").strip()
    unidade = (request.args.get("unidade") or "").strip()
    objeto = (request.args.get("objeto") or "").strip()

    if not etapa or not componente:
        return jsonify({"unidades": [], "objetos": [], "habilidades": []})

    try:
        return jsonify(consultar_bncc(
            DB_PATH, etapa=etapa, componente=componente, ano_serie=ano_serie,
            unidade=unidade, objeto=objeto
        ))
    except sqlite3.Error as erro:
        print("ERRO AO CONSULTAR BNCC:", erro)
        return jsonify({"erro": "Não foi possível consultar o catálogo da BNCC."}), 500


def _mapear_bncc_dev_item(item):
    """Normaliza registros da API pública bncc.dev para o formato usado pela ARK EDUS."""
    if not isinstance(item, dict):
        return None
    codigo = str(item.get("codigo") or item.get("code") or "").strip().upper()
    descricao = str(item.get("texto") or item.get("descricao") or item.get("habilidade") or "").strip()
    if not codigo or not descricao:
        return None

    organizacao = item.get("organizacao") or {}
    nomes = organizacao.get("nomes") if isinstance(organizacao, dict) else {}
    nomes = nomes if isinstance(nomes, dict) else {}
    unidade = (
        item.get("unidade_tematica") or item.get("unidadeTematica") or
        item.get("pratica_linguagem") or item.get("praticaLinguagem") or
        nomes.get("unidadeTematica") or nomes.get("praticaLinguagem") or
        nomes.get("praticaDeLinguagem") or nomes.get("campoAtuacao") or ""
    )

    objetos = (
        item.get("objetosConhecimento") or item.get("objetos_conhecimento") or
        item.get("objeto_conhecimento") or item.get("objetoConhecimento") or []
    )
    if isinstance(objetos, dict):
        objetos = [objetos]
    elif isinstance(objetos, str):
        objetos = [objetos]
    nomes_objetos = []
    for objeto in objetos if isinstance(objetos, list) else []:
        if isinstance(objeto, dict):
            nome = objeto.get("nome") or objeto.get("texto") or objeto.get("descricao")
        else:
            nome = str(objeto or "")
        if nome and str(nome).strip():
            nomes_objetos.append(str(nome).strip())

    componente = item.get("componente") or {}
    if isinstance(componente, dict):
        componente_nome = componente.get("nome") or ""
    else:
        componente_nome = str(componente or "")

    return {
        "codigo": codigo,
        "descricao": descricao,
        "valor": f"{codigo} — {descricao}",
        "unidade_tematica": str(unidade or "").strip(),
        "objeto_conhecimento": " • ".join(nomes_objetos),
        "componente": str(componente_nome or "").strip(),
    }


def _buscar_bncc_remota(termo, limite=30):
    """Consulta bncc.dev quando o catálogo local ainda não foi importado no banco persistente."""
    termo = str(termo or "").strip()
    if len(termo) < 2:
        return []
    try:
        codigo = re.sub(r"[^A-Za-z0-9]", "", termo).upper()
        if re.fullmatch(r"(?:EF|EM)[A-Z0-9]{5,}", codigo):
            resposta = requests.get(
                f"https://api.bncc.dev/v1/aprendizagens/{codigo}", timeout=5
            )
            if resposta.ok:
                normalizado = _mapear_bncc_dev_item(resposta.json())
                return [normalizado] if normalizado else []

        resposta = requests.get(
            "https://api.bncc.dev/v1/busca",
            params={"q": termo}, timeout=5
        )
        if not resposta.ok:
            return []
        dados = resposta.json()
        if isinstance(dados, dict):
            itens = (
                dados.get("resultados") or dados.get("itens") or dados.get("items") or
                dados.get("dados") or dados.get("habilidades") or []
            )
        else:
            itens = dados
        resultado = []
        for item in itens if isinstance(itens, list) else []:
            normalizado = _mapear_bncc_dev_item(item)
            if normalizado:
                resultado.append(normalizado)
            if len(resultado) >= limite:
                break
        return resultado
    except Exception as erro:
        print("AVISO BNCC REMOTA:", erro)
        return []


@app.route("/api/bncc/buscar")
def api_bncc_buscar():
    if not permissao_modulo("Questões"):
        return jsonify({"erro": "Acesso negado."}), 403

    termo = (request.args.get("q") or "").strip()
    etapa = (request.args.get("etapa") or "").strip()
    componente = (request.args.get("componente") or "").strip()
    ano_serie = (request.args.get("ano_serie") or "").strip()
    if len(termo) < 2:
        return jsonify({"habilidades": []})

    def normalizar_busca(valor):
        txt = unicodedata.normalize("NFD", str(valor or "").casefold())
        return "".join(c for c in txt if unicodedata.category(c) != "Mn")

    chave = normalizar_busca(termo)
    habilidades = []
    vistos = set()

    # 1) Pesquisa ampla no catálogo local. Não exige que etapa/componente estejam
    # perfeitamente iguais ao cadastro, evitando o caso em que o professor digita
    # EF05LP06 e a lista fica vazia por diferença de nomenclatura.
    try:
        banco_bncc = sqlite3.connect(DB_PATH)
        banco_bncc.row_factory = sqlite3.Row
        cur_bncc = banco_bncc.cursor()
        cur_bncc.execute("""
            SELECT etapa_ensino, ano_serie, area_conhecimento, componente,
                   unidade_tematica, objeto_conhecimento, codigo, descricao
            FROM bncc_habilidades
            WHERE ativo = 1
              AND (LOWER(codigo) LIKE LOWER(?) OR LOWER(descricao) LIKE LOWER(?))
            ORDER BY
              CASE WHEN LOWER(codigo) = LOWER(?) THEN 0
                   WHEN LOWER(codigo) LIKE LOWER(?) THEN 1 ELSE 2 END,
              codigo COLLATE NOCASE
            LIMIT 80
        """, (f"%{termo}%", f"%{termo}%", termo, f"{termo}%"))
        linhas = cur_bncc.fetchall()
        banco_bncc.close()

        for r in linhas:
            alvo_contexto = True
            # Contexto é apenas preferência, nunca bloqueio absoluto.
            contexto_score = 0
            if componente and normalizar_busca(componente) in normalizar_busca(r["componente"] or r["area_conhecimento"]):
                contexto_score += 2
            if etapa and normalizar_busca(etapa).split("-")[0].strip() in normalizar_busca(r["etapa_ensino"]):
                contexto_score += 1
            if ano_serie and _ano_compativel_bncc_seguro(r["ano_serie"], ano_serie, etapa):
                contexto_score += 1
            item = {
                "codigo": r["codigo"],
                "descricao": r["descricao"],
                "valor": f"{r['codigo']} — {r['descricao']}",
                "unidade_tematica": r["unidade_tematica"] or "",
                "objeto_conhecimento": r["objeto_conhecimento"] or "",
                "componente": r["componente"] or "",
                "contexto_score": contexto_score,
            }
            codigo = str(item["codigo"]).upper()
            if codigo not in vistos:
                vistos.add(codigo)
                habilidades.append(item)
        habilidades.sort(key=lambda h: (-int(h.get("contexto_score", 0)), str(h.get("codigo", ""))))
        habilidades = habilidades[:30]
    except sqlite3.Error as erro:
        print("ERRO AO PESQUISAR BNCC LOCAL:", erro)

    # 2) Complementa com consulta contextual do catálogo local quando disponível.
    if len(habilidades) < 10 and etapa and componente:
        try:
            dados = consultar_bncc(DB_PATH, etapa=etapa, componente=componente, ano_serie=ano_serie)
            for h in dados.get("habilidades", []):
                alvo = normalizar_busca(f"{h.get('codigo','')} {h.get('descricao','')}")
                codigo = str(h.get("codigo") or "").upper()
                if chave in alvo and codigo and codigo not in vistos:
                    vistos.add(codigo)
                    habilidades.append(h)
                    if len(habilidades) >= 30:
                        break
        except sqlite3.Error as erro:
            print("ERRO AO COMPLEMENTAR BNCC LOCAL:", erro)

    # 3) Se o banco persistente ainda não recebeu o catálogo, usa a API pública
    # da BNCC como fallback. O resultado continua sendo apenas sugestão ao professor.
    if not habilidades:
        habilidades = _buscar_bncc_remota(termo, 30)

    for h in habilidades:
        h.pop("contexto_score", None)
    return jsonify({"habilidades": habilidades[:30]})


def _ano_compativel_bncc_seguro(ano_registro, ano_selecionado, etapa):
    """Comparação tolerante usada apenas para ordenar sugestões BNCC."""
    ar = re.sub(r"\D", "", str(ano_registro or ""))
    asel = re.sub(r"\D", "", str(ano_selecionado or ""))
    if not ar or not asel or "medio" in normalizar_sem_acentos(str(etapa or "")):
        return True
    return ar == asel or asel in ar


def normalizar_sem_acentos(valor):
    txt = unicodedata.normalize("NFD", str(valor or "").casefold())
    return "".join(c for c in txt if unicodedata.category(c) != "Mn")


@app.route("/api/matrizes/opcoes")
def api_matrizes_opcoes():
    if not permissao_modulo("Questões"):
        return jsonify({"erro": "Acesso negado."}), 403

    sistema = (request.args.get("sistema") or "").strip().upper()
    etapa = (request.args.get("etapa") or "").strip()
    componente = (request.args.get("componente") or "").strip()
    ano_serie = (request.args.get("ano_serie") or "").strip()
    matriz_id = request.args.get("matriz_id", type=int)

    if not sistema or not etapa or not componente or not ano_serie:
        return jsonify({
            "matrizes": [], "descritores": [], "disponivel": False,
            "mensagem": "Selecione matriz, componente, etapa e ano/série."
        })

    try:
        return jsonify(consultar_matrizes(
            DB_PATH,
            sistema=sistema,
            etapa=etapa,
            componente=componente,
            ano_serie=ano_serie,
            matriz_id=matriz_id
        ))
    except sqlite3.Error as erro:
        print("ERRO AO CONSULTAR MATRIZES:", erro)
        return jsonify({"erro": "Não foi possível consultar os descritores da matriz selecionada."}), 500


@app.route("/questoes/nova")
def nova_questao():
    if not permissao_modulo("Questões"):
        return redirect("/acesso_negado")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    componentes = []
    valores_anteriores = None
    prova_id = request.args.get("prova_id", type=int)
    escola_filtro_id = obter_escola_usuario()

    try:
        # Quando o cadastro foi aberto durante a montagem de uma avaliação,
        # valida o acesso e utiliza os dados da própria avaliação para
        # preencher automaticamente componente, etapa e ano/série.
        if prova_id:
            if not _pode_gerenciar_prova(
                cursor,
                prova_id,
                exigir_edicao=True
            ):
                return _redirecionar_acesso_negado_prova()

            cursor.execute("""
                SELECT
                    p.id,
                    p.disciplina,
                    COALESCE(p.escola_id, t.escola_id) AS escola_id,
                    t.etapa,
                    t.ano AS turma_ano
                FROM provas AS p
                INNER JOIN turmas AS t
                    ON t.id = p.turma_id
                WHERE p.id = ?
                LIMIT 1
            """, (prova_id,))
            prova_origem = cursor.fetchone()

            if not prova_origem:
                flash("Avaliação não encontrada.", "erro")
                return redirect("/provas")

            escola_filtro_id = prova_origem["escola_id"]
            valores_anteriores = {
                "disciplina": prova_origem["disciplina"] or "",
                "etapa_ensino": prova_origem["etapa"] or "",
                "ano_serie": prova_origem["turma_ano"] or ""
            }

        elif request.args.get("repetir") == "1":
            valores_anteriores = session.get("ultima_classificacao_questao")

        if escola_filtro_id:
            cursor.execute("""
                SELECT DISTINCT nome
                FROM componentes_curriculares
                WHERE ativo = 1
                  AND escola_id = ?
                  AND nome IS NOT NULL
                  AND TRIM(nome) <> ''
                ORDER BY nome COLLATE NOCASE
            """, (escola_filtro_id,))
        else:
            cursor.execute("""
                SELECT DISTINCT nome
                FROM componentes_curriculares
                WHERE ativo = 1
                  AND nome IS NOT NULL
                  AND TRIM(nome) <> ''
                ORDER BY nome COLLATE NOCASE
            """)

        componentes = [linha["nome"] for linha in cursor.fetchall()]

        # Garante que o componente da avaliação apareça mesmo em cadastros
        # antigos que ainda não estejam na tabela de componentes curriculares.
        if valores_anteriores:
            componente_prova = (valores_anteriores.get("disciplina") or "").strip()
            if componente_prova and componente_prova not in componentes:
                componentes.insert(0, componente_prova)

    except sqlite3.Error as erro:
        print("ERRO AO CARREGAR NOVA QUESTÃO:", erro)
        flash("Não foi possível carregar o cadastro da questão.", "erro")

    finally:
        banco.close()

    return render_template(
        "questoes/nova.html",
        componentes=componentes,
        valores_anteriores=valores_anteriores,
        prova_id=prova_id
    )


@app.route("/questoes/<int:questao_id>/excluir", methods=["POST"])
def excluir_questao(questao_id):
    if not permissao_modulo("Questões"):
        return redirect("/acesso_negado")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()
    try:
        cursor.execute("SELECT escola_id, criado_por FROM questoes WHERE id = ?", (questao_id,))
        questao = cursor.fetchone()
        if not questao:
            flash("Questão não encontrada.", "erro")
            return redirect("/questoes")

        cargo = (session.get("usuario_cargo") or "").strip()
        escola_id = obter_escola_usuario()
        usuario_id = session.get("usuario_id")
        pode_excluir = cargo in {
            "Administrador Geral",
            "Administrador da Instituição"
        }
        pode_excluir = pode_excluir or questao["criado_por"] == usuario_id

        if (
            cargo != "Administrador Geral"
            and questao["escola_id"] not in (None, escola_id)
        ):
            pode_excluir = False

        if not pode_excluir:
            return redirect("/acesso_negado")

        cursor.execute("DELETE FROM questoes WHERE id = ?", (questao_id,))
        banco.commit()
        flash("Questão excluída com sucesso.", "sucesso")
    except sqlite3.IntegrityError:
        banco.rollback()
        flash("Esta questão está vinculada a uma avaliação e não pode ser excluída.", "erro")
    except sqlite3.Error as erro:
        banco.rollback()
        flash(f"Não foi possível excluir a questão: {erro}", "erro")
    finally:
        banco.close()

    return redirect(request.referrer or "/questoes")


@app.route("/cadastrar_questao", methods=["POST"])
def cadastrar_questao():

    if not cargo_permitido([
        "Administrador Geral",
        "Administrador da Instituição",
        "Coordenador",
        "Professor"
    ]):
        return redirect("/login")

    if not permissao_modulo("Questões"):
        return redirect("/acesso_negado")

    disciplina = (request.form.get("disciplina") or "").strip()
    etapa_ensino = (request.form.get("etapa_ensino") or "").strip()
    ano_serie = (request.form.get("ano_serie") or "").strip()
    # Campos antigos mantidos vazios apenas para compatibilidade com o banco.
    assunto = ""
    assunto_temporario = False
    subassunto = ""
    tipo_questao = (request.form.get("tipo_questao") or "multipla_escolha").strip()
    enunciado = (request.form.get("enunciado") or "").strip()
    enunciado_html = (request.form.get("enunciado_html") or "").strip()
    dificuldade = (request.form.get("dificuldade") or "").strip()
    taxonomia_bloom = (request.form.get("taxonomia_bloom") or "").strip()
    unidade_tematica = (request.form.get("unidade_tematica") or "").strip()
    objeto_conhecimento = (request.form.get("objeto_conhecimento") or "").strip()
    habilidade_bncc = (request.form.get("habilidade_bncc") or "").strip()
    sistema_matriz = (request.form.get("sistema_matriz") or "").strip().upper()
    matriz_referencia = (request.form.get("matriz_referencia") or "").strip()
    descritor_saeb = (request.form.get("descritor_saeb") or "").strip()
    if sistema_matriz not in {"", "SAETO", "SAEB"}:
        sistema_matriz = ""
    fonte = (request.form.get("fonte") or "").strip()
    tags = (request.form.get("tags") or "").strip()
    observacoes = (request.form.get("observacoes") or "").strip()
    resposta_esperada = (request.form.get("resposta_esperada") or "").strip()
    criterios_correcao = (request.form.get("criterios_correcao") or "").strip()
    manter_classificacao = request.form.get("manter_classificacao") == "1"
    prova_id = request.form.get("prova_id", type=int)
    url_nova_questao = (
        f"/questoes/nova?prova_id={prova_id}"
        if prova_id
        else "/questoes/nova"
    )

    try:
        linhas_resposta = int(request.form.get("linhas_resposta") or 5)
    except (TypeError, ValueError):
        linhas_resposta = 5
    linhas_resposta = max(1, min(linhas_resposta, 30))

    try:
        ano_fonte = int(request.form.get("ano_fonte")) if request.form.get("ano_fonte") else None
    except (TypeError, ValueError):
        ano_fonte = None

    try:
        tempo_estimado = int(request.form.get("tempo_estimado")) if request.form.get("tempo_estimado") else None
    except (TypeError, ValueError):
        tempo_estimado = None

    # Mantém o campo legado para filtros e telas antigas.
    habilidade = " | ".join(
        item for item in [habilidade_bncc, descritor_saeb] if item
    )

    tipos_validos = {
        "multipla_escolha", "verdadeiro_falso", "multiplas_respostas",
        "discursiva", "resposta_curta", "numerica"
    }

    if tipo_questao not in tipos_validos:
        flash("Tipo de questão inválido.", "erro")
        return redirect(url_nova_questao)

    if not disciplina or not enunciado or not dificuldade:
        flash("Preencha o componente, o enunciado e a dificuldade.", "erro")
        return redirect(url_nova_questao)

    extensoes_permitidas = {".png", ".jpg", ".jpeg", ".webp"}

    def salvar_imagem(arquivo, prefixo):
        if not arquivo or not arquivo.filename:
            return ""
        nome_seguro = secure_filename(arquivo.filename)
        extensao = os.path.splitext(nome_seguro)[1].lower()
        if extensao not in extensoes_permitidas:
            raise ValueError("Envie apenas imagens PNG, JPG, JPEG ou WEBP.")
        nome_arquivo = f"{prefixo}_{uuid.uuid4().hex}{extensao}"
        arquivo.save(os.path.join(app.config["UPLOAD_FOLDER"], nome_arquivo))
        return nome_arquivo

    def salvar_imagens_embutidas(conteudo_html):
        """Converte imagens data URL do editor em arquivos reais no servidor."""
        if not conteudo_html:
            return ""

        padrao = re.compile(
            r'src=["\']data:image/(?P<tipo>png|jpeg|jpg|webp);base64,(?P<dados>[^"\']+)["\']',
            re.IGNORECASE
        )

        def substituir(match):
            tipo = match.group("tipo").lower()
            extensao = ".jpg" if tipo in {"jpg", "jpeg"} else f".{tipo}"
            try:
                dados = base64.b64decode(match.group("dados"), validate=True)
            except Exception as erro:
                raise ValueError("Uma das imagens inseridas no enunciado é inválida.") from erro

            if len(dados) > 8 * 1024 * 1024:
                raise ValueError("Cada imagem do enunciado deve ter no máximo 8 MB.")

            nome_arquivo = f"enunciado_{uuid.uuid4().hex}{extensao}"
            with open(os.path.join(app.config["UPLOAD_FOLDER"], nome_arquivo), "wb") as destino:
                destino.write(dados)
            return f'src="/uploads/{nome_arquivo}"'

        conteudo_html = padrao.sub(substituir, conteudo_html)
        conteudo_html = re.sub(r'<\s*(script|iframe|object|embed)[^>]*>.*?<\s*/\s*\1\s*>', '', conteudo_html, flags=re.I | re.S)
        conteudo_html = re.sub(r'\son[a-z]+\s*=\s*(["\']).*?\1', '', conteudo_html, flags=re.I | re.S)
        conteudo_html = re.sub(r'javascript\s*:', '', conteudo_html, flags=re.I)
        return conteudo_html

    try:
        enunciado_html = salvar_imagens_embutidas(enunciado_html)
        nome_imagem = salvar_imagem(request.files.get("imagem"), "questao")
    except ValueError as erro:
        flash(str(erro), "erro")
        return redirect(url_nova_questao)

    alternativas = []
    respostas_corretas = []

    if tipo_questao in {"multipla_escolha", "multiplas_respostas"}:
        textos = request.form.getlist("alternativas[]")
        imagens = request.files.getlist("imagens_alternativas[]")
        indices_corretos = set(request.form.getlist("corretas[]"))
        total_linhas = max(len(textos), len(imagens))

        for indice in range(total_linhas):
            texto = textos[indice].strip() if indice < len(textos) else ""
            arquivo_imagem = imagens[indice] if indice < len(imagens) else None
            try:
                imagem_alternativa = salvar_imagem(arquivo_imagem, "alternativa")
            except ValueError as erro:
                flash(f"Alternativa {indice + 1}: {erro}", "erro")
                return redirect(url_nova_questao)

            if not texto and not imagem_alternativa:
                continue

            letra = chr(65 + len(alternativas))
            alternativas.append({"letra": letra, "texto": texto, "imagem": imagem_alternativa})
            if str(indice) in indices_corretos:
                respostas_corretas.append(letra)

        if len(alternativas) < 2:
            flash("Cadastre pelo menos duas alternativas com texto ou imagem.", "erro")
            return redirect(url_nova_questao)
        if tipo_questao == "multipla_escolha" and len(respostas_corretas) != 1:
            flash("Marque exatamente uma alternativa correta.", "erro")
            return redirect(url_nova_questao)
        if tipo_questao == "multiplas_respostas" and not respostas_corretas:
            flash("Marque pelo menos uma alternativa correta.", "erro")
            return redirect(url_nova_questao)

    elif tipo_questao == "verdadeiro_falso":
        resposta_vf = (request.form.get("resposta_vf") or "").strip()
        if resposta_vf not in {"V", "F"}:
            flash("Selecione Verdadeiro ou Falso.", "erro")
            return redirect(url_nova_questao)
        alternativas = [
            {"letra": "V", "texto": "Verdadeiro", "imagem": ""},
            {"letra": "F", "texto": "Falso", "imagem": ""}
        ]
        respostas_corretas = [resposta_vf]

    elif tipo_questao in {"resposta_curta", "numerica"} and not resposta_esperada:
        flash("Informe a resposta esperada.", "erro")
        return redirect(url_nova_questao)

    textos_legados = []
    for item in alternativas[:4]:
        texto_legado = item.get("texto") or ""
        if not texto_legado and item.get("imagem"):
            texto_legado = "[Alternativa com imagem]"
        textos_legados.append(texto_legado)
    textos_legados += [""] * (4 - len(textos_legados))
    correta_legada = respostas_corretas[0] if respostas_corretas else ""

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        escola_questao_id = obter_escola_usuario()
        prova_destino = None

        if prova_id:
            if not _pode_gerenciar_prova(cursor, prova_id, exigir_edicao=True):
                return _redirecionar_acesso_negado_prova()

            cursor.execute("""
                SELECT
                    p.id,
                    p.disciplina,
                    COALESCE(p.escola_id, t.escola_id) AS escola_id
                FROM provas AS p
                INNER JOIN turmas AS t
                    ON t.id = p.turma_id
                WHERE p.id = ?
                LIMIT 1
            """, (prova_id,))
            prova_destino = cursor.fetchone()

            if not prova_destino:
                flash("A avaliação de destino não foi encontrada.", "erro")
                return redirect("/provas")

            if (prova_destino["disciplina"] or "").strip().lower() != disciplina.lower():
                flash(
                    "O componente da questão precisa ser o mesmo da avaliação.",
                    "erro"
                )
                return redirect(url_nova_questao)

            escola_questao_id = prova_destino["escola_id"]

        cursor.execute("""
            INSERT INTO questoes (
                disciplina, etapa_ensino, ano_serie, assunto,
                assunto_temporario, subassunto,
                tipo_questao, enunciado, enunciado_html, imagem,
                alternativa_a, alternativa_b, alternativa_c, alternativa_d,
                correta, alternativas_json, respostas_corretas,
                resposta_esperada, criterios_correcao, habilidade,
                habilidade_bncc, unidade_tematica, objeto_conhecimento,
                sistema_matriz, matriz_referencia, descritor_saeb, taxonomia_bloom,
                dificuldade, fonte, ano_fonte, tags, tempo_estimado,
                linhas_resposta, observacoes, escola_id, criado_por, criado_em
            )
            VALUES (
                ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?,
                ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?
            )
        """, (
            disciplina, etapa_ensino, ano_serie, assunto,
            1 if assunto_temporario else 0, subassunto,
            tipo_questao, enunciado, enunciado_html, nome_imagem,
            textos_legados[0], textos_legados[1], textos_legados[2], textos_legados[3],
            correta_legada, json.dumps(alternativas, ensure_ascii=False),
            json.dumps(respostas_corretas, ensure_ascii=False), resposta_esperada,
            criterios_correcao, habilidade, habilidade_bncc, unidade_tematica,
            objeto_conhecimento, sistema_matriz, matriz_referencia, descritor_saeb,
            taxonomia_bloom, dificuldade, fonte, ano_fonte, tags,
            tempo_estimado, linhas_resposta if tipo_questao == "discursiva" else None,
            observacoes, escola_questao_id, session.get("usuario_id"),
            datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        ))

        questao_id = cursor.lastrowid

        if prova_id:
            cursor.execute("""
                SELECT COALESCE(MAX(ordem), 0) + 1 AS proxima_ordem
                FROM prova_questoes
                WHERE prova_id = ?
            """, (prova_id,))
            proxima_ordem = cursor.fetchone()["proxima_ordem"]

            cursor.execute("""
                INSERT OR IGNORE INTO prova_questoes (
                    prova_id,
                    questao_id,
                    peso,
                    ordem
                )
                VALUES (?, ?, 0, ?)
            """, (prova_id, questao_id, proxima_ordem))

            cursor.execute("""
                UPDATE provas
                SET quantidade = (
                    SELECT COUNT(*) FROM prova_questoes WHERE prova_id = ?
                ), atualizado_em = ?
                WHERE id = ?
            """, (
                prova_id,
                datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                prova_id
            ))

        banco.commit()
        flash(
            "Questão cadastrada e adicionada à avaliação!" if prova_id
            else "Questão cadastrada com sucesso!",
            "sucesso"
        )

        if manter_classificacao:
            session["ultima_classificacao_questao"] = {
                "disciplina": disciplina,
                "etapa_ensino": etapa_ensino,
                "ano_serie": ano_serie,
                "dificuldade": dificuldade,
                "taxonomia_bloom": taxonomia_bloom,
                "unidade_tematica": unidade_tematica,
                "objeto_conhecimento": objeto_conhecimento,
                "habilidade_bncc": habilidade_bncc,
                "sistema_matriz": sistema_matriz,
                "matriz_referencia": matriz_referencia,
                "descritor_saeb": descritor_saeb,
                "fonte": fonte,
                "ano_fonte": ano_fonte or "",
                "tags": tags,
                "tempo_estimado": tempo_estimado or "",
                "abrir_avancado": bool(
                    unidade_tematica or objeto_conhecimento or habilidade_bncc or
                    matriz_referencia or descritor_saeb or fonte or tags
                )
            }
        else:
            session.pop("ultima_classificacao_questao", None)

    except sqlite3.Error as erro:
        banco.rollback()
        print("ERRO AO CADASTRAR QUESTÃO:", erro)
        flash(f"Não foi possível cadastrar a questão: {erro}", "erro")

    finally:
        banco.close()

    if prova_id:
        return redirect(f"/provas/{prova_id}/montar")

    return redirect("/questoes/nova?repetir=1" if manter_classificacao else "/questoes")


# =========================================================
# PROVAS — LISTAGEM, CADASTRO E MONTAGEM
# =========================================================

def _contexto_provas():
    """Retorna o contexto acadêmico e de acesso usado no módulo Provas."""
    contexto = obter_contexto_plataforma()

    return {
        "usuario_id": contexto.get("usuario_id"),
        "cargo": (contexto.get("cargo") or "").strip(),
        "escola_id": contexto.get("escola_id"),
        "ano_letivo_id": contexto.get("ano_letivo_id"),
        "ano": contexto.get("ano")
    }


def _professor_legado_do_usuario(cursor, usuario_id):
    """
    Compatibiliza o cadastro atual de usuários com a tabela professores,
    que ainda é utilizada pela tabela provas.
    """
    if not usuario_id:
        return None

    cursor.execute("""
        SELECT
            u.id AS usuario_id,
            u.nome,
            u.email,
            u.escola_id
        FROM usuarios AS u
        INNER JOIN cargos AS c
            ON c.id = u.cargo_id
        WHERE u.id = ?
          AND c.nome IN ('Professor', 'Professor Autônomo')
        LIMIT 1
    """, (usuario_id,))

    usuario = cursor.fetchone()

    if not usuario:
        return None

    cursor.execute("""
        SELECT id
        FROM professores
        WHERE escola_id = ?
          AND (
                LOWER(TRIM(COALESCE(email, ''))) =
                    LOWER(TRIM(COALESCE(?, '')))
                OR LOWER(TRIM(nome)) = LOWER(TRIM(?))
          )
        ORDER BY
            CASE
                WHEN LOWER(TRIM(COALESCE(email, ''))) =
                     LOWER(TRIM(COALESCE(?, '')))
                THEN 0
                ELSE 1
            END,
            id
        LIMIT 1
    """, (
        usuario["escola_id"],
        usuario["email"],
        usuario["nome"],
        usuario["email"]
    ))

    professor = cursor.fetchone()

    if professor:
        return professor["id"]

    cursor.execute("""
        INSERT INTO professores (
            nome,
            email,
            disciplina,
            escola_id
        )
        VALUES (?, ?, ?, ?)
    """, (
        usuario["nome"],
        usuario["email"],
        "",
        usuario["escola_id"]
    ))

    return cursor.lastrowid



def _sincronizar_professores_da_escola(cursor, escola_id=None):
    """Sincroniza usuários com cargo Professor com a tabela professores."""
    parametros = []
    filtro_escola = ""

    if escola_id:
        filtro_escola = " AND u.escola_id = ? "
        parametros.append(escola_id)

    cursor.execute(f"""
        SELECT u.id, u.nome, u.email, u.escola_id
        FROM usuarios AS u
        INNER JOIN cargos AS c ON c.id = u.cargo_id
        WHERE c.nome IN ('Professor', 'Professor Autônomo')
          AND u.ativo = 1
          AND u.escola_id IS NOT NULL
          {filtro_escola}
        ORDER BY u.nome COLLATE NOCASE
    """, parametros)

    for usuario in cursor.fetchall():
        cursor.execute("""
            SELECT id
            FROM professores
            WHERE escola_id = ?
              AND (
                    (
                        COALESCE(TRIM(?), '') <> ''
                        AND LOWER(TRIM(COALESCE(email, ''))) = LOWER(TRIM(?))
                    )
                    OR LOWER(TRIM(nome)) = LOWER(TRIM(?))
              )
            ORDER BY id
            LIMIT 1
        """, (
            usuario["escola_id"],
            usuario["email"],
            usuario["email"],
            usuario["nome"]
        ))

        existente = cursor.fetchone()

        if existente:
            cursor.execute("""
                UPDATE professores
                SET nome = ?, email = ?, escola_id = ?
                WHERE id = ?
            """, (
                usuario["nome"],
                usuario["email"],
                usuario["escola_id"],
                existente["id"]
            ))
        else:
            cursor.execute("""
                INSERT INTO professores (nome, email, disciplina, escola_id)
                VALUES (?, ?, '', ?)
            """, (
                usuario["nome"],
                usuario["email"],
                usuario["escola_id"]
            ))


def _pode_criar_prova(cargo):
    return cargo in {
        "Administrador Geral",
        "Administrador da Instituição",
        "Coordenador",
        "Professor",
        "Professor Autônomo"
    }


def _pode_gerenciar_prova(
    cursor,
    prova_id,
    exigir_edicao=False,
    permitir_finalizada=False
):
    """
    Valida o acesso no backend. Assim, alterar manualmente a URL não
    permite acessar uma avaliação de outra instituição ou professor.
    """
    contexto = _contexto_provas()
    cargo = contexto["cargo"]

    cursor.execute("""
        SELECT
            p.id,
            p.professor_id,
            p.escola_id,
            p.ano_letivo_id,
            p.status,
            t.escola_id AS turma_escola_id,
            t.ano_letivo_id AS turma_ano_letivo_id
        FROM provas AS p
        INNER JOIN turmas AS t
            ON t.id = p.turma_id
        WHERE p.id = ?
        LIMIT 1
    """, (prova_id,))

    prova = cursor.fetchone()

    if not prova:
        return False

    # A finalização bloqueia alterações na estrutura da avaliação, mas não
    # impede criar, corrigir, importar ou excluir uma aplicação vinculada.
    if (
        exigir_edicao
        and not permitir_finalizada
        and (prova["status"] or "rascunho").strip().lower() == "finalizada"
    ):
        return False

    if cargo == "Administrador Geral":
        return True

    escola_prova = prova["escola_id"] or prova["turma_escola_id"]

    if (
        not contexto["escola_id"]
        or int(escola_prova or 0) != int(contexto["escola_id"])
    ):
        return False

    ano_prova = prova["ano_letivo_id"] or prova["turma_ano_letivo_id"]

    if contexto["ano_letivo_id"] and ano_prova:
        if int(ano_prova) != int(contexto["ano_letivo_id"]):
            return False

    if cargo == "Professor":
        professor_id = _professor_legado_do_usuario(
            cursor,
            contexto["usuario_id"]
        )
        return bool(
            professor_id
            and int(prova["professor_id"] or 0) == int(professor_id)
        )

    if exigir_edicao:
        return cargo in {
            "Administrador da Instituição",
            "Coordenador",
            "Professor Autônomo"
        }

    return cargo in {
        "Administrador da Instituição",
        "Coordenador",
        "Secretaria",
        "Professor Autônomo"
    }


def _redirecionar_acesso_negado_prova():
    flash(
        "Você não possui permissão para acessar esta avaliação.",
        "erro"
    )
    return redirect("/provas")


@app.route("/provas")
def provas():
    if not permissao_modulo("Provas"):
        return redirect("/acesso_negado")

    contexto = _contexto_provas()
    cargo = contexto["cargo"]
    usuario_id = contexto["usuario_id"]
    escola_id = contexto["escola_id"]
    ano_letivo_id = contexto["ano_letivo_id"]
    ano_visualizado = contexto["ano"]

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    escolas = []
    turmas = []
    professores = []
    registros = []
    professor_logado_id = None
    vinculos_duplicacao = []

    try:
        if cargo == "Administrador Geral":
            _sincronizar_professores_da_escola(cursor)
            banco.commit()

            cursor.execute("""
                SELECT id, nome_instituicao
                FROM escolas
                WHERE COALESCE(status, 1) = 1
                ORDER BY nome_instituicao COLLATE NOCASE
            """)
            escolas = cursor.fetchall()

            if ano_visualizado is not None:
                cursor.execute("""
                    SELECT DISTINCT
                        t.id,
                        t.nome,
                        t.etapa,
                        t.ano,
                        t.turno,
                        t.escola_id,
                        e.nome_instituicao
                    FROM turmas AS t
                    INNER JOIN escolas AS e
                        ON e.id = t.escola_id
                    INNER JOIN anos_letivos AS al
                        ON al.id = t.ano_letivo_id
                       AND al.escola_id = t.escola_id
                    WHERE al.ano = ?
                    ORDER BY
                        e.nome_instituicao COLLATE NOCASE,
                        t.nome COLLATE NOCASE
                """, (ano_visualizado,))
                turmas = cursor.fetchall()

                cursor.execute("""
                    SELECT
                        pr.id,
                        pr.nome,
                        pr.escola_id,
                        e.nome_instituicao
                    FROM professores AS pr
                    INNER JOIN escolas AS e
                        ON e.id = pr.escola_id
                    WHERE COALESCE(e.status, 1) = 1
                    ORDER BY
                        e.nome_instituicao COLLATE NOCASE,
                        pr.nome COLLATE NOCASE
                """)
                professores = cursor.fetchall()

                cursor.execute("""
                    SELECT
                        p.id,
                        p.nome,
                        p.disciplina,
                        p.data_geracao,
                        p.data_aplicacao,
                        p.status,
                        p.media_ativa,
                        p.media_aprovacao,
                        p.escola_id,
                        t.nome AS turma_nome,
                        e.nome_instituicao,
                        COALESCE(pr.nome, 'Não informado') AS professor_nome,
                        COUNT(pq.id) AS quantidade_real
                    FROM provas AS p
                    INNER JOIN turmas AS t
                        ON t.id = p.turma_id
                    INNER JOIN escolas AS e
                        ON e.id = COALESCE(p.escola_id, t.escola_id)
                    INNER JOIN anos_letivos AS al
                        ON al.id = COALESCE(
                            p.ano_letivo_id,
                            t.ano_letivo_id
                        )
                    LEFT JOIN professores AS pr
                        ON pr.id = p.professor_id
                    LEFT JOIN prova_questoes AS pq
                        ON pq.prova_id = p.id
                    WHERE al.ano = ?
                    GROUP BY p.id
                    ORDER BY p.id DESC
                """, (ano_visualizado,))
                registros = cursor.fetchall()

        else:
            if not escola_id:
                flash(
                    "Não foi possível identificar sua instituição.",
                    "erro"
                )
            else:
                _sincronizar_professores_da_escola(cursor, escola_id)
                banco.commit()

                if not ano_letivo_id:
                    ano = atualizar_ano_letivo_na_sessao(escola_id)
                    ano_letivo_id = ano["id"] if ano else None

                if cargo == "Professor":
                    professor_logado_id = _professor_legado_do_usuario(
                        cursor,
                        usuario_id
                    )

                    cursor.execute("""
                        SELECT DISTINCT
                            t.id,
                            t.nome,
                            t.etapa,
                            t.ano,
                            t.turno,
                            t.escola_id
                        FROM turmas AS t
                        INNER JOIN professor_vinculos AS pv
                            ON pv.turma_id = t.id
                        WHERE pv.professor_id = ?
                          AND t.escola_id = ?
                          AND t.ano_letivo_id = ?
                        ORDER BY t.nome COLLATE NOCASE
                    """, (
                        usuario_id,
                        escola_id,
                        ano_letivo_id
                    ))
                    turmas = cursor.fetchall()

                    if professor_logado_id:
                        cursor.execute("""
                            SELECT id, nome, escola_id
                            FROM professores
                            WHERE id = ?
                        """, (professor_logado_id,))
                        professores = cursor.fetchall()

                        cursor.execute("""
                            SELECT
                                p.id,
                                p.nome,
                                p.disciplina,
                                p.data_geracao,
                                p.data_aplicacao,
                                p.status,
                                p.media_ativa,
                                p.media_aprovacao,
                                p.escola_id,
                                t.nome AS turma_nome,
                                e.nome_instituicao,
                                COALESCE(
                                    pr.nome,
                                    'Não informado'
                                ) AS professor_nome,
                                COUNT(pq.id) AS quantidade_real
                            FROM provas AS p
                            INNER JOIN turmas AS t
                                ON t.id = p.turma_id
                            INNER JOIN escolas AS e
                                ON e.id = t.escola_id
                            LEFT JOIN professores AS pr
                                ON pr.id = p.professor_id
                            LEFT JOIN prova_questoes AS pq
                                ON pq.prova_id = p.id
                            WHERE p.professor_id = ?
                              AND t.escola_id = ?
                              AND COALESCE(
                                  p.ano_letivo_id,
                                  t.ano_letivo_id
                              ) = ?
                            GROUP BY p.id
                            ORDER BY p.id DESC
                        """, (
                            professor_logado_id,
                            escola_id,
                            ano_letivo_id
                        ))
                        registros = cursor.fetchall()

                else:
                    cursor.execute("""
                        SELECT
                            id,
                            nome,
                            etapa,
                            ano,
                            turno,
                            escola_id
                        FROM turmas
                        WHERE escola_id = ?
                          AND ano_letivo_id = ?
                        ORDER BY nome COLLATE NOCASE
                    """, (escola_id, ano_letivo_id))
                    turmas = cursor.fetchall()

                    cursor.execute("""
                        SELECT id, nome, escola_id
                        FROM professores
                        WHERE escola_id = ?
                        ORDER BY nome COLLATE NOCASE
                    """, (escola_id,))
                    professores = cursor.fetchall()

                    cursor.execute("""
                        SELECT
                            p.id,
                            p.nome,
                            p.disciplina,
                            p.data_geracao,
                            p.data_aplicacao,
                            p.status,
                            p.media_ativa,
                            p.media_aprovacao,
                            p.escola_id,
                            t.nome AS turma_nome,
                            e.nome_instituicao,
                            COALESCE(
                                pr.nome,
                                'Não informado'
                            ) AS professor_nome,
                            COUNT(pq.id) AS quantidade_real
                        FROM provas AS p
                        INNER JOIN turmas AS t
                            ON t.id = p.turma_id
                        INNER JOIN escolas AS e
                            ON e.id = t.escola_id
                        LEFT JOIN professores AS pr
                            ON pr.id = p.professor_id
                        LEFT JOIN prova_questoes AS pq
                            ON pq.prova_id = p.id
                        WHERE t.escola_id = ?
                          AND COALESCE(
                              p.ano_letivo_id,
                              t.ano_letivo_id
                          ) = ?
                        GROUP BY p.id
                        ORDER BY p.id DESC
                    """, (escola_id, ano_letivo_id))
                    registros = cursor.fetchall()

        # Opções válidas para o modal de duplicação.
        # Relaciona o professor legado usado em provas ao usuário professor
        # e aos vínculos oficiais de turma/componente.
        if cargo == "Administrador Geral":
            filtro_vinculos = ""
            parametros_vinculos = ()
        elif escola_id:
            filtro_vinculos = " AND t.escola_id = ? "
            parametros_vinculos = (escola_id,)
        else:
            filtro_vinculos = " AND 1 = 0 "
            parametros_vinculos = ()

        cursor.execute(f"""
            SELECT DISTINCT
                dados.professor_id,
                dados.turma_id,
                dados.componente_id,
                dados.escola_id,
                dados.componente_nome
            FROM (
                SELECT
                    pr.id AS professor_id,
                    pv.turma_id,
                    pv.componente_id,
                    t.escola_id,
                    cc.nome AS componente_nome
                FROM professor_vinculos AS pv
                INNER JOIN usuarios AS u
                    ON u.id = pv.professor_id
                INNER JOIN professores AS pr
                    ON pr.escola_id = u.escola_id
                   AND (
                        (
                            TRIM(COALESCE(pr.email, '')) <> ''
                            AND TRIM(COALESCE(u.email, '')) <> ''
                            AND LOWER(TRIM(pr.email)) = LOWER(TRIM(u.email))
                        )
                        OR LOWER(TRIM(pr.nome)) = LOWER(TRIM(u.nome))
                   )
                INNER JOIN turmas AS t
                    ON t.id = pv.turma_id
                   AND t.escola_id = u.escola_id
                INNER JOIN componentes_curriculares AS cc
                    ON cc.id = pv.componente_id
                   AND cc.escola_id = t.escola_id
                   AND COALESCE(cc.ativo, 1) = 1
                WHERE COALESCE(u.ativo, 1) = 1

                UNION

                SELECT
                    pc.professor_id,
                    pc.turma_id,
                    pc.componente_id,
                    t.escola_id,
                    cc.nome AS componente_nome
                FROM professor_componentes AS pc
                INNER JOIN professores AS pr
                    ON pr.id = pc.professor_id
                INNER JOIN turmas AS t
                    ON t.id = pc.turma_id
                   AND t.escola_id = pr.escola_id
                INNER JOIN componentes_curriculares AS cc
                    ON cc.id = pc.componente_id
                   AND cc.escola_id = t.escola_id
                   AND COALESCE(cc.ativo, 1) = 1
            ) AS dados
            INNER JOIN professores AS pr_ordem
                ON pr_ordem.id = dados.professor_id
            INNER JOIN turmas AS t
                ON t.id = dados.turma_id
            WHERE 1 = 1
              {filtro_vinculos}
            ORDER BY
                dados.componente_nome COLLATE NOCASE,
                pr_ordem.nome COLLATE NOCASE
        """, parametros_vinculos)

        vinculos_duplicacao = [
            dict(linha)
            for linha in cursor.fetchall()
        ]

        hoje = datetime.now().date()
        lista_provas = []

        for registro in registros:
            data_aplicacao = (
                registro["data_aplicacao"] or ""
            ).strip()
            data_objeto = None

            for formato in ("%Y-%m-%d", "%d/%m/%Y"):
                if not data_aplicacao:
                    break
                try:
                    data_objeto = datetime.strptime(
                        data_aplicacao,
                        formato
                    ).date()
                    break
                except ValueError:
                    continue

            status_banco = _sincronizar_status_prova(
                cursor,
                registro["id"]
            ) or (registro["status"] or "rascunho").strip().lower()

            mapa_status = {
                "rascunho": ("Rascunho", "rascunho"),
                "agendada": ("Agendada", "agendada"),
                "em_correcao": ("Em correção", "em-correcao"),
                "em correção": ("Em correção", "em-correcao"),
                "aguardando correção": ("Em correção", "em-correcao"),
                "finalizada": ("Finalizada", "finalizada")
            }
            status, status_slug = mapa_status.get(
                status_banco,
                ("Rascunho", "rascunho")
            )

            lista_provas.append({
                "id": registro["id"],
                "nome": registro["nome"],
                "turma": registro["turma_nome"],
                "professor": registro["professor_nome"],
                "disciplina": registro["disciplina"],
                "quantidade": registro["quantidade_real"] or 0,
                "data_geracao": (
                    registro["data_geracao"]
                    or "Não informada"
                ),
                "data_aplicacao": (
                    data_objeto.strftime("%d/%m/%Y")
                    if data_objeto
                    else data_aplicacao or "Não definida"
                ),
                "status": status,
                "status_slug": status_slug,
                "media_ativa": int(registro["media_ativa"] or 0),
                "media_aprovacao": registro["media_aprovacao"],
                "escola_id": registro["escola_id"],
                "instituicao": registro["nome_instituicao"]
            })

        banco.commit()

        indicadores = {
            "total": len(lista_provas),
            "finalizadas": sum(
                item["status_slug"] == "finalizada"
                for item in lista_provas
            ),
            "em_correcao": sum(
                item["status_slug"] == "em-correcao"
                for item in lista_provas
            ),
            "agendadas": sum(
                item["status_slug"] == "agendada"
                for item in lista_provas
            ),
            "rascunhos": sum(
                item["status_slug"] == "rascunho"
                for item in lista_provas
            )
        }

        disciplinas = sorted({
            item["disciplina"]
            for item in lista_provas
            if item["disciplina"]
        })

        pode_criar = _pode_criar_prova(cargo)
        pode_editar = _pode_criar_prova(cargo)
        pode_excluir = _pode_criar_prova(cargo)

        return render_template(
            "provas.html",
            provas=lista_provas,
            turmas=turmas,
            professores=professores,
            escolas=escolas,
            disciplinas=disciplinas,
            indicadores=indicadores,
            pode_criar=pode_criar,
            pode_editar=pode_editar,
            pode_excluir=pode_excluir,
            exibir_instituicao=(
                cargo == "Administrador Geral"
            ),
            professor_logado_id=professor_logado_id,
            vinculos_duplicacao=vinculos_duplicacao,
            cargo=cargo
        )

    except sqlite3.Error as erro:
        import traceback
        traceback.print_exc()

        flash(
            f"Não foi possível carregar as avaliações: {erro}",
            "erro"
        )

        return render_template(
            "provas.html",
            provas=[],
            turmas=[],
            professores=[],
            escolas=[],
            disciplinas=[],
            indicadores={
                "total": 0,
                "aplicadas": 0,
                "agendadas": 0,
                "rascunhos": 0
            },
            pode_criar=False,
            pode_editar=False,
            pode_excluir=False,
            exibir_instituicao=False,
            vinculos_duplicacao=[],
            cargo=cargo
        )

    finally:
        banco.close()



@app.route("/provas/<int:prova_id>/duplicar", methods=["POST"])
def duplicar_prova(prova_id):
    """Duplica a avaliação e seus vínculos com questões como um novo rascunho."""
    if not permissao_modulo("Provas"):
        return redirect("/acesso_negado")

    contexto = _contexto_provas()
    cargo = contexto["cargo"]
    usuario_id = contexto["usuario_id"]

    if not _pode_criar_prova(cargo):
        return _redirecionar_acesso_negado_prova()

    nome = (request.form.get("nome") or "").strip()
    disciplina = (request.form.get("disciplina") or "").strip()
    data_aplicacao = (request.form.get("data_aplicacao") or "").strip()

    try:
        turma_id = int(request.form.get("turma_id") or 0)
        professor_id = int(request.form.get("professor_id") or 0)
    except (TypeError, ValueError):
        flash("Selecione turma e professor válidos.", "erro")
        return redirect("/provas")

    if not nome:
        flash("Informe o nome da nova avaliação.", "erro")
        return redirect("/provas")

    if not disciplina:
        flash("Selecione o componente curricular.", "erro")
        return redirect("/provas")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        if not _pode_gerenciar_prova(cursor, prova_id, exigir_edicao=False):
            return _redirecionar_acesso_negado_prova()

        cursor.execute("""
            SELECT *
            FROM provas
            WHERE id = ?
            LIMIT 1
        """, (prova_id,))
        original = cursor.fetchone()

        if not original:
            flash("A avaliação original não foi encontrada.", "erro")
            return redirect("/provas")

        modo_media = (request.form.get("modo_media") or "manter").strip().lower()

        if modo_media == "nova":
            media_texto = (
                request.form.get("nova_media") or ""
            ).strip().replace(",", ".")

            try:
                media_aprovacao = float(media_texto)
            except (TypeError, ValueError):
                flash("Informe uma nova média válida.", "erro")
                return redirect("/provas")

            if not 0 <= media_aprovacao <= 10:
                flash("A média deve estar entre 0 e 10.", "erro")
                return redirect("/provas")

            media_ativa = 1
        else:
            media_ativa = int(original["media_ativa"] or 0)
            media_aprovacao = (
                original["media_aprovacao"]
                if media_ativa
                else None
            )

        cursor.execute("""
            SELECT
                t.id,
                t.escola_id,
                t.ano_letivo_id
            FROM turmas AS t
            WHERE t.id = ?
            LIMIT 1
        """, (turma_id,))
        turma = cursor.fetchone()

        if not turma:
            flash("A turma selecionada não foi encontrada.", "erro")
            return redirect("/provas")

        if cargo == "Administrador Geral":
            try:
                escola_id = int(request.form.get("escola_id") or 0)
            except (TypeError, ValueError):
                flash("Selecione uma instituição.", "erro")
                return redirect("/provas")

            if int(turma["escola_id"] or 0) != escola_id:
                flash("A turma não pertence à instituição selecionada.", "erro")
                return redirect("/provas")
        else:
            escola_id = contexto["escola_id"]

            if not escola_id or int(turma["escola_id"] or 0) != int(escola_id):
                return _redirecionar_acesso_negado_prova()

        cursor.execute("""
            SELECT
                pr.id,
                pr.nome,
                pr.email,
                pr.escola_id
            FROM professores AS pr
            WHERE pr.id = ?
              AND pr.escola_id = ?
            LIMIT 1
        """, (professor_id, escola_id))
        professor = cursor.fetchone()

        if not professor:
            flash("O professor selecionado não pertence à instituição.", "erro")
            return redirect("/provas")

        # Confirma o vínculo entre professor, turma e componente.
        # Aceita tanto o vínculo oficial atual quanto registros legados.
        cursor.execute("""
            SELECT 1
            WHERE EXISTS (
                SELECT 1
                FROM professores AS pr
                INNER JOIN usuarios AS u
                    ON u.escola_id = pr.escola_id
                   AND (
                        (
                            TRIM(COALESCE(u.email, '')) <> ''
                            AND TRIM(COALESCE(pr.email, '')) <> ''
                            AND LOWER(TRIM(u.email)) = LOWER(TRIM(pr.email))
                        )
                        OR LOWER(TRIM(u.nome)) = LOWER(TRIM(pr.nome))
                   )
                INNER JOIN professor_vinculos AS pv
                    ON pv.professor_id = u.id
                   AND pv.turma_id = ?
                INNER JOIN turmas AS t
                    ON t.id = pv.turma_id
                   AND t.escola_id = ?
                   AND t.escola_id = u.escola_id
                INNER JOIN componentes_curriculares AS cc
                    ON cc.id = pv.componente_id
                   AND cc.escola_id = t.escola_id
                   AND COALESCE(cc.ativo, 1) = 1
                WHERE pr.id = ?
                  AND LOWER(TRIM(cc.nome)) = LOWER(TRIM(?))
            )
            OR EXISTS (
                SELECT 1
                FROM professor_componentes AS pc
                INNER JOIN professores AS pr
                    ON pr.id = pc.professor_id
                INNER JOIN turmas AS t
                    ON t.id = pc.turma_id
                   AND t.escola_id = pr.escola_id
                INNER JOIN componentes_curriculares AS cc
                    ON cc.id = pc.componente_id
                   AND cc.escola_id = t.escola_id
                   AND COALESCE(cc.ativo, 1) = 1
                WHERE pc.professor_id = ?
                  AND pc.turma_id = ?
                  AND t.escola_id = ?
                  AND LOWER(TRIM(cc.nome)) = LOWER(TRIM(?))
            )
            LIMIT 1
        """, (
            turma_id,
            escola_id,
            professor_id,
            disciplina,
            professor_id,
            turma_id,
            escola_id,
            disciplina
        ))

        if not cursor.fetchone():
            flash(
                "O professor não possui vínculo com essa turma e componente.",
                "erro"
            )
            return redirect("/provas")

        if cargo == "Professor":
            professor_logado_id = _professor_legado_do_usuario(
                cursor,
                usuario_id
            )
            if int(professor_logado_id or 0) != professor_id:
                return _redirecionar_acesso_negado_prova()

        data_geracao = datetime.now().strftime("%d/%m/%Y")
        agora = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        cursor.execute("""
            INSERT INTO provas (
                nome,
                turma_id,
                professor_id,
                disciplina,
                quantidade,
                data_geracao,
                data_aplicacao,
                escola_id,
                ano_letivo_id,
                status,
                atualizado_em,
                media_ativa,
                media_aprovacao,
                peso_total,
                tipo_peso,
                tem_nota
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, 'rascunho', ?, ?, ?, ?, ?, ?)
        """, (
            nome,
            turma_id,
            professor_id,
            disciplina,
            int(original["quantidade"] or 0),
            data_geracao,
            data_aplicacao,
            escola_id,
            turma["ano_letivo_id"],
            agora,
            media_ativa,
            media_aprovacao,
            float(original["peso_total"] or 10),
            original["tipo_peso"] or "automatico",
            int(original["tem_nota"] if "tem_nota" in original.keys() else 1)
        ))

        nova_prova_id = cursor.lastrowid

        cursor.execute("""
            INSERT INTO prova_questoes (
                prova_id,
                questao_id,
                peso,
                ordem
            )
            SELECT
                ?,
                questao_id,
                COALESCE(peso, 0),
                COALESCE(ordem, 0)
            FROM prova_questoes
            WHERE prova_id = ?
            ORDER BY COALESCE(ordem, 0), id
        """, (nova_prova_id, prova_id))

        cursor.execute("""
            UPDATE provas
            SET quantidade = (
                SELECT COUNT(*)
                FROM prova_questoes
                WHERE prova_id = ?
            )
            WHERE id = ?
        """, (nova_prova_id, nova_prova_id))

        banco.commit()

        flash(
            "Avaliação duplicada com sucesso. A nova cópia foi salva como rascunho.",
            "sucesso"
        )
        return redirect(f"/provas/{nova_prova_id}/montar")

    except sqlite3.Error as erro:
        banco.rollback()
        import traceback
        traceback.print_exc()
        flash(f"Não foi possível duplicar a avaliação: {erro}", "erro")
        return redirect("/provas")

    finally:
        banco.close()



@app.route("/gerar_prova", methods=["POST"])
def gerar_prova():
    """
    Cria apenas os dados básicos da avaliação.

    As questões são escolhidas depois, na página de montagem.
    """
    if not permissao_modulo("Provas"):
        return redirect("/acesso_negado")

    contexto = _contexto_provas()
    cargo = contexto["cargo"]
    usuario_id = contexto["usuario_id"]

    if not _pode_criar_prova(cargo):
        return _redirecionar_acesso_negado_prova()

    nome = request.form.get("nome", "").strip()
    disciplina = request.form.get("disciplina", "").strip()
    data_aplicacao = request.form.get(
        "data_aplicacao",
        ""
    ).strip()

    tem_nota = 1 if request.form.get("tem_nota") == "1" else 0
    media_ativa = 1 if request.form.get("media_ativa") == "1" else 0
    media_aprovacao = None

    if media_ativa:
        media_texto = request.form.get("media_aprovacao", "").strip().replace(",", ".")
        try:
            media_aprovacao = float(media_texto)
        except (TypeError, ValueError):
            flash("Informe uma média válida entre 0,0 e 10,0.", "erro")
            return redirect("/provas")

        if not 0 <= media_aprovacao <= 10:
            flash("A média deve estar entre 0,0 e 10,0.", "erro")
            return redirect("/provas")

    try:
        turma_id = int(
            request.form.get("turma_id", "").strip()
        )
    except (TypeError, ValueError):
        flash("Selecione uma turma válida.", "erro")
        return redirect("/provas")

    if not nome:
        flash("Informe o nome da avaliação.", "erro")
        return redirect("/provas")

    if not disciplina:
        flash(
            "Informe o componente curricular da avaliação.",
            "erro"
        )
        return redirect("/provas")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        cursor.execute("""
            SELECT
                t.id,
                t.escola_id,
                t.ano_letivo_id
            FROM turmas AS t
            WHERE t.id = ?
            LIMIT 1
        """, (turma_id,))
        turma = cursor.fetchone()

        if not turma:
            flash("A turma selecionada não foi encontrada.", "erro")
            return redirect("/provas")

        if cargo == "Administrador Geral":
            try:
                escola_id = int(
                    request.form.get("escola_id", "").strip()
                )
            except (TypeError, ValueError):
                flash("Selecione uma instituição.", "erro")
                return redirect("/provas")

            if int(turma["escola_id"]) != escola_id:
                flash(
                    "A turma não pertence à instituição selecionada.",
                    "erro"
                )
                return redirect("/provas")
        else:
            escola_id = contexto["escola_id"]

            if (
                not escola_id
                or int(turma["escola_id"]) != int(escola_id)
            ):
                return _redirecionar_acesso_negado_prova()

        ano_letivo_id = turma["ano_letivo_id"]

        if cargo == "Professor":
            cursor.execute("""
                SELECT 1
                FROM professor_vinculos
                WHERE professor_id = ?
                  AND turma_id = ?
                LIMIT 1
            """, (usuario_id, turma_id))

            if not cursor.fetchone():
                flash(
                    "Você não possui vínculo com essa turma.",
                    "erro"
                )
                return redirect("/provas")

            professor_id = _professor_legado_do_usuario(
                cursor,
                usuario_id
            )
        else:
            try:
                professor_id = int(
                    request.form.get(
                        "professor_id",
                        ""
                    ).strip()
                )
            except (TypeError, ValueError):
                flash("Selecione um professor.", "erro")
                return redirect("/provas")

            cursor.execute("""
                SELECT id
                FROM professores
                WHERE id = ?
                  AND escola_id = ?
                LIMIT 1
            """, (professor_id, escola_id))

            if not cursor.fetchone():
                flash(
                    "O professor selecionado não pertence à instituição.",
                    "erro"
                )
                return redirect("/provas")

        data_geracao = datetime.now().strftime("%d/%m/%Y")
        agora = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        cursor.execute("""
            INSERT INTO provas (
                nome,
                turma_id,
                professor_id,
                disciplina,
                quantidade,
                data_geracao,
                data_aplicacao,
                escola_id,
                ano_letivo_id,
                status,
                atualizado_em,
                media_ativa,
                media_aprovacao,
                tem_nota
            )
            VALUES (?, ?, ?, ?, 0, ?, ?, ?, ?, 'rascunho', ?, ?, ?, ?)
        """, (
            nome,
            turma_id,
            professor_id,
            disciplina,
            data_geracao,
            data_aplicacao,
            escola_id,
            ano_letivo_id,
            agora,
            media_ativa,
            media_aprovacao,
            tem_nota
        ))

        prova_id = cursor.lastrowid
        banco.commit()

        flash(
            "Dados da avaliação salvos. Agora adicione as questões.",
            "sucesso"
        )

        return redirect(f"/provas/{prova_id}/montar")

    except sqlite3.Error as erro:
        banco.rollback()
        import traceback
        traceback.print_exc()

        flash(
            f"Não foi possível criar a avaliação: {erro}",
            "erro"
        )
        return redirect("/provas")

    finally:
        banco.close()




@app.route("/questoes/<int:questao_id>/editar", methods=["GET", "POST"])
def editar_questao(questao_id):
    if not permissao_modulo("Questões"):
        return redirect("/acesso_negado")

    prova_id = request.args.get("prova_id", type=int)
    if request.method == "POST":
        prova_id = request.form.get("prova_id", type=int)

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        cursor.execute("SELECT * FROM questoes WHERE id = ? LIMIT 1", (questao_id,))
        questao = cursor.fetchone()

        if not questao:
            flash("Questão não encontrada.", "erro")
            return redirect(f"/provas/{prova_id}/montar" if prova_id else "/questoes")

        escola_usuario = obter_escola_usuario()
        cargo = (session.get("usuario_cargo") or "").strip()

        if (
            cargo != "Administrador Geral"
            and questao["escola_id"] is not None
            and escola_usuario is not None
            and int(questao["escola_id"]) != int(escola_usuario)
        ):
            return redirect("/acesso_negado")

        if prova_id and not _pode_gerenciar_prova(cursor, prova_id, exigir_edicao=True):
            return _redirecionar_acesso_negado_prova()

        if request.method == "GET":
            try:
                alternativas = json.loads(questao["alternativas_json"] or "[]")
            except (TypeError, ValueError, json.JSONDecodeError):
                alternativas = []

            if not alternativas:
                alternativas = []
                for indice, campo in enumerate(
                    ["alternativa_a", "alternativa_b", "alternativa_c", "alternativa_d"]
                ):
                    texto = questao[campo] or ""
                    if texto:
                        alternativas.append({
                            "letra": chr(65 + indice),
                            "texto": texto,
                            "imagem": ""
                        })

            try:
                respostas = json.loads(questao["respostas_corretas"] or "[]")
            except (TypeError, ValueError, json.JSONDecodeError):
                respostas = []

            if not respostas and questao["correta"]:
                respostas = [questao["correta"]]

            cursor.execute("""
                SELECT DISTINCT nome
                FROM componentes_curriculares
                WHERE ativo = 1
                  AND (escola_id = ? OR ? IS NULL)
                ORDER BY nome
            """, (escola_usuario, escola_usuario))
            componentes = [linha["nome"] for linha in cursor.fetchall()]

            return render_template(
                "editar_questao.html",
                questao=questao,
                alternativas=alternativas,
                respostas=respostas,
                componentes=componentes,
                prova_id=prova_id,
                valores_anteriores=dict(questao)
            )

        disciplina = (request.form.get("disciplina") or "").strip()
        assunto = (request.form.get("assunto") or "").strip()
        dificuldade = (request.form.get("dificuldade") or "").strip()
        tipo_questao = (request.form.get("tipo_questao") or "multipla_escolha").strip()
        enunciado = (request.form.get("enunciado") or "").strip()
        enunciado_html = (request.form.get("enunciado_html") or "").strip()
        resposta_esperada = (request.form.get("resposta_esperada") or "").strip()
        criterios_correcao = (request.form.get("criterios_correcao") or "").strip()
        etapa_ensino = (request.form.get("etapa_ensino") or questao["etapa_ensino"] or "").strip()
        ano_serie = (request.form.get("ano_serie") or questao["ano_serie"] or "").strip()
        unidade_tematica = (request.form.get("unidade_tematica") or "").strip()
        objeto_conhecimento = (request.form.get("objeto_conhecimento") or "").strip()
        habilidade_bncc = (request.form.get("habilidade_bncc") or "").strip()
        sistema_matriz = (request.form.get("sistema_matriz") or "").strip().upper()
        matriz_referencia = (request.form.get("matriz_referencia") or "").strip()
        descritor_saeb = (request.form.get("descritor_saeb") or "").strip()
        if sistema_matriz not in {"", "SAETO", "SAEB"}:
            sistema_matriz = ""

        try:
            linhas_resposta = int(request.form.get("linhas_resposta") or 5)
        except (TypeError, ValueError):
            linhas_resposta = 5
        linhas_resposta = max(1, min(linhas_resposta, 30))

        if not disciplina or not assunto or not dificuldade or not enunciado:
            flash("Preencha componente, assunto, dificuldade e enunciado.", "erro")
            return redirect(
                f"/questoes/{questao_id}/editar"
                + (f"?prova_id={prova_id}" if prova_id else "")
            )

        tipos_validos = {
            "multipla_escolha", "multiplas_respostas", "verdadeiro_falso",
            "discursiva", "resposta_curta", "numerica"
        }
        if tipo_questao not in tipos_validos:
            flash("Tipo de questão inválido.", "erro")
            return redirect(
                f"/questoes/{questao_id}/editar"
                + (f"?prova_id={prova_id}" if prova_id else "")
            )

        extensoes_permitidas = {".png", ".jpg", ".jpeg", ".webp"}

        def salvar_imagem_edicao(arquivo, prefixo):
            if not arquivo or not arquivo.filename:
                return ""
            nome_seguro = secure_filename(arquivo.filename)
            extensao = os.path.splitext(nome_seguro)[1].lower()
            if extensao not in extensoes_permitidas:
                raise ValueError("Envie apenas imagens PNG, JPG, JPEG ou WEBP.")
            nome_arquivo = f"{prefixo}_{uuid.uuid4().hex}{extensao}"
            arquivo.save(os.path.join(app.config["UPLOAD_FOLDER"], nome_arquivo))
            return nome_arquivo

        def salvar_imagens_embutidas_edicao(conteudo_html):
            if not conteudo_html:
                return ""
            padrao = re.compile(r'src=["\']data:image/(?P<tipo>png|jpeg|jpg|webp);base64,(?P<dados>[^"\']+)["\']', re.IGNORECASE)
            def substituir(match):
                tipo=match.group("tipo").lower(); ext=".jpg" if tipo in {"jpg","jpeg"} else f".{tipo}"
                try: dados=base64.b64decode(match.group("dados"), validate=True)
                except Exception as erro: raise ValueError("Uma das imagens inseridas no enunciado é inválida.") from erro
                if len(dados)>8*1024*1024: raise ValueError("Cada imagem do enunciado deve ter no máximo 8 MB.")
                nome=f"enunciado_{uuid.uuid4().hex}{ext}"
                with open(os.path.join(app.config["UPLOAD_FOLDER"],nome),"wb") as destino: destino.write(dados)
                return f'src="/uploads/{nome}"'
            conteudo_html=padrao.sub(substituir,conteudo_html)
            conteudo_html=re.sub(r'<\s*(script|iframe|object|embed)[^>]*>.*?<\s*/\s*\1\s*>','',conteudo_html,flags=re.I|re.S)
            conteudo_html=re.sub(r'\son[a-z]+\s*=\s*(["\']).*?\1','',conteudo_html,flags=re.I|re.S)
            conteudo_html=re.sub(r'javascript\s*:','',conteudo_html,flags=re.I)
            return conteudo_html


        try:
            enunciado_html = salvar_imagens_embutidas_edicao(enunciado_html)
        except ValueError as erro:
            flash(str(erro), "erro")
            return redirect(f"/questoes/{questao_id}/editar" + (f"?prova_id={prova_id}" if prova_id else ""))

        # A imagem do enunciado agora vive dentro de enunciado_html.
        # O campo legado fica vazio nas novas cópias para não duplicar a figura.
        imagem_questao = ""

        alternativas = []
        respostas_corretas = []

        if tipo_questao in {"multipla_escolha", "multiplas_respostas"}:
            textos = request.form.getlist("alternativas[]")
            imagens_novas = request.files.getlist("imagens_alternativas[]")
            imagens_atuais = request.form.getlist("imagens_atuais[]")
            remover_imagens = set(request.form.getlist("remover_imagens[]"))
            corretas = set(request.form.getlist("corretas[]"))

            for indice, texto in enumerate(textos):
                texto = (texto or "").strip()
                imagem_alt = imagens_atuais[indice] if indice < len(imagens_atuais) else ""
                if str(indice) in remover_imagens:
                    imagem_alt = ""
                arquivo = imagens_novas[indice] if indice < len(imagens_novas) else None
                try:
                    nova_alt = salvar_imagem_edicao(arquivo, "alternativa")
                    if nova_alt:
                        imagem_alt = nova_alt
                except ValueError as erro:
                    flash(f"Alternativa {indice + 1}: {erro}", "erro")
                    return redirect(f"/questoes/{questao_id}/editar" + (f"?prova_id={prova_id}" if prova_id else ""))
                if not texto and not imagem_alt:
                    continue
                letra = chr(65 + len(alternativas))
                alternativas.append({"letra": letra, "texto": texto, "imagem": imagem_alt})
                if str(indice) in corretas:
                    respostas_corretas.append(letra)

            if len(alternativas) < 2:
                flash("Cadastre pelo menos duas alternativas.", "erro")
                return redirect(
                    f"/questoes/{questao_id}/editar"
                    + (f"?prova_id={prova_id}" if prova_id else "")
                )
            if tipo_questao == "multipla_escolha" and len(respostas_corretas) != 1:
                flash("Marque exatamente uma alternativa correta.", "erro")
                return redirect(
                    f"/questoes/{questao_id}/editar"
                    + (f"?prova_id={prova_id}" if prova_id else "")
                )
            if tipo_questao == "multiplas_respostas" and not respostas_corretas:
                flash("Marque pelo menos uma alternativa correta.", "erro")
                return redirect(
                    f"/questoes/{questao_id}/editar"
                    + (f"?prova_id={prova_id}" if prova_id else "")
                )

        elif tipo_questao == "verdadeiro_falso":
            resposta_vf = (request.form.get("resposta_vf") or "").strip()
            if resposta_vf not in {"V", "F"}:
                flash("Selecione Verdadeiro ou Falso.", "erro")
                return redirect(
                    f"/questoes/{questao_id}/editar"
                    + (f"?prova_id={prova_id}" if prova_id else "")
                )
            alternativas = [
                {"letra": "V", "texto": "Verdadeiro", "imagem": ""},
                {"letra": "F", "texto": "Falso", "imagem": ""}
            ]
            respostas_corretas = [resposta_vf]

        elif tipo_questao in {"resposta_curta", "numerica"} and not resposta_esperada:
            flash("Informe a resposta esperada.", "erro")
            return redirect(
                f"/questoes/{questao_id}/editar"
                + (f"?prova_id={prova_id}" if prova_id else "")
            )

        textos_legados = [(item.get("texto") or "") for item in alternativas[:4]]
        textos_legados += [""] * (4 - len(textos_legados))
        correta_legada = respostas_corretas[0] if respostas_corretas else ""

        agora = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        escola_copia_id = (
            questao["escola_id"]
            if cargo == "Administrador Geral"
            else escola_usuario
        )

        # A edição nunca altera a questão original. Ela cria uma nova questão
        # com o usuário logado como autor, preservando os campos não exibidos
        # neste formulário (imagem, fonte, tags, observações etc.).
        cursor.execute("""
            INSERT INTO questoes (
                disciplina, etapa_ensino, ano_serie, assunto,
                assunto_temporario, subassunto,
                tipo_questao, enunciado, enunciado_html, imagem,
                alternativa_a, alternativa_b, alternativa_c, alternativa_d,
                correta, alternativas_json, respostas_corretas,
                resposta_esperada, criterios_correcao, habilidade,
                habilidade_bncc, unidade_tematica, objeto_conhecimento,
                sistema_matriz, matriz_referencia, descritor_saeb, taxonomia_bloom,
                dificuldade, fonte, ano_fonte, tags, tempo_estimado,
                linhas_resposta, observacoes, escola_id, criado_por, criado_em,
                atualizado_em
            )
            VALUES (
                ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?,
                ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?
            )
        """, (
            disciplina, etapa_ensino, ano_serie, assunto,
            questao["assunto_temporario"], questao["subassunto"],
            tipo_questao, enunciado, enunciado_html, imagem_questao,
            textos_legados[0], textos_legados[1],
            textos_legados[2], textos_legados[3],
            correta_legada,
            json.dumps(alternativas, ensure_ascii=False),
            json.dumps(respostas_corretas, ensure_ascii=False),
            resposta_esperada, criterios_correcao,
            " | ".join(item for item in [habilidade_bncc, descritor_saeb] if item),
            habilidade_bncc, unidade_tematica, objeto_conhecimento,
            sistema_matriz, matriz_referencia, descritor_saeb,
            questao["taxonomia_bloom"], dificuldade,
            questao["fonte"], questao["ano_fonte"], questao["tags"],
            questao["tempo_estimado"],
            linhas_resposta if tipo_questao == "discursiva" else None,
            questao["observacoes"], escola_copia_id,
            session.get("usuario_id"), agora, agora
        ))

        nova_questao_id = cursor.lastrowid

        # Quando a edição foi aberta dentro de uma avaliação, troca somente
        # o vínculo daquela avaliação para a nova cópia. A questão original
        # continua intacta e permanece disponível no banco de questões.
        if prova_id:
            cursor.execute("""
                UPDATE prova_questoes
                SET questao_id = ?
                WHERE prova_id = ?
                  AND questao_id = ?
            """, (nova_questao_id, prova_id, questao_id))

        banco.commit()
        flash("Cópia da questão criada com sucesso. A original foi preservada.", "sucesso")
        return redirect(f"/provas/{prova_id}/montar" if prova_id else "/questoes")

    except sqlite3.Error as erro:
        banco.rollback()
        flash(f"Não foi possível editar a questão: {erro}", "erro")
        return redirect(f"/provas/{prova_id}/montar" if prova_id else "/questoes")
    finally:
        banco.close()


def normalizar_ordem_questoes_prova(cursor, prova_id):
    """Garante uma sequência de ordem 1, 2, 3... para a avaliação."""
    cursor.execute("""
        SELECT id
        FROM prova_questoes
        WHERE prova_id = ?
        ORDER BY
            CASE WHEN COALESCE(ordem, 0) <= 0 THEN 1 ELSE 0 END,
            ordem,
            id
    """, (prova_id,))

    for indice, registro in enumerate(cursor.fetchall(), start=1):
        cursor.execute("""
            UPDATE prova_questoes
            SET ordem = ?
            WHERE id = ? AND prova_id = ?
        """, (indice, registro["id"], prova_id))



# =========================================================
# IMPORTAÇÃO EM LOTE DE QUESTÕES
# =========================================================

def _normalizar_cabecalho_importacao(valor):
    texto = unicodedata.normalize("NFKD", str(valor or ""))
    texto = "".join(c for c in texto if not unicodedata.combining(c))
    texto = re.sub(r"[^a-zA-Z0-9]+", "_", texto.lower()).strip("_")
    return texto


def _valor_linha_importacao(linha, *aliases):
    for alias in aliases:
        chave = _normalizar_cabecalho_importacao(alias)
        if chave in linha and linha[chave] not in (None, ""):
            return str(linha[chave]).strip()
    return ""


def _tipo_importacao(valor, tem_alternativas=False):
    v = _normalizar_cabecalho_importacao(valor)
    mapa = {
        "multipla_escolha": "multipla_escolha", "multipla_escolha_unica": "multipla_escolha",
        "objetiva": "multipla_escolha", "objetiva_unica": "multipla_escolha",
        "multiplas_respostas": "multiplas_respostas", "multipla_resposta": "multiplas_respostas",
        "verdadeiro_falso": "verdadeiro_falso", "vf": "verdadeiro_falso",
        "discursiva": "discursiva", "dissertativa": "discursiva", "aberta": "discursiva",
        "resposta_curta": "resposta_curta", "curta": "resposta_curta",
        "numerica": "numerica", "numerico": "numerica",
    }
    if v in mapa:
        return mapa[v]
    return "multipla_escolha" if tem_alternativas else "discursiva"


def _dificuldade_importacao(valor, padrao="Média"):
    v = _normalizar_cabecalho_importacao(valor)
    if v in {"facil", "fácil"}: return "Fácil"
    if v in {"dificil", "difícil"}: return "Difícil"
    if v in {"media", "medio", "média", "médio"}: return "Média"
    return padrao or "Média"


def _questao_de_linha_importacao(linha, defaults):
    linha = {_normalizar_cabecalho_importacao(k): ("" if v is None else v) for k, v in linha.items()}
    enunciado = _valor_linha_importacao(linha, "enunciado", "questao", "pergunta", "texto")
    if not enunciado:
        return None

    alternativas = []
    for letra in "ABCDEFGH":
        texto = _valor_linha_importacao(
            linha, f"alternativa_{letra.lower()}", f"alternativa {letra}", letra.lower()
        )
        if texto:
            alternativas.append({"letra": letra, "texto": texto, "imagem": ""})

    correta_raw = _valor_linha_importacao(linha, "gabarito", "correta", "resposta_correta", "respostas_corretas")
    corretas = []
    if correta_raw:
        for pedaco in re.split(r"[,;/|\s]+", correta_raw.upper()):
            pedaco = re.sub(r"[^A-HVF]", "", pedaco)
            if pedaco and pedaco not in corretas:
                corretas.append(pedaco)

    tipo = _tipo_importacao(_valor_linha_importacao(linha, "tipo", "tipo_questao"), bool(alternativas))
    if tipo == "verdadeiro_falso" and not alternativas:
        alternativas = [
            {"letra": "V", "texto": "Verdadeiro", "imagem": ""},
            {"letra": "F", "texto": "Falso", "imagem": ""},
        ]

    return {
        "disciplina": _valor_linha_importacao(linha, "disciplina", "componente", "componente_curricular", "componente curricular") or defaults.get("disciplina", ""),
        "turma": _valor_linha_importacao(linha, "turma", "classe"),
        "etapa_ensino": _valor_linha_importacao(linha, "etapa", "etapa_ensino") or defaults.get("etapa_ensino", ""),
        "ano_serie": _valor_linha_importacao(linha, "ano_serie", "ano", "serie", "ano série") or _valor_linha_importacao(linha, "turma", "classe") or defaults.get("ano_serie", ""),
        "tipo_questao": tipo,
        "enunciado": enunciado,
        "alternativas": alternativas,
        "respostas_corretas": corretas,
        "resposta_esperada": _valor_linha_importacao(linha, "resposta_esperada", "resposta esperada"),
        "criterios_correcao": _valor_linha_importacao(linha, "criterios_correcao", "criterios", "critério", "criterio"),
        "dificuldade": _dificuldade_importacao(_valor_linha_importacao(linha, "dificuldade", "nivel", "nivel_de_dificuldade", "nível de dificuldade"), defaults.get("dificuldade", "Média")),
        "habilidade_bncc": _valor_linha_importacao(linha, "habilidade_bncc", "bncc", "habilidade"),
        "unidade_tematica": _valor_linha_importacao(linha, "unidade_tematica", "unidade temática"),
        "objeto_conhecimento": _valor_linha_importacao(linha, "objeto_conhecimento", "objeto de conhecimento", "objeto do conhecimento"),
        "sistema_matriz": (_valor_linha_importacao(linha, "sistema_matriz", "matriz_sistema") or "").upper(),
        "matriz_referencia": _valor_linha_importacao(linha, "matriz_referencia", "matriz de referencia"),
        "descritor_saeb": _valor_linha_importacao(linha, "descritor_saeb", "descritor", "saeb", "saeto"),
        "taxonomia_bloom": _valor_linha_importacao(linha, "taxonomia_bloom", "bloom"),
        "fonte": _valor_linha_importacao(linha, "fonte"),
        "tags": _valor_linha_importacao(linha, "tags", "etiquetas"),
        "observacoes": _valor_linha_importacao(linha, "observacoes", "observação", "observacao"),
    }


def _linhas_planilha_importacao(caminho, extensao):
    if extensao == ".csv":
        dados = Path(caminho).read_bytes()
        texto = None
        for enc in ("utf-8-sig", "utf-8", "latin-1"):
            try:
                texto = dados.decode(enc)
                break
            except UnicodeDecodeError:
                pass
        if texto is None:
            raise ValueError("Não foi possível identificar a codificação do CSV.")
        amostra = texto[:4096]
        try:
            dialeto = csv.Sniffer().sniff(amostra, delimiters=",;\t|")
        except csv.Error:
            dialeto = csv.excel
            dialeto.delimiter = ";"
        leitor = csv.DictReader(texto.splitlines(), dialect=dialeto)
        return list(leitor)

    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise ValueError("Para importar Excel, instale a dependência openpyxl.") from exc
    wb = load_workbook(caminho, read_only=True, data_only=True)
    ws = wb.active
    linhas = list(ws.iter_rows(values_only=True))
    if not linhas:
        return []
    cabecalhos = [str(v or "").strip() for v in linhas[0]]
    return [dict(zip(cabecalhos, valores)) for valores in linhas[1:] if any(v not in (None, "") for v in valores)]


def _texto_arquivo_importacao(caminho, extensao):
    if extensao == ".txt":
        dados = Path(caminho).read_bytes()
        for enc in ("utf-8-sig", "utf-8", "latin-1"):
            try:
                return dados.decode(enc)
            except UnicodeDecodeError:
                continue
        return dados.decode("utf-8", errors="replace")
    if extensao == ".pdf":
        if PdfReader is None:
            raise ValueError("A biblioteca pypdf não está instalada.")
        leitor = PdfReader(caminho)
        return "\n".join((pagina.extract_text() or "") for pagina in leitor.pages)
    if extensao == ".docx":
        try:
            from docx import Document
        except ImportError as exc:
            raise ValueError("Para importar Word, instale a dependência python-docx.") from exc
        doc = Document(caminho)
        partes = [p.text for p in doc.paragraphs]
        for tabela in doc.tables:
            for linha in tabela.rows:
                partes.append("\t".join(c.text for c in linha.cells))
        return "\n".join(partes)
    raise ValueError("Formato de arquivo não suportado.")


def _questoes_texto_importacao(texto, defaults):
    """Reconhece questões em Word/PDF/TXT usando uma estrutura pedagógica explícita.

    Estrutura recomendada por questão:
      Tipo de questão: Objetiva
      Componente curricular: História
      Turma: 8º ano B
      Questão: ...
      Alternativas:
      A) ...
      B) ...
      C) ...
      D) ...
      Gabarito: A
      Nível de dificuldade: Média
      BNCC
      Unidade temática: ...
      Objeto do conhecimento: ...
      Habilidade: EF08HI...

    Cada bloco pode ter um componente curricular diferente, inclusive quando
    todas as questões serão adicionadas à mesma avaliação.
    """
    texto = (texto or "").replace("\r\n", "\n").replace("\r", "\n")
    texto = re.sub(r"\u00a0", " ", texto)

    # Um novo bloco começa em "Questão 1", "QUESTÃO 02" ou em "Tipo de questão:"
    # quando o arquivo segue o modelo estruturado sem numeração.
    inicio_num = re.compile(r"(?im)^\s*(?:quest[aã]o\s*)?(\d{1,3})\s*[\.)\-:]\s*(?=\S|$)")
    matches = list(inicio_num.finditer(texto))
    blocos = []
    if matches:
        for i, m in enumerate(matches):
            ini = m.end()
            fim = matches[i + 1].start() if i + 1 < len(matches) else len(texto)
            bloco = texto[ini:fim].strip()
            if bloco:
                blocos.append(bloco)
    else:
        # Em documentos estruturados, cada ocorrência de "Tipo de questão" pode
        # marcar o início de uma nova questão.
        partes = re.split(r"(?im)(?=^\s*tipo\s+de\s+quest[aã]o\s*[:\-])", texto)
        blocos = [b.strip() for b in partes if b.strip()]
        if len(blocos) <= 1:
            blocos = [b.strip() for b in re.split(r"\n\s*\n+", texto) if b.strip()]

    aliases_meta = {
        "tipo": ["tipo", "tipo de questao", "tipo da questao"],
        "disciplina": ["componente curricular", "componente", "disciplina"],
        "turma": ["turma", "classe"],
        "gabarito": ["gabarito", "resposta correta", "resposta"],
        "dificuldade": ["nivel de dificuldade", "nível de dificuldade", "dificuldade", "nivel"],
        "unidade_tematica": ["unidade tematica", "unidade temática"],
        "objeto_conhecimento": ["objeto do conhecimento", "objeto de conhecimento"],
        "habilidade": ["habilidade", "habilidade bncc", "codigo da habilidade", "código da habilidade"],
        "bncc": ["bncc"],
        "fonte": ["fonte"],
        "tags": ["tag", "tags"],
    }
    alias_para_chave = {}
    for chave, nomes in aliases_meta.items():
        for nome in nomes:
            alias_para_chave[_normalizar_cabecalho_importacao(nome)] = chave

    resultado = []
    for bloco in blocos:
        linhas = [re.sub(r"\s+", " ", l).strip() for l in bloco.split("\n") if l.strip()]
        if not linhas:
            continue

        metadados = {}
        alternativas = []
        enunciado_linhas = []
        secao = None
        ultima_alt = None
        bncc_ativo = False

        for linha in linhas:
            # Títulos de seção que não carregam valor.
            if re.fullmatch(r"(?i)alternativas?\s*:?", linha):
                secao = "alternativas"
                ultima_alt = None
                continue
            if re.fullmatch(r"(?i)bncc\s*:?", linha):
                bncc_ativo = True
                secao = "bncc"
                continue

            alt = re.match(r"^([A-H])\s*[\)\.\-:]\s*(.+)$", linha, flags=re.I)
            if alt:
                secao = "alternativas"
                ultima_alt = {"letra": alt.group(1).upper(), "texto": alt.group(2).strip(), "imagem": ""}
                alternativas.append(ultima_alt)
                continue

            # Qualquer rótulo "Nome: valor" é normalizado e comparado aos aliases.
            meta = re.match(r"^([^:]{2,60})\s*[:\-]\s*(.*)$", linha)
            if meta:
                rotulo = _normalizar_cabecalho_importacao(meta.group(1))
                valor = meta.group(2).strip()
                chave = alias_para_chave.get(rotulo)
                if chave:
                    if chave == "bncc" and not valor:
                        bncc_ativo = True
                        secao = "bncc"
                    else:
                        metadados[chave] = valor
                        if chave in {"unidade_tematica", "objeto_conhecimento", "habilidade"}:
                            bncc_ativo = True
                    secao = "meta"
                    ultima_alt = None
                    continue
                if rotulo in {"questao", "enunciado", "pergunta"}:
                    if valor:
                        enunciado_linhas.append(valor)
                    secao = "enunciado"
                    ultima_alt = None
                    continue

            # Caso "Questão" venha em uma linha e o enunciado na linha seguinte.
            if re.fullmatch(r"(?i)(quest[aã]o|enunciado|pergunta)\s*:?", linha):
                secao = "enunciado"
                ultima_alt = None
                continue

            # Continuação de alternativa quebrada em mais de uma linha.
            if secao == "alternativas" and ultima_alt is not None:
                ultima_alt["texto"] += " " + linha
                continue

            # Texto livre após "Questão:" faz parte do enunciado.
            if secao == "enunciado" or not secao:
                enunciado_linhas.append(linha)
                secao = "enunciado"
                continue

            # Evita perder texto de enunciado em PDFs cuja ordem de extração varie.
            if not bncc_ativo and secao != "meta":
                enunciado_linhas.append(linha)

        enunciado = " ".join(enunciado_linhas).strip()
        if not enunciado:
            continue

        gabarito = metadados.get("gabarito", "")
        corretas = []
        for pedaco in re.split(r"[,;/|\s]+", gabarito.upper()):
            letra = re.sub(r"[^A-HVF]", "", pedaco)
            if letra and letra not in corretas:
                corretas.append(letra)

        tipo = _tipo_importacao(metadados.get("tipo", ""), bool(alternativas))
        resposta_esperada = ""
        if tipo in {"discursiva", "resposta_curta", "numerica"} and not alternativas:
            resposta_esperada = gabarito
            corretas = []

        turma = metadados.get("turma", "")
        resultado.append({
            "disciplina": metadados.get("disciplina") or defaults.get("disciplina", ""),
            "turma": turma,
            "etapa_ensino": defaults.get("etapa_ensino", ""),
            "ano_serie": turma or defaults.get("ano_serie", ""),
            "tipo_questao": tipo,
            "enunciado": enunciado,
            "alternativas": alternativas,
            "respostas_corretas": corretas,
            "resposta_esperada": resposta_esperada,
            "criterios_correcao": "",
            "dificuldade": _dificuldade_importacao(metadados.get("dificuldade", ""), defaults.get("dificuldade", "Média")),
            "habilidade_bncc": metadados.get("habilidade", ""),
            "unidade_tematica": metadados.get("unidade_tematica", ""),
            "objeto_conhecimento": metadados.get("objeto_conhecimento", ""),
            "sistema_matriz": "BNCC" if any(metadados.get(k) for k in ("habilidade", "unidade_tematica", "objeto_conhecimento")) else "",
            "matriz_referencia": "",
            "descritor_saeb": "",
            "taxonomia_bloom": "",
            "fonte": metadados.get("fonte", ""),
            "tags": metadados.get("tags", ""),
            "observacoes": "",
        })
    return resultado


def _extrair_texto_resposta_openai(payload):
    """Obtém o texto final da Responses API sem depender do SDK."""
    if isinstance(payload, dict) and isinstance(payload.get("output_text"), str):
        return payload["output_text"].strip()
    partes = []
    for item in (payload or {}).get("output", []) if isinstance(payload, dict) else []:
        for conteudo in item.get("content", []) if isinstance(item, dict) else []:
            if not isinstance(conteudo, dict):
                continue
            if conteudo.get("type") in {"output_text", "text"} and conteudo.get("text"):
                partes.append(str(conteudo["text"]))
    return "\n".join(partes).strip()


def _salvar_recorte_figura_importacao(caminho_imagem, caixa):
    """Recorta apenas a figura/ilustração da questão, evitando duplicar o print inteiro."""
    if not isinstance(caixa, dict):
        return ""
    try:
        x1 = max(0.0, min(1.0, float(caixa.get("x1", 0))))
        y1 = max(0.0, min(1.0, float(caixa.get("y1", 0))))
        x2 = max(0.0, min(1.0, float(caixa.get("x2", 0))))
        y2 = max(0.0, min(1.0, float(caixa.get("y2", 0))))
        if x2 <= x1 or y2 <= y1 or (x2-x1) < .08 or (y2-y1) < .05:
            return ""
        with Image.open(caminho_imagem) as img:
            img = img.convert("RGB")
            w, h = img.size
            margem_x = int(w * .01)
            margem_y = int(h * .01)
            box = (
                max(0, int(x1*w)-margem_x), max(0, int(y1*h)-margem_y),
                min(w, int(x2*w)+margem_x), min(h, int(y2*h)+margem_y),
            )
            recorte = img.crop(box)
            nome = f"questao_ia_{uuid.uuid4().hex}.jpg"
            destino = os.path.join(app.config["UPLOAD_FOLDER"], nome)
            recorte.save(destino, "JPEG", quality=92, optimize=True)
            return nome
    except Exception:
        return ""


def _questoes_de_imagem_com_ia(caminho_imagem, extensao, defaults):
    """Lê print/foto de livro com visão e devolve questões estruturadas para revisão."""
    import requests

    chave = (os.environ.get("OPENAI_API_KEY") or "").strip()
    if not chave:
        raise ValueError(
            "A importação por foto precisa da variável OPENAI_API_KEY configurada no servidor. "
            "Nenhum dado do banco é alterado por essa configuração."
        )

    modelo = (os.environ.get("OPENAI_VISION_MODEL") or "gpt-5.6-luna").strip()
    mime = {
        ".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".webp": "image/webp"
    }.get(extensao, "image/jpeg")
    dados = Path(caminho_imagem).read_bytes()
    if len(dados) > 15 * 1024 * 1024:
        raise ValueError("A imagem deve ter no máximo 15 MB.")
    data_url = f"data:{mime};base64,{base64.b64encode(dados).decode('ascii')}"

    instrucao = f"""
Você está importando uma questão escolar a partir de uma FOTO ou PRINT de livro para a plataforma ARK EDUS.
Leia apenas o conteúdo visível. Retorne SOMENTE JSON válido.

Regras:
- Identifique uma ou mais questões presentes na imagem.
- Preserve o enunciado e as alternativas, corrigindo apenas quebras de linha causadas pela foto.
- Não invente gabarito. Se ele não estiver visível, use respostas_corretas=[] e gabarito_confianca="nao_visivel".
- Se houver uma figura, tirinha, gráfico, mapa, tabela ou ilustração que FAÇA PARTE da questão e deva aparecer na prova, informe figura_caixa com coordenadas normalizadas de 0 a 1: x1,y1,x2,y2. A caixa deve conter SOMENTE a figura necessária, sem incluir enunciado nem alternativas. Se não houver figura necessária, use null.
- Sugira disciplina, dificuldade, BNCC, unidade temática, objeto do conhecimento e habilidade apenas quando houver base suficiente; caso contrário deixe string vazia. Campo sugerido deve poder ser revisado pelo professor.
- tipo_questao: multipla_escolha, multiplas_respostas, verdadeiro_falso, discursiva, resposta_curta ou numerica.
- alternativas: lista de objetos {{"letra":"A","texto":"...","imagem":""}}.
- respostas_corretas: lista de letras.

Padrões informados pelo professor, use apenas se a imagem não indicar algo melhor:
Disciplina: {defaults.get('disciplina','')}
Etapa: {defaults.get('etapa_ensino','')}
Ano/série: {defaults.get('ano_serie','')}
Dificuldade: {defaults.get('dificuldade','Média')}

Formato JSON obrigatório:
{{"questoes":[{{"disciplina":"","turma":"","etapa_ensino":"","ano_serie":"","tipo_questao":"multipla_escolha","enunciado":"","alternativas":[],"respostas_corretas":[],"gabarito_confianca":"nao_visivel","resposta_esperada":"","criterios_correcao":"","dificuldade":"Média","habilidade_bncc":"","unidade_tematica":"","objeto_conhecimento":"","sistema_matriz":"","matriz_referencia":"","descritor_saeb":"","taxonomia_bloom":"","fonte":"","tags":"","observacoes":"","figura_caixa":null}}]}}
""".strip()

    corpo = {
        "model": modelo,
        "input": [{
            "role": "user",
            "content": [
                {"type": "input_text", "text": instrucao},
                {"type": "input_image", "image_url": data_url, "detail": "high"},
            ],
        }],
        "text": {"format": {"type": "json_object"}},
    }
    try:
        resposta = requests.post(
            "https://api.openai.com/v1/responses",
            headers={"Authorization": f"Bearer {chave}", "Content-Type": "application/json"},
            json=corpo,
            timeout=90,
        )
    except requests.RequestException as exc:
        raise ValueError("Não foi possível conectar ao serviço de leitura por IA. Tente novamente.") from exc

    if resposta.status_code >= 400:
        detalhe = ""
        try:
            detalhe = (resposta.json().get("error") or {}).get("message", "")
        except Exception:
            pass
        if resposta.status_code in {401, 403}:
            raise ValueError("A chave da IA não foi aceita. Confira OPENAI_API_KEY no Render.")
        if resposta.status_code == 429:
            raise ValueError("O limite da IA foi atingido no momento. Aguarde e tente novamente.")
        raise ValueError("A IA não conseguiu processar a imagem." + (f" Detalhe: {detalhe[:180]}" if detalhe else ""))

    try:
        texto = _extrair_texto_resposta_openai(resposta.json())
        resultado = json.loads(texto)
    except Exception as exc:
        raise ValueError("A leitura da imagem retornou um formato inesperado. Tente uma foto mais nítida.") from exc

    questoes = resultado.get("questoes", []) if isinstance(resultado, dict) else []
    saida = []
    for q in questoes:
        if not isinstance(q, dict):
            continue
        alternativas = q.get("alternativas") if isinstance(q.get("alternativas"), list) else []
        alternativas_ok = []
        for i, alt in enumerate(alternativas[:8]):
            if not isinstance(alt, dict):
                continue
            letra = (str(alt.get("letra") or chr(65+i)).strip().upper()[:1] or chr(65+i))
            alternativas_ok.append({"letra": letra, "texto": str(alt.get("texto") or "").strip(), "imagem": ""})
        corretas = [str(x).strip().upper()[:1] for x in (q.get("respostas_corretas") or []) if str(x).strip()]
        imagem = _salvar_recorte_figura_importacao(caminho_imagem, q.get("figura_caixa"))
        saida.append({
            "disciplina": str(q.get("disciplina") or defaults.get("disciplina", "")).strip(),
            "turma": str(q.get("turma") or "").strip(),
            "etapa_ensino": str(q.get("etapa_ensino") or defaults.get("etapa_ensino", "")).strip(),
            "ano_serie": str(q.get("ano_serie") or defaults.get("ano_serie", "")).strip(),
            "tipo_questao": _tipo_importacao(q.get("tipo_questao", ""), bool(alternativas_ok)),
            "enunciado": str(q.get("enunciado") or "").strip(),
            "imagem": imagem,
            "alternativas": alternativas_ok,
            "respostas_corretas": corretas,
            "resposta_esperada": str(q.get("resposta_esperada") or "").strip(),
            "criterios_correcao": str(q.get("criterios_correcao") or "").strip(),
            "dificuldade": _dificuldade_importacao(q.get("dificuldade", ""), defaults.get("dificuldade", "Média")),
            "habilidade_bncc": str(q.get("habilidade_bncc") or "").strip(),
            "unidade_tematica": str(q.get("unidade_tematica") or "").strip(),
            "objeto_conhecimento": str(q.get("objeto_conhecimento") or "").strip(),
            "sistema_matriz": str(q.get("sistema_matriz") or "").strip().upper(),
            "matriz_referencia": str(q.get("matriz_referencia") or "").strip(),
            "descritor_saeb": str(q.get("descritor_saeb") or "").strip(),
            "taxonomia_bloom": str(q.get("taxonomia_bloom") or "").strip(),
            "fonte": str(q.get("fonte") or "").strip(),
            "tags": str(q.get("tags") or "").strip(),
            "observacoes": str(q.get("observacoes") or "").strip(),
            "gabarito_confianca": str(q.get("gabarito_confianca") or "").strip(),
        })
    return saida


def _validar_questoes_importadas(questoes):
    validas, avisos = [], []
    for i, q in enumerate(questoes, 1):
        if not (q.get("enunciado") or "").strip():
            avisos.append(f"Item {i}: ignorado porque não possui enunciado.")
            continue
        if not (q.get("disciplina") or "").strip():
            avisos.append(f"Item {i}: ignorado porque não possui componente curricular.")
            continue
        tipo = q.get("tipo_questao") or "multipla_escolha"
        alternativas = q.get("alternativas") or []
        corretas = q.get("respostas_corretas") or []
        if tipo in {"multipla_escolha", "multiplas_respostas"}:
            if len(alternativas) < 2:
                avisos.append(f"Item {i}: ignorado porque uma questão objetiva precisa de pelo menos 2 alternativas.")
                continue
            if not corretas:
                avisos.append(f"Item {i}: objetiva sem gabarito. Ela será importada, mas revise antes de aplicar.")
        validas.append(q)
    return validas, avisos


def _salvar_preview_importacao_questoes(questoes, prova_id):
    pasta = os.path.join(UPLOAD_FOLDER, "importacoes_questoes")
    os.makedirs(pasta, exist_ok=True)
    token = uuid.uuid4().hex
    caminho = os.path.join(pasta, f"{token}.json")
    payload = {
        "usuario_id": session.get("usuario_id"),
        "prova_id": prova_id,
        "criado_em": datetime.now().isoformat(),
        "questoes": questoes,
    }
    with open(caminho, "w", encoding="utf-8") as arq:
        json.dump(payload, arq, ensure_ascii=False)
    return token


def _carregar_preview_importacao_questoes(token):
    if not re.fullmatch(r"[a-f0-9]{32}", token or ""):
        return None, None
    caminho = os.path.join(UPLOAD_FOLDER, "importacoes_questoes", f"{token}.json")
    if not os.path.isfile(caminho):
        return None, None
    try:
        with open(caminho, "r", encoding="utf-8") as arq:
            payload = json.load(arq)
        if payload.get("usuario_id") != session.get("usuario_id"):
            return None, None
        return payload, caminho
    except (OSError, json.JSONDecodeError):
        return None, None


@app.route("/questoes/importar", methods=["GET", "POST"])
def importar_questoes():
    if not cargo_permitido(["Administrador Geral", "Administrador da Instituição", "Coordenador", "Professor"]):
        return redirect("/login")
    if not permissao_modulo("Questões"):
        return redirect("/acesso_negado")

    prova_id = request.values.get("prova_id", type=int)
    prova = None
    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()
    try:
        if prova_id:
            if not _pode_gerenciar_prova(cursor, prova_id, exigir_edicao=True):
                return _redirecionar_acesso_negado_prova()
            cursor.execute("""
                SELECT p.id, p.nome, p.disciplina, COALESCE(p.escola_id,t.escola_id) escola_id,
                       t.etapa, t.nome AS turma_nome
                FROM provas p INNER JOIN turmas t ON t.id=p.turma_id
                WHERE p.id=?
            """, (prova_id,))
            prova = cursor.fetchone()
            if not prova:
                flash("Avaliação não encontrada.", "erro")
                return redirect("/provas")

        if request.method == "GET":
            return render_template("questoes/importar.html", prova=prova, prova_id=prova_id, preview=None)

        arquivo = request.files.get("arquivo")
        if not arquivo or not arquivo.filename:
            flash("Selecione um arquivo para importar.", "erro")
            return render_template("questoes/importar.html", prova=prova, prova_id=prova_id, preview=None)

        nome = secure_filename(arquivo.filename)
        extensao = os.path.splitext(nome)[1].lower()
        permitidas = {".xlsx", ".csv", ".docx", ".txt", ".pdf", ".png", ".jpg", ".jpeg", ".webp"}
        if extensao not in permitidas:
            flash("Formato não suportado. Use XLSX, CSV, DOCX, TXT, PDF, PNG, JPG, JPEG ou WEBP.", "erro")
            return render_template("questoes/importar.html", prova=prova, prova_id=prova_id, preview=None)

        disciplina_padrao = (request.form.get("disciplina_padrao") or (prova["disciplina"] if prova else "")).strip()
        defaults = {
            "disciplina": disciplina_padrao,
            "etapa_ensino": (request.form.get("etapa_ensino") or (prova["etapa"] if prova else "")).strip(),
            "ano_serie": (request.form.get("ano_serie") or "").strip(),
            "dificuldade": (request.form.get("dificuldade_padrao") or "Média").strip(),
        }
        with tempfile.NamedTemporaryFile(delete=False, suffix=extensao) as temp:
            arquivo.save(temp.name)
            caminho_temp = temp.name
        try:
            if extensao in {".xlsx", ".csv"}:
                linhas = _linhas_planilha_importacao(caminho_temp, extensao)
                questoes = []
                for linha in linhas:
                    q = _questao_de_linha_importacao(linha, defaults)
                    if q:
                        questoes.append(q)
            elif extensao in {".png", ".jpg", ".jpeg", ".webp"}:
                questoes = _questoes_de_imagem_com_ia(caminho_temp, extensao, defaults)
            else:
                texto = _texto_arquivo_importacao(caminho_temp, extensao)
                if not texto.strip():
                    raise ValueError("Não foi encontrado texto legível no arquivo. Se for um PDF escaneado, envie um print/foto em PNG ou JPG para leitura por IA.")
                questoes = _questoes_texto_importacao(texto, defaults)
        finally:
            try: os.remove(caminho_temp)
            except OSError: pass

        questoes, avisos = _validar_questoes_importadas(questoes)
        if not questoes:
            flash("Nenhuma questão válida foi encontrada. Confira o formato do arquivo.", "erro")
            return render_template("questoes/importar.html", prova=prova, prova_id=prova_id, preview=None, avisos=avisos)

        token = _salvar_preview_importacao_questoes(questoes, prova_id)
        return render_template(
            "questoes/importar.html", prova=prova, prova_id=prova_id,
            preview=questoes, token=token, avisos=avisos, nome_arquivo=arquivo.filename
        )
    except ValueError as erro:
        flash(str(erro), "erro")
        return render_template("questoes/importar.html", prova=prova, prova_id=prova_id, preview=None)
    finally:
        banco.close()


@app.route("/questoes/importar/confirmar", methods=["POST"])
def confirmar_importacao_questoes():
    if not cargo_permitido(["Administrador Geral", "Administrador da Instituição", "Coordenador", "Professor"]):
        return redirect("/login")
    if not permissao_modulo("Questões"):
        return redirect("/acesso_negado")

    token = request.form.get("token", "")
    payload, caminho_preview = _carregar_preview_importacao_questoes(token)
    if not payload:
        flash("A prévia da importação expirou ou não é válida. Envie o arquivo novamente.", "erro")
        return redirect("/questoes/importar")

    prova_id = payload.get("prova_id")
    questoes = payload.get("questoes") or []
    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()
    importadas = 0
    try:
        escola_id = obter_escola_usuario()
        proxima_ordem = 1
        if prova_id:
            if not _pode_gerenciar_prova(cursor, int(prova_id), exigir_edicao=True):
                return _redirecionar_acesso_negado_prova()
            cursor.execute("""
                SELECT p.disciplina, COALESCE(p.escola_id,t.escola_id) escola_id
                FROM provas p INNER JOIN turmas t ON t.id=p.turma_id WHERE p.id=?
            """, (prova_id,))
            prova = cursor.fetchone()
            if not prova:
                flash("Avaliação de destino não encontrada.", "erro")
                return redirect("/provas")
            escola_id = prova["escola_id"]
            cursor.execute("SELECT COALESCE(MAX(ordem),0)+1 AS ordem FROM prova_questoes WHERE prova_id=?", (prova_id,))
            proxima_ordem = cursor.fetchone()["ordem"]

        for q in questoes:
            alternativas = q.get("alternativas") or []
            corretas = q.get("respostas_corretas") or []
            textos = [(a.get("texto") or "") for a in alternativas[:4]] + [""] * 4
            correta_legada = corretas[0] if corretas else ""
            habilidade = " | ".join(x for x in [q.get("habilidade_bncc", ""), q.get("descritor_saeb", "")] if x)
            cursor.execute("""
                INSERT INTO questoes (
                    disciplina, etapa_ensino, ano_serie, assunto, assunto_temporario, subassunto,
                    tipo_questao, enunciado, enunciado_html, imagem,
                    alternativa_a, alternativa_b, alternativa_c, alternativa_d,
                    correta, alternativas_json, respostas_corretas, resposta_esperada,
                    criterios_correcao, habilidade, habilidade_bncc, unidade_tematica,
                    objeto_conhecimento, sistema_matriz, matriz_referencia, descritor_saeb,
                    taxonomia_bloom, dificuldade, fonte, tags, observacoes,
                    linhas_resposta, escola_id, criado_por, criado_em
                ) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
            """, (
                q.get("disciplina", ""), q.get("etapa_ensino", ""), q.get("ano_serie", ""), "", 0, "",
                q.get("tipo_questao", "multipla_escolha"), q.get("enunciado", ""), "", q.get("imagem", ""),
                textos[0], textos[1], textos[2], textos[3], correta_legada,
                json.dumps(alternativas, ensure_ascii=False), json.dumps(corretas, ensure_ascii=False),
                q.get("resposta_esperada", ""), q.get("criterios_correcao", ""), habilidade,
                q.get("habilidade_bncc", ""), q.get("unidade_tematica", ""), q.get("objeto_conhecimento", ""),
                q.get("sistema_matriz", ""), q.get("matriz_referencia", ""), q.get("descritor_saeb", ""),
                q.get("taxonomia_bloom", ""), q.get("dificuldade", "Média"), q.get("fonte", ""),
                q.get("tags", ""), q.get("observacoes", ""), 5 if q.get("tipo_questao") == "discursiva" else None,
                escola_id, session.get("usuario_id"), datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            ))
            questao_id = cursor.lastrowid
            if prova_id:
                cursor.execute("INSERT INTO prova_questoes (prova_id,questao_id,peso,ordem) VALUES (?,?,0,?)", (prova_id, questao_id, proxima_ordem))
                proxima_ordem += 1
            importadas += 1

        if prova_id:
            cursor.execute("""
                UPDATE provas SET quantidade=(SELECT COUNT(*) FROM prova_questoes WHERE prova_id=?), atualizado_em=? WHERE id=?
            """, (prova_id, datetime.now().strftime("%Y-%m-%d %H:%M:%S"), prova_id))
        banco.commit()
        try: os.remove(caminho_preview)
        except OSError: pass
        flash(f"{importadas} questão(ões) importada(s) com sucesso!", "sucesso")
        return redirect(f"/provas/{prova_id}/montar" if prova_id else "/questoes")
    except Exception as erro:
        banco.rollback()
        app.logger.exception("Erro ao confirmar importação de questões")
        flash(f"Não foi possível concluir a importação: {erro}", "erro")
        return redirect(f"/questoes/importar?prova_id={prova_id}" if prova_id else "/questoes/importar")
    finally:
        banco.close()

@app.route("/provas/<int:prova_id>/montar")
def montar_prova(prova_id):
    if not permissao_modulo("Provas"):
        return redirect("/acesso_negado")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        if not _pode_gerenciar_prova(
            cursor,
            prova_id,
            exigir_edicao=False
        ):
            return _redirecionar_acesso_negado_prova()

        cursor.execute("""
            SELECT
                p.id,
                p.nome,
                p.disciplina,
                p.data_aplicacao,
                p.status,
                COALESCE(p.peso_total, 10) AS peso_total,
                COALESCE(p.tipo_peso, 'automatico') AS tipo_peso,
                p.escola_id,
                p.ano_letivo_id,
                t.id AS turma_id,
                t.nome AS turma_nome,
                t.etapa,
                t.ano AS turma_ano,
                t.turno,
                e.nome_instituicao,
                COALESCE(
                    pr.nome,
                    'Não informado'
                ) AS professor_nome
            FROM provas AS p
            INNER JOIN turmas AS t
                ON t.id = p.turma_id
            INNER JOIN escolas AS e
                ON e.id = COALESCE(
                    p.escola_id,
                    t.escola_id
                )
            LEFT JOIN professores AS pr
                ON pr.id = p.professor_id
            WHERE p.id = ?
            LIMIT 1
        """, (prova_id,))
        prova = cursor.fetchone()

        if not prova:
            flash("Avaliação não encontrada.", "erro")
            return redirect("/provas")

        if (prova["status"] or "rascunho").strip().lower() == "finalizada":
            flash("Esta avaliação está finalizada e disponível somente para visualização.", "aviso")
            return redirect(f"/prova/{prova_id}")

        normalizar_ordem_questoes_prova(cursor, prova_id)
        banco.commit()

        cursor.execute("""
            SELECT
                pq.id AS vinculo_id,
                COALESCE(pq.peso, 0) AS peso,
                COALESCE(pq.ordem, 0) AS ordem,
                q.*
            FROM prova_questoes AS pq
            INNER JOIN questoes AS q
                ON q.id = pq.questao_id
            WHERE pq.prova_id = ?
            ORDER BY pq.ordem, pq.id
        """, (prova_id,))
        questoes_adicionadas_raw = cursor.fetchall()
        questoes_adicionadas = []
        for registro in questoes_adicionadas_raw:
            item = dict(registro)
            try:
                alternativas_render = json.loads(item.get("alternativas_json") or "[]")
            except (TypeError, ValueError, json.JSONDecodeError):
                alternativas_render = []
            if not alternativas_render and item.get("tipo_questao") in {"multipla_escolha", "multiplas_respostas"}:
                for indice, campo in enumerate(("alternativa_a", "alternativa_b", "alternativa_c", "alternativa_d")):
                    texto = item.get(campo) or ""
                    if texto:
                        alternativas_render.append({"letra": chr(65 + indice), "texto": texto, "imagem": ""})
            item["alternativas_render"] = alternativas_render
            questoes_adicionadas.append(item)

        cursor.execute("""
            SELECT q.*
            FROM questoes AS q
            WHERE q.disciplina = ?
              AND q.id NOT IN (
                    SELECT questao_id
                    FROM prova_questoes
                    WHERE prova_id = ?
              )
            ORDER BY q.id DESC
        """, (
            prova["disciplina"],
            prova_id
        ))
        banco_questoes = cursor.fetchall()

        return render_template(
            "montar_prova.html",
            prova=prova,
            questoes=questoes_adicionadas,
            banco_questoes=banco_questoes,
            total_questoes=len(questoes_adicionadas)
        )

    except sqlite3.Error as erro:
        import traceback
        traceback.print_exc()

        flash(
            f"Não foi possível abrir a montagem: {erro}",
            "erro"
        )
        return redirect("/provas")

    finally:
        banco.close()


@app.route("/provas/<int:prova_id>/banco-questoes")
def selecionar_questoes_prova(prova_id):
    if not permissao_modulo("Provas"):
        return redirect("/acesso_negado")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        if not _pode_gerenciar_prova(cursor, prova_id, exigir_edicao=True):
            flash("Esta avaliação não pode mais ser editada.", "aviso")
            return redirect(f"/provas/{prova_id}/montar")

        cursor.execute("""
            SELECT p.id, p.nome, p.disciplina, p.escola_id, p.status,
                   t.escola_id AS turma_escola_id, t.nome AS turma_nome
            FROM provas p
            INNER JOIN turmas t ON t.id = p.turma_id
            WHERE p.id = ?
            LIMIT 1
        """, (prova_id,))
        prova = cursor.fetchone()

        if not prova:
            flash("Avaliação não encontrada.", "erro")
            return redirect("/provas")

        busca = request.args.get("busca", "").strip()
        assunto = request.args.get("assunto", "").strip()
        dificuldade = request.args.get("dificuldade", "").strip()
        tipo = request.args.get("tipo", "").strip()
        etapa = request.args.get("etapa", "").strip()
        ano_serie = request.args.get("ano_serie", "").strip()

        filtros = [prova["disciplina"], prova_id]
        condicoes = [
            "q.disciplina = ?",
            "q.id NOT IN (SELECT questao_id FROM prova_questoes WHERE prova_id = ?)"
        ]

        if busca:
            condicoes.append("(q.enunciado LIKE ? OR COALESCE(q.assunto, '') LIKE ?)")
            termo = f"%{busca}%"
            filtros.extend([termo, termo])
        if assunto:
            condicoes.append("q.assunto = ?")
            filtros.append(assunto)
        if dificuldade:
            condicoes.append("q.dificuldade = ?")
            filtros.append(dificuldade)
        if tipo:
            condicoes.append("q.tipo_questao = ?")
            filtros.append(tipo)
        if etapa:
            condicoes.append("q.etapa_ensino = ?")
            filtros.append(etapa)
        if ano_serie:
            condicoes.append("q.ano_serie = ?")
            filtros.append(ano_serie)

        cursor.execute(f"""
            SELECT q.*
            FROM questoes q
            WHERE {' AND '.join(condicoes)}
            ORDER BY q.id DESC
        """, filtros)

        questoes = []
        for registro in cursor.fetchall():
            questao = dict(registro)

            try:
                alternativas = json.loads(questao.get("alternativas_json") or "[]")
            except (TypeError, ValueError, json.JSONDecodeError):
                alternativas = []

            if not alternativas:
                alternativas = []
                for indice, campo in enumerate(
                    ["alternativa_a", "alternativa_b", "alternativa_c", "alternativa_d"]
                ):
                    texto = (questao.get(campo) or "").strip()
                    if texto:
                        alternativas.append({
                            "letra": chr(65 + indice),
                            "texto": texto,
                            "imagem": ""
                        })

            questao["alternativas"] = alternativas
            questoes.append(questao)

        cursor.execute("""
            SELECT DISTINCT assunto FROM questoes
            WHERE disciplina = ? AND TRIM(COALESCE(assunto, '')) <> ''
            ORDER BY assunto
        """, (prova["disciplina"],))
        assuntos = [r["assunto"] for r in cursor.fetchall()]

        return render_template(
            "selecionar_questoes_prova.html",
            prova=prova, questoes=questoes, assuntos=assuntos,
            filtros={"busca": busca, "assunto": assunto, "dificuldade": dificuldade,
                     "tipo": tipo, "etapa": etapa, "ano_serie": ano_serie}
        )
    except sqlite3.Error as erro:
        flash(f"Não foi possível abrir o banco de questões: {erro}", "erro")
        return redirect(f"/provas/{prova_id}/montar")
    finally:
        banco.close()


@app.route("/provas/<int:prova_id>/banco-questoes/adicionar", methods=["POST"])
def adicionar_questoes_selecionadas(prova_id):
    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        if not _pode_gerenciar_prova(cursor, prova_id, exigir_edicao=True):
            flash("Esta avaliação não pode mais ser editada.", "aviso")
            return redirect(f"/provas/{prova_id}/montar")

        ids_brutos = request.form.getlist("questoes_ids")
        questoes_ids = []
        for valor in ids_brutos:
            try:
                questoes_ids.append(int(valor))
            except (TypeError, ValueError):
                continue

        if not questoes_ids:
            flash("Selecione pelo menos uma questão.", "aviso")
            return redirect(f"/provas/{prova_id}/banco-questoes")

        cursor.execute("""
            SELECT p.disciplina, p.escola_id, t.escola_id AS turma_escola_id
            FROM provas p INNER JOIN turmas t ON t.id = p.turma_id
            WHERE p.id = ? LIMIT 1
        """, (prova_id,))
        prova = cursor.fetchone()
        escola_id = prova["escola_id"] or prova["turma_escola_id"]

        cursor.execute("SELECT COALESCE(MAX(ordem), 0) AS ordem FROM prova_questoes WHERE prova_id = ?", (prova_id,))
        ordem = cursor.fetchone()["ordem"]
        adicionadas = 0

        for questao_id in dict.fromkeys(questoes_ids):
            cursor.execute("""
                SELECT id FROM questoes
                WHERE id = ? AND disciplina = ?
                LIMIT 1
            """, (questao_id, prova["disciplina"]))
            if not cursor.fetchone():
                continue

            cursor.execute("SELECT 1 FROM prova_questoes WHERE prova_id = ? AND questao_id = ?", (prova_id, questao_id))
            if cursor.fetchone():
                continue

            ordem += 1
            cursor.execute("""
                INSERT INTO prova_questoes (prova_id, questao_id, peso, ordem)
                VALUES (?, ?, 0, ?)
            """, (prova_id, questao_id, ordem))
            adicionadas += 1

        cursor.execute("""
            UPDATE provas SET quantidade = (SELECT COUNT(*) FROM prova_questoes WHERE prova_id = ?),
                atualizado_em = ? WHERE id = ?
        """, (prova_id, datetime.now().strftime("%Y-%m-%d %H:%M:%S"), prova_id))
        banco.commit()

        if adicionadas:
            flash(f"{adicionadas} questão(ões) adicionada(s) à avaliação.", "sucesso")
        else:
            flash("Nenhuma questão nova foi adicionada.", "aviso")
        return redirect(f"/provas/{prova_id}/montar")
    except sqlite3.Error as erro:
        banco.rollback()
        flash(f"Não foi possível adicionar as questões: {erro}", "erro")
        return redirect(f"/provas/{prova_id}/banco-questoes")
    finally:
        banco.close()


@app.route(
    "/provas/<int:prova_id>/questoes/adicionar",
    methods=["POST"]
)
def adicionar_questao_prova(prova_id):
    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        if not _pode_gerenciar_prova(
            cursor,
            prova_id,
            exigir_edicao=True
        ):
            return _redirecionar_acesso_negado_prova()

        try:
            questao_id = int(
                request.form.get("questao_id", "")
            )
        except (TypeError, ValueError):
            flash("Selecione uma questão válida.", "erro")
            return redirect(f"/provas/{prova_id}/montar")

        cursor.execute("""
            SELECT
                p.escola_id,
                p.disciplina,
                t.escola_id AS turma_escola_id
            FROM provas AS p
            INNER JOIN turmas AS t
                ON t.id = p.turma_id
            WHERE p.id = ?
            LIMIT 1
        """, (prova_id,))
        prova = cursor.fetchone()

        cursor.execute("""
            SELECT id, escola_id, disciplina
            FROM questoes
            WHERE id = ?
            LIMIT 1
        """, (questao_id,))
        questao = cursor.fetchone()

        if not prova or not questao:
            flash("Questão não encontrada.", "erro")
            return redirect(f"/provas/{prova_id}/montar")

        escola_prova = (
            prova["escola_id"]
            or prova["turma_escola_id"]
        )

        if (
            questao["escola_id"] is not None
            and int(questao["escola_id"]) != int(escola_prova)
        ):
            return _redirecionar_acesso_negado_prova()

        if (
            (questao["disciplina"] or "").strip().lower()
            != (prova["disciplina"] or "").strip().lower()
        ):
            flash(
                "A questão não pertence ao componente da avaliação.",
                "erro"
            )
            return redirect(f"/provas/{prova_id}/montar")

        cursor.execute("""
            SELECT id
            FROM prova_questoes
            WHERE prova_id = ?
              AND questao_id = ?
            LIMIT 1
        """, (prova_id, questao_id))

        if cursor.fetchone():
            flash(
                "Essa questão já foi adicionada à avaliação.",
                "aviso"
            )
            return redirect(f"/provas/{prova_id}/montar")

        cursor.execute("""
            SELECT COALESCE(MAX(ordem), 0) + 1 AS proxima_ordem
            FROM prova_questoes
            WHERE prova_id = ?
        """, (prova_id,))
        proxima_ordem = cursor.fetchone()["proxima_ordem"]

        cursor.execute("""
            INSERT INTO prova_questoes (
                prova_id,
                questao_id,
                peso,
                ordem
            )
            VALUES (?, ?, 0, ?)
        """, (prova_id, questao_id, proxima_ordem))

        cursor.execute("""
            UPDATE provas
            SET
                quantidade = (
                    SELECT COUNT(*)
                    FROM prova_questoes
                    WHERE prova_id = ?
                ),
                atualizado_em = ?
            WHERE id = ?
        """, (
            prova_id,
            datetime.now().strftime(
                "%Y-%m-%d %H:%M:%S"
            ),
            prova_id
        ))

        banco.commit()
        flash("Questão adicionada à avaliação.", "sucesso")

    except sqlite3.Error as erro:
        banco.rollback()
        flash(
            f"Não foi possível adicionar a questão: {erro}",
            "erro"
        )

    finally:
        banco.close()

    return redirect(f"/provas/{prova_id}/montar")




@app.route(
    "/provas/<int:prova_id>/pesos",
    methods=["POST"]
)
def salvar_pesos_prova(prova_id):
    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        if not _pode_gerenciar_prova(
            cursor,
            prova_id,
            exigir_edicao=True
        ):
            return _redirecionar_acesso_negado_prova()

        tipo_peso = (request.form.get("tipo_peso") or "automatico").strip()
        if tipo_peso not in {"automatico", "manual"}:
            tipo_peso = "automatico"

        try:
            peso_total = float(
                (request.form.get("peso_total") or "10").replace(",", ".")
            )
        except (TypeError, ValueError):
            peso_total = 10.0

        if peso_total <= 0:
            flash("O peso total deve ser maior que zero.", "erro")
            return redirect(f"/provas/{prova_id}/montar")

        cursor.execute("""
            SELECT id
            FROM prova_questoes
            WHERE prova_id = ?
            ORDER BY ordem, id
        """, (prova_id,))
        vinculos = cursor.fetchall()

        if not vinculos:
            flash("Adicione questões antes de configurar os pesos.", "aviso")
            return redirect(f"/provas/{prova_id}/montar")

        if tipo_peso == "automatico":
            quantidade = len(vinculos)
            valor_base = round(peso_total / quantidade, 2)
            acumulado = 0.0

            for indice, vinculo in enumerate(vinculos):
                if indice == quantidade - 1:
                    peso = round(peso_total - acumulado, 2)
                else:
                    peso = valor_base
                    acumulado = round(acumulado + peso, 2)

                cursor.execute("""
                    UPDATE prova_questoes
                    SET peso = ?
                    WHERE id = ? AND prova_id = ?
                """, (peso, vinculo["id"], prova_id))

        else:
            soma = 0.0
            pesos = []

            for vinculo in vinculos:
                bruto = request.form.get(f"peso_{vinculo['id']}", "0")
                try:
                    peso = float(str(bruto).replace(",", "."))
                except (TypeError, ValueError):
                    peso = 0.0

                if peso < 0:
                    flash("Os pesos não podem ser negativos.", "erro")
                    return redirect(f"/provas/{prova_id}/montar")

                peso = round(peso, 2)
                pesos.append((peso, vinculo["id"]))
                soma = round(soma + peso, 2)

            if abs(soma - peso_total) > 0.009:
                flash(
                    f"A soma dos pesos ({soma:.2f}) precisa ser igual ao peso total ({peso_total:.2f}).",
                    "erro"
                )
                return redirect(f"/provas/{prova_id}/montar")

            for peso, vinculo_id in pesos:
                cursor.execute("""
                    UPDATE prova_questoes
                    SET peso = ?
                    WHERE id = ? AND prova_id = ?
                """, (peso, vinculo_id, prova_id))

        cursor.execute("""
            UPDATE provas
            SET peso_total = ?,
                tipo_peso = ?,
                atualizado_em = ?
            WHERE id = ?
        """, (
            round(peso_total, 2),
            tipo_peso,
            datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            prova_id
        ))

        banco.commit()
        flash("Pontuação da avaliação salva com sucesso.", "sucesso")

    except sqlite3.Error as erro:
        banco.rollback()
        flash(f"Não foi possível salvar os pesos: {erro}", "erro")

    finally:
        banco.close()

    return redirect(f"/provas/{prova_id}/montar")




@app.route(
    "/provas/<int:prova_id>/questoes/<int:vinculo_id>/duplicar",
    methods=["POST"]
)
def duplicar_questao_prova(prova_id, vinculo_id):
    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        if not _pode_gerenciar_prova(
            cursor,
            prova_id,
            exigir_edicao=True
        ):
            return _redirecionar_acesso_negado_prova()

        cursor.execute("""
            SELECT q.*
            FROM prova_questoes AS pq
            INNER JOIN questoes AS q
                ON q.id = pq.questao_id
            WHERE pq.id = ?
              AND pq.prova_id = ?
            LIMIT 1
        """, (vinculo_id, prova_id))

        original = cursor.fetchone()

        if not original:
            flash("Questão não encontrada na avaliação.", "erro")
            return redirect(f"/provas/{prova_id}/montar")

        # Duplica todas as colunas existentes da questão, exceto o ID.
        # Os campos de criação/atualização são ajustados para o novo registro.
        dados = dict(original)
        dados.pop("id", None)

        agora = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        if "criado_por" in dados:
            dados["criado_por"] = session.get("usuario_id")

        if "criado_em" in dados:
            dados["criado_em"] = agora

        if "atualizado_em" in dados:
            dados["atualizado_em"] = agora

        colunas = list(dados.keys())
        marcadores = ", ".join(["?"] * len(colunas))
        nomes_colunas = ", ".join(colunas)

        cursor.execute(
            f"""
            INSERT INTO questoes ({nomes_colunas})
            VALUES ({marcadores})
            """,
            [dados[coluna] for coluna in colunas]
        )

        nova_questao_id = cursor.lastrowid

        cursor.execute("""
            SELECT COALESCE(MAX(ordem), 0) + 1 AS proxima_ordem
            FROM prova_questoes
            WHERE prova_id = ?
        """, (prova_id,))

        proxima_ordem = cursor.fetchone()["proxima_ordem"]

        cursor.execute("""
            INSERT INTO prova_questoes (
                prova_id,
                questao_id,
                peso,
                ordem
            )
            VALUES (?, ?, 0, ?)
        """, (
            prova_id,
            nova_questao_id,
            proxima_ordem
        ))

        cursor.execute("""
            UPDATE provas
            SET quantidade = (
                    SELECT COUNT(*)
                    FROM prova_questoes
                    WHERE prova_id = ?
                ),
                atualizado_em = ?
            WHERE id = ?
        """, (
            prova_id,
            agora,
            prova_id
        ))

        banco.commit()
        flash(
            "Questão duplicada e adicionada ao final da avaliação.",
            "sucesso"
        )

    except sqlite3.Error as erro:
        banco.rollback()
        flash(
            f"Não foi possível duplicar a questão: {erro}",
            "erro"
        )

    finally:
        banco.close()

    return redirect(f"/provas/{prova_id}/montar")


@app.route(
    "/provas/<int:prova_id>/questoes/<int:vinculo_id>/mover/<direcao>",
    methods=["POST"]
)
def mover_questao_prova(prova_id, vinculo_id, direcao):
    if direcao not in {"cima", "baixo"}:
        flash("Direção de movimentação inválida.", "erro")
        return redirect(f"/provas/{prova_id}/montar")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        if not _pode_gerenciar_prova(cursor, prova_id, exigir_edicao=True):
            return _redirecionar_acesso_negado_prova()

        normalizar_ordem_questoes_prova(cursor, prova_id)

        cursor.execute("""
            SELECT id, ordem
            FROM prova_questoes
            WHERE id = ? AND prova_id = ?
            LIMIT 1
        """, (vinculo_id, prova_id))
        atual = cursor.fetchone()

        if not atual:
            flash("Questão não encontrada na avaliação.", "erro")
            return redirect(f"/provas/{prova_id}/montar")

        operador = "<" if direcao == "cima" else ">"
        ordenacao = "DESC" if direcao == "cima" else "ASC"

        cursor.execute(f"""
            SELECT id, ordem
            FROM prova_questoes
            WHERE prova_id = ? AND ordem {operador} ?
            ORDER BY ordem {ordenacao}, id {ordenacao}
            LIMIT 1
        """, (prova_id, atual["ordem"]))
        vizinha = cursor.fetchone()

        if vizinha:
            cursor.execute("""
                UPDATE prova_questoes SET ordem = ?
                WHERE id = ? AND prova_id = ?
            """, (vizinha["ordem"], atual["id"], prova_id))
            cursor.execute("""
                UPDATE prova_questoes SET ordem = ?
                WHERE id = ? AND prova_id = ?
            """, (atual["ordem"], vizinha["id"], prova_id))
            cursor.execute("""
                UPDATE provas SET atualizado_em = ? WHERE id = ?
            """, (datetime.now().strftime("%Y-%m-%d %H:%M:%S"), prova_id))
            banco.commit()
            flash("Ordem das questões atualizada.", "sucesso")

    except sqlite3.Error as erro:
        banco.rollback()
        flash(f"Não foi possível reordenar a questão: {erro}", "erro")
    finally:
        banco.close()

    return redirect(f"/provas/{prova_id}/montar")


@app.route(
    "/provas/<int:prova_id>/questoes/<int:vinculo_id>/remover",
    methods=["POST"]
)
def remover_questao_prova(prova_id, vinculo_id):
    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        if not _pode_gerenciar_prova(
            cursor,
            prova_id,
            exigir_edicao=True
        ):
            return _redirecionar_acesso_negado_prova()

        cursor.execute("""
            DELETE FROM prova_questoes
            WHERE id = ?
              AND prova_id = ?
        """, (vinculo_id, prova_id))

        normalizar_ordem_questoes_prova(cursor, prova_id)

        cursor.execute("""
            UPDATE provas
            SET
                quantidade = (
                    SELECT COUNT(*)
                    FROM prova_questoes
                    WHERE prova_id = ?
                ),
                atualizado_em = ?
            WHERE id = ?
        """, (
            prova_id,
            datetime.now().strftime(
                "%Y-%m-%d %H:%M:%S"
            ),
            prova_id
        ))

        banco.commit()
        flash("Questão removida da avaliação.", "sucesso")

    except sqlite3.Error as erro:
        banco.rollback()
        flash(
            f"Não foi possível remover a questão: {erro}",
            "erro"
        )

    finally:
        banco.close()

    return redirect(f"/provas/{prova_id}/montar")


@app.route(
    "/provas/<int:prova_id>/finalizar",
    methods=["POST"]
)
def finalizar_prova(prova_id):
    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        if not _pode_gerenciar_prova(
            cursor,
            prova_id,
            exigir_edicao=True
        ):
            return _redirecionar_acesso_negado_prova()

        cursor.execute("""
            SELECT COUNT(*) AS total
            FROM prova_questoes
            WHERE prova_id = ?
        """, (prova_id,))
        total = cursor.fetchone()["total"]

        if total <= 0:
            flash(
                "Adicione pelo menos uma questão antes de finalizar.",
                "aviso"
            )
            return redirect(f"/provas/{prova_id}/montar")

        cursor.execute("""
            UPDATE provas
            SET
                quantidade = ?,
                status = 'agendada',
                atualizado_em = ?
            WHERE id = ?
        """, (
            total,
            datetime.now().strftime(
                "%Y-%m-%d %H:%M:%S"
            ),
            prova_id
        ))

        banco.commit()
        flash("Avaliação finalizada e agendada com sucesso.", "sucesso")

        return redirect(f"/prova/{prova_id}")

    except sqlite3.Error as erro:
        banco.rollback()
        flash(
            f"Não foi possível finalizar a avaliação: {erro}",
            "erro"
        )
        return redirect(f"/provas/{prova_id}/montar")

    finally:
        banco.close()

@app.route("/prova/<int:prova_id>")
@app.route("/provas/<int:prova_id>")
def visualizar_prova(prova_id):
    if not permissao_modulo("Provas"):
        return redirect("/acesso_negado")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        # Verifica se o usuário pode visualizar a avaliação.
        if not _pode_gerenciar_prova(
            cursor,
            prova_id,
            exigir_edicao=False
        ):
            return _redirecionar_acesso_negado_prova()

        # =====================================================
        # DADOS PRINCIPAIS DA AVALIAÇÃO
        # =====================================================
        cursor.execute("""
            SELECT
                p.id,
                p.nome,
                p.turma_id,
                t.nome AS turma_nome,
                t.ano AS turma_ano,
                t.turno AS turma_turno,
                p.disciplina,
                COUNT(pq.id) AS quantidade,
                COALESCE(pr.nome, 'Não informado') AS professor_nome,
                p.data_geracao,
                p.data_aplicacao,
                e.id AS escola_id,
                e.nome_instituicao,
                e.cidade,
                e.estado,
                e.logo
            FROM provas AS p

            INNER JOIN turmas AS t
                ON t.id = p.turma_id

            INNER JOIN escolas AS e
                ON e.id = COALESCE(
                    p.escola_id,
                    t.escola_id
                )

            LEFT JOIN professores AS pr
                ON pr.id = p.professor_id

            LEFT JOIN prova_questoes AS pq
                ON pq.prova_id = p.id

            WHERE p.id = ?

            GROUP BY
                p.id,
                p.nome,
                p.turma_id,
                t.nome,
                t.ano,
                t.turno,
                p.disciplina,
                pr.nome,
                p.data_geracao,
                p.data_aplicacao,
                e.id,
                e.nome_instituicao,
                e.cidade,
                e.estado,
                e.logo

            LIMIT 1
        """, (prova_id,))

        prova = cursor.fetchone()

        if prova is None:
            flash(
                "Avaliação não encontrada.",
                "erro"
            )
            return redirect(url_for("provas"))

        # =====================================================
        # QUESTÕES DA AVALIAÇÃO
        # =====================================================
        cursor.execute("""
            SELECT
                q.*,
                pq.id AS prova_questao_id,
                COALESCE(pq.peso, 0) AS peso,
                COALESCE(pq.anulada, 0) AS anulada,
                COALESCE(pq.ordem, pq.id) AS ordem_prova

            FROM prova_questoes AS pq

            INNER JOIN questoes AS q
                ON q.id = pq.questao_id

            WHERE pq.prova_id = ?

            ORDER BY
                COALESCE(
                    NULLIF(pq.ordem, 0),
                    pq.id
                ),
                pq.id
        """, (prova_id,))

        questoes_banco = cursor.fetchall()
        questoes = []
        for registro in questoes_banco:
            item = dict(registro)
            try:
                alternativas_render = json.loads(item.get("alternativas_json") or "[]")
            except (TypeError, ValueError, json.JSONDecodeError):
                alternativas_render = []
            if not alternativas_render and item.get("tipo_questao") in {"multipla_escolha", "multiplas_respostas"}:
                for indice, campo in enumerate(("alternativa_a", "alternativa_b", "alternativa_c", "alternativa_d")):
                    texto = item.get(campo) or ""
                    if texto:
                        alternativas_render.append({"letra": chr(65 + indice), "texto": texto, "imagem": ""})
            item["alternativas_render"] = alternativas_render
            questoes.append(item)

        # =====================================================
        # RESUMO DOS PESOS
        # =====================================================
        peso_total = sum(
            float(questao["peso"] or 0)
            for questao in questoes
        )

        quantidade_anuladas = sum(
            1
            for questao in questoes
            if int(questao["anulada"] or 0) == 1
        )

        # =====================================================
        # DADOS DA INSTITUIÇÃO
        # =====================================================
        instituicao = {
            "id": prova["escola_id"],
            "nome": (
                prova["nome_instituicao"]
                or "ARK EDUS"
            ),
            "cidade": (
                prova["cidade"]
                or "Não informada"
            ),
            "estado": (
                prova["estado"]
                or "Não informado"
            ),
            "logo": prova["logo"]
        }

        return render_template(
            "visualizar_prova.html",
            prova=prova,
            questoes=questoes,
            instituicao=instituicao,
            peso_total=peso_total,
            quantidade_anuladas=quantidade_anuladas
        )

    except sqlite3.Error as erro:
        print(
            "ERRO AO VISUALIZAR AVALIAÇÃO:",
            erro
        )

        flash(
            f"Não foi possível abrir a avaliação: {erro}",
            "erro"
        )

        return redirect(url_for("provas"))

    except (TypeError, ValueError) as erro:
        print(
            "ERRO AO CALCULAR OS DADOS DA AVALIAÇÃO:",
            erro
        )

        flash(
            "Não foi possível calcular os dados da avaliação.",
            "erro"
        )

        return redirect(url_for("provas"))

    finally:
        banco.close()

@app.route(
    "/prova/<int:prova_id>/questao/<int:prova_questao_id>/alternar-anulacao",
    methods=["POST"]
)
def alternar_anulacao_questao(prova_id, prova_questao_id):
    if not permissao_modulo("Provas"):
        return redirect("/acesso_negado")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        if not _pode_gerenciar_prova(
            cursor,
            prova_id,
            exigir_edicao=False
        ):
            return _redirecionar_acesso_negado_prova()

        cursor.execute("""
            SELECT
                id,
                COALESCE(anulada, 0) AS anulada
            FROM prova_questoes
            WHERE id = ?
              AND prova_id = ?
            LIMIT 1
        """, (prova_questao_id, prova_id))

        prova_questao = cursor.fetchone()

        if prova_questao is None:
            flash("Questão não encontrada nesta avaliação.", "erro")
            return redirect(url_for("visualizar_prova", prova_id=prova_id))

        novo_estado = 0 if int(prova_questao["anulada"] or 0) == 1 else 1

        cursor.execute("""
            UPDATE prova_questoes
            SET anulada = ?
            WHERE id = ?
              AND prova_id = ?
        """, (novo_estado, prova_questao_id, prova_id))

        banco.commit()

        if novo_estado == 1:
            flash("Questão anulada com sucesso.", "success")
        else:
            flash("Questão reativada com sucesso.", "success")

        return redirect(
            url_for("visualizar_prova", prova_id=prova_id)
            + f"#questao-{prova_questao_id}"
        )

    except sqlite3.Error as erro:
        banco.rollback()
        flash(
            f"Não foi possível alterar a questão: {erro}",
            "erro"
        )
        return redirect(url_for("visualizar_prova", prova_id=prova_id))

    finally:
        banco.close()


@app.route("/cartao_resposta/<int:prova_id>")
def cartao_resposta(prova_id):
    if "usuario_id" not in session:
        return redirect("/login")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        cursor.execute("""
            SELECT
                p.id,
                p.nome,
                p.disciplina,
                p.turma_id,
                p.data_aplicacao,
                p.escola_id,
                t.nome AS turma_nome,
                t.ano_letivo_id,
                COALESCE(pr.nome, 'Não informado') AS professor_nome,
                e.nome_instituicao,
                e.cidade,
                e.estado,
                e.logo,
                COALESCE(al.ano, e.ano_letivo) AS ano_letivo
            FROM provas AS p
            INNER JOIN turmas AS t
                ON t.id = p.turma_id
            LEFT JOIN professores AS pr
                ON pr.id = p.professor_id
            LEFT JOIN escolas AS e
                ON e.id = COALESCE(p.escola_id, t.escola_id)
            LEFT JOIN anos_letivos AS al
                ON al.id = t.ano_letivo_id
            WHERE p.id = ?
            LIMIT 1
        """, (prova_id,))

        prova = cursor.fetchone()

        if not prova:
            flash("Avaliação não encontrada.", "erro")
            return redirect("/provas")

        # Recupera a ordem real das questões da avaliação. A posição exibida
        # no cartão é a mesma posição usada na prova impressa.
        cursor.execute("""
            SELECT
                q.id,
                LOWER(TRIM(COALESCE(q.tipo_questao, 'multipla_escolha'))) AS tipo_questao,
                COALESCE(q.linhas_resposta, 5) AS linhas_resposta,
                COALESCE(NULLIF(pq.ordem, 0), pq.id) AS ordem_exibicao
            FROM prova_questoes AS pq
            INNER JOIN questoes AS q
                ON q.id = pq.questao_id
            WHERE pq.prova_id = ?
            ORDER BY
                CASE WHEN COALESCE(pq.ordem, 0) > 0 THEN pq.ordem ELSE pq.id END,
                pq.id
        """, (prova_id,))

        questoes_banco = cursor.fetchall()
        questoes_cartao = []
        questoes_discursivas = []

        tipos_discursivos = {
            "discursiva",
            "dissertativa",
            "resposta_aberta",
            "resposta aberta"
        }

        for numero, questao in enumerate(questoes_banco, start=1):
            tipo_original = (questao["tipo_questao"] or "").strip().lower()
            eh_discursiva = tipo_original in tipos_discursivos

            try:
                linhas = int(questao["linhas_resposta"] or 5)
            except (TypeError, ValueError):
                linhas = 5

            linhas = max(1, min(linhas, 30))

            item = {
                "id": questao["id"],
                "numero": numero,
                "tipo_questao": "discursiva" if eh_discursiva else "objetiva",
                "linhas_resposta": linhas
            }
            questoes_cartao.append(item)

            if eh_discursiva:
                questoes_discursivas.append(item)

        quantidade = len(questoes_cartao)

        # Usa a matrícula do ano letivo atual quando ela existir e mantém
        # compatibilidade com alunos cadastrados no modelo antigo.
        cursor.execute("""
            SELECT DISTINCT
                a.id,
                a.nome,
                COALESCE(a.matricula, '') AS matricula
            FROM alunos AS a
            LEFT JOIN aluno_matriculas AS am
                ON am.aluno_id = a.id
               AND am.turma_id = ?
               AND (? IS NULL OR am.ano_letivo_id = ?)
               AND COALESCE(am.situacao, 'Cursando') = 'Cursando'
            WHERE am.id IS NOT NULL
               OR a.turma_id = ?
            ORDER BY a.nome COLLATE NOCASE
        """, (
            prova["turma_id"],
            prova["ano_letivo_id"],
            prova["ano_letivo_id"],
            prova["turma_id"]
        ))
        alunos = cursor.fetchall()

        instituicao = {
            "nome": prova["nome_instituicao"] or "ARK EDUS",
            "cidade": prova["cidade"] or "",
            "estado": prova["estado"] or "",
            "logo": prova["logo"],
            "ano_letivo": prova["ano_letivo"] or ""
        }

        def gerar_qr_base64(conteudo):
            qr = qrcode.QRCode(
                version=2,
                error_correction=qrcode.constants.ERROR_CORRECT_M,
                box_size=5,
                border=2
            )
            qr.add_data(conteudo)
            qr.make(fit=True)

            imagem_qr = qr.make_image(
                fill_color="black",
                back_color="white"
            )
            buffer = BytesIO()
            imagem_qr.save(buffer, format="PNG")
            return base64.b64encode(buffer.getvalue()).decode("utf-8")

        cartoes = []

        for aluno in alunos:
            codigo_base = (
                f"PROVA:{prova['id']}|ALUNO:{aluno['id']}|"
                f"TURMA:{prova['turma_nome']}"
            )

            cartoes.append({
                "aluno": aluno,
                "qr_base64": gerar_qr_base64(
                    f"{codigo_base}|FOLHA:OBJETIVAS"
                ),
                "qr_discursiva_base64": gerar_qr_base64(
                    f"{codigo_base}|FOLHA:DISCURSIVAS"
                )
            })

        return render_template(
            "cartao_resposta.html",
            prova=prova,
            quantidade=quantidade,
            instituicao=instituicao,
            cartoes=cartoes,
            questoes_cartao=questoes_cartao,
            questoes_discursivas=questoes_discursivas
        )

    except sqlite3.Error as erro:
        print("ERRO AO GERAR CARTÕES-RESPOSTA:", erro)
        flash(f"Não foi possível gerar os cartões-resposta: {erro}", "erro")
        return redirect("/provas")

    finally:
        banco.close()

@app.route("/instituicao")
def instituicao():

    if not cargo_permitido([
        "Administrador Geral",
        "Administrador da Instituição",
        "Coordenador",
        "Secretaria"
    ]):
        return redirect("/login")

    banco = conectar_banco()
    cursor = banco.cursor()

    cursor.execute("""
        SELECT *
        FROM instituicao
        WHERE id = 1
    """)

    dados = cursor.fetchone()

    banco.close()

    return render_template(
        "instituicao.html",
        instituicao=dados
    )

@app.route("/salvar_instituicao", methods=["POST"])
def salvar_instituicao():

    if not cargo_permitido([
        "Administrador Geral",
        "Administrador da Instituição",
        "Coordenador",
        "Secretaria"
    ]):
        return redirect("/login")

    nome = request.form["nome"]
    cidade = request.form["cidade"]
    estado = request.form["estado"]
    diretor = request.form["diretor"]
    coordenador = request.form["coordenador"]
    ano_letivo = request.form["ano_letivo"]

    logo = request.files.get("logo")
    nome_logo = ""

    if logo and logo.filename != "":
        nome_logo = secure_filename(logo.filename)

        logo.save(
            os.path.join(
                app.config["UPLOAD_FOLDER"],
                nome_logo
            )
        )

    banco = conectar_banco()
    cursor = banco.cursor()

    cursor.execute(
        "SELECT * FROM instituicao WHERE id = 1"
    )

    existe = cursor.fetchone()

    if existe:

        if nome_logo:

            cursor.execute("""
                UPDATE instituicao
                SET nome = ?,
                    cidade = ?,
                    estado = ?,
                    diretor = ?,
                    coordenador = ?,
                    ano_letivo = ?,
                    logo = ?
                WHERE id = 1
            """, (
                nome,
                cidade,
                estado,
                diretor,
                coordenador,
                ano_letivo,
                nome_logo
            ))

        else:

            cursor.execute("""
                UPDATE instituicao
                SET nome = ?,
                    cidade = ?,
                    estado = ?,
                    diretor = ?,
                    coordenador = ?,
                    ano_letivo = ?
                WHERE id = 1
            """, (
                nome,
                cidade,
                estado,
                diretor,
                coordenador,
                ano_letivo
            ))

    else:

        cursor.execute("""
            INSERT INTO instituicao (
                id,
                nome,
                cidade,
                estado,
                diretor,
                coordenador,
                ano_letivo,
                logo
            )
            VALUES (1, ?, ?, ?, ?, ?, ?, ?)
        """, (
            nome,
            cidade,
            estado,
            diretor,
            coordenador,
            ano_letivo,
            nome_logo
        ))

    banco.commit()
    banco.close()

    return redirect("/instituicao")

@app.route(
    "/excluir_prova/<int:prova_id>",
    methods=["GET", "POST"]
)
def excluir_prova(prova_id):
    """Exclui uma avaliação após validar sessão, módulo e vínculo do usuário."""

    if not cargo_permitido([
        "Administrador Geral",
        "Administrador da Instituição",
        "Coordenador",
        "Professor"
    ]):
        return redirect("/login")

    if not permissao_modulo("Provas"):
        return redirect("/acesso_negado")

    cargo = (session.get("usuario_cargo") or "").strip()

    if cargo not in {
        "Administrador Geral",
        "Administrador da Instituição",
        "Coordenador",
        "Professor"
    }:
        flash(
            "Você não possui permissão para excluir avaliações.",
            "erro"
        )
        return redirect("/provas")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        # Usa a mesma validação de instituição, ano letivo e professor
        # aplicada nas demais ações do módulo de avaliações.
        if not _pode_gerenciar_prova(
            cursor,
            prova_id,
            exigir_edicao=False
        ):
            flash(
                "Avaliação não encontrada ou você não possui permissão para excluí-la.",
                "erro"
            )
            return redirect("/provas")

        cursor.execute("""
            SELECT nome
            FROM provas
            WHERE id = ?
            LIMIT 1
        """, (prova_id,))

        prova = cursor.fetchone()

        if prova is None:
            flash("Avaliação não encontrada.", "erro")
            return redirect("/provas")

        # O projeto não ativa PRAGMA foreign_keys em todas as conexões.
        # Por isso, os registros dependentes são removidos explicitamente.
        cursor.execute(
            "DELETE FROM respostas_alunos WHERE prova_id = ?",
            (prova_id,)
        )

        cursor.execute(
            "DELETE FROM resultados WHERE prova_id = ?",
            (prova_id,)
        )

        cursor.execute(
            "DELETE FROM prova_questoes WHERE prova_id = ?",
            (prova_id,)
        )

        cursor.execute(
            "DELETE FROM provas WHERE id = ?",
            (prova_id,)
        )

        if cursor.rowcount == 0:
            banco.rollback()
            flash("A avaliação não foi encontrada.", "erro")
            return redirect("/provas")

        banco.commit()

        flash(
            f'A avaliação “{prova["nome"]}” foi excluída com sucesso.',
            "sucesso"
        )

    except sqlite3.Error as erro:
        banco.rollback()
        print("ERRO AO EXCLUIR AVALIAÇÃO:", erro)
        flash(
            f"Não foi possível excluir a avaliação: {erro}",
            "erro"
        )

    finally:
        banco.close()

    return redirect("/provas")



@app.route("/editar_prova/<int:prova_id>")
def editar_prova(prova_id):
    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    cursor.execute("SELECT status FROM provas WHERE id = ?", (prova_id,))
    status_prova = cursor.fetchone()
    if status_prova and (status_prova["status"] or "rascunho").lower() == "finalizada":
        banco.close()
        flash("Esta avaliação já foi finalizada e não pode mais ser editada.", "aviso")
        return redirect(f"/provas/{prova_id}/montar")

    if not _pode_gerenciar_prova(cursor, prova_id, exigir_edicao=True):
        banco.close()
        return _redirecionar_acesso_negado_prova()

    cursor.execute("SELECT * FROM provas WHERE id = ?", (prova_id,))
    prova = cursor.fetchone()

    cursor.execute("""
        SELECT * FROM turmas
        WHERE escola_id = ? AND ano_letivo_id = ?
        ORDER BY nome
    """, (prova["escola_id"], prova["ano_letivo_id"]))
    turmas = cursor.fetchall()

    cursor.execute("""
        SELECT * FROM professores
        WHERE escola_id = ?
        ORDER BY nome
    """, (prova["escola_id"],))
    professores = cursor.fetchall()

    banco.close()

    return render_template(
        "editar_prova.html",
        prova=prova,
        turmas=turmas,
        professores=professores
    )


@app.route("/atualizar_prova/<int:prova_id>", methods=["POST"])
def atualizar_prova(prova_id):
    nome = request.form.get("nome", "").strip()
    disciplina = request.form.get("disciplina", "").strip()
    data_aplicacao = request.form.get("data_aplicacao", "").strip()
    tem_nota = 1 if request.form.get("tem_nota") == "1" else 0
    media_ativa = 1 if request.form.get("media_ativa") == "1" else 0
    media_aprovacao = None

    if not nome or not disciplina:
        flash("Preencha os dados obrigatórios da avaliação.", "erro")
        return redirect(f"/editar_prova/{prova_id}")

    try:
        turma_id = int(request.form.get("turma_id", ""))
        professor_id = int(request.form.get("professor_id", ""))
    except (TypeError, ValueError):
        flash("Selecione uma turma e um professor válidos.", "erro")
        return redirect(f"/editar_prova/{prova_id}")

    if media_ativa:
        try:
            media_aprovacao = float(
                request.form.get("media_aprovacao", "").strip().replace(",", ".")
            )
        except (TypeError, ValueError):
            flash("Informe uma média válida entre 0,0 e 10,0.", "erro")
            return redirect(f"/editar_prova/{prova_id}")

        if not 0 <= media_aprovacao <= 10:
            flash("A média deve estar entre 0,0 e 10,0.", "erro")
            return redirect(f"/editar_prova/{prova_id}")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        if not _pode_gerenciar_prova(cursor, prova_id, exigir_edicao=True):
            return _redirecionar_acesso_negado_prova()

        cursor.execute("SELECT escola_id, ano_letivo_id FROM provas WHERE id = ?", (prova_id,))
        prova_atual = cursor.fetchone()

        cursor.execute("""
            SELECT 1 FROM turmas
            WHERE id = ? AND escola_id = ? AND ano_letivo_id = ?
        """, (turma_id, prova_atual["escola_id"], prova_atual["ano_letivo_id"]))
        if not cursor.fetchone():
            flash("A turma selecionada não pertence à avaliação.", "erro")
            return redirect(f"/editar_prova/{prova_id}")

        cursor.execute("""
            SELECT 1 FROM professores
            WHERE id = ? AND escola_id = ?
        """, (professor_id, prova_atual["escola_id"]))
        if not cursor.fetchone():
            flash("O professor selecionado não pertence à instituição.", "erro")
            return redirect(f"/editar_prova/{prova_id}")

        cursor.execute("""
            UPDATE provas
            SET nome = ?, turma_id = ?, professor_id = ?, disciplina = ?,
                data_aplicacao = ?, media_ativa = ?, media_aprovacao = ?, tem_nota = ?,
                atualizado_em = ?
            WHERE id = ?
        """, (
            nome, turma_id, professor_id, disciplina, data_aplicacao,
            media_ativa, media_aprovacao, tem_nota,
            datetime.now().strftime("%Y-%m-%d %H:%M:%S"), prova_id
        ))
        banco.commit()
        flash("Avaliação atualizada com sucesso.", "sucesso")
        return redirect("/provas")

    except sqlite3.Error as erro:
        banco.rollback()
        flash(f"Erro ao atualizar a avaliação: {erro}", "erro")
        return redirect(f"/editar_prova/{prova_id}")
    finally:
        banco.close()

# ==========================
# EXCLUIR PROFESSOR
# ==========================

@app.route("/excluir_professor/<int:professor_id>")
def excluir_professor(professor_id):

    if not cargo_permitido([
        "Administrador Geral",
        "Administrador da Instituição",
        "Coordenador"
    ]):
        return redirect("/login")

    banco = conectar_banco()
    cursor = banco.cursor()

    cursor.execute(
        "DELETE FROM professor_disciplinas WHERE professor_id = ?",
        (professor_id,)
    )

    cursor.execute(
        "DELETE FROM professor_turmas WHERE professor_id = ?",
        (professor_id,)
    )

    cursor.execute(
        "UPDATE provas SET professor_id = NULL WHERE professor_id = ?",
        (professor_id,)
    )

    cursor.execute(
        "DELETE FROM professores WHERE id = ?",
        (professor_id,)
    )

    banco.commit()
    banco.close()

    return redirect("/professores")

# ==========================
# TELA EDITAR PROFESSOR
# ==========================

@app.route("/editar_professor/<int:professor_id>")
def editar_professor(professor_id):

    if not cargo_permitido([
        "Administrador Geral",
        "Administrador da Instituição",
        "Coordenador"
    ]):
        return redirect("/login")

    banco = conectar_banco()
    cursor = banco.cursor()

    cursor.execute(
        "SELECT * FROM professores WHERE id = ?",
        (professor_id,)
    )
    professor = cursor.fetchone()

    cursor.execute(
        "SELECT disciplina FROM professor_disciplinas WHERE professor_id = ?",
        (professor_id,)
    )
    disciplinas = [d[0] for d in cursor.fetchall()]

    cursor.execute(
        "SELECT turma_id FROM professor_turmas WHERE professor_id = ?",
        (professor_id,)
    )
    turmas_vinculadas = [t[0] for t in cursor.fetchall()]

    cursor.execute("SELECT * FROM turmas ORDER BY nome")
    turmas = cursor.fetchall()

    banco.close()

    return render_template(
        "editar_professor.html",
        professor=professor,
        disciplinas=disciplinas,
        turmas=turmas,
        turmas_vinculadas=turmas_vinculadas
    )

# ==========================
# SALVAR EDIÇÃO PROFESSOR
# ==========================

@app.route("/atualizar_professor/<int:professor_id>", methods=["POST"])
def atualizar_professor(professor_id):

    nome = request.form["nome"]
    email = request.form["email"]

    disciplinas = request.form.getlist("disciplinas")
    turmas = request.form.getlist("turmas")

    banco = conectar_banco()
    cursor = banco.cursor()

    cursor.execute("""
        UPDATE professores
        SET nome = ?, email = ?, disciplina = ?
        WHERE id = ?
    """, (
        nome,
        email,
        ", ".join(disciplinas),
        professor_id
    ))

    cursor.execute(
        "DELETE FROM professor_disciplinas WHERE professor_id = ?",
        (professor_id,)
    )

    for disciplina in disciplinas:
        cursor.execute("""
            INSERT INTO professor_disciplinas
            (professor_id, disciplina)
            VALUES (?, ?)
        """, (professor_id, disciplina))

    cursor.execute(
        "DELETE FROM professor_turmas WHERE professor_id = ?",
        (professor_id,)
    )

    for turma_id in turmas:
        cursor.execute("""
            INSERT INTO professor_turmas
            (professor_id, turma_id)
            VALUES (?, ?)
        """, (professor_id, turma_id))

    banco.commit()
    banco.close()

    return redirect("/professores")

# ==========================
# AÇÕES EM LOTE - ALUNOS
# ==========================

@app.route("/alunos/acoes-lote", methods=["POST"])
def acoes_lote_alunos():
    if not cargo_permitido(["Administrador Geral", "Administrador da Instituição", "Coordenador", "Secretaria"]):
        flash("Você não possui permissão para alterar alunos em lote.", "erro")
        return redirect("/acesso_negado")

    ids = []
    for valor in request.form.getlist("aluno_ids"):
        try:
            ids.append(int(valor))
        except (TypeError, ValueError):
            pass
    ids = list(dict.fromkeys(ids))
    if not ids:
        flash("Selecione pelo menos um aluno.", "erro")
        return redirect("/alunos")

    acao = (request.form.get("acao") or "").strip()
    cargo = session.get("usuario_cargo", "").strip()
    escola_usuario_id = obter_escola_usuario()
    banco = conectar_banco(); banco.row_factory = sqlite3.Row; cursor = banco.cursor()
    try:
        placeholders = ",".join("?" for _ in ids)
        cursor.execute(f"SELECT id, escola_id FROM alunos WHERE id IN ({placeholders})", ids)
        encontrados = cursor.fetchall()
        permitidos = [r["id"] for r in encontrados if cargo == "Administrador Geral" or r["escola_id"] == escola_usuario_id]
        if len(permitidos) != len(ids):
            flash("Há alunos selecionados fora da sua instituição.", "erro")
            return redirect("/alunos")

        ph = ",".join("?" for _ in permitidos)
        if acao == "mover":
            try: turma_id = int(request.form.get("turma_destino_id") or 0)
            except ValueError: turma_id = 0
            cursor.execute("SELECT id, escola_id, ano_letivo_id FROM turmas WHERE id = ? LIMIT 1", (turma_id,))
            turma = cursor.fetchone()
            if not turma or (cargo != "Administrador Geral" and turma["escola_id"] != escola_usuario_id):
                flash("Turma de destino inválida.", "erro"); return redirect("/alunos")
            # Evita mover cadastros entre instituições por engano.
            if any(r["escola_id"] != turma["escola_id"] for r in encontrados):
                flash("A turma de destino deve pertencer à mesma instituição dos alunos selecionados.", "erro"); return redirect("/alunos")
            cursor.execute(f"UPDATE alunos SET turma_id = ?, escola_id = ?, ano_letivo_id = ? WHERE id IN ({ph})", [turma_id, turma["escola_id"], turma["ano_letivo_id"], *permitidos])
            flash(f"{len(permitidos)} aluno(s) movido(s) com sucesso.", "sucesso")
        elif acao == "excluir":
            cursor.execute("PRAGMA foreign_keys = ON")
            cursor.execute(f"DELETE FROM alunos WHERE id IN ({ph})", permitidos)
            flash(f"{len(permitidos)} aluno(s) excluído(s) com sucesso.", "sucesso")
        else:
            flash("Ação em lote inválida.", "erro"); return redirect("/alunos")
        banco.commit()
    except sqlite3.IntegrityError:
        banco.rollback(); flash("Não foi possível concluir a ação porque existem registros vinculados a um ou mais alunos.", "erro")
    except Exception:
        banco.rollback(); app.logger.exception("Erro em ação em lote de alunos"); flash("Não foi possível concluir a ação em lote.", "erro")
    finally:
        banco.close()
    return redirect("/alunos")

# ==========================
# EXCLUIR ALUNO
# ==========================

@app.route(
    "/excluir_aluno/<int:aluno_id>",
    methods=["GET", "POST"]
)
def excluir_aluno(aluno_id):

    if not cargo_permitido([
        "Administrador Geral",
        "Administrador da Instituição",
        "Coordenador",
        "Secretaria"
    ]):
        flash(
            "Você não possui permissão para excluir alunos.",
            "erro"
        )
        return redirect("/acesso_negado")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    cargo = session.get("usuario_cargo", "").strip()
    escola_usuario_id = obter_escola_usuario()

    try:
        cursor.execute("PRAGMA foreign_keys = ON")

        cursor.execute("""
            SELECT
                id,
                nome,
                escola_id
            FROM alunos
            WHERE id = ?
            LIMIT 1
        """, (aluno_id,))

        aluno = cursor.fetchone()

        if aluno is None:
            flash(
                "O aluno informado não foi encontrado.",
                "erro"
            )
            return redirect("/alunos")

        if (
            cargo != "Administrador Geral"
            and aluno["escola_id"] != escola_usuario_id
        ):
            flash(
                "Você não possui permissão para excluir este aluno.",
                "erro"
            )
            return redirect("/alunos")

        nome_aluno = aluno["nome"] or "Aluno"

        cursor.execute("""
            DELETE FROM alunos
            WHERE id = ?
        """, (aluno_id,))

        if cursor.rowcount == 0:
            banco.rollback()
            flash(
                "Não foi possível localizar o aluno para exclusão.",
                "erro"
            )
            return redirect("/alunos")

        banco.commit()

        flash(
            f'Aluno "{nome_aluno}" excluído com sucesso.',
            "success"
        )

        return redirect("/alunos")

    except sqlite3.IntegrityError as erro:
        banco.rollback()

        print(
            "ERRO DE INTEGRIDADE AO EXCLUIR ALUNO:",
            erro
        )

        flash(
            "Não foi possível excluir o aluno porque existem "
            "registros relacionados que impedem a exclusão.",
            "erro"
        )

        return redirect("/alunos")

    except sqlite3.Error as erro:
        banco.rollback()

        import traceback
        traceback.print_exc()

        print("ERRO AO EXCLUIR ALUNO:", erro)

        flash(
            f"Não foi possível excluir o aluno: {erro}",
            "erro"
        )

        return redirect("/alunos")

    finally:
        banco.close()

# ==========================
# TELA EDITAR ALUNO
# ==========================

@app.route("/editar_aluno/<int:aluno_id>")
def editar_aluno(aluno_id):

    if not cargo_permitido([
        "Administrador Geral",
        "Administrador da Instituição",
        "Coordenador",
        "Secretaria"
    ]):
        return redirect("/login")

    banco = conectar_banco()
    cursor = banco.cursor()

    cursor.execute(
        "SELECT * FROM alunos WHERE id = ?",
        (aluno_id,)
    )
    aluno = cursor.fetchone()

    cursor.execute(
        "SELECT * FROM turmas ORDER BY nome"
    )
    turmas = cursor.fetchall()

    banco.close()

    return render_template(
        "editar_aluno.html",
        aluno=aluno,
        turmas=turmas
    )

# ==========================
# SALVAR EDIÇÃO ALUNO
# ==========================

@app.route("/atualizar_aluno/<int:aluno_id>", methods=["POST"])
def atualizar_aluno(aluno_id):

    nome = request.form["nome"]
    matricula = request.form["matricula"]
    turma_id = request.form["turma_id"]

    banco = conectar_banco()
    cursor = banco.cursor()

    cursor.execute("""
        UPDATE alunos
        SET nome = ?, matricula = ?, turma_id = ?
        WHERE id = ?
    """, (
        nome,
        matricula,
        turma_id,
        aluno_id
    ))

    banco.commit()
    banco.close()

    return redirect("/alunos")

# ==========================
# ATUALIZAR TURMA
# ==========================

@app.route("/atualizar_turma/<int:turma_id>", methods=["POST"])
def atualizar_turma(turma_id):

    nome = request.form["nome"]
    ano = request.form["ano"]
    turno = request.form["turno"]

    banco = conectar_banco()
    cursor = banco.cursor()

    cursor.execute("""
        UPDATE turmas
        SET nome = ?, ano = ?, turno = ?
        WHERE id = ?
    """, (
        nome,
        ano,
        turno,
        turma_id
    ))

    banco.commit()
    banco.close()

    return redirect("/turmas")


# ==========================
# ALUNOS DA TURMA
# ==========================

@app.route("/turma/<int:turma_id>/alunos")
def alunos_turma(turma_id):

    banco = conectar_banco()
    cursor = banco.cursor()

    cursor.execute(
        "SELECT * FROM turmas WHERE id = ?",
        (turma_id,)
    )

    turma = cursor.fetchone()

    cursor.execute("""
        SELECT id, nome, matricula
        FROM alunos
        WHERE turma_id = ?
        ORDER BY nome
    """, (turma_id,))

    alunos = cursor.fetchall()

    banco.close()

    return render_template(
        "alunos_turma.html",
        turma=turma,
        alunos=alunos
    )

@app.route("/turma_professores/<int:turma_id>")
def turma_professores(turma_id):

    banco = conectar_banco()
    cursor = banco.cursor()

    # Dados da turma
    cursor.execute("""
        SELECT *
        FROM turmas
        WHERE id = ?
    """, (turma_id,))

    turma = cursor.fetchone()

    # Professores vinculados
    cursor.execute("""
        SELECT
            professores.nome,
            professores.email,
            professores.disciplina
        FROM professor_turmas
        JOIN professores
            ON professor_turmas.professor_id = professores.id
        WHERE professor_turmas.turma_id = ?
    """, (turma_id,))

    professores = cursor.fetchall()

    banco.close()

    return render_template(
        "turma_professores.html",
        turma=turma,
        professores=professores
    )

# =====================================
# SELEÇÃO MANUAL DE QUESTÕES
# =====================================

@app.route("/selecionar_questoes")
def selecionar_questoes():

    professor_id = request.args.get("professor_id", "")
    disciplina = request.args.get("disciplina", "")
    habilidade = request.args.get("habilidade", "")
    descritor = request.args.get("descritor", "")

    banco = conectar_banco()
    cursor = banco.cursor()

    sql = """
        SELECT *
        FROM questoes
        WHERE 1=1
    """

    parametros = []

    if disciplina:
        sql += " AND disciplina = ? "
        parametros.append(disciplina)

    if habilidade:
        sql += " AND habilidade LIKE ? "
        parametros.append(f"%{habilidade}%")

    if descritor:
        sql += " AND enunciado LIKE ? "
        parametros.append(f"%{descritor}%")

    sql += " ORDER BY disciplina, id "

    cursor.execute(sql, parametros)
    questoes = cursor.fetchall()

    cursor.execute("SELECT * FROM turmas")
    turmas = cursor.fetchall()

    cursor.execute("""
        SELECT
            professores.id,
            professores.nome,
            professores.email,
            professores.disciplina
        FROM professores
        ORDER BY professores.nome
    """)
    professores = cursor.fetchall()

    banco.close()

    return render_template(
        "selecionar_questoes.html",
        questoes=questoes,
        turmas=turmas,
        professores=professores,
        professor_id=professor_id,
        disciplina=disciplina,
        habilidade=habilidade,
        descritor=descritor
    )


@app.route("/gerar_prova_manual", methods=["POST"])
def gerar_prova_manual():

    nome = request.form["nome"]
    turma_id = request.form["turma_id"]
    professor_id = request.form["professor_id"]
    disciplina = request.form["disciplina"]

    questoes_selecionadas = request.form.getlist("questoes")

    banco = conectar_banco()
    cursor = banco.cursor()

    cursor.execute("""
        INSERT INTO provas
        (nome, turma_id, professor_id, disciplina, quantidade)
        VALUES (?, ?, ?, ?, ?)
    """, (
        nome,
        turma_id,
        professor_id,
        disciplina,
        len(questoes_selecionadas)
    ))

    prova_id = cursor.lastrowid

    for questao_id in questoes_selecionadas:
        cursor.execute("""
            INSERT INTO prova_questoes
            (prova_id, questao_id)
            VALUES (?, ?)
        """, (
            prova_id,
            questao_id
        ))

    banco.commit()
    banco.close()

    return redirect("/provas")

@app.route("/importar_cartoes/<int:prova_id>")
def importar_cartoes(prova_id):

    banco = conectar_banco()
    cursor = banco.cursor()

    cursor.execute("""
        SELECT
            provas.id,
            provas.nome,
            turmas.nome
        FROM provas
        JOIN turmas ON provas.turma_id = turmas.id
        WHERE provas.id = ?
    """, (prova_id,))

    prova = cursor.fetchone()

    banco.close()

    return render_template(
        "importar_cartoes.html",
        prova=prova
    )

def ler_respostas_cartao(caminho_imagem, quantidade):
    imagem = cv2.imread(caminho_imagem)

    if imagem is None:
        return {}

    cinza = cv2.cvtColor(imagem, cv2.COLOR_BGR2GRAY)
    blur = cv2.GaussianBlur(cinza, (5, 5), 0)

    _, thresh = cv2.threshold(
        blur, 0, 255,
        cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU
    )

    contornos, _ = cv2.findContours(
        thresh, cv2.RETR_EXTERNAL,
        cv2.CHAIN_APPROX_SIMPLE
    )

    marcadores = []

    for c in contornos:
        area = cv2.contourArea(c)
        x, y, w, h = cv2.boundingRect(c)

        if area > 250 and area < 2000:
            proporcao = w / float(h)

            if 0.75 <= proporcao <= 1.25:
                marcadores.append((x, y, w, h))

    if len(marcadores) < 4:
        return {}

    marcadores = sorted(marcadores, key=lambda m: m[2] * m[3], reverse=True)[:4]

    pontos = []

    for x, y, w, h in marcadores:
        pontos.append([x + w // 2, y + h // 2])

    pontos = np.array(pontos, dtype="float32")

    soma = pontos.sum(axis=1)
    diferenca = np.diff(pontos, axis=1)

    ordenados = np.zeros((4, 2), dtype="float32")
    ordenados[0] = pontos[np.argmin(soma)]
    ordenados[2] = pontos[np.argmax(soma)]
    ordenados[1] = pontos[np.argmin(diferenca)]
    ordenados[3] = pontos[np.argmax(diferenca)]

    largura = 900
    altura = 350

    destino = np.array([
        [0, 0],
        [largura - 1, 0],
        [largura - 1, altura - 1],
        [0, altura - 1]
    ], dtype="float32")

    matriz = cv2.getPerspectiveTransform(ordenados, destino)
    recorte = cv2.warpPerspective(imagem, matriz, (largura, altura))

    cinza = cv2.cvtColor(recorte, cv2.COLOR_BGR2GRAY)
    blur = cv2.GaussianBlur(cinza, (5, 5), 0)

    _, thresh = cv2.threshold(
        blur, 0, 255,
        cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU
    )

    contornos, _ = cv2.findContours(
        thresh, cv2.RETR_EXTERNAL,
        cv2.CHAIN_APPROX_SIMPLE
    )

    bolhas = []

    for c in contornos:
        x, y, w, h = cv2.boundingRect(c)
        area = cv2.contourArea(c)

        if 12 <= w <= 35 and 12 <= h <= 35 and area > 80:
            proporcao = w / float(h)

            if 0.75 <= proporcao <= 1.25:
                bolhas.append((x, y, w, h))

    respostas = {}
    alternativas = ["A", "B", "C", "D"]

    colunas = 1

    if quantidade > 40:
        colunas = 3
    elif quantidade > 20:
        colunas = 2

    largura_coluna = largura // colunas

    for coluna in range(colunas):

        inicio_x = coluna * largura_coluna
        fim_x = inicio_x + largura_coluna

        bolhas_coluna = [
            b for b in bolhas
            if inicio_x < b[0] < fim_x
        ]

        bolhas_coluna = sorted(bolhas_coluna, key=lambda b: b[1])

        linhas = []

        for bolha in bolhas_coluna:
            x, y, w, h = bolha
            adicionou = False

            for linha in linhas:
                if abs(linha[0][1] - y) < 15:
                    linha.append(bolha)
                    adicionou = True
                    break

            if not adicionou:
                linhas.append([bolha])

        linhas = sorted(linhas, key=lambda l: l[0][1])

        for indice_linha, linha in enumerate(linhas):

            linha = sorted(linha, key=lambda b: b[0])

            if len(linha) < 4:
                continue

            linha = linha[:4]

            questao = coluna * 20 + indice_linha + 1

            if questao > quantidade:
                continue

            preenchimentos = []

            for x, y, w, h in linha:
                roi = thresh[y:y+h, x:x+w]
                total = cv2.countNonZero(roi)
                preenchimentos.append(total)

            maior = max(preenchimentos)
            indice = preenchimentos.index(maior)

            if maior > 180:
                respostas[questao] = alternativas[indice]

    return respostas


def _taxa_escura_melhor_circulo(cinza, centro_x, centro_y, raio=14, busca=10):
    """Mede uma bolha tolerando pequenos deslocamentos de scan/foto."""
    altura, largura = cinza.shape[:2]
    melhor = 0.0
    # limiar 165 é mais tolerante a impressão/scanner do que 120.
    for dy in range(-busca, busca + 1, 2):
        for dx in range(-busca, busca + 1, 2):
            x0 = int(centro_x + dx)
            y0 = int(centro_y + dy)
            x1, x2 = max(0, x0 - raio), min(largura, x0 + raio + 1)
            y1, y2 = max(0, y0 - raio), min(altura, y0 + raio + 1)
            roi = cinza[y1:y2, x1:x2]
            if roi.size == 0:
                continue
            yy, xx = np.ogrid[:roi.shape[0], :roi.shape[1]]
            cx, cy = x0 - x1, y0 - y1
            mascara = (xx - cx) ** 2 + (yy - cy) ** 2 <= raio ** 2
            pixels = roi[mascara]
            if pixels.size:
                taxa = float(np.mean(pixels < 165))
                melhor = max(melhor, taxa)
    return melhor


def ler_modelo_cartao(caminho_imagem, quantidade_modelos=4):
    """Lê o modelo com tolerância a PDF escaneado, foto, escala e deslocamento."""
    imagem = cv2.imread(caminho_imagem)
    if imagem is None:
        return None

    quantidade_modelos = max(1, min(4, int(quantidade_modelos or 1)))
    largura_base, altura_base = 1449, 2048
    normalizada = cv2.resize(imagem, (largura_base, altura_base))
    cinza = cv2.cvtColor(normalizada, cv2.COLOR_BGR2GRAY)
    cinza = cv2.createCLAHE(clipLimit=1.6, tileGridSize=(8, 8)).apply(cinza)

    centros_x = [1134, 1204, 1274, 1344]
    centro_y = 515
    preenchimentos = [
        _taxa_escura_melhor_circulo(cinza, x, centro_y, raio=15, busca=12)
        for x in centros_x[:quantidade_modelos]
    ]
    if not preenchimentos:
        return None

    ordem = np.argsort(preenchimentos)[::-1]
    melhor = preenchimentos[int(ordem[0])]
    segundo = preenchimentos[int(ordem[1])] if len(ordem) > 1 else 0.0

    # Em scanner, preenchimento sólido costuma ficar entre 0.25 e 0.90.
    if melhor < 0.24 or (len(preenchimentos) > 1 and (melhor - segundo) < 0.10):
        return None
    return int(ordem[0]) + 1


def ler_respostas_cartao_detalhado(caminho_imagem, quantidade):
    """Lê A/B/C/D também em cartões impressos e posteriormente escaneados.

    Em vez de exigir coordenada perfeita e preto absoluto, procura o centro mais
    escuro perto da posição esperada e compara cada alternativa com as demais.
    """
    imagem = cv2.imread(caminho_imagem)
    if imagem is None:
        raise ValueError("Não foi possível abrir a imagem enviada.")

    quantidade = max(0, int(quantidade or 0))
    if quantidade == 0:
        return {}

    largura_base, altura_base = 1449, 2048
    normalizada = cv2.resize(imagem, (largura_base, altura_base))
    cinza = cv2.cvtColor(normalizada, cv2.COLOR_BGR2GRAY)
    cinza = cv2.createCLAHE(clipLimit=1.6, tileGridSize=(8, 8)).apply(cinza)

    centros_x = [636, 739, 843, 946]
    primeiro_y = 1032
    passo_y = 49
    letras = ["A", "B", "C", "D"]
    resultado = {}

    for indice in range(quantidade):
        centro_y = primeiro_y + (indice * passo_y)
        preenchimentos = [
            _taxa_escura_melhor_circulo(cinza, x, centro_y, raio=14, busca=11)
            for x in centros_x
        ]
        ordem = np.argsort(preenchimentos)[::-1]
        maior = preenchimentos[int(ordem[0])]
        segundo = preenchimentos[int(ordem[1])]

        # Limiares deliberadamente conservadores: dúvida vai para revisão,
        # nunca é descartada nem transformada silenciosamente em resposta.
        marcadas_indices = [i for i, v in enumerate(preenchimentos) if v >= 0.28]
        if maior < 0.20:
            situacao, resposta, marcadas = "em_branco", "", []
        elif len(marcadas_indices) >= 2 and segundo >= 0.25:
            situacao, resposta = "dupla_marcacao", ""
            marcadas = [letras[i] for i in marcadas_indices]
        elif (maior - segundo) < 0.09:
            situacao, resposta = "dupla_marcacao", ""
            marcadas = [letras[int(ordem[0])], letras[int(ordem[1])]]
        else:
            situacao = "respondida"
            resposta = letras[int(ordem[0])]
            marcadas = [resposta]

        resultado[indice + 1] = {
            "resposta": resposta,
            "situacao": situacao,
            "marcadas": marcadas,
            "preenchimentos": [round(v, 4) for v in preenchimentos],
        }
    return resultado


@app.route("/corrigir_cartoes/<int:prova_id>", methods=["POST"])
def corrigir_cartoes(prova_id):

    if not cargo_permitido([
        "Administrador Geral",
        "Administrador da Instituição",
        "Coordenador",
        "Professor"
    ]):
        return redirect("/login")

    if "arquivo" not in request.files:
        return "Nenhum arquivo enviado."

    arquivo = request.files["arquivo"]

    if arquivo.filename == "":
        return "Nenhum arquivo enviado."

    os.makedirs("uploads_cartoes", exist_ok=True)

    nome_arquivo = secure_filename(arquivo.filename)

    caminho = os.path.join(
        "uploads_cartoes",
        nome_arquivo
    )

    arquivo.save(caminho)

    try:
        imagem_pil = Image.open(caminho).convert("RGB")
        imagem = np.array(imagem_pil)

    except Exception as erro:
        print("Erro ao cadastrar usuário:", erro)
        traceback.print_exc()

    qr_texto = _decodificar_qr_cartao(caminho)

    if not qr_texto:
        return """
        <h2>QR Code não encontrado</h2>
        <p>Envie uma foto mais nítida, bem iluminada e sem cortes.</p>
        <a href='/provas'>Voltar para provas</a>
        """

    try:
        partes = qr_texto.split("|")
        prova_qr = int(partes[0].replace("PROVA:", ""))
        aluno_id = int(partes[1].replace("ALUNO:", ""))

    except Exception:
        return f"""
        <h2>Erro ao ler dados do QR Code</h2>
        <p>QR Code lido: {qr_texto}</p>
        <a href='/provas'>Voltar</a>
        """

    if prova_qr != prova_id:
        return """
        <h2>Cartão não pertence a esta prova.</h2>
        <p>Confira se você enviou o cartão-resposta correto.</p>
        <a href='/provas'>Voltar</a>
        """

    banco = conectar_banco()
    cursor = banco.cursor()

    cursor.execute("""
        SELECT COUNT(*)
        FROM prova_questoes
        WHERE prova_id = ?
    """, (prova_qr,))

    total_questoes = cursor.fetchone()[0]

    respostas_lidas = ler_respostas_cartao(
        caminho,
        total_questoes
    )

    cursor.execute("""
        SELECT questoes.correta
        FROM prova_questoes
        JOIN questoes
            ON prova_questoes.questao_id = questoes.id
        WHERE prova_questoes.prova_id = ?
        ORDER BY COALESCE(NULLIF(prova_questoes.ordem, 0), prova_questoes.id), prova_questoes.id
    """, (prova_qr,))

    gabarito = cursor.fetchall()

    cursor.execute("""
        DELETE FROM respostas_alunos
        WHERE prova_id = ? AND aluno_id = ?
    """, (prova_qr, aluno_id))

    acertos = 0

    for i, g in enumerate(gabarito, start=1):

        correta = g[0]
        resposta_aluno = respostas_lidas.get(i)

        acertou = 0

        if resposta_aluno == correta:
            acertos += 1
            acertou = 1

        cursor.execute("""
            INSERT INTO respostas_alunos
            (
                prova_id,
                aluno_id,
                numero_questao,
                resposta_aluno,
                resposta_correta,
                acertou
            )
            VALUES (?, ?, ?, ?, ?, ?)
        """, (
            prova_qr,
            aluno_id,
            i,
            resposta_aluno,
            correta,
            acertou
        ))

    erros = total_questoes - acertos

    if total_questoes > 0:
        nota = round((acertos / total_questoes) * 10, 1)
    else:
        nota = 0

    cursor.execute("""
        DELETE FROM resultados
        WHERE prova_id = ? AND aluno_id = ?
    """, (prova_qr, aluno_id))

    cursor.execute("""
        INSERT INTO resultados
        (
            prova_id,
            aluno_id,
            acertos,
            erros,
            nota
        )
        VALUES (?, ?, ?, ?, ?)
    """, (
        prova_qr,
        aluno_id,
        acertos,
        erros,
        nota
    ))

    banco.commit()
    banco.close()

    return f"""
    <h2>Cartão corrigido com sucesso!</h2>

    <p><strong>Prova:</strong> {prova_qr}</p>
    <p><strong>Aluno ID:</strong> {aluno_id}</p>
    <p><strong>Arquivo:</strong> {nome_arquivo}</p>
    <p><strong>Total de questões:</strong> {total_questoes}</p>

    <p><strong>Acertos:</strong> {acertos}</p>
    <p><strong>Erros:</strong> {erros}</p>
    <p><strong>Nota:</strong> {nota}</p>

    <br>

    <a href='/resultados/{prova_qr}'>Ver resultados</a>
    """

@app.route("/resultados/<int:prova_id>")
def resultados(prova_id):
    """Centro completo de resultados de uma avaliação."""
    if not cargo_permitido([
        "Administrador Geral",
        "Administrador da Instituição",
        "Coordenador",
        "Professor"
    ]):
        return redirect("/login")

    _garantir_tabelas_aplicacoes()
    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    def colunas(tabela):
        cursor.execute(f"PRAGMA table_info({tabela})")
        return {linha["name"] for linha in cursor.fetchall()}

    try:
        cols_q = colunas("questoes")
        cols_pq = colunas("prova_questoes")
        cols_aa = colunas("aplicacao_alunos")

        cursor.execute("""
            SELECT
                p.*,
                t.nome AS turma_nome,
                t.ano AS turma_ano,
                t.turno AS turma_turno,
                e.nome_instituicao,
                e.logo,
                e.cidade,
                e.estado,
                e.diretor,
                e.coordenador1,
                e.coordenador2,
                e.coordenador3,
                COALESCE(al.ano, p.ano_letivo_id) AS ano_letivo,
                COALESCE(prof.nome, u.nome, '—') AS professor_nome
            FROM provas p
            LEFT JOIN turmas t ON t.id = p.turma_id
            LEFT JOIN escolas e ON e.id = p.escola_id
            LEFT JOIN anos_letivos al ON al.id = p.ano_letivo_id
            LEFT JOIN usuarios u ON u.id = p.professor_id
            LEFT JOIN professores prof ON prof.id = p.professor_id
            WHERE p.id = ?
            LIMIT 1
        """, (prova_id,))
        prova = cursor.fetchone()

        if not prova:
            flash("Avaliação não encontrada.", "erro")
            return redirect("/provas")

        if not _pode_gerenciar_prova(cursor, prova_id, exigir_edicao=False, permitir_finalizada=True):
            return redirect("/acesso_negado")

        _recalcular_notas_aplicacoes_por_peso(cursor, prova_id=prova_id)
        banco.commit()

        nota_final_expr = "aa.nota_final" if "nota_final" in cols_aa else "NULL"
        nota_disc_expr = "aa.nota_discursiva" if "nota_discursiva" in cols_aa else "NULL"

        cursor.execute(f"""
            WITH atuais AS (
                SELECT
                    aa.aluno_id,
                    aa.aplicacao_id,
                    aa.status,
                    COALESCE(aa.acertos_objetivos, 0) AS acertos,
                    MAX(COALESCE(aa.total_objetivas, 0) - COALESCE(aa.acertos_objetivos, 0), 0) AS erros,
                    COALESCE(aa.total_objetivas, 0) AS total_objetivas,
                    aa.nota_objetiva,
                    {nota_disc_expr} AS nota_discursiva,
                    COALESCE({nota_final_expr},
                        CASE WHEN aa.objetiva_corrigida = 1 AND COALESCE(aa.discursiva_pendente, 0) = 0
                             THEN aa.nota_objetiva ELSE NULL END) AS nota,
                    ap.data_aplicacao,
                    1 AS origem_nova
                FROM aplicacoes ap
                JOIN aplicacao_alunos aa ON aa.aplicacao_id = ap.id
                WHERE ap.prova_id = ?
            ), legados AS (
                SELECT r.aluno_id, NULL aplicacao_id, 'Corrigido' status,
                       COALESCE(r.acertos,0) acertos, COALESCE(r.erros,0) erros,
                       COALESCE(r.acertos,0)+COALESCE(r.erros,0) total_objetivas,
                       r.nota nota_objetiva, NULL nota_discursiva, r.nota nota,
                       NULL data_aplicacao, 0 origem_nova
                FROM resultados r
                WHERE r.prova_id = ?
                  AND NOT EXISTS (
                    SELECT 1 FROM atuais a WHERE a.aluno_id = r.aluno_id
                  )
            )
            SELECT a.id AS aluno_id, a.nome, a.matricula,
                   c.aplicacao_id, c.status, c.acertos, c.erros,
                   c.total_objetivas, c.nota_objetiva, c.nota_discursiva,
                   ROUND(c.nota,2) nota, c.data_aplicacao, c.origem_nova
            FROM (SELECT * FROM atuais UNION ALL SELECT * FROM legados) c
            JOIN alunos a ON a.id = c.aluno_id
            ORDER BY CASE WHEN c.nota IS NULL THEN 1 ELSE 0 END, c.nota DESC, a.nome COLLATE NOCASE
        """, (prova_id, prova_id))
        todos_registros = [dict(r) for r in cursor.fetchall()]

        ausentes = [r for r in todos_registros if (r.get("status") or "").strip().lower() in {"ausente", "faltou"}]
        resultados_lista = [r for r in todos_registros if r.get("nota") is not None and r not in ausentes]
        pendentes = [r for r in todos_registros if r.get("nota") is None and r not in ausentes]

        notas = [float(r["nota"]) for r in resultados_lista]
        total_alunos = len(resultados_lista)
        total_turma = len(todos_registros)
        presentes = total_turma - len(ausentes)
        media_turma = round(sum(notas) / len(notas), 2) if notas else 0
        maior_nota = round(max(notas), 2) if notas else 0
        menor_nota = round(min(notas), 2) if notas else 0

        notas_ordenadas = sorted(notas)
        if notas_ordenadas:
            meio = len(notas_ordenadas) // 2
            mediana = (notas_ordenadas[meio] if len(notas_ordenadas) % 2 else
                       (notas_ordenadas[meio - 1] + notas_ordenadas[meio]) / 2)
            variancia = sum((n - media_turma) ** 2 for n in notas) / len(notas)
            desvio_padrao = variancia ** 0.5
        else:
            mediana = desvio_padrao = 0

        media_ativa = bool(prova["media_ativa"]) if "media_ativa" in prova.keys() else False
        media_aprovacao = float(prova["media_aprovacao"]) if media_ativa and prova["media_aprovacao"] is not None else None
        aprovados = sum(1 for n in notas if media_aprovacao is not None and n >= media_aprovacao)
        reprovados = sum(1 for n in notas if media_aprovacao is not None and n < media_aprovacao)
        taxa_aprovacao = round(aprovados / len(notas) * 100, 1) if notas and media_aprovacao is not None else 0

        total_acertos = sum(int(r.get("acertos") or 0) for r in resultados_lista)
        total_itens = sum(int(r.get("total_objetivas") or 0) for r in resultados_lista)
        percentual_medio = round(total_acertos / total_itens * 100, 1) if total_itens else 0

        # Questões e indicadores pedagógicos
        habilidade_expr = "q.habilidade" if "habilidade" in cols_q else "NULL"
        descritor_expr = "q.descritor" if "descritor" in cols_q else "NULL"
        peso_expr = "pq.peso" if "peso" in cols_pq else ("pq.valor" if "valor" in cols_pq else "1.0")
        anulada_expr = "COALESCE(pq.anulada,0)" if "anulada" in cols_pq else "0"
        ordem_expr = "COALESCE(NULLIF(pq.ordem,0), pq.id)" if "ordem" in cols_pq else "pq.id"

        cursor.execute(f"""
            SELECT pq.questao_id, {ordem_expr} ordem,
                   q.tipo_questao, q.enunciado, {habilidade_expr} habilidade,
                   {descritor_expr} descritor, {peso_expr} peso, {anulada_expr} anulada
            FROM prova_questoes pq
            JOIN questoes q ON q.id = pq.questao_id
            WHERE pq.prova_id = ?
            ORDER BY {ordem_expr}, pq.id
        """, (prova_id,))
        questoes = [dict(q) for q in cursor.fetchall()]

        cursor.execute("""
            SELECT aro.questao_id,
                   COUNT(*) respondentes,
                   SUM(CASE WHEN aro.acertou = 1 THEN 1 ELSE 0 END) acertos,
                   SUM(CASE WHEN aro.situacao = 'em_branco' OR TRIM(COALESCE(aro.resposta,'')) = '' THEN 1 ELSE 0 END) em_branco
            FROM aplicacao_respostas_objetivas aro
            JOIN aplicacoes ap ON ap.id = aro.aplicacao_id
            WHERE ap.prova_id = ?
            GROUP BY aro.questao_id
        """, (prova_id,))
        obj_stats = {r["questao_id"]: dict(r) for r in cursor.fetchall()}

        relatorio_questoes = []
        for numero, q in enumerate(questoes, 1):
            st = obj_stats.get(q["questao_id"], {})
            respondentes = int(st.get("respondentes") or 0)
            acertos_q = int(st.get("acertos") or 0)
            brancos = int(st.get("em_branco") or 0)
            erros_q = max(respondentes - acertos_q - brancos, 0)
            percentual = round(acertos_q / respondentes * 100, 1) if respondentes else 0
            q.update(numero=numero, respondentes=respondentes, acertos=acertos_q,
                     erros=erros_q, em_branco=brancos, percentual=percentual)
            relatorio_questoes.append(q)

        validas = [q for q in relatorio_questoes if q["respondentes"] > 0 and not q["anulada"]]
        questao_mais_facil = max(validas, key=lambda x: x["percentual"]) if validas else None
        questao_mais_dificil = min(validas, key=lambda x: x["percentual"]) if validas else None
        anuladas = [q for q in relatorio_questoes if q["anulada"]]

        def agrupar_indicador(campo):
            grupos = {}
            for q in relatorio_questoes:
                valor = (q.get(campo) or "").strip()
                if not valor:
                    continue
                item = grupos.setdefault(valor, {"codigo": valor, "questoes": [], "acertos": 0, "respondentes": 0})
                item["questoes"].append(q["numero"])
                item["acertos"] += q["acertos"]
                item["respondentes"] += q["respondentes"]
            saida = []
            for item in grupos.values():
                item["percentual"] = round(item["acertos"] / item["respondentes"] * 100, 1) if item["respondentes"] else 0
                item["nivel"] = "Avançado" if item["percentual"] >= 80 else "Adequado" if item["percentual"] >= 60 else "Básico" if item["percentual"] >= 40 else "Abaixo do básico"
                saida.append(item)
            return sorted(saida, key=lambda x: x["percentual"], reverse=True)

        habilidades = agrupar_indicador("habilidade")
        descritores = agrupar_indicador("descritor")

        # Acrescenta a descrição oficial da habilidade BNCC ao relatório.
        # A questão continua guardando o código, mas a impressão exibe código + descrição.
        codigos_habilidades = sorted({
            (q.get("habilidade") or "").strip()
            for q in relatorio_questoes
            if (q.get("habilidade") or "").strip()
        })
        descricoes_habilidades = {}
        if codigos_habilidades:
            try:
                placeholders = ",".join("?" for _ in codigos_habilidades)
                cursor.execute(
                    f"SELECT codigo, descricao FROM bncc_habilidades WHERE codigo IN ({placeholders})",
                    codigos_habilidades
                )
                descricoes_habilidades = {
                    (r["codigo"] or "").strip(): (r["descricao"] or "").strip()
                    for r in cursor.fetchall()
                }
            except sqlite3.Error:
                descricoes_habilidades = {}

        # Fallback remoto: bancos antigos podem conter apenas o código na questão
        # sem a linha correspondente em bncc_habilidades. Nesse caso, consulta a
        # fonte BNCC já utilizada pela própria tela de cadastro para obter a
        # descrição, sem alterar os registros da avaliação.
        codigos_sem_descricao = [
            codigo for codigo in codigos_habilidades
            if not descricoes_habilidades.get(codigo)
        ]
        for codigo in codigos_sem_descricao:
            try:
                itens_remotos = _buscar_bncc_remota(codigo, limite=3)
                item_exato = next(
                    (item for item in itens_remotos
                     if str(item.get("codigo") or "").strip().upper() == codigo.upper()),
                    itens_remotos[0] if itens_remotos else None
                )
                if item_exato and (item_exato.get("descricao") or "").strip():
                    descricoes_habilidades[codigo] = (item_exato.get("descricao") or "").strip()
            except Exception as erro:
                print(f"AVISO DESCRICAO BNCC {codigo}: {erro}")

        for q in relatorio_questoes:
            codigo = (q.get("habilidade") or "").strip()
            q["habilidade_descricao"] = descricoes_habilidades.get(codigo, "")
        for h in habilidades:
            h["descricao"] = descricoes_habilidades.get((h.get("codigo") or "").strip(), "")

        # Resumo das discursivas
        cursor.execute("""
            SELECT rda.questao_id, COUNT(*) total,
                   SUM(CASE WHEN rda.corrigida = 1 THEN 1 ELSE 0 END) corrigidas,
                   AVG(CASE WHEN rda.corrigida = 1 THEN rda.nota END) media
            FROM respostas_discursivas_aplicacao rda
            JOIN aplicacoes ap ON ap.id = rda.aplicacao_id
            WHERE ap.prova_id = ?
            GROUP BY rda.questao_id
        """, (prova_id,))
        disc_stats = {r["questao_id"]: dict(r) for r in cursor.fetchall()}
        discursivas = []
        for q in relatorio_questoes:
            tipo = (q.get("tipo_questao") or "").lower()
            if tipo in {"discursiva", "dissertativa", "resposta_aberta", "resposta aberta"}:
                st = disc_stats.get(q["questao_id"], {})
                total = int(st.get("total") or 0)
                corrigidas_q = int(st.get("corrigidas") or 0)
                discursivas.append({**q, "total_respostas": total, "corrigidas": corrigidas_q,
                                    "pendentes": max(total-corrigidas_q,0), "media_discursiva": round(float(st.get("media") or 0),2)})

        # =========================================================
        # MAPA DE RESPOSTAS
        # =========================================================
        # A matriz reúne as respostas objetivas e discursivas por aluno.
        # As posições seguem a numeração exibida no cartão/prova de cada modelo.
        tipos_discursivos = {
            "discursiva",
            "dissertativa",
            "resposta_aberta",
            "resposta aberta"
        }

        total_objetivas_mapa = sum(
            1
            for questao in relatorio_questoes
            if (questao.get("tipo_questao") or "").strip().lower()
            not in tipos_discursivos
        )
        total_discursivas_mapa = sum(
            1
            for questao in relatorio_questoes
            if (questao.get("tipo_questao") or "").strip().lower()
            in tipos_discursivos
        )
        total_questoes_mapa = (
            total_objetivas_mapa + total_discursivas_mapa
        )

        mapa_colunas = []
        for numero in range(1, total_questoes_mapa + 1):
            mapa_colunas.append({
                "numero": numero,
                "tipo": (
                    "objetiva"
                    if numero <= total_objetivas_mapa
                    else "discursiva"
                )
            })

        respostas_objetivas_mapa = {}
        respostas_discursivas_mapa = {}

        cols_aro = colunas("aplicacao_respostas_objetivas")
        anulada_aro_expr = (
            "COALESCE(aro.anulada, 0)"
            if "anulada" in cols_aro
            else "0"
        )

        cursor.execute(f"""
            SELECT
                aro.aplicacao_id,
                aro.aluno_id,
                aro.numero_questao,
                aro.resposta,
                aro.situacao,
                aro.acertou,
                {anulada_aro_expr} AS anulada
            FROM aplicacao_respostas_objetivas aro
            INNER JOIN aplicacoes ap
                ON ap.id = aro.aplicacao_id
            WHERE ap.prova_id = ?
            ORDER BY
                aro.aplicacao_id,
                aro.aluno_id,
                aro.numero_questao
        """, (prova_id,))

        for registro in cursor.fetchall():
            chave = (
                registro["aplicacao_id"],
                registro["aluno_id"]
            )
            respostas_objetivas_mapa.setdefault(chave, {})[
                int(registro["numero_questao"])
            ] = dict(registro)

        # Compatibilidade com correções antigas salvas em respostas_alunos.
        cursor.execute("""
            SELECT name
            FROM sqlite_master
            WHERE type = 'table'
              AND name = 'respostas_alunos'
            LIMIT 1
        """)
        if cursor.fetchone():
            cols_ra = colunas("respostas_alunos")
            resposta_ra_expr = (
                "ra.resposta_aluno"
                if "resposta_aluno" in cols_ra
                else "ra.resposta"
                if "resposta" in cols_ra
                else "NULL"
            )
            correta_ra_expr = (
                "ra.resposta_correta"
                if "resposta_correta" in cols_ra
                else "NULL"
            )

            cursor.execute(f"""
                SELECT
                    ra.aluno_id,
                    ra.numero_questao,
                    {resposta_ra_expr} AS resposta,
                    {correta_ra_expr} AS resposta_correta,
                    COALESCE(ra.acertou, 0) AS acertou
                FROM respostas_alunos ra
                WHERE ra.prova_id = ?
                ORDER BY
                    ra.aluno_id,
                    ra.numero_questao
            """, (prova_id,))

            for registro in cursor.fetchall():
                chave = (None, registro["aluno_id"])
                resposta = (
                    registro["resposta"]
                    if "resposta" in registro.keys()
                    else None
                )
                respostas_objetivas_mapa.setdefault(chave, {})[
                    int(registro["numero_questao"])
                ] = {
                    "resposta": resposta,
                    "situacao": (
                        "respondida"
                        if resposta not in (None, "")
                        else "em_branco"
                    ),
                    "acertou": int(registro["acertou"] or 0),
                    "anulada": 0
                }

        cursor.execute("""
            SELECT
                rd.aplicacao_id,
                rd.aluno_id,
                rd.numero_exibicao,
                rd.nota,
                rd.corrigida,
                ROUND(COALESCE(pq.peso, 0), 2) AS peso,
                COALESCE(pq.anulada, 0) AS anulada
            FROM respostas_discursivas_aplicacao rd
            INNER JOIN aplicacoes ap
                ON ap.id = rd.aplicacao_id
            LEFT JOIN prova_questoes pq
                ON pq.prova_id = ap.prova_id
               AND pq.questao_id = rd.questao_id
            WHERE ap.prova_id = ?
            ORDER BY
                rd.aplicacao_id,
                rd.aluno_id,
                rd.numero_exibicao
        """, (prova_id,))

        for registro in cursor.fetchall():
            chave = (
                registro["aplicacao_id"],
                registro["aluno_id"]
            )
            respostas_discursivas_mapa.setdefault(chave, {})[
                int(registro["numero_exibicao"])
            ] = dict(registro)

        mapa_respostas = []

        for registro_aluno in todos_registros:
            chave = (
                registro_aluno.get("aplicacao_id"),
                registro_aluno["aluno_id"]
            )
            respostas_objetivas = respostas_objetivas_mapa.get(
                chave,
                {}
            )
            respostas_discursivas = respostas_discursivas_mapa.get(
                chave,
                {}
            )

            celulas = []
            total_acertos_mapa = 0
            aluno_ausente = (
                (registro_aluno.get("status") or "")
                .strip()
                .lower()
                in {"ausente", "faltou"}
            )

            for coluna in mapa_colunas:
                numero = coluna["numero"]

                if aluno_ausente:
                    celulas.append({
                        "numero": numero,
                        "tipo": coluna["tipo"],
                        "valor": "AUS",
                        "status": "ausente",
                        "titulo": "Aluno ausente"
                    })
                    continue

                if coluna["tipo"] == "objetiva":
                    resposta = respostas_objetivas.get(numero)

                    if not resposta:
                        celulas.append({
                            "numero": numero,
                            "tipo": "objetiva",
                            "valor": "—",
                            "status": "sem_resposta",
                            "titulo": "Resposta não registrada"
                        })
                        continue

                    valor = str(
                        resposta.get("resposta") or ""
                    ).strip().upper()
                    situacao = str(
                        resposta.get("situacao") or ""
                    ).strip().lower()
                    anulada = int(
                        resposta.get("anulada") or 0
                    ) == 1
                    acertou = int(
                        resposta.get("acertou") or 0
                    ) == 1

                    if anulada:
                        status = "anulada"
                        valor = "AN"
                        titulo = "Questão anulada"
                        total_acertos_mapa += 1
                    elif situacao == "dupla_marcacao":
                        status = "errada"
                        valor = "DM"
                        titulo = "Dupla marcação"
                    elif situacao == "em_branco" or not valor:
                        status = "em_branco"
                        valor = "—"
                        titulo = "Resposta em branco"
                    elif acertou:
                        status = "correta"
                        titulo = f"Resposta correta: {valor}"
                        total_acertos_mapa += 1
                    else:
                        status = "errada"
                        titulo = f"Resposta incorreta: {valor}"

                    celulas.append({
                        "numero": numero,
                        "tipo": "objetiva",
                        "valor": valor,
                        "status": status,
                        "titulo": titulo
                    })

                else:
                    resposta = respostas_discursivas.get(numero)

                    if not resposta:
                        celulas.append({
                            "numero": numero,
                            "tipo": "discursiva",
                            "valor": "—",
                            "status": "sem_resposta",
                            "titulo": "Resposta discursiva não registrada"
                        })
                        continue

                    anulada = int(
                        resposta.get("anulada") or 0
                    ) == 1
                    corrigida = int(
                        resposta.get("corrigida") or 0
                    ) == 1
                    nota_discursiva = float(
                        resposta.get("nota") or 0
                    )
                    peso_discursiva = float(
                        resposta.get("peso") or 0
                    )

                    if anulada:
                        status = "anulada"
                        valor = "AN"
                        titulo = "Questão anulada"
                        total_acertos_mapa += 1
                    elif not corrigida:
                        status = "pendente"
                        valor = "—"
                        titulo = "Correção discursiva pendente"
                    elif (
                        peso_discursiva > 0
                        and nota_discursiva >= peso_discursiva - 0.001
                    ):
                        status = "correta"
                        valor = "C"
                        titulo = (
                            "Resposta discursiva totalmente correta "
                            f"({nota_discursiva:.2f}/{peso_discursiva:.2f})"
                        )
                        total_acertos_mapa += 1
                    elif nota_discursiva > 0:
                        status = "parcial"
                        valor = "P"
                        titulo = (
                            "Resposta discursiva parcialmente correta "
                            f"({nota_discursiva:.2f}/{peso_discursiva:.2f})"
                        )
                    else:
                        status = "errada"
                        valor = "E"
                        titulo = (
                            "Resposta discursiva incorreta "
                            f"({nota_discursiva:.2f}/{peso_discursiva:.2f})"
                        )

                    celulas.append({
                        "numero": numero,
                        "tipo": "discursiva",
                        "valor": valor,
                        "status": status,
                        "titulo": titulo
                    })

            mapa_respostas.append({
                "aluno_id": registro_aluno["aluno_id"],
                "nome": registro_aluno["nome"],
                "matricula": registro_aluno.get("matricula"),
                "status": registro_aluno.get("status"),
                "ausente": aluno_ausente,
                "celulas": celulas,
                "total_acertos": (
                    0 if aluno_ausente else total_acertos_mapa
                )
            })

        mapa_respostas.sort(
            key=lambda item: (
                item["ausente"],
                item["nome"].casefold()
            )
        )

        # Ranking e situação
        for posicao, item in enumerate(resultados_lista, 1):
            item["posicao"] = posicao
            item["percentual"] = round((item["acertos"] / item["total_objetivas"] * 100), 1) if item["total_objetivas"] else 0
            item["situacao_final"] = ("Aprovado" if media_aprovacao is not None and item["nota"] >= media_aprovacao else
                                      "Abaixo da média" if media_aprovacao is not None else "Resultado disponível")

        return render_template(
            "resultados.html", prova=prova, resultados=resultados_lista,
            todos_registros=todos_registros, total_alunos=total_alunos,
            total_turma=total_turma, presentes=presentes, ausentes=ausentes,
            pendentes=pendentes, media_turma=media_turma, maior_nota=maior_nota,
            menor_nota=menor_nota, mediana=round(mediana,2), desvio_padrao=round(desvio_padrao,2),
            media_ativa=media_ativa, media_aprovacao=media_aprovacao,
            aprovados=aprovados, reprovados=reprovados, taxa_aprovacao=taxa_aprovacao,
            percentual_medio=percentual_medio, relatorio_questoes=relatorio_questoes,
            habilidades=habilidades, descritores=descritores, discursivas=discursivas,
            anuladas=anuladas, questao_mais_facil=questao_mais_facil,
            questao_mais_dificil=questao_mais_dificil,
            mapa_colunas=mapa_colunas,
            mapa_respostas=mapa_respostas,
            total_questoes_mapa=total_questoes_mapa,
            gerado_em=agora_local().strftime("%d/%m/%Y às %H:%M")
        )

    except sqlite3.Error as erro:
        import traceback
        traceback.print_exc()
        flash(f"Não foi possível carregar os resultados: {erro}", "erro")
        return redirect("/provas")
    finally:
        banco.close()


@app.route("/questao_relatorio/<int:prova_id>/<int:numero>")
def questao_relatorio(prova_id, numero):

    if not cargo_permitido([
        "Administrador Geral",
        "Administrador da Instituição",
        "Coordenador",
        "Professor"
    ]):
        return redirect("/login")

    banco = conectar_banco()
    cursor = banco.cursor()

    cursor.execute("""
        SELECT
            questoes.enunciado,
            questoes.imagem,
            questoes.alternativa_a,
            questoes.alternativa_b,
            questoes.alternativa_c,
            questoes.alternativa_d,
            questoes.correta
        FROM prova_questoes
        JOIN questoes
            ON prova_questoes.questao_id = questoes.id
        WHERE prova_questoes.prova_id = ?
        ORDER BY COALESCE(NULLIF(prova_questoes.ordem, 0), prova_questoes.id), prova_questoes.id
    """, (prova_id,))

    questoes = cursor.fetchall()

    if numero < 1 or numero > len(questoes):
        banco.close()
        return "Questão não encontrada."

    questao = questoes[numero - 1]

    cursor.execute("""
        SELECT COUNT(*), SUM(acertou)
        FROM respostas_alunos
        WHERE prova_id = ?
        AND numero_questao = ?
    """, (prova_id, numero))

    dados = cursor.fetchone()

    respondentes = dados[0] or 0
    acertos = dados[1] or 0
    erros = respondentes - acertos

    banco.close()

    return render_template(
        "questao_relatorio.html",
        numero=numero,
        questao=questao,
        respondentes=respondentes,
        acertos=acertos,
        erros=erros
    )

@app.route("/relatorio_questoes/<int:prova_id>")
def relatorio_questoes(prova_id):

    if not permissao_modulo("Relatorios_questoes"):
        return redirect("/acesso_negado")

    banco = conectar_banco()
    cursor = banco.cursor()

    cursor.execute("""
        SELECT nome
        FROM provas
        WHERE id = ?
    """, (prova_id,))

    prova = cursor.fetchone()

    if not prova:
        banco.close()
        return "Prova não encontrada."

    cursor.execute("""
        SELECT
            numero_questao,
            COUNT(*),
            SUM(acertou)
        FROM respostas_alunos
        WHERE prova_id = ?
        GROUP BY numero_questao
        ORDER BY numero_questao
    """, (prova_id,))

    dados = cursor.fetchall()

    relatorio = []

    for questao, total, acertos in dados:

        acertos = acertos or 0
        erros = total - acertos

        percentual = (
            round((acertos / total) * 100, 1)
            if total > 0
            else 0
        )

        relatorio.append((
            questao,
            total,
            acertos,
            erros,
            percentual
        ))

    banco.close()

    return render_template(
        "relatorio_questoes.html",
        prova=prova,
        prova_id=prova_id,
        relatorio=relatorio
    )

# =========================================================
# LISTAR INSTITUIÇÕES
# =========================================================

@app.route("/gestao/instituicoes")
def gestao_instituicoes():

    # Somente o Administrador Geral pode acessar.
    if not cargo_permitido([
        "Administrador Geral"
    ]):
        return redirect("/acesso_negado")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:

        cursor.execute("""
            SELECT
                escolas.*,

                anos_letivos.id AS ano_letivo_id_ativo,
                anos_letivos.ano AS ano_letivo_ativo,
                anos_letivos.data_inicio AS ano_data_inicio,
                anos_letivos.data_fim AS ano_data_fim,
                anos_letivos.encerrado AS ano_encerrado,

                (
                    SELECT COUNT(*)
                    FROM usuarios
                    WHERE usuarios.escola_id = escolas.id
                      AND usuarios.ativo = 1
                ) AS total_usuarios,

                (
                    SELECT COUNT(*)
                    FROM turmas
                    WHERE turmas.escola_id = escolas.id
                      AND turmas.ano_letivo_id = anos_letivos.id
                ) AS total_turmas_ano_ativo,

                (
                    SELECT COUNT(*)
                    FROM alunos
                    WHERE alunos.escola_id = escolas.id
                      AND alunos.ano_letivo_id = anos_letivos.id
                ) AS total_alunos_ano_ativo,

                (
                    SELECT COUNT(*)
                    FROM provas
                    WHERE provas.escola_id = escolas.id
                      AND provas.ano_letivo_id = anos_letivos.id
                ) AS total_provas_ano_ativo

            FROM escolas

            LEFT JOIN anos_letivos
                ON anos_letivos.escola_id = escolas.id
               AND anos_letivos.ativo = 1
               AND anos_letivos.encerrado = 0

            ORDER BY
                escolas.nome_instituicao COLLATE NOCASE ASC
        """)

        escolas = cursor.fetchall()

        return render_template(
            "gestao/gestao_instituicoes.html",
            escolas=escolas
        )

    except sqlite3.Error as erro:

        import traceback
        traceback.print_exc()

        print(
            "ERRO AO LISTAR INSTITUIÇÕES:",
            erro
        )

        flash(
            f"Erro ao carregar as instituições: {erro}",
            "erro"
        )

        return render_template(
            "gestao/gestao_instituicoes.html",
            escolas=[]
        )

    finally:
        banco.close()

@app.route("/gestao/instituicoes/nova", methods=["GET", "POST"])
def nova_instituicao():

    print("ROTA NOVA INSTITUIÇÃO ACESSADA")
    print("MÉTODO RECEBIDO:", request.method)

    if not cargo_permitido(["Administrador Geral"]):
        return redirect("/login")

    if request.method == "POST":

        print("FORMULÁRIO RECEBIDO:")
        print(request.form.to_dict(flat=False))

    if not cargo_permitido(["Administrador Geral"]):
        return redirect("/login")

    if request.method == "POST":

        # =====================================================
        # DADOS DA INSTITUIÇÃO
        # =====================================================

        nome_instituicao = request.form.get(
            "nome_instituicao",
            ""
        ).strip()

        codigo_inep = request.form.get(
            "codigo_inep",
            ""
        ).strip()

        cnpj = request.form.get(
            "cnpj",
            ""
        ).strip()

        cep = request.form.get(
            "cep",
            ""
        ).strip()

        endereco = request.form.get(
            "endereco",
            ""
        ).strip()

        cidade = request.form.get(
            "cidade",
            ""
        ).strip()

        estado = request.form.get(
            "estado",
            ""
        ).strip()

        telefone = request.form.get(
            "telefone",
            ""
        ).strip()

        whatsapp = request.form.get(
            "whatsapp",
            ""
        ).strip()

        email_institucional = request.form.get(
            "email",
            ""
        ).strip().lower()

        site = request.form.get(
            "site",
            ""
        ).strip()

        diretor = request.form.get(
            "diretor",
            ""
        ).strip()

        coordenador1 = request.form.get(
            "coordenador1",
            ""
        ).strip()

        coordenador2 = request.form.get(
            "coordenador2",
            ""
        ).strip()

        coordenador3 = request.form.get(
            "coordenador3",
            ""
        ).strip()

        secretario = request.form.get(
            "secretario",
            ""
        ).strip()

        # =====================================================
        # PLANO E COBRANÇA DA INSTITUIÇÃO
        # =====================================================
        plano_instituicao = (request.form.get("plano_codigo") or "start").strip().lower()
        if plano_instituicao not in {"start", "essencial"}:
            plano_instituicao = "start"
        assinatura_status = (request.form.get("assinatura_status") or "ativa").strip().lower()
        if assinatura_status not in {"ativa", "pendente", "trial"}:
            assinatura_status = "ativa"
        metodo_pagamento = (request.form.get("metodo_pagamento") or "manual").strip().lower()
        proxima_cobranca = (request.form.get("proxima_cobranca") or "").strip() or None
        renovacao_automatica = 1 if request.form.get("renovacao_automatica") == "1" else 0
        observacao_cobranca = (request.form.get("observacao_cobranca") or "").strip() or None
        isento_instituicao = 1 if request.form.get("isento_cobranca") == "1" else 0
        motivo_isencao = (request.form.get("motivo_isencao") or "").strip() or None
        isencao_fim = (request.form.get("isencao_fim") or "").strip() or None
        if isento_instituicao:
            assinatura_status = "isenta"

        if not nome_instituicao:
            flash(
                "Informe o nome da instituição.",
                "erro"
            )

            return render_template(
                "gestao/nova_instituicao.html"
            )

        # =====================================================
        # DADOS DO ADMINISTRADOR
        # =====================================================

        admin_nome = request.form.get(
            "admin_nome",
            ""
        ).strip()

        admin_email = request.form.get(
            "admin_email",
            ""
        ).strip().lower()

        admin_cpf = request.form.get(
            "admin_cpf",
            ""
        ).strip()

        admin_senha = request.form.get(
            "admin_senha",
            ""
        ).strip()

        admin_senha2 = request.form.get(
            "admin_senha2",
            ""
        ).strip()

        if not admin_nome:
            flash(
                "Informe o nome do administrador.",
                "erro"
            )

            return render_template(
                "gestao/nova_instituicao.html"
            )

        if not admin_email:
            flash(
                "Informe o e-mail do administrador.",
                "erro"
            )

            return render_template(
                "gestao/nova_instituicao.html"
            )

        if not admin_senha:
            flash(
                "Informe a senha do administrador.",
                "erro"
            )

            return render_template(
                "gestao/nova_instituicao.html"
            )

        if admin_senha != admin_senha2:
            flash(
                "As senhas do administrador não conferem.",
                "erro"
            )

            return render_template(
                "gestao/nova_instituicao.html"
            )

        if len(admin_senha) < 6:
            flash(
                "A senha do administrador deve possuir pelo menos 6 caracteres.",
                "erro"
            )

            return render_template(
                "gestao/nova_instituicao.html"
            )

        if not re.fullmatch(r"[^\s@]+@[^\s@]+\.[^\s@]+", admin_email):
            flash("Informe um e-mail válido para o administrador.", "erro")
            return render_template("gestao/nova_instituicao.html")

        if admin_cpf:
            cpf_digitos = somente_digitos(admin_cpf)
            if len(cpf_digitos) != 11:
                flash("O CPF do administrador deve possuir 11 dígitos.", "erro")
                return render_template("gestao/nova_instituicao.html")
            if not validar_cpf(cpf_digitos):
                flash("Informe um CPF válido para o administrador.", "erro")
                return render_template("gestao/nova_instituicao.html")
            admin_cpf = cpf_digitos

        # =====================================================
        # DADOS ACADÊMICOS
        # =====================================================

        tipo_instituicao = request.form.get(
            "tipo_instituicao",
            ""
        ).strip()

        ano_letivo = request.form.get(
            "ano_letivo",
            ""
        ).strip()

        modalidades = request.form.getlist(
            "modalidade_ensino"
        )

        etapas = request.form.getlist(
            "etapas_ensino"
        )

        componentes_recebidos = request.form.getlist(
            "componentes_curriculares"
        )

        if not tipo_instituicao:
            flash(
                "Selecione o tipo da instituição.",
                "erro"
            )

            return render_template(
                "gestao/nova_instituicao.html"
            )

        if not ano_letivo:
            flash(
                "Selecione o ano letivo.",
                "erro"
            )

            return render_template(
                "gestao/nova_instituicao.html"
            )

        try:
            ano_letivo_int = int(ano_letivo)
        except (TypeError, ValueError):
            flash("O ano letivo informado é inválido.", "erro")
            return render_template("gestao/nova_instituicao.html")

        if ano_letivo_int < 2000 or ano_letivo_int > datetime.now().year + 5:
            flash("Selecione um ano letivo válido.", "erro")
            return render_template("gestao/nova_instituicao.html")
        ano_letivo = str(ano_letivo_int)

        if not etapas:
            flash(
                "Selecione pelo menos uma etapa de ensino.",
                "erro"
            )

            return render_template(
                "gestao/nova_instituicao.html"
            )

        if not componentes_recebidos:
            flash(
                "Selecione pelo menos um componente curricular.",
                "erro"
            )

            return render_template(
                "gestao/nova_instituicao.html"
            )

        modalidade_ensino = ", ".join(modalidades)
        etapas_ensino = ", ".join(etapas)

        # =====================================================
        # ORGANIZA OS COMPONENTES RECEBIDOS DO HTML
        # =====================================================

        componentes_processados = []
        componentes_repetidos = set()

        for componente_json in componentes_recebidos:

            try:
                componente = json.loads(componente_json)

                etapa = str(
                    componente.get("etapa", "")
                ).strip()

                nome = str(
                    componente.get("nome", "")
                ).strip()

                tipo = str(
                    componente.get("tipo", "padrao")
                ).strip().lower()

            except (
                json.JSONDecodeError,
                TypeError,
                AttributeError
            ):
                continue

            if not etapa or not nome:
                continue

            # Impede componente de uma etapa não selecionada
            if etapa not in etapas:
                continue

            # Aceita somente os dois tipos previstos
            if tipo not in ["padrao", "manual"]:
                tipo = "padrao"

            chave_componente = (
                etapa.lower(),
                nome.lower()
            )

            # Impede componentes repetidos
            if chave_componente in componentes_repetidos:
                continue

            componentes_repetidos.add(
                chave_componente
            )

            componentes_processados.append({
                "etapa": etapa,
                "nome": nome,
                "tipo": tipo
            })

        if not componentes_processados:
            flash(
                "Não foi possível identificar os componentes curriculares selecionados.",
                "erro"
            )

            return render_template(
                "gestao/nova_instituicao.html"
            )

        # =====================================================
        # LOGO
        # =====================================================

        logo = request.files.get("logo")
        nome_logo = ""

        if logo and logo.filename:

            nome_logo = secure_filename(
                logo.filename
            )

            caminho_logo = os.path.join(
                app.config["UPLOAD_FOLDER"],
                nome_logo
            )

            logo.save(caminho_logo)

        # =====================================================
        # BANCO DE DADOS
        # =====================================================

        banco = conectar_banco()
        banco.row_factory = sqlite3.Row
        cursor = banco.cursor()

        try:

            # -------------------------------------------------
            # Evita duplicidade de instituição por INEP/CNPJ
            # -------------------------------------------------
            if codigo_inep:
                cursor.execute("SELECT id FROM escolas WHERE TRIM(codigo_inep) = TRIM(?) LIMIT 1", (codigo_inep,))
                if cursor.fetchone():
                    flash("Já existe uma instituição cadastrada com esse código INEP.", "erro")
                    return render_template("gestao/nova_instituicao.html")

            if cnpj:
                cnpj_digitos = re.sub(r"\D", "", cnpj)
                if len(cnpj_digitos) != 14:
                    flash("O CNPJ deve possuir 14 dígitos.", "erro")
                    return render_template("gestao/nova_instituicao.html")
                cursor.execute("""
                    SELECT id FROM escolas
                    WHERE REPLACE(REPLACE(REPLACE(REPLACE(TRIM(cnpj), '.', ''), '/', ''), '-', ''), ' ', '') = ?
                    LIMIT 1
                """, (cnpj_digitos,))
                if cursor.fetchone():
                    flash("Já existe uma instituição cadastrada com esse CNPJ.", "erro")
                    return render_template("gestao/nova_instituicao.html")

            # -------------------------------------------------
            # Verifica se o e-mail do administrador já existe
            # -------------------------------------------------

            cursor.execute("""
                SELECT id
                FROM usuarios
                WHERE LOWER(email) = LOWER(?)
                LIMIT 1
            """, (
                admin_email,
            ))

            usuario_existente = cursor.fetchone()

            if usuario_existente:

                flash(
                    "Já existe um usuário cadastrado com esse e-mail.",
                    "erro"
                )

                banco.close()

                return render_template(
                    "gestao/nova_instituicao.html"
                )

            # -------------------------------------------------
            # Busca o cargo Administrador da Instituição
            # -------------------------------------------------

            cursor.execute("""
                SELECT id
                FROM cargos
                WHERE nome = ?
                LIMIT 1
            """, (
                "Administrador da Instituição",
            ))

            cargo = cursor.fetchone()

            if cargo is None:

                cursor.execute("""
                    INSERT INTO cargos (nome)
                    VALUES (?)
                """, (
                    "Administrador da Instituição",
                ))

                cargo_id = cursor.lastrowid

            else:
                cargo_id = cargo["id"]

            # -------------------------------------------------
            # Cadastra a instituição
            # -------------------------------------------------

            cursor.execute("""
                INSERT INTO escolas (
                    nome_instituicao,
                    codigo_inep,
                    cnpj,
                    cep,
                    endereco,
                    cidade,
                    estado,
                    telefone,
                    whatsapp,
                    email,
                    site,
                    diretor,
                    coordenador1,
                    coordenador2,
                    coordenador3,
                    secretario,
                    tipo_instituicao,
                    ano_letivo,
                    modalidade_ensino,
                    etapas_ensino,
                    logo,
                    status,
                    criado_em
                )
                VALUES (
                    ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?,
                    ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?
                )
            """, (
                nome_instituicao,
                codigo_inep,
                cnpj,
                cep,
                endereco,
                cidade,
                estado,
                telefone,
                whatsapp,
                email_institucional,
                site,
                diretor,
                coordenador1,
                coordenador2,
                coordenador3,
                secretario,
                tipo_instituicao,
                ano_letivo,
                modalidade_ensino,
                etapas_ensino,
                nome_logo,
                1,
                datetime.now().strftime(
                    "%Y-%m-%d %H:%M:%S"
                )
            ))

            escola_id = cursor.lastrowid

            # -------------------------------------------------
            # Cria e ativa o primeiro ano letivo oficial
            # -------------------------------------------------
            sincronizar_ano_letivo_instituicao(
                cursor,
                escola_id,
                ano_letivo,
                tornar_ativo=True
            )

            # -------------------------------------------------
            # Salva os componentes curriculares da instituição
            # -------------------------------------------------

            for componente in componentes_processados:

                cursor.execute("""
                    INSERT INTO componentes_curriculares (
                        escola_id,
                        etapa_ensino,
                        nome,
                        tipo,
                        ativo
                    )
                    VALUES (?, ?, ?, ?, 1)
                """, (
                    escola_id,
                    componente["etapa"],
                    componente["nome"],
                    componente["tipo"]
                ))

            # -------------------------------------------------
            # Cria o administrador da instituição
            # -------------------------------------------------

            cursor.execute("""
                INSERT INTO usuarios (
                    nome,
                    email,
                    senha,
                    cargo_id,
                    ativo,
                    escola_id,
                    cpf,
                    tipo_conta,
                    plano_codigo,
                    assinatura_status
                )
                VALUES (?, ?, ?, ?, 1, ?, ?, 'institucional', ?, ?)
            """, (
                admin_nome,
                admin_email,
                generate_password_hash(admin_senha),
                cargo_id,
                escola_id,
                admin_cpf,
                plano_instituicao,
                assinatura_status
            ))
            admin_usuario_id = cursor.lastrowid

            cursor.execute("UPDATE escolas SET plano_codigo=? WHERE id=?", (plano_instituicao, escola_id))
            cursor.execute("""
                INSERT INTO assinaturas_saas (
                    usuario_id, escola_id, plano_codigo, status, gateway,
                    proxima_cobranca, isento, motivo_isencao, isencao_inicio,
                    isencao_fim, concedida_por, metodo_pagamento,
                    renovacao_automatica, observacao_cobranca
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, DATE('now'), ?, ?, ?, ?, ?)
            """, (
                admin_usuario_id, escola_id, plano_instituicao, assinatura_status,
                'cortesia' if isento_instituicao else 'manual', proxima_cobranca,
                isento_instituicao, motivo_isencao, isencao_fim,
                session.get('usuario_id'), metodo_pagamento, renovacao_automatica,
                observacao_cobranca
            ))
            assinatura_id = cursor.lastrowid
            if not isento_instituicao:
                valor_plano = 199.0 if plano_instituicao == 'start' else 299.0
                cursor.execute("""INSERT INTO pagamentos_saas
                    (assinatura_id, gateway, forma_pagamento, valor, status, vencimento)
                    VALUES (?, 'manual', ?, ?, ?, ?)""",
                    (assinatura_id, metodo_pagamento, valor_plano,
                     'pendente' if assinatura_status == 'pendente' else 'pago', proxima_cobranca))

            banco.commit()

            flash(
                "Instituição, administrador e componentes curriculares cadastrados com sucesso!",
                "success"
            )

            return redirect(
                "/gestao/instituicoes"
            )

        except Exception as erro:

            banco.rollback()

            import traceback
            traceback.print_exc()

            print(
                "ERRO COMPLETO AO CADASTRAR INSTITUIÇÃO:",
                repr(erro)
            )

            mensagem_erro = str(erro)
            if "user-defined function raised exception" in mensagem_erro.lower():
                mensagem_erro = (
                    "Falha ao proteger os dados pessoais do administrador. "
                    "Verifique o CPF informado e tente novamente."
                )
            flash(
                f"Erro ao cadastrar a instituição: {mensagem_erro}",
                "erro"
            )

            return render_template(
                "gestao/nova_instituicao.html"
            )

        finally:
            banco.close()

    return render_template(
        "gestao/nova_instituicao.html"
    )

@app.route("/usuarios")
def usuarios():

    if not cargo_permitido([
        "Administrador Geral",
        "Administrador da Instituição"
    ]):
        return redirect("/login")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    usuario_cargo = session.get("usuario_cargo", "").strip()
    usuario_id = session.get("usuario_id")
    escola_id = session.get("escola_id")

    # Se for administrador da instituição e o escola_id
    # ainda não estiver na sessão, busca no banco.
    if (
        usuario_cargo == "Administrador da Instituição"
        and not escola_id
        and usuario_id
    ):
        cursor.execute("""
            SELECT escola_id
            FROM usuarios
            WHERE id = ?
            LIMIT 1
        """, (usuario_id,))

        usuario_logado = cursor.fetchone()

        if usuario_logado and usuario_logado["escola_id"]:
            escola_id = usuario_logado["escola_id"]
            session["escola_id"] = escola_id

    if usuario_cargo == "Administrador Geral":

        cursor.execute("""
            SELECT
                usuarios.id,
                usuarios.nome,
                usuarios.email,
                usuarios.cpf,
                usuarios.ativo,
                usuarios.escola_id,
                cargos.nome AS cargo,
                escolas.nome_instituicao
            FROM usuarios
            LEFT JOIN cargos
                ON usuarios.cargo_id = cargos.id
            LEFT JOIN escolas
                ON usuarios.escola_id = escolas.id
            ORDER BY usuarios.nome COLLATE NOCASE ASC
        """)

    elif usuario_cargo == "Administrador da Instituição":

        if not escola_id:
            banco.close()

            flash(
                "Seu usuário não está vinculado a uma instituição.",
                "erro"
            )

            return redirect("/")

        cursor.execute("""
            SELECT
                usuarios.id,
                usuarios.nome,
                usuarios.email,
                usuarios.cpf,
                usuarios.ativo,
                usuarios.escola_id,
                cargos.nome AS cargo,
                escolas.nome_instituicao
            FROM usuarios
            LEFT JOIN cargos
                ON usuarios.cargo_id = cargos.id
            LEFT JOIN escolas
                ON usuarios.escola_id = escolas.id
            WHERE usuarios.escola_id = ?
            ORDER BY usuarios.nome COLLATE NOCASE ASC
        """, (escola_id,))

    else:
        banco.close()
        return redirect("/")

    lista_usuarios = cursor.fetchall()

    banco.close()

    return render_template(
        "gestao/usuarios.html",
        usuarios=lista_usuarios
    )

@app.route("/gestao")
def gestao():

    if not cargo_permitido(["Administrador"]):
        return redirect("/login")

    banco = conectar_banco()
    cursor = banco.cursor()

    cursor.execute("""
        SELECT
            usuarios.id,
            usuarios.nome,
            usuarios.email,
            cargos.nome
        FROM usuarios
        JOIN cargos
        ON usuarios.cargo_id = cargos.id
        ORDER BY usuarios.nome
    """)

    usuarios = cursor.fetchall()

    banco.close()

    return render_template(
        "gestao.html",
        usuarios=usuarios
    )

@app.route("/cargos")
def cargos():

    if not cargo_permitido(["Administrador"]):
        return redirect("/login")

    banco = conectar_banco()
    cursor = banco.cursor()

    cursor.execute("""
        SELECT id, nome
        FROM cargos
        ORDER BY nome
    """)

    cargos = cursor.fetchall()

    banco.close()

    return render_template(
        "cargos.html",
        cargos=cargos
    )

import sqlite3
import traceback

from flask import (
    request,
    redirect,
    render_template,
    session,
    flash
)


def coluna_existe(cursor, tabela, coluna):
    """
    Verifica se determinada coluna existe em uma tabela.
    """
    cursor.execute(f"PRAGMA table_info({tabela})")
    colunas = cursor.fetchall()

    return any(
        registro[1] == coluna
        for registro in colunas
    )


def garantir_estrutura_usuarios(cursor):
    """
    Garante que as tabelas e colunas utilizadas no cadastro
    de usuários estejam disponíveis no banco.
    """

    # ==============================
    # COLUNA escola_id EM USUÁRIOS
    # ==============================

    if not coluna_existe(cursor, "usuarios", "escola_id"):
        cursor.execute("""
            ALTER TABLE usuarios
            ADD COLUMN escola_id INTEGER
        """)

    # ==============================
    # COLUNA cpf EM USUÁRIOS
    # ==============================

    if not coluna_existe(cursor, "usuarios", "cpf"):
        cursor.execute("""
            ALTER TABLE usuarios
            ADD COLUMN cpf TEXT
        """)

    # ==============================
    # DADOS PESSOAIS DO PERFIL
    # ==============================

    if not coluna_existe(cursor, "usuarios", "telefone"):
        cursor.execute("""
            ALTER TABLE usuarios
            ADD COLUMN telefone TEXT
        """)

    if not coluna_existe(cursor, "usuarios", "data_nascimento"):
        cursor.execute("""
            ALTER TABLE usuarios
            ADD COLUMN data_nascimento TEXT
        """)

    # ==============================
    # COLUNA escola_id EM TURMAS
    # ==============================

    if not coluna_existe(cursor, "turmas", "escola_id"):
        cursor.execute("""
            ALTER TABLE turmas
            ADD COLUMN escola_id INTEGER
        """)

    # ==============================
    # PERMISSÕES DOS USUÁRIOS
    # ==============================

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS usuario_permissoes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            usuario_id INTEGER NOT NULL,
            modulo TEXT NOT NULL,
            pode_acessar INTEGER DEFAULT 0,
            UNIQUE(usuario_id, modulo),
            FOREIGN KEY (usuario_id)
                REFERENCES usuarios(id)
                ON DELETE CASCADE
        )
    """)

    # ==============================
    # TURMAS DOS COORDENADORES
    # ==============================

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS coordenador_turmas (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            usuario_id INTEGER NOT NULL,
            turma_id INTEGER NOT NULL,
            UNIQUE(usuario_id, turma_id),
            FOREIGN KEY (usuario_id)
                REFERENCES usuarios(id)
                ON DELETE CASCADE,
            FOREIGN KEY (turma_id)
                REFERENCES turmas(id)
                ON DELETE CASCADE
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS componentes_curriculares (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            escola_id INTEGER NOT NULL,
            etapa_ensino TEXT NOT NULL,
            nome TEXT NOT NULL,
            tipo TEXT NOT NULL DEFAULT 'padrao',
            ativo INTEGER NOT NULL DEFAULT 1,

            FOREIGN KEY (escola_id)
                REFERENCES escolas(id)
                ON DELETE CASCADE,

            UNIQUE (
                escola_id,
                etapa_ensino,
                nome
            )
        )
    """)

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS professor_vinculos (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        professor_id INTEGER NOT NULL,
        turma_id INTEGER NOT NULL,
        componente_id INTEGER NOT NULL,

        FOREIGN KEY (professor_id)
            REFERENCES usuarios(id)
            ON DELETE CASCADE,

        FOREIGN KEY (turma_id)
            REFERENCES turmas(id)
            ON DELETE CASCADE,

        FOREIGN KEY (componente_id)
            REFERENCES componentes_curriculares(id)
            ON DELETE CASCADE,

        UNIQUE (
            professor_id,
            turma_id,
            componente_id
        )
    )
""")

@app.route("/cadastrar_usuario", methods=["GET", "POST"])
def cadastrar_usuario():

    # ==============================
    # CONTROLE DE ACESSO
    # ==============================

    if not cargo_permitido([
        "Administrador Geral",
        "Administrador da Instituição"
    ]):
        flash(
            "Você não possui permissão para acessar esta página.",
            "erro"
        )
        return redirect("/login")

    banco = None

    try:
        banco = conectar_banco()
        banco.row_factory = sqlite3.Row

        cursor = banco.cursor()

        # Ativa o suporte a chaves estrangeiras
        cursor.execute("PRAGMA foreign_keys = ON")

        # Corrige automaticamente a estrutura do banco
        garantir_estrutura_usuarios(cursor)
        banco.commit()

        cargo_logado = session.get(
            "usuario_cargo",
            ""
        ).strip()

        escola_logada_id = session.get("escola_id")

        modulos_plataforma = [
            "Dashboard",
            "Instituições",
            "Usuários",
            "Turmas",
            "Alunos",
            "Questões",
            "Provas",
            "Relatórios"
        ]

        # ==============================
        # CARGOS DISPONÍVEIS
        # ==============================

        if cargo_logado == "Administrador Geral":

            cursor.execute("""
                SELECT
                    id,
                    nome
                FROM cargos
                ORDER BY nome
            """)

        else:

            cursor.execute("""
                SELECT
                    id,
                    nome
                FROM cargos
                WHERE nome != 'Administrador Geral'
                ORDER BY nome
            """)

        cargos = cursor.fetchall()

        # ==============================
        # INSTITUIÇÕES DISPONÍVEIS
        # ==============================

        if cargo_logado == "Administrador Geral":

            cursor.execute("""
                SELECT
                    id,
                    nome_instituicao
                FROM escolas
                WHERE COALESCE(status, 1) = 1
                ORDER BY nome_instituicao
            """)

            escolas = cursor.fetchall()

        else:
            escolas = []

        # ==============================
        # TURMAS DISPONÍVEIS
        # ==============================

        if cargo_logado == "Administrador Geral":

            cursor.execute("""
                SELECT
                    id,
                    nome,
                    ano,
                    turno,
                    escola_id
                FROM turmas
                ORDER BY
                    escola_id,
                    ano,
                    nome,
                    turno
            """)

        else:

            cursor.execute("""
                SELECT
                    id,
                    nome,
                    ano,
                    turno,
                    escola_id
                FROM turmas
                WHERE escola_id = ?
                ORDER BY
                    ano,
                    nome,
                    turno
            """, (
                escola_logada_id,
            ))

        turmas = cursor.fetchall()

        # ==============================
        # NOME DA INSTITUIÇÃO LOGADA
        # ==============================

        nome_instituicao = ""

        if (
            cargo_logado == "Administrador da Instituição"
            and escola_logada_id
        ):

            cursor.execute("""
                SELECT nome_instituicao
                FROM escolas
                WHERE id = ?
                LIMIT 1
            """, (
                escola_logada_id,
            ))

            escola = cursor.fetchone()

            if escola:
                nome_instituicao = escola["nome_instituicao"]

        # ==============================
        # SALVAMENTO DO USUÁRIO
        # ==============================

        if request.method == "POST":

            nome = request.form.get(
                "nome",
                ""
            ).strip()

            email = request.form.get(
                "email",
                ""
            ).strip().lower()

            cpf = request.form.get(
                "cpf",
                ""
            ).strip()

            senha = request.form.get(
                "senha",
                ""
            ).strip()

            confirmar_senha = request.form.get(
                "confirmar_senha",
                ""
            ).strip()

            cargo_id = request.form.get(
                "cargo_id",
                ""
            ).strip()

            modulos_selecionados = request.form.getlist(
                "modulos_permitidos"
            )

            turmas_vinculadas = request.form.getlist(
                "turmas_vinculadas"
            )

            # ==============================
            # VALIDAÇÕES BÁSICAS
            # ==============================

            if not nome:
                flash(
                    "Informe o nome do usuário.",
                    "erro"
                )
                return redirect("/cadastrar_usuario")

            if not email:
                flash(
                    "Informe o e-mail do usuário.",
                    "erro"
                )
                return redirect("/cadastrar_usuario")

            if "@" not in email:
                flash(
                    "Informe um endereço de e-mail válido.",
                    "erro"
                )
                return redirect("/cadastrar_usuario")

            if not senha:
                flash(
                    "Informe uma senha.",
                    "erro"
                )
                return redirect("/cadastrar_usuario")

            if len(senha) < 6:
                flash(
                    "A senha deve possuir pelo menos 6 caracteres.",
                    "erro"
                )
                return redirect("/cadastrar_usuario")

            if senha != confirmar_senha:
                flash(
                    "As senhas não conferem.",
                    "erro"
                )
                return redirect("/cadastrar_usuario")

            if not cargo_id:
                flash(
                    "Selecione um cargo.",
                    "erro"
                )
                return redirect("/cadastrar_usuario")


                foto.stream.seek(0, os.SEEK_END)
                tamanho_foto = foto.stream.tell()
                foto.stream.seek(0)

                limite_foto = 5 * 1024 * 1024

                if tamanho_foto > limite_foto:
                    flash(
                        "A foto deve ter no máximo 5 MB.",
                        "erro"
                    )
                    return redirect("/cadastrar_usuario")

            # ==============================
            # VERIFICA O CARGO SELECIONADO
            # ==============================

            cursor.execute("""
                SELECT
                    id,
                    nome
                FROM cargos
                WHERE id = ?
                LIMIT 1
            """, (
                cargo_id,
            ))

            cargo_selecionado = cursor.fetchone()

            if cargo_selecionado is None:
                flash(
                    "O cargo selecionado é inválido.",
                    "erro"
                )
                return redirect("/cadastrar_usuario")

            nome_cargo = cargo_selecionado["nome"].strip()

            # Administrador da Instituição não pode
            # criar Administrador Geral
            if (
                cargo_logado != "Administrador Geral"
                and nome_cargo == "Administrador Geral"
            ):
                flash(
                    "Você não pode criar um Administrador Geral.",
                    "erro"
                )
                return redirect("/cadastrar_usuario")

            # ==============================
            # INSTITUIÇÃO DO NOVO USUÁRIO
            # ==============================

            if cargo_logado == "Administrador Geral":

                escola_id = request.form.get(
                    "escola_id",
                    ""
                ).strip()

                escola_id = escola_id or None

            else:
                escola_id = escola_logada_id

            # Administrador Geral não precisa ficar
            # vinculado a uma instituição
            if nome_cargo == "Administrador Geral":
                escola_id = None

            # Demais cargos precisam de instituição
            if (
                nome_cargo != "Administrador Geral"
                and not escola_id
            ):
                flash(
                    "Selecione a instituição do usuário.",
                    "erro"
                )
                return redirect("/cadastrar_usuario")

            # Confirma se a instituição existe
            if escola_id:

                cursor.execute("""
                    SELECT id
                    FROM escolas
                    WHERE id = ?
                      AND COALESCE(status, 1) = 1
                    LIMIT 1
                """, (
                    escola_id,
                ))

                if cursor.fetchone() is None:
                    flash(
                        "A instituição selecionada é inválida ou está inativa.",
                        "erro"
                    )
                    return redirect("/cadastrar_usuario")

            # ==============================
            # CONTROLE DE PERMISSÕES
            # ==============================

            if cargo_logado == "Administrador da Instituição":

                modulos_selecionados = [
                    modulo
                    for modulo in modulos_selecionados
                    if modulo != "Instituições"
                ]

            # ==============================
            # VERIFICA E-MAIL DUPLICADO
            # ==============================

            cursor.execute("""
                SELECT id
                FROM usuarios
                WHERE LOWER(email) = LOWER(?)
                LIMIT 1
            """, (
                email,
            ))

            if cursor.fetchone():
                flash(
                    "Já existe um usuário cadastrado com este e-mail.",
                    "erro"
                )
                return redirect("/cadastrar_usuario")

            # ==============================
            # VERIFICA CPF DUPLICADO
            # ==============================

            if cpf:

                cursor.execute("""
                    SELECT id
                    FROM usuarios
                    WHERE REPLACE(
                        REPLACE(
                            REPLACE(cpf, '.', ''),
                            '-',
                            ''
                        ),
                        ' ',
                        ''
                    ) = REPLACE(
                        REPLACE(
                            REPLACE(?, '.', ''),
                            '-',
                            ''
                        ),
                        ' ',
                        ''
                    )
                    LIMIT 1
                """, (
                    cpf,
                ))

                if cursor.fetchone():
                    flash(
                        "Já existe um usuário cadastrado com este CPF.",
                        "erro"
                    )
                    return redirect("/cadastrar_usuario")

            # ==============================
            # INSERE O USUÁRIO
            # ==============================

            cursor.execute("""
                INSERT INTO usuarios (
                    nome,
                    email,
                    cpf,
                    senha,
                    cargo_id,
                    escola_id,
                    ativo
                )
                VALUES (?, ?, ?, ?, ?, ?, 1)
            """, (
                nome,
                email,
                cpf or None,
                senha,
                cargo_id,
                escola_id
            ))

            usuario_id = cursor.lastrowid

            # ==============================
            # PERMISSÕES INDIVIDUAIS
            # ==============================

            for modulo in modulos_plataforma:

                pode_acessar = (
                    1
                    if modulo in modulos_selecionados
                    else 0
                )

                # Administrador Geral recebe acesso completo
                if nome_cargo == "Administrador Geral":
                    pode_acessar = 1

                # Usuários de instituição não acessam
                # o cadastro geral de instituições
                if (
                    cargo_logado == "Administrador da Instituição"
                    and modulo == "Instituições"
                ):
                    pode_acessar = 0

                cursor.execute("""
                    INSERT OR REPLACE INTO usuario_permissoes (
                        usuario_id,
                        modulo,
                        pode_acessar
                    )
                    VALUES (?, ?, ?)
                """, (
                    usuario_id,
                    modulo,
                    pode_acessar
                ))

            # ==============================
            # TURMAS DA COORDENAÇÃO
            # ==============================

            cargo_coordenacao = (
                "coordenador" in nome_cargo.lower()
                or "coordenação" in nome_cargo.lower()
                or "coordenacao" in nome_cargo.lower()
            )

            if cargo_coordenacao and escola_id:

                for turma_id in turmas_vinculadas:

                    cursor.execute("""
                        SELECT id
                        FROM turmas
                        WHERE id = ?
                          AND escola_id = ?
                        LIMIT 1
                    """, (
                        turma_id,
                        escola_id
                    ))

                    turma_valida = cursor.fetchone()

                    if turma_valida:

                        cursor.execute("""
                            INSERT OR IGNORE INTO coordenador_turmas (
                                usuario_id,
                                turma_id
                            )
                            VALUES (?, ?)
                        """, (
                            usuario_id,
                            turma_id
                        ))

            banco.commit()

            flash(
                "Usuário cadastrado com sucesso.",
                "success"
            )

            return redirect("/usuarios")

        # ==============================
        # ABERTURA DO FORMULÁRIO
        # ==============================

        return render_template(
            "gestao/cadastrar_usuario.html",
            cargos=cargos,
            escolas=escolas,
            turmas=turmas,
            nome_instituicao=nome_instituicao,
            modulos_plataforma=modulos_plataforma
        )

    except sqlite3.IntegrityError as erro:

        if banco:
            banco.rollback()

        traceback.print_exc()

        print(
            f"Erro de integridade ao cadastrar usuário: {erro}"
        )

        flash(
            "Não foi possível cadastrar o usuário. "
            "Verifique se o e-mail ou CPF já está cadastrado.",
            "erro"
        )

        return redirect("/cadastrar_usuario")

    except sqlite3.OperationalError as erro:

        if banco:
            banco.rollback()

        traceback.print_exc()

        print(
            f"Erro na estrutura do banco de dados: {erro}"
        )

        flash(
            f"Erro na estrutura do banco de dados: {erro}",
            "erro"
        )

        return redirect("/usuarios")

    except Exception as erro:

        if banco:
            banco.rollback()

        traceback.print_exc()

        print(
            f"Erro ao cadastrar usuário: {erro}"
        )

        flash(
            "Ocorreu um erro ao cadastrar o usuário.",
            "erro"
        )

        return redirect("/usuarios")

    finally:

        if banco:
            banco.close()

@app.route("/editar_usuario/<int:id>", methods=["GET", "POST"])
def editar_usuario(id):

    if not cargo_permitido([
        "Administrador Geral",
        "Administrador da Instituição"
    ]):
        flash(
            "Você não possui permissão para editar usuários.",
            "erro"
        )
        return redirect("/usuarios")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    cargo_logado = session.get(
        "usuario_cargo",
        ""
    ).strip()

    escola_logada_id = session.get("escola_id")

    modulos_plataforma = [
        "Dashboard",
        "Instituições",
        "Usuários",
        "Turmas",
        "Professores",
        "Alunos",
        "Questões",
        "Provas",
        "Relatórios"
    ]

    try:

        # ==========================================
        # BUSCA O USUÁRIO COM SEGURANÇA
        # ==========================================

        if cargo_logado == "Administrador Geral":

            cursor.execute("""
                SELECT
                    usuarios.*,
                    escolas.nome_instituicao
                FROM usuarios
                LEFT JOIN escolas
                    ON usuarios.escola_id = escolas.id
                WHERE usuarios.id = ?
                LIMIT 1
            """, (
                id,
            ))

        else:

            cursor.execute("""
                SELECT
                    usuarios.*,
                    escolas.nome_instituicao
                FROM usuarios
                LEFT JOIN escolas
                    ON usuarios.escola_id = escolas.id
                WHERE usuarios.id = ?
                  AND usuarios.escola_id = ?
                LIMIT 1
            """, (
                id,
                escola_logada_id
            ))

        usuario = cursor.fetchone()

        if usuario is None:

            flash(
                "Usuário não encontrado ou sem permissão de acesso.",
                "erro"
            )

            return redirect("/usuarios")

        # ==========================================
        # CARGOS DISPONÍVEIS
        # ==========================================

        if cargo_logado == "Administrador Geral":

            cursor.execute("""
                SELECT
                    id,
                    nome
                FROM cargos
                ORDER BY nome
            """)

        else:

            cursor.execute("""
                SELECT
                    id,
                    nome
                FROM cargos
                WHERE nome != 'Administrador Geral'
                ORDER BY nome
            """)

        cargos = cursor.fetchall()

        # ==========================================
        # INSTITUIÇÕES DISPONÍVEIS
        # ==========================================

        if cargo_logado == "Administrador Geral":

            cursor.execute("""
                SELECT
                    id,
                    nome_instituicao
                FROM escolas
                WHERE COALESCE(status, 1) = 1
                ORDER BY nome_instituicao
            """)

            escolas = cursor.fetchall()

        else:
            escolas = []

        # ==========================================
        # SALVAMENTO
        # ==========================================

        if request.method == "POST":

            nome = request.form.get(
                "nome",
                ""
            ).strip()

            email = request.form.get(
                "email",
                ""
            ).strip().lower()

            cpf = request.form.get(
                "cpf",
                ""
            ).strip()

            nova_senha = request.form.get(
                "senha",
                ""
            ).strip()

            cargo_id = request.form.get(
                "cargo_id",
                ""
            ).strip()

            modulos_selecionados = request.form.getlist(
                "modulos_permitidos"
            )

            turmas_selecionadas = request.form.getlist(
                "turmas_vinculadas"
            )

            # ======================================
            # VALIDAÇÕES
            # ======================================

            if not nome:
                flash(
                    "Informe o nome do usuário.",
                    "erro"
                )
                return redirect(f"/editar_usuario/{id}")

            if not email:
                flash(
                    "Informe o e-mail do usuário.",
                    "erro"
                )
                return redirect(f"/editar_usuario/{id}")

            if "@" not in email:
                flash(
                    "Informe um endereço de e-mail válido.",
                    "erro"
                )
                return redirect(f"/editar_usuario/{id}")

            if not cargo_id:
                flash(
                    "Selecione um cargo.",
                    "erro"
                )
                return redirect(f"/editar_usuario/{id}")

            if nova_senha and len(nova_senha) < 6:
                flash(
                    "A nova senha deve possuir pelo menos 6 caracteres.",
                    "erro"
                )
                return redirect(f"/editar_usuario/{id}")

            # ======================================
            # VERIFICA E-MAIL DUPLICADO
            # ======================================

            cursor.execute("""
                SELECT id
                FROM usuarios
                WHERE LOWER(email) = LOWER(?)
                  AND id != ?
                LIMIT 1
            """, (
                email,
                id
            ))

            if cursor.fetchone():

                flash(
                    "Este e-mail já está sendo utilizado por outro usuário.",
                    "erro"
                )

                return redirect(f"/editar_usuario/{id}")

            # ======================================
            # VERIFICA O CARGO
            # ======================================

            cursor.execute("""
                SELECT
                    id,
                    nome
                FROM cargos
                WHERE id = ?
                LIMIT 1
            """, (
                cargo_id,
            ))

            cargo_selecionado = cursor.fetchone()

            if cargo_selecionado is None:

                flash(
                    "O cargo selecionado não existe.",
                    "erro"
                )

                return redirect(f"/editar_usuario/{id}")

            nome_cargo = cargo_selecionado["nome"].strip()

            if (
                cargo_logado != "Administrador Geral"
                and nome_cargo == "Administrador Geral"
            ):

                flash(
                    "Você não pode atribuir o cargo de Administrador Geral.",
                    "erro"
                )

                return redirect(f"/editar_usuario/{id}")

            # ======================================
            # DEFINE A INSTITUIÇÃO
            # ======================================

            if cargo_logado == "Administrador Geral":

                escola_id = request.form.get(
                    "escola_id",
                    ""
                ).strip()

                escola_id = escola_id or None

            else:
                escola_id = escola_logada_id

            # Administrador Geral pode não possuir instituição
            if nome_cargo == "Administrador Geral":
                escola_id = None

            # Outros cargos precisam estar vinculados
            if (
                nome_cargo != "Administrador Geral"
                and not escola_id
            ):

                flash(
                    "Selecione a instituição do usuário.",
                    "erro"
                )

                return redirect(f"/editar_usuario/{id}")

            # Confirma se a instituição existe e está ativa
            if escola_id:

                cursor.execute("""
                    SELECT id
                    FROM escolas
                    WHERE id = ?
                      AND COALESCE(status, 1) = 1
                    LIMIT 1
                """, (
                    escola_id,
                ))

                if cursor.fetchone() is None:

                    flash(
                        "A instituição selecionada é inválida ou está inativa.",
                        "erro"
                    )

                    return redirect(f"/editar_usuario/{id}")

            # Administrador de instituição não pode
            # liberar acesso ao módulo Instituições
            if cargo_logado == "Administrador da Instituição":

                modulos_selecionados = [
                    modulo
                    for modulo in modulos_selecionados
                    if modulo != "Instituições"
                ]

            # ======================================
            # ATUALIZA O USUÁRIO
            # ======================================

            if nova_senha:

                cursor.execute("""
                    UPDATE usuarios
                    SET
                        nome = ?,
                        email = ?,
                        cpf = ?,
                        senha = ?,
                        cargo_id = ?,
                        escola_id = ?
                    WHERE id = ?
                """, (
                    nome,
                    email,
                    cpf or None,
                    nova_senha,
                    cargo_id,
                    escola_id,
                    id
                ))

            else:

                cursor.execute("""
                    UPDATE usuarios
                    SET
                        nome = ?,
                        email = ?,
                        cpf = ?,
                        cargo_id = ?,
                        escola_id = ?
                    WHERE id = ?
                """, (
                    nome,
                    email,
                    cpf or None,
                    cargo_id,
                    escola_id,
                    id
                ))

            # ======================================
            # ATUALIZA AS PERMISSÕES
            # ======================================

            cursor.execute("""
                DELETE FROM usuario_permissoes
                WHERE usuario_id = ?
            """, (
                id,
            ))

            for modulo in modulos_plataforma:

                pode_acessar = (
                    1
                    if modulo in modulos_selecionados
                    else 0
                )

                # Administrador Geral recebe acesso completo
                if nome_cargo == "Administrador Geral":
                    pode_acessar = 1

                # Usuários administrados pela instituição
                # nunca recebem o módulo Instituições
                if (
                    cargo_logado == "Administrador da Instituição"
                    and modulo == "Instituições"
                ):
                    pode_acessar = 0

                cursor.execute("""
                    INSERT INTO usuario_permissoes (
                        usuario_id,
                        modulo,
                        pode_acessar
                    )
                    VALUES (?, ?, ?)
                """, (
                    id,
                    modulo,
                    pode_acessar
                ))

            # ======================================
            # ATUALIZA AS TURMAS DA COORDENAÇÃO
            # ======================================

            cursor.execute("""
                DELETE FROM coordenador_turmas
                WHERE usuario_id = ?
            """, (
                id,
            ))

            nome_cargo_minusculo = nome_cargo.lower()

            cargo_coordenacao = (
                "coordenador" in nome_cargo_minusculo
                or "coordenação" in nome_cargo_minusculo
                or "coordenacao" in nome_cargo_minusculo
            )

            if cargo_coordenacao and escola_id:

                for turma_id in turmas_selecionadas:

                    cursor.execute("""
                        SELECT id
                        FROM turmas
                        WHERE id = ?
                          AND escola_id = ?
                        LIMIT 1
                    """, (
                        turma_id,
                        escola_id
                    ))

                    turma_valida = cursor.fetchone()

                    if turma_valida:

                        cursor.execute("""
                            INSERT OR IGNORE INTO coordenador_turmas (
                                usuario_id,
                                turma_id
                            )
                            VALUES (?, ?)
                        """, (
                            id,
                            turma_id
                        ))

            banco.commit()

            flash(
                "Usuário atualizado com sucesso.",
                "success"
            )

            return redirect("/usuarios")

        # ==========================================
        # PERMISSÕES JÁ MARCADAS
        # ==========================================

        cursor.execute("""
            SELECT modulo
            FROM usuario_permissoes
            WHERE usuario_id = ?
              AND pode_acessar = 1
        """, (
            id,
        ))

        permissoes_marcadas = [
            linha["modulo"]
            for linha in cursor.fetchall()
        ]

        # ==========================================
        # TURMAS DISPONÍVEIS
        # ==========================================

        if cargo_logado == "Administrador Geral":

            cursor.execute("""
                SELECT
                    id,
                    nome,
                    ano,
                    turno,
                    escola_id
                FROM turmas
                ORDER BY escola_id, ano, nome, turno
            """)

        else:

            cursor.execute("""
                SELECT
                    id,
                    nome,
                    ano,
                    turno,
                    escola_id
                FROM turmas
                WHERE escola_id = ?
                ORDER BY ano, nome, turno
            """, (
                escola_logada_id,
            ))

        turmas = cursor.fetchall()

        # ==========================================
        # TURMAS JÁ MARCADAS
        # ==========================================

        cursor.execute("""
            SELECT turma_id
            FROM coordenador_turmas
            WHERE usuario_id = ?
        """, (
            id,
        ))

        turmas_marcadas = [
            linha["turma_id"]
            for linha in cursor.fetchall()
        ]

        # ==========================================
        # ABRE A TELA
        # ==========================================

        return render_template(
            "gestao/editar_usuario.html",
            usuario=usuario,
            cargos=cargos,
            escolas=escolas,
            turmas=turmas,
            turmas_marcadas=turmas_marcadas,
            modulos_plataforma=modulos_plataforma,
            permissoes_marcadas=permissoes_marcadas
        )

    except sqlite3.IntegrityError as erro:

        banco.rollback()

        print(
            f"Erro de integridade ao editar usuário: {erro}"
        )

        flash(
            "Não foi possível salvar. Verifique se o e-mail ou CPF já está cadastrado.",
            "erro"
        )

        return redirect(f"/editar_usuario/{id}")

    except Exception as erro:

        banco.rollback()

        import traceback
        traceback.print_exc()

        print(
            f"Erro ao editar usuário: {erro}"
        )

        flash(
            "Ocorreu um erro ao editar o usuário.",
            "erro"
        )

        return redirect("/usuarios")

    finally:
        banco.close()

@app.route("/ativar_inativar_usuario/<int:id>")
def ativar_inativar_usuario(id):

    if not cargo_permitido([
        "Administrador Geral",
        "Administrador da Instituição"
    ]):
        return redirect("/login")

    # Impede o usuário conectado de inativar a própria conta
    if id == session.get("usuario_id"):
        flash(
            "Você não pode inativar o usuário que está conectado.",
            "erro"
        )
        return redirect("/usuarios")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    cargo_logado = session.get("usuario_cargo")
    escola_id = session.get("escola_id")

    # Administrador Geral pode alterar qualquer usuário.
    # Administrador da Instituição altera apenas usuários da própria escola.
    if cargo_logado == "Administrador Geral":

        cursor.execute("""
            SELECT id, ativo
            FROM usuarios
            WHERE id = ?
            LIMIT 1
        """, (id,))

    else:

        cursor.execute("""
            SELECT id, ativo
            FROM usuarios
            WHERE id = ?
              AND escola_id = ?
            LIMIT 1
        """, (
            id,
            escola_id
        ))

    usuario = cursor.fetchone()

    if usuario is None:
        banco.close()

        flash(
            "Usuário não encontrado ou sem permissão para esta ação.",
            "erro"
        )

        return redirect("/usuarios")

    novo_status = 0 if usuario["ativo"] == 1 else 1

    cursor.execute("""
        UPDATE usuarios
        SET ativo = ?
        WHERE id = ?
    """, (
        novo_status,
        id
    ))

    banco.commit()
    banco.close()

    if novo_status == 1:
        flash("Usuário ativado com sucesso.", "success")
    else:
        flash("Usuário inativado com sucesso.", "success")

    return redirect("/usuarios")

@app.route("/meu-perfil", methods=["GET", "POST"])
def meu_perfil():
    """Perfil pessoal e, para autônomos, configuração do próprio espaço."""
    usuario_id = session.get("usuario_id")
    if not usuario_id:
        flash("Faça login para acessar seu perfil.", "aviso")
        return redirect("/login")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        garantir_estrutura_usuarios(cursor)
        # Campos extras de identidade visual do espaço autônomo.
        colunas_escola = {r[1] for r in cursor.execute("PRAGMA table_info(escolas)").fetchall()}
        for coluna, definicao in {
            "cor_primaria": "TEXT",
            "cor_secundaria": "TEXT",
            "cabecalho_documentos": "TEXT",
            "rodape_documentos": "TEXT"
        }.items():
            if coluna not in colunas_escola:
                cursor.execute(f"ALTER TABLE escolas ADD COLUMN {coluna} {definicao}")
        banco.commit()

        cursor.execute("""
            SELECT u.*, c.nome AS cargo
            FROM usuarios u
            LEFT JOIN cargos c ON c.id = u.cargo_id
            WHERE u.id = ? LIMIT 1
        """, (usuario_id,))
        conta = cursor.fetchone()
        if not conta:
            session.clear()
            flash("Usuário não encontrado. Entre novamente.", "erro")
            return redirect("/login")

        eh_autonomo = (conta["tipo_conta"] == "autonomo" or conta["cargo"] == "Professor Autônomo")
        escola_id = conta["escola_id"]

        if request.method == "POST":
            acao = request.form.get("acao", "dados").strip()

            if acao == "dados":
                nome = request.form.get("nome", "").strip()
                email = request.form.get("email", "").strip().lower()
                cpf = request.form.get("cpf", "").strip()
                telefone = request.form.get("telefone", "").strip()
                data_nascimento = request.form.get("data_nascimento", "").strip()
                if not nome or not email:
                    flash("Informe seu nome e e-mail.", "erro")
                    return redirect(url_for("meu_perfil"))
                duplicado = cursor.execute(
                    "SELECT id FROM usuarios WHERE LOWER(email)=LOWER(?) AND id<>? LIMIT 1",
                    (email, usuario_id)
                ).fetchone()
                if duplicado:
                    flash("Este e-mail já está sendo utilizado.", "erro")
                    return redirect(url_for("meu_perfil"))
                cursor.execute("""UPDATE usuarios SET nome=?,email=?,cpf=?,telefone=?,data_nascimento=? WHERE id=?""",
                    (nome,email,cpf or None,telefone or None,data_nascimento or None,usuario_id))
                banco.commit(); session["usuario_nome"] = nome
                flash("Informações pessoais atualizadas com sucesso.", "success")
                return redirect(url_for("meu_perfil"))

            if acao == "institucional":
                if not eh_autonomo or not escola_id:
                    flash("Esta configuração é exclusiva do professor autônomo.", "erro")
                    return redirect(url_for("meu_perfil"))
                nome_instituicao = request.form.get("nome_instituicao", "").strip()
                if not nome_instituicao:
                    flash("Informe o nome da instituição ou identificação profissional.", "erro")
                    return redirect(url_for("meu_perfil") + "#institucional")
                ano_letivo = request.form.get("ano_letivo", "").strip()
                etapas = [x.strip() for x in request.form.getlist("etapas_ensino") if x.strip()]
                modalidades = [x.strip() for x in request.form.getlist("modalidade_ensino") if x.strip()]
                componentes_raw = request.form.getlist("componentes_curriculares")
                componentes=[]; vistos=set()
                for item in componentes_raw:
                    try:
                        obj=json.loads(item); etapa=str(obj.get("etapa","")).strip(); nome=str(obj.get("nome","")).strip(); tipo=str(obj.get("tipo","padrao")).strip()
                    except Exception:
                        continue
                    chave=(etapa.lower(),nome.lower())
                    if etapa and nome and etapa in etapas and chave not in vistos:
                        vistos.add(chave); componentes.append((etapa,nome,tipo if tipo in ("padrao","manual") else "padrao"))
                if etapas and not componentes:
                    flash("Selecione ao menos um componente curricular para as etapas marcadas.", "erro")
                    return redirect(url_for("meu_perfil") + "#academico")

                logo_atual = cursor.execute("SELECT logo FROM escolas WHERE id=?", (escola_id,)).fetchone()
                nome_logo = logo_atual["logo"] if logo_atual else ""
                logo=request.files.get("logo")
                if logo and logo.filename:
                    extensao=Path(secure_filename(logo.filename)).suffix.lower()
                    if extensao not in {".png",".jpg",".jpeg",".webp"}:
                        flash("A logo deve estar em PNG, JPG, JPEG ou WEBP.", "erro")
                        return redirect(url_for("meu_perfil") + "#identidade")
                    nome_logo=f"autonomo_{escola_id}_{int(datetime.now().timestamp())}{extensao}"
                    logo.save(os.path.join(app.config["UPLOAD_FOLDER"], nome_logo))

                cursor.execute("""
                    UPDATE escolas SET nome_instituicao=?,codigo_inep=?,cnpj=?,cep=?,endereco=?,cidade=?,estado=?,
                    telefone=?,whatsapp=?,email=?,site=?,tipo_instituicao=?,ano_letivo=?,modalidade_ensino=?,etapas_ensino=?,
                    logo=?,cor_primaria=?,cor_secundaria=?,cabecalho_documentos=?,rodape_documentos=? WHERE id=?
                """, (
                    nome_instituicao,request.form.get("codigo_inep","").strip() or None,
                    request.form.get("cnpj","").strip() or None,request.form.get("cep","").strip() or None,
                    request.form.get("endereco","").strip() or None,request.form.get("cidade","").strip() or None,
                    request.form.get("estado","").strip() or None,request.form.get("telefone_institucional","").strip() or None,
                    request.form.get("whatsapp","").strip() or None,request.form.get("email_institucional","").strip().lower() or None,
                    request.form.get("site","").strip() or None,request.form.get("tipo_instituicao","").strip() or None,
                    ano_letivo or None,", ".join(modalidades),", ".join(etapas),nome_logo,
                    request.form.get("cor_primaria","").strip() or None,request.form.get("cor_secundaria","").strip() or None,
                    request.form.get("cabecalho_documentos","").strip() or None,request.form.get("rodape_documentos","").strip() or None,
                    escola_id
                ))
                if ano_letivo:
                    sincronizar_ano_letivo_instituicao(cursor, escola_id, ano_letivo, tornar_ativo=True)
                cursor.execute("DELETE FROM componentes_curriculares WHERE escola_id=?", (escola_id,))
                for etapa,nome,tipo in componentes:
                    cursor.execute("INSERT INTO componentes_curriculares(escola_id,etapa_ensino,nome,tipo,ativo) VALUES(?,?,?,?,1)", (escola_id,etapa,nome,tipo))
                banco.commit()
                atualizar_ano_letivo_na_sessao(escola_id)
                flash("Informações profissionais, acadêmicas e identidade visual salvas com sucesso.", "success")
                return redirect(url_for("meu_perfil") + "#institucional")

            if acao == "senha":
                senha_atual=request.form.get("senha_atual",""); nova=request.form.get("nova_senha",""); confirmar=request.form.get("confirmar_senha","")
                senha_salva=conta["senha"] or ""
                try: valida=check_password_hash(senha_salva,senha_atual)
                except (ValueError,TypeError): valida=(senha_salva==senha_atual)
                if not valida:
                    flash("A senha atual está incorreta.", "erro"); return redirect(url_for("meu_perfil")+"#seguranca")
                if len(nova)<8:
                    flash("A nova senha deve ter pelo menos 8 caracteres.", "erro"); return redirect(url_for("meu_perfil")+"#seguranca")
                if nova!=confirmar:
                    flash("A confirmação da nova senha não confere.", "erro"); return redirect(url_for("meu_perfil")+"#seguranca")
                cursor.execute("UPDATE usuarios SET senha=? WHERE id=?", (generate_password_hash(nova),usuario_id)); banco.commit()
                flash("Senha alterada com sucesso.", "success"); return redirect(url_for("meu_perfil")+"#seguranca")

        cursor.execute("""
            SELECT u.id,u.nome,u.email,u.cpf,u.telefone,u.data_nascimento,u.tipo_conta,c.nome AS cargo,
                   e.*
            FROM usuarios u LEFT JOIN cargos c ON c.id=u.cargo_id LEFT JOIN escolas e ON e.id=u.escola_id
            WHERE u.id=? LIMIT 1
        """, (usuario_id,))
        usuario=cursor.fetchone()
        componentes=[]
        if eh_autonomo and escola_id:
            componentes=[dict(r) for r in cursor.execute("SELECT etapa_ensino,nome,tipo FROM componentes_curriculares WHERE escola_id=? AND ativo=1 ORDER BY etapa_ensino,nome", (escola_id,)).fetchall()]
        return render_template("meu_perfil.html", usuario=usuario, eh_autonomo=eh_autonomo, componentes_salvos=componentes,
            modalidades_marcadas=[x.strip() for x in (usuario["modalidade_ensino"] or "").split(",") if x.strip()] if usuario else [],
            etapas_marcadas=[x.strip() for x in (usuario["etapas_ensino"] or "").split(",") if x.strip()] if usuario else [])
    except sqlite3.IntegrityError:
        banco.rollback(); flash("Não foi possível salvar. Verifique dados duplicados.", "erro"); return redirect(url_for("meu_perfil"))
    except sqlite3.Error as erro:
        banco.rollback(); app.logger.exception("Erro no perfil"); flash(f"Não foi possível atualizar seu perfil: {erro}", "erro"); return redirect("/")
    finally:
        banco.close()


@app.route("/login", methods=["GET", "POST"])
def login():

    # O retorno do logout deve sempre mostrar a tela de login, mesmo que o
    # navegador envie por engano um cookie antigo na primeira requisição.
    if request.method == "GET" and request.args.get("saiu") == "1":
        session.clear()
        session.modified = True
        return render_template("login.html")

    if request.method == "GET" and session.get("usuario_id"):
        return redirect(url_for("index"))

    if request.method == "POST":
        email = request.form.get("email", "").strip()
        senha = request.form.get("senha", "").strip()

        if not email or not senha:
            flash("Informe o usuário/e-mail e a senha.", "erro")
            return render_template("login.html")

        banco = conectar_banco()
        banco.row_factory = sqlite3.Row
        cursor = banco.cursor()

        try:
            cursor.execute("""
                SELECT
                    usuarios.*,
                    cargos.nome AS cargo
                FROM usuarios
                LEFT JOIN cargos
                    ON usuarios.cargo_id = cargos.id
                WHERE LOWER(TRIM(usuarios.email)) = LOWER(TRIM(?))
                  AND usuarios.ativo = 1
                LIMIT 1
            """, (email,))

            usuario = cursor.fetchone()

            senha_valida = False

            if usuario:
                senha_salva = usuario["senha"] or ""

                try:
                    senha_valida = check_password_hash(senha_salva, senha)
                except (ValueError, TypeError):
                    senha_valida = False

                # Compatibilidade temporária com senhas antigas em texto puro.
                if not senha_valida and senha_salva == senha:
                    senha_valida = True
                    cursor.execute(
                        "UPDATE usuarios SET senha = ? WHERE id = ?",
                        (generate_password_hash(senha), usuario["id"])
                    )
                    banco.commit()

            if not usuario or not senha_valida:
                flash("Usuário, senha ou status inválido.", "erro")
                return render_template("login.html")

            session.clear()
            session["usuario_id"] = usuario["id"]
            session["usuario_nome"] = usuario["nome"]
            session["usuario_cargo"] = (
                "Professor Autônomo"
                if (usuario["tipo_conta"] or "").strip().lower() == "autonomo"
                else (usuario["cargo"] or "")
            )
            session["tipo_conta"] = usuario["tipo_conta"] or "institucional"
            session["escola_id"] = usuario["escola_id"]

            if usuario["escola_id"]:
                atualizar_ano_letivo_na_sessao(usuario["escola_id"])

            if usuario["escola_id"]:
                session["escola_id"] = int(usuario["escola_id"])
                garantir_ano_atual_para_escola(usuario["escola_id"])
                atualizar_ano_letivo_na_sessao(usuario["escola_id"])
            else:
                obter_ano_global_administrador()

            token_sessao = uuid.uuid4().hex
            session["token_sessao"] = token_sessao
            ip_login = request.headers.get("X-Forwarded-For", request.remote_addr)
            cursor.execute("""
                INSERT INTO sessoes_usuario(token, usuario_id, ip, user_agent)
                VALUES (?, ?, ?, ?)
            """, (token_sessao, usuario["id"], ip_login, request.headers.get("User-Agent", "")[:500]))
            cursor.execute("""
                UPDATE usuarios SET ultimo_login_em=CURRENT_TIMESTAMP, ultimo_login_ip=? WHERE id=?
            """, (ip_login, usuario["id"]))
            banco.commit()
            registrar_auditoria("LOGIN", recurso="autenticacao", recurso_id=usuario["id"])

            termos_ok = usuario["termos_versao"] == TERMOS_VERSAO and bool(usuario["termos_aceitos_em"])
            privacidade_ok = usuario["privacidade_versao"] == PRIVACIDADE_VERSAO and bool(usuario["privacidade_aceita_em"])
            if not (termos_ok and privacidade_ok):
                return redirect(url_for("aceite_legal"))
            return redirect(url_for("index"))

        except sqlite3.Error as erro:
            print("ERRO NO LOGIN:", erro)
            flash("Não foi possível realizar o login.", "erro")
            return render_template("login.html")

        finally:
            banco.close()

    return render_template("login.html")

@app.route("/logout", methods=["GET", "POST"], strict_slashes=False)
@app.route("/sair", methods=["GET", "POST"], strict_slashes=False)
def logout():
    """Encerra totalmente a sessão atual e volta ao login."""
    # Remove todos os dados do usuário. Não preserva nome, escola, plano ou
    # qualquer chave usada pelos módulos acadêmico e SaaS.
    token_sessao = session.get("token_sessao")
    if token_sessao:
        try:
            banco = conectar_banco()
            banco.execute("UPDATE sessoes_usuario SET ativa=0, encerrado_em=CURRENT_TIMESTAMP WHERE token=?", (token_sessao,))
            banco.commit(); banco.close()
        except Exception:
            app.logger.exception("Falha ao encerrar registro de sessão")
    registrar_auditoria("LOGOUT", recurso="autenticacao")
    for chave in list(session.keys()):
        session.pop(chave, None)
    session.clear()
    session.modified = True

    resposta = redirect(url_for("login", saiu="1"), code=303)
    resposta.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0, private"
    resposta.headers["Pragma"] = "no-cache"
    resposta.headers["Expires"] = "0"
    resposta.headers["Clear-Site-Data"] = '"cache"'

    # Expira o cookie no mesmo caminho em que o Flask o cria. O próprio
    # Flask também salvará a sessão vazia ao finalizar a resposta.
    nome_cookie = app.config.get("SESSION_COOKIE_NAME", "session")
    resposta.delete_cookie(nome_cookie, path="/")
    dominio_cookie = app.config.get("SESSION_COOKIE_DOMAIN")
    if dominio_cookie:
        resposta.delete_cookie(nome_cookie, path="/", domain=dominio_cookie)

    return resposta


@app.after_request
def impedir_cache_de_paginas_autenticadas(resposta):
    """Evita que o navegador reapresente o dashboard depois do logout."""
    if session.get("usuario_id") or request.endpoint in {"logout", "login"}:
        resposta.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0, private"
        resposta.headers["Pragma"] = "no-cache"
        resposta.headers["Expires"] = "0"
    return resposta

@app.route("/permissoes/<int:cargo_id>")
def permissoes(cargo_id):

    if not cargo_permitido(["Administrador"]):
        return redirect("/login")

    modulos = [
        "Dashboard",
        "Turmas",
        "Alunos",
        "Professores",
        "Questões",
        "Provas",
        "Correção",
        "Relatórios",
        "Gestão"
    ]

    banco = conectar_banco()
    cursor = banco.cursor()

    cursor.execute("SELECT nome FROM cargos WHERE id = ?", (cargo_id,))
    cargo = cursor.fetchone()

    cursor.execute("""
        SELECT modulo
        FROM permissoes
        WHERE cargo_id = ?
        AND pode_acessar = 1
    """, (cargo_id,))

    permissoes_salvas = [linha[0] for linha in cursor.fetchall()]

    banco.close()

    return render_template(
        "permissoes.html",
        cargo=cargo,
        cargo_id=cargo_id,
        modulos=modulos,
        permissoes_salvas=permissoes_salvas
    )

def permissao_modulo(modulo):

    if "usuario_id" not in session:
        return False

    usuario_id = session.get("usuario_id")
    cargo = session.get("usuario_cargo", "").strip()

    # Administrador Geral possui acesso completo
    if cargo == "Administrador Geral":
        return True

    # Professor autônomo acessa os módulos pedagógicos do próprio espaço.
    if cargo == "Professor Autônomo":
        return modulo in {"Dashboard", "Turmas", "Alunos", "Banco de Questões", "Questões", "Provas", "Relatórios", "Aplicações"}

    # Administrador da Instituição possui acesso completo,
    # exceto ao gerenciamento geral das instituições
    if cargo == "Administrador da Instituição":
        return modulo != "Instituições"

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:

        # Verifica se o usuário possui permissões individuais cadastradas
        cursor.execute("""
            SELECT COUNT(*) AS total
            FROM usuario_permissoes
            WHERE usuario_id = ?
        """, (usuario_id,))

        possui_permissoes_individuais = (
            cursor.fetchone()["total"] > 0
        )

        # Se existem permissões individuais, elas são definitivas:
        # marcado = acessa; desmarcado = não acessa.
        if possui_permissoes_individuais:

            cursor.execute("""
                SELECT pode_acessar
                FROM usuario_permissoes
                WHERE usuario_id = ?
                  AND modulo = ?
                LIMIT 1
            """, (
                usuario_id,
                modulo
            ))

            permissao_individual = cursor.fetchone()

            if permissao_individual is None:
                return False

            return permissao_individual["pode_acessar"] == 1

        # Só usa as permissões do cargo quando o usuário ainda
        # não possui nenhuma configuração individual.
        cursor.execute("""
            SELECT cargo_id
            FROM usuarios
            WHERE id = ?
            LIMIT 1
        """, (usuario_id,))

        usuario = cursor.fetchone()

        if usuario is None:
            return False

        cursor.execute("""
            SELECT pode_acessar
            FROM permissoes
            WHERE cargo_id = ?
              AND modulo = ?
            LIMIT 1
        """, (
            usuario["cargo_id"],
            modulo
        ))

        permissao_cargo = cursor.fetchone()

        if permissao_cargo is None:
            return False

        return permissao_cargo["pode_acessar"] == 1

    finally:
        banco.close()

@app.context_processor
def inject_permissoes():
    return dict(
        permissao_modulo=permissao_modulo
    )

@app.route("/salvar_permissoes/<int:cargo_id>", methods=["POST"])
def salvar_permissoes(cargo_id):

    if not cargo_permitido(["Administrador"]):
        return redirect("/login")

    modulos_marcados = request.form.getlist("modulos")

    todos_modulos = [
        "Dashboard",
        "Turmas",
        "Alunos",
        "Professores",
        "Questões",
        "Provas",
        "Correção",
        "Relatórios",
        "Gestão"
    ]

    banco = conectar_banco()
    cursor = banco.cursor()

    cursor.execute(
        "DELETE FROM permissoes WHERE cargo_id = ?",
        (cargo_id,)
    )

    for modulo in todos_modulos:
        pode_acessar = 1 if modulo in modulos_marcados else 0

        cursor.execute("""
            INSERT INTO permissoes
            (cargo_id, modulo, pode_acessar)
            VALUES (?, ?, ?)
        """, (
            cargo_id,
            modulo,
            pode_acessar
        ))

    banco.commit()

    cursor.execute(
        "SELECT nome FROM cargos WHERE id = ?",
        (cargo_id,)
    )

    nome_cargo = cursor.fetchone()[0]

    flash(
        f"Permissões do cargo '{nome_cargo}' salvas com sucesso!",
        "success"
    )

    banco.close()

    return redirect(f"/permissoes/{cargo_id}")

@app.route("/acesso_negado")
def acesso_negado():
    return render_template("acesso_negado.html")

@app.route("/recuperar_senha", methods=["POST"])
def recuperar_senha():
    email = (request.form.get("email") or "").strip().lower()

    if not email:
        flash("Informe o e-mail cadastrado.", "erro")
        return redirect("/login")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        cursor.execute("""
            SELECT id, nome, email
            FROM usuarios
            WHERE LOWER(TRIM(email)) = LOWER(TRIM(?))
              AND ativo = 1
            LIMIT 1
        """, (email,))

        usuario = cursor.fetchone()

        # Por segurança, a tela seguinte é exibida mesmo quando o e-mail
        # não existe. Isso evita informar quais endereços estão cadastrados.
        if not usuario:
            flash(
                "Se o e-mail estiver cadastrado, um código será enviado.",
                "sucesso"
            )
            return render_template("verificar_codigo.html", email=email)

        if not app.config.get("MAIL_USERNAME") or not app.config.get("MAIL_PASSWORD"):
            print(
                "ERRO DE E-MAIL: MAIL_USERNAME ou MAIL_PASSWORD não configurados."
            )
            flash(
                "O serviço de recuperação de senha ainda não está configurado. "
                "Entre em contato com o administrador.",
                "erro"
            )
            return redirect("/login")

        import secrets

        codigo = f"{secrets.randbelow(1_000_000):06d}"
        criado_em = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        # Invalida códigos anteriores ainda não utilizados.
        cursor.execute("""
            UPDATE codigos_recuperacao
            SET usado = 1
            WHERE usuario_id = ?
              AND usado = 0
        """, (usuario["id"],))

        cursor.execute("""
            INSERT INTO codigos_recuperacao (
                usuario_id,
                codigo,
                usado,
                criado_em
            )
            VALUES (?, ?, 0, ?)
        """, (
            usuario["id"],
            codigo,
            criado_em
        ))

        mensagem = Message(
            subject="Código de recuperação de senha | ARK EDUS",
            recipients=[usuario["email"]],
            sender=app.config.get("MAIL_DEFAULT_SENDER")
        )

        mensagem.html = f"""
        <div style="font-family:Arial,sans-serif;max-width:560px;margin:auto;
                    color:#0f172a;border:1px solid #dbe4f0;border-radius:18px;
                    overflow:hidden">
            <div style="background:linear-gradient(135deg,#17468f,#2563eb);
                        color:#fff;padding:26px">
                <h2 style="margin:0">Recuperação de senha</h2>
            </div>

            <div style="padding:26px">
                <p>Olá, <strong>{usuario["nome"]}</strong>!</p>

                <p>Use o código abaixo para redefinir sua senha:</p>

                <div style="margin:24px 0;padding:18px;text-align:center;
                            border-radius:14px;background:#eff6ff;
                            color:#17468f;font-size:32px;font-weight:800;
                            letter-spacing:8px">
                    {codigo}
                </div>

                <p>O código expira em 30 minutos e só poderá ser usado uma vez.</p>

                <p style="color:#64748b;font-size:13px">
                    Caso você não tenha solicitado a alteração, ignore esta mensagem.
                </p>
            </div>
        </div>
        """

        try:
            mail.send(mensagem)
        except Exception as erro_email:
            banco.rollback()
            print("ERRO AO ENVIAR E-MAIL DE RECUPERAÇÃO:", repr(erro_email))
            flash(
                "Não foi possível enviar o código agora. "
                "Verifique as configurações de e-mail da plataforma.",
                "erro"
            )
            return redirect("/login")

        banco.commit()

        flash("Enviamos um código de 6 dígitos para o seu e-mail.", "sucesso")
        return render_template(
            "verificar_codigo.html",
            email=usuario["email"]
        )

    except sqlite3.Error as erro:
        banco.rollback()
        print("ERRO NA RECUPERAÇÃO DE SENHA:", erro)
        flash(
            "Não foi possível iniciar a recuperação de senha.",
            "erro"
        )
        return redirect("/login")

    finally:
        banco.close()


@app.route("/verificar_codigo", methods=["POST"])
def verificar_codigo():
    email = (request.form.get("email") or "").strip().lower()
    codigo = (request.form.get("codigo") or "").strip()

    if not email or not codigo:
        flash("Informe o e-mail e o código recebido.", "erro")
        return render_template("verificar_codigo.html", email=email)

    if not codigo.isdigit() or len(codigo) != 6:
        flash("O código deve possuir exatamente 6 números.", "erro")
        return render_template("verificar_codigo.html", email=email)

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        cursor.execute("""
            SELECT
                u.id AS usuario_id,
                cr.id AS codigo_id,
                cr.criado_em
            FROM usuarios AS u
            INNER JOIN codigos_recuperacao AS cr
                ON cr.usuario_id = u.id
            WHERE LOWER(TRIM(u.email)) = LOWER(TRIM(?))
              AND cr.codigo = ?
              AND cr.usado = 0
              AND u.ativo = 1
            ORDER BY cr.id DESC
            LIMIT 1
        """, (email, codigo))

        registro = cursor.fetchone()

        if not registro:
            flash("Código inválido ou já utilizado.", "erro")
            return render_template("verificar_codigo.html", email=email)

        try:
            criado_em = datetime.strptime(
                registro["criado_em"],
                "%Y-%m-%d %H:%M:%S"
            )
        except (TypeError, ValueError):
            criado_em = None

        if not criado_em or (datetime.now() - criado_em).total_seconds() > 1800:
            cursor.execute("""
                UPDATE codigos_recuperacao
                SET usado = 1
                WHERE id = ?
            """, (registro["codigo_id"],))
            banco.commit()

            flash("Este código expirou. Solicite um novo código.", "erro")
            return redirect("/login")

        return render_template(
            "nova_senha.html",
            email=email,
            codigo=codigo
        )

    except sqlite3.Error as erro:
        print("ERRO AO VERIFICAR CÓDIGO:", erro)
        flash("Não foi possível verificar o código.", "erro")
        return render_template("verificar_codigo.html", email=email)

    finally:
        banco.close()


@app.route("/salvar_senha_usuario", methods=["POST"])
def salvar_senha_usuario():
    email = (request.form.get("email") or "").strip().lower()
    codigo = (request.form.get("codigo") or "").strip()
    nova_senha = request.form.get("nova_senha") or ""
    confirmar_senha = request.form.get("confirmar_senha") or ""

    if not email or not codigo:
        flash("A solicitação de recuperação é inválida.", "erro")
        return redirect("/login")

    if nova_senha != confirmar_senha:
        flash("As senhas informadas não conferem.", "erro")
        return render_template(
            "nova_senha.html",
            email=email,
            codigo=codigo
        )

    if len(nova_senha) < 6:
        flash("A nova senha deve ter pelo menos 6 caracteres.", "erro")
        return render_template(
            "nova_senha.html",
            email=email,
            codigo=codigo
        )

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        cursor.execute("""
            SELECT
                u.id AS usuario_id,
                cr.id AS codigo_id,
                cr.criado_em
            FROM usuarios AS u
            INNER JOIN codigos_recuperacao AS cr
                ON cr.usuario_id = u.id
            WHERE LOWER(TRIM(u.email)) = LOWER(TRIM(?))
              AND cr.codigo = ?
              AND cr.usado = 0
              AND u.ativo = 1
            ORDER BY cr.id DESC
            LIMIT 1
        """, (email, codigo))

        registro = cursor.fetchone()

        if not registro:
            flash("Código inválido ou já utilizado.", "erro")
            return redirect("/login")

        try:
            criado_em = datetime.strptime(
                registro["criado_em"],
                "%Y-%m-%d %H:%M:%S"
            )
        except (TypeError, ValueError):
            criado_em = None

        if not criado_em or (datetime.now() - criado_em).total_seconds() > 1800:
            cursor.execute("""
                UPDATE codigos_recuperacao
                SET usado = 1
                WHERE id = ?
            """, (registro["codigo_id"],))
            banco.commit()

            flash("O código expirou. Solicite uma nova recuperação.", "erro")
            return redirect("/login")

        cursor.execute("""
            UPDATE usuarios
            SET senha = ?
            WHERE id = ?
        """, (
            generate_password_hash(nova_senha),
            registro["usuario_id"]
        ))

        cursor.execute("""
            UPDATE codigos_recuperacao
            SET usado = 1
            WHERE usuario_id = ?
        """, (registro["usuario_id"],))

        banco.commit()

        flash(
            "Senha alterada com sucesso. Entre com a nova senha.",
            "sucesso"
        )
        return redirect("/login")

    except sqlite3.Error as erro:
        banco.rollback()
        print("ERRO AO SALVAR NOVA SENHA:", erro)
        flash("Não foi possível alterar a senha.", "erro")
        return render_template(
            "nova_senha.html",
            email=email,
            codigo=codigo
        )

    finally:
        banco.close()

def garantir_estrutura_lgpd():
    banco = conectar_banco()
    cursor = banco.cursor()
    try:
        colunas = {linha[1] for linha in cursor.execute("PRAGMA table_info(usuarios)").fetchall()}
        for nome, definicao in (
            ("cpf_criptografado", "TEXT"),
            ("cpf_hash", "TEXT"),
            ("cpf_final", "TEXT"),
            ("termos_aceitos_em", "TEXT"),
            ("termos_versao", "TEXT"),
            ("privacidade_aceita_em", "TEXT"),
            ("privacidade_versao", "TEXT"),
            ("ultimo_login_em", "TEXT"),
            ("ultimo_login_ip", "TEXT"),
        ):
            if nome not in colunas:
                cursor.execute(f"ALTER TABLE usuarios ADD COLUMN {nome} {definicao}")

        cursor.executescript("""
            CREATE UNIQUE INDEX IF NOT EXISTS idx_usuarios_cpf_hash
                ON usuarios(cpf_hash) WHERE cpf_hash IS NOT NULL;

            CREATE TABLE IF NOT EXISTS logs_auditoria (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                usuario_id INTEGER,
                escola_id INTEGER,
                acao TEXT NOT NULL,
                recurso TEXT,
                recurso_id TEXT,
                metodo TEXT,
                rota TEXT,
                ip TEXT,
                user_agent TEXT,
                detalhes TEXT,
                criado_em TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY(usuario_id) REFERENCES usuarios(id) ON DELETE SET NULL
            );

            CREATE TABLE IF NOT EXISTS solicitacoes_lgpd (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                usuario_id INTEGER,
                escola_id INTEGER,
                tipo TEXT NOT NULL,
                descricao TEXT,
                status TEXT NOT NULL DEFAULT 'aberta',
                criado_em TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
                concluido_em TEXT,
                FOREIGN KEY(usuario_id) REFERENCES usuarios(id) ON DELETE SET NULL
            );

            CREATE TABLE IF NOT EXISTS aceites_legais (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                usuario_id INTEGER NOT NULL,
                documento TEXT NOT NULL,
                versao TEXT NOT NULL,
                aceito_em TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
                ip TEXT,
                user_agent TEXT,
                FOREIGN KEY(usuario_id) REFERENCES usuarios(id) ON DELETE CASCADE
            );

            CREATE TABLE IF NOT EXISTS sessoes_usuario (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                token TEXT NOT NULL UNIQUE,
                usuario_id INTEGER NOT NULL,
                ip TEXT,
                user_agent TEXT,
                criado_em TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
                ultimo_acesso_em TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
                encerrado_em TEXT,
                ativa INTEGER NOT NULL DEFAULT 1,
                FOREIGN KEY(usuario_id) REFERENCES usuarios(id) ON DELETE CASCADE
            );

            CREATE TABLE IF NOT EXISTS incidentes_seguranca (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                titulo TEXT NOT NULL,
                descricao TEXT NOT NULL,
                gravidade TEXT NOT NULL DEFAULT 'baixa',
                status TEXT NOT NULL DEFAULT 'aberto',
                escola_id INTEGER,
                criado_por INTEGER,
                criado_em TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
                resolvido_em TEXT,
                FOREIGN KEY(criado_por) REFERENCES usuarios(id) ON DELETE SET NULL
            );

            DROP TRIGGER IF EXISTS trg_usuarios_proteger_cpf_insert;
            DROP TRIGGER IF EXISTS trg_usuarios_proteger_cpf_update;

            CREATE TRIGGER trg_usuarios_proteger_cpf_insert
            AFTER INSERT ON usuarios
            WHEN NEW.cpf IS NOT NULL AND instr(NEW.cpf, '*') = 0
            BEGIN
                UPDATE usuarios
                SET cpf_criptografado = cpf_encrypt(NEW.cpf),
                    cpf_hash = cpf_hash(NEW.cpf),
                    cpf_final = cpf_final(NEW.cpf),
                    cpf = cpf_mask(NEW.cpf)
                WHERE id = NEW.id;
            END;

            CREATE TRIGGER trg_usuarios_proteger_cpf_update
            AFTER UPDATE OF cpf ON usuarios
            WHEN NEW.cpf IS NOT NULL AND instr(NEW.cpf, '*') = 0
            BEGIN
                UPDATE usuarios
                SET cpf_criptografado = cpf_encrypt(NEW.cpf),
                    cpf_hash = cpf_hash(NEW.cpf),
                    cpf_final = cpf_final(NEW.cpf),
                    cpf = cpf_mask(NEW.cpf)
                WHERE id = NEW.id;
            END;
        """)

        # Migra CPFs antigos em texto aberto.
        antigos = cursor.execute("""
            SELECT id, cpf FROM usuarios
            WHERE cpf IS NOT NULL AND TRIM(cpf) <> '' AND instr(cpf, '*') = 0
        """).fetchall()
        for usuario_id, cpf in antigos:
            try:
                criptografado, assinatura, final = preparar_cpf(cpf)
                cursor.execute("""
                    UPDATE usuarios SET cpf_criptografado=?, cpf_hash=?, cpf_final=?, cpf=? WHERE id=?
                """, (criptografado, assinatura, final, mascarar_cpf(cpf), usuario_id))
            except (ValueError, RuntimeError):
                cursor.execute("UPDATE usuarios SET cpf=NULL WHERE id=?", (usuario_id,))
        banco.commit()
    finally:
        banco.close()


def registrar_auditoria(acao, recurso=None, recurso_id=None, detalhes=None):
    try:
        banco = conectar_banco()
        banco.execute("""
            INSERT INTO logs_auditoria
            (usuario_id, escola_id, acao, recurso, recurso_id, metodo, rota, ip, user_agent, detalhes)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            session.get("usuario_id"), session.get("escola_id"), acao, recurso,
            str(recurso_id) if recurso_id is not None else None, request.method,
            request.path, request.headers.get("X-Forwarded-For", request.remote_addr),
            request.headers.get("User-Agent", "")[:500],
            json.dumps(detalhes, ensure_ascii=False) if detalhes else None,
        ))
        banco.commit(); banco.close()
    except Exception:
        app.logger.exception("Falha ao registrar auditoria")


@app.before_request
def aplicar_seguranca_sessao():
    if not session.get("usuario_id"):
        return None
    session.permanent = True
    rotas_livres = {"aceite_legal", "logout", "site_termos", "site_privacidade", "static"}
    if request.endpoint in rotas_livres or request.path.startswith("/static/"):
        return None
    try:
        banco = conectar_banco(); banco.row_factory = sqlite3.Row
        usuario = banco.execute("SELECT termos_aceitos_em,termos_versao,privacidade_aceita_em,privacidade_versao FROM usuarios WHERE id=?", (session.get("usuario_id"),)).fetchone()
        token = session.get("token_sessao")
        if token:
            banco.execute("UPDATE sessoes_usuario SET ultimo_acesso_em=CURRENT_TIMESTAMP WHERE token=? AND ativa=1", (token,))
            banco.commit()
        banco.close()
        if usuario and not (usuario["termos_aceitos_em"] and usuario["termos_versao"] == TERMOS_VERSAO and usuario["privacidade_aceita_em"] and usuario["privacidade_versao"] == PRIVACIDADE_VERSAO):
            return redirect(url_for("aceite_legal"))
    except sqlite3.Error:
        app.logger.exception("Falha ao validar aceite legal")
    return None


@app.after_request
def aplicar_cabecalhos_seguranca(resposta):
    resposta.headers["X-Content-Type-Options"] = "nosniff"
    resposta.headers["X-Frame-Options"] = "SAMEORIGIN"
    resposta.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"
    resposta.headers["Permissions-Policy"] = "camera=(), microphone=(), geolocation=()"
    resposta.headers["Content-Security-Policy"] = (
        "default-src 'self'; img-src 'self' data: blob:; "
        "style-src 'self' 'unsafe-inline' https://cdnjs.cloudflare.com; "
        "script-src 'self' 'unsafe-inline' https://cdnjs.cloudflare.com; "
        "font-src 'self' data: https://cdnjs.cloudflare.com; connect-src 'self'"
    )
    if request.is_secure:
        resposta.headers["Strict-Transport-Security"] = "max-age=31536000; includeSubDomains"
    if session.get("usuario_id") and request.method in {"POST", "PUT", "PATCH", "DELETE"}:
        registrar_auditoria("ALTERACAO", recurso=request.endpoint)
    return resposta


criar_tabelas()

# Migração defensiva para bancos persistentes criados por versões antigas.
# IMPORTANTE: apenas cria estruturas ausentes; nunca apaga tabelas nem dados.
def garantir_schema_persistente_critico():
    banco = conectar_banco()
    try:
        cursor = banco.cursor()
        cursor.execute("PRAGMA foreign_keys = ON")
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS componentes_curriculares (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                escola_id INTEGER NOT NULL,
                etapa_ensino TEXT NOT NULL,
                nome TEXT NOT NULL,
                tipo TEXT NOT NULL DEFAULT 'padrao',
                ativo INTEGER NOT NULL DEFAULT 1,
                FOREIGN KEY (escola_id) REFERENCES escolas(id) ON DELETE CASCADE,
                UNIQUE (escola_id, etapa_ensino, nome)
            )
        """)
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS professor_vinculos (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                professor_id INTEGER NOT NULL,
                turma_id INTEGER NOT NULL,
                componente_id INTEGER NOT NULL,
                FOREIGN KEY (professor_id) REFERENCES usuarios(id) ON DELETE CASCADE,
                FOREIGN KEY (turma_id) REFERENCES turmas(id) ON DELETE CASCADE,
                FOREIGN KEY (componente_id) REFERENCES componentes_curriculares(id) ON DELETE CASCADE,
                UNIQUE (professor_id, turma_id, componente_id)
            )
        """)
        banco.commit()
        print("[ARK EDUS] Schema persistente crítico verificado sem apagar dados.")
    finally:
        banco.close()

garantir_schema_persistente_critico()
garantir_estrutura_lgpd()
sincronizar_anos_letivos_legados()
corrigir_ano_corrente_inconsistente()

@app.route("/gestao/instituicoes/editar/<int:id>", methods=["GET", "POST"])
def editar_instituicao(id):

    if not cargo_permitido(["Administrador Geral"]):
        return redirect("/login")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        cursor.execute("""
            SELECT *
            FROM escolas
            WHERE id = ?
            LIMIT 1
        """, (id,))

        escola = cursor.fetchone()

        if escola is None:
            flash("Instituição não encontrada.", "erro")
            return redirect("/gestao/instituicoes")

        cursor.execute("""
            SELECT usuarios.*
            FROM usuarios
            LEFT JOIN cargos
                ON cargos.id = usuarios.cargo_id
            WHERE usuarios.escola_id = ?
              AND cargos.nome = 'Administrador da Instituição'
            ORDER BY usuarios.id
            LIMIT 1
        """, (id,))

        administrador = cursor.fetchone()

        if request.method == "POST":

            nome_instituicao = request.form.get("nome_instituicao", "").strip()
            codigo_inep = request.form.get("codigo_inep", "").strip()
            cnpj = request.form.get("cnpj", "").strip()
            cep = request.form.get("cep", "").strip()
            endereco = request.form.get("endereco", "").strip()
            cidade = request.form.get("cidade", "").strip()
            estado = request.form.get("estado", "").strip()
            telefone = request.form.get("telefone", "").strip()
            whatsapp = request.form.get("whatsapp", "").strip()
            email_institucional = request.form.get("email", "").strip().lower()
            site = request.form.get("site", "").strip()
            diretor = request.form.get("diretor", "").strip()
            coordenador1 = request.form.get("coordenador1", "").strip()
            coordenador2 = request.form.get("coordenador2", "").strip()
            coordenador3 = request.form.get("coordenador3", "").strip()
            secretario = request.form.get("secretario", "").strip()
            tipo_instituicao = request.form.get("tipo_instituicao", "").strip()
            ano_letivo = request.form.get("ano_letivo", "").strip()

            modalidades = request.form.getlist("modalidade_ensino")
            etapas = request.form.getlist("etapas_ensino")
            componentes_recebidos = request.form.getlist("componentes_curriculares")

            admin_nome = request.form.get("admin_nome", "").strip()
            admin_email = request.form.get("admin_email", "").strip().lower()
            admin_cpf = request.form.get("admin_cpf", "").strip()
            admin_senha = request.form.get("admin_senha", "").strip()

            if not nome_instituicao:
                flash("Informe o nome da instituição.", "erro")
                return redirect(f"/gestao/instituicoes/editar/{id}")

            if not tipo_instituicao:
                flash("Selecione o tipo da instituição.", "erro")
                return redirect(f"/gestao/instituicoes/editar/{id}")

            if not ano_letivo:
                flash("Selecione o ano letivo.", "erro")
                return redirect(f"/gestao/instituicoes/editar/{id}")

            if not etapas:
                flash("Selecione pelo menos uma etapa de ensino.", "erro")
                return redirect(f"/gestao/instituicoes/editar/{id}")

            if not componentes_recebidos:
                flash("Selecione pelo menos um componente curricular.", "erro")
                return redirect(f"/gestao/instituicoes/editar/{id}")

            if not admin_nome:
                flash("Informe o nome do administrador.", "erro")
                return redirect(f"/gestao/instituicoes/editar/{id}")

            if not admin_email:
                flash("Informe o e-mail do administrador.", "erro")
                return redirect(f"/gestao/instituicoes/editar/{id}")

            if admin_senha and len(admin_senha) < 6:
                flash("A nova senha deve possuir pelo menos 6 caracteres.", "erro")
                return redirect(f"/gestao/instituicoes/editar/{id}")

            if not administrador and not admin_senha:
                flash("Informe uma senha para criar o administrador da instituição.", "erro")
                return redirect(f"/gestao/instituicoes/editar/{id}")

            if administrador:
                cursor.execute("""
                    SELECT id
                    FROM usuarios
                    WHERE LOWER(email) = LOWER(?)
                      AND id != ?
                    LIMIT 1
                """, (admin_email, administrador["id"]))
            else:
                cursor.execute("""
                    SELECT id
                    FROM usuarios
                    WHERE LOWER(email) = LOWER(?)
                    LIMIT 1
                """, (admin_email,))

            if cursor.fetchone():
                flash("Este e-mail já está sendo utilizado por outro usuário.", "erro")
                return redirect(f"/gestao/instituicoes/editar/{id}")

            componentes_processados = []
            componentes_repetidos = set()

            for componente_json in componentes_recebidos:
                try:
                    componente = json.loads(componente_json)
                    etapa = str(componente.get("etapa", "")).strip()
                    nome = str(componente.get("nome", "")).strip()
                    tipo = str(componente.get("tipo", "padrao")).strip().lower()
                except (json.JSONDecodeError, TypeError, AttributeError):
                    continue

                if not etapa or not nome:
                    continue

                if etapa not in etapas:
                    continue

                if tipo not in ["padrao", "manual"]:
                    tipo = "padrao"

                chave = (etapa.lower(), nome.lower())

                if chave in componentes_repetidos:
                    continue

                componentes_repetidos.add(chave)
                componentes_processados.append({
                    "etapa": etapa,
                    "nome": nome,
                    "tipo": tipo
                })

            if not componentes_processados:
                flash("Não foi possível identificar os componentes curriculares selecionados.", "erro")
                return redirect(f"/gestao/instituicoes/editar/{id}")

            modalidade_ensino = ", ".join(modalidades)
            etapas_ensino = ", ".join(etapas)

            logo = request.files.get("logo")
            nome_logo = escola["logo"] or ""

            if logo and logo.filename:
                nome_logo = secure_filename(logo.filename)
                logo.save(os.path.join(app.config["UPLOAD_FOLDER"], nome_logo))

            cursor.execute("""
                UPDATE escolas
                SET
                    nome_instituicao = ?,
                    codigo_inep = ?,
                    cnpj = ?,
                    cep = ?,
                    endereco = ?,
                    cidade = ?,
                    estado = ?,
                    telefone = ?,
                    whatsapp = ?,
                    email = ?,
                    site = ?,
                    diretor = ?,
                    coordenador1 = ?,
                    coordenador2 = ?,
                    coordenador3 = ?,
                    secretario = ?,
                    tipo_instituicao = ?,
                    ano_letivo = ?,
                    modalidade_ensino = ?,
                    etapas_ensino = ?,
                    logo = ?
                WHERE id = ?
            """, (
                nome_instituicao,
                codigo_inep,
                cnpj,
                cep,
                endereco,
                cidade,
                estado,
                telefone,
                whatsapp,
                email_institucional,
                site,
                diretor,
                coordenador1,
                coordenador2,
                coordenador3,
                secretario,
                tipo_instituicao,
                ano_letivo,
                modalidade_ensino,
                etapas_ensino,
                nome_logo,
                id
            ))

            # Mantém a tabela oficial de anos letivos sincronizada.
            # O ano selecionado na edição passa a ser o ano ativo da escola.
            sincronizar_ano_letivo_instituicao(
                cursor,
                id,
                ano_letivo,
                tornar_ativo=True
            )

            cursor.execute("""
                DELETE FROM componentes_curriculares
                WHERE escola_id = ?
            """, (id,))

            for componente in componentes_processados:
                cursor.execute("""
                    INSERT INTO componentes_curriculares (
                        escola_id,
                        etapa_ensino,
                        nome,
                        tipo,
                        ativo
                    )
                    VALUES (?, ?, ?, ?, 1)
                """, (
                    id,
                    componente["etapa"],
                    componente["nome"],
                    componente["tipo"]
                ))

            if administrador:
                if admin_senha:
                    cursor.execute("""
                        UPDATE usuarios
                        SET nome = ?, email = ?, cpf = ?, senha = ?, escola_id = ?, ativo = 1
                        WHERE id = ?
                    """, (
                        admin_nome,
                        admin_email,
                        admin_cpf,
                        admin_senha,
                        id,
                        administrador["id"]
                    ))
                else:
                    cursor.execute("""
                        UPDATE usuarios
                        SET nome = ?, email = ?, cpf = ?, escola_id = ?, ativo = 1
                        WHERE id = ?
                    """, (
                        admin_nome,
                        admin_email,
                        admin_cpf,
                        id,
                        administrador["id"]
                    ))
            else:
                cursor.execute("""
                    SELECT id
                    FROM cargos
                    WHERE nome = ?
                    LIMIT 1
                """, ("Administrador da Instituição",))

                cargo = cursor.fetchone()

                if cargo is None:
                    cursor.execute("""
                        INSERT INTO cargos (nome)
                        VALUES (?)
                    """, ("Administrador da Instituição",))
                    cargo_id = cursor.lastrowid
                else:
                    cargo_id = cargo["id"]

                cursor.execute("""
                    INSERT INTO usuarios (
                        nome,
                        email,
                        senha,
                        cargo_id,
                        ativo,
                        escola_id,
                        cpf
                    )
                    VALUES (?, ?, ?, ?, 1, ?, ?)
                """, (
                    admin_nome,
                    admin_email,
                    admin_senha,
                    cargo_id,
                    id,
                    admin_cpf
                ))

            banco.commit()

            flash(
                "Instituição, administrador e componentes curriculares atualizados com sucesso!",
                "success"
            )

            return redirect("/gestao/instituicoes")

        modalidades_marcadas = [
            item.strip()
            for item in (escola["modalidade_ensino"] or "").split(",")
            if item.strip()
        ]

        etapas_marcadas = [
            item.strip()
            for item in (escola["etapas_ensino"] or "").split(",")
            if item.strip()
        ]

        cursor.execute("""
            SELECT
                id,
                escola_id,
                etapa_ensino,
                nome,
                tipo,
                ativo
            FROM componentes_curriculares
            WHERE escola_id = ?
              AND ativo = 1
            ORDER BY etapa_ensino, nome
        """, (id,))

        componentes_banco = cursor.fetchall()

        componentes_salvos = [
            {
                "id": componente["id"],
                "escola_id": componente["escola_id"],
                "etapa": componente["etapa_ensino"],
                "nome": componente["nome"],
                "tipo": componente["tipo"] or "padrao",
                "ativo": componente["ativo"]
            }
            for componente in componentes_banco
        ]

        return render_template(
            "gestao/editar_instituicao.html",
            escola=escola,
            administrador=administrador,
            modalidades_marcadas=modalidades_marcadas,
            etapas_marcadas=etapas_marcadas,
            componentes_salvos=componentes_salvos
        )

    except sqlite3.IntegrityError as erro:
        banco.rollback()

        import traceback
        traceback.print_exc()

        print(
            "ERRO DE INTEGRIDADE AO EDITAR INSTITUIÇÃO:",
            repr(erro)
        )

        flash(
            f"Não foi possível salvar as alterações: {erro}",
            "erro"
        )

        return redirect(f"/gestao/instituicoes/editar/{id}")

    except Exception as erro:
        banco.rollback()

        import traceback
        traceback.print_exc()

        print(
            "ERRO COMPLETO AO EDITAR INSTITUIÇÃO:",
            repr(erro)
        )

        flash(
            f"Ocorreu um erro ao salvar as alterações: {erro}",
            "erro"
        )

        return redirect(f"/gestao/instituicoes/editar/{id}")

    finally:
        banco.close()

@app.route("/gestao/instituicoes/ver/<int:id>")
def ver_instituicao(id):

    if not cargo_permitido(["Administrador Geral"]):
        return redirect("/login")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    # ======================================================
    # INSTITUIÇÃO
    # ======================================================

    cursor.execute("""
        SELECT *
        FROM escolas
        WHERE id = ?
    """, (id,))

    escola = cursor.fetchone()

    if escola is None:
        banco.close()
        return redirect("/gestao/instituicoes")

    # ======================================================
    # ADMINISTRADOR DA INSTITUIÇÃO
    # ======================================================

    cursor.execute("""
        SELECT
            usuarios.id,
            usuarios.nome,
            usuarios.email,
            usuarios.cpf,
            usuarios.ativo,
            cargos.nome AS cargo
        FROM usuarios
        LEFT JOIN cargos
            ON usuarios.cargo_id = cargos.id
        WHERE usuarios.escola_id = ?
          AND cargos.nome = 'Administrador da Instituição'
        ORDER BY usuarios.id
        LIMIT 1
    """, (id,))

    administrador = cursor.fetchone()

    # ======================================================
    # TODOS OS USUÁRIOS VINCULADOS À INSTITUIÇÃO
    # ======================================================

    cursor.execute("""
        SELECT
            usuarios.id,
            usuarios.nome,
            usuarios.email,
            usuarios.cpf,
            usuarios.ativo,
            cargos.nome AS cargo
        FROM usuarios
        LEFT JOIN cargos
            ON usuarios.cargo_id = cargos.id
        WHERE usuarios.escola_id = ?
        ORDER BY
            usuarios.ativo DESC,
            cargos.nome,
            usuarios.nome
    """, (id,))

    usuarios_instituicao = cursor.fetchall()

    # ======================================================
    # COMPONENTES CURRICULARES
    # ======================================================

    cursor.execute("""
        SELECT
            id,
            escola_id,
            etapa_ensino,
            nome,
            tipo,
            ativo
        FROM componentes_curriculares
        WHERE escola_id = ?
          AND ativo = 1
        ORDER BY etapa_ensino, nome
    """, (id,))

    componentes = cursor.fetchall()

    componentes_salvos = []

    for componente in componentes:

        componentes_salvos.append({
            "id": componente["id"],
            "escola_id": componente["escola_id"],
            "etapa": componente["etapa_ensino"],
            "nome": componente["nome"],
            "tipo": componente["tipo"],
            "ativo": componente["ativo"]
        })

    banco.close()

    return render_template(
        "gestao/ver_instituicao.html",
        escola=escola,
        administrador=administrador,
        usuarios_instituicao=usuarios_instituicao,
        componentes_salvos=componentes_salvos
    )

@app.route("/gestao/instituicoes/inativar/<int:id>")
def inativar_instituicao(id):

    if not cargo_permitido(["Administrador Geral"]):
        return redirect("/login")

    banco = conectar_banco()
    cursor = banco.cursor()

    cursor.execute("""
        UPDATE escolas
        SET status = 0
        WHERE id = ?
    """, (id,))

    cursor.execute("""
        UPDATE usuarios
        SET ativo = 0
        WHERE escola_id = ?
    """, (id,))

    banco.commit()
    banco.close()

    flash("Instituição inativada com sucesso.", "success")

    return redirect("/gestao/instituicoes")

@app.route("/gestao/instituicoes/ativar/<int:id>")
def ativar_instituicao(id):

    if not cargo_permitido(["Administrador Geral"]):
        return redirect("/login")

    banco = conectar_banco()
    cursor = banco.cursor()

    cursor.execute("""
        UPDATE escolas
        SET status = 1
        WHERE id = ?
    """, (id,))

    cursor.execute("""
        UPDATE usuarios
        SET ativo = 1
        WHERE escola_id = ?
    """, (id,))

    banco.commit()
    banco.close()

    flash("Instituição ativada com sucesso.", "success")

    return redirect("/gestao/instituicoes")

@app.route("/gestao/instituicoes/excluir/<int:id>")
def excluir_instituicao(id):

    if not cargo_permitido(["Administrador Geral"]):
        return redirect("/login")

    banco = conectar_banco()
    cursor = banco.cursor()

    try:
        cursor.execute("""
            DELETE FROM usuarios
            WHERE escola_id = ?
        """, (id,))

        cursor.execute("""
            DELETE FROM escolas
            WHERE id = ?
        """, (id,))

        banco.commit()

        flash("Instituição excluída com sucesso.", "success")

    except Exception as erro:
        banco.rollback()

        print(f"Erro ao excluir instituição: {erro}")

        flash(
            "Não foi possível excluir a instituição.",
            "erro"
        )

    finally:
        banco.close()

    return redirect("/gestao/instituicoes")

@app.route("/excluir_usuario/<int:id>")
def excluir_usuario(id):

    if not cargo_permitido([
        "Administrador Geral",
        "Administrador da Instituição"
    ]):
        return redirect("/login")

    if id == session.get("usuario_id"):
        flash(
            "Você não pode excluir o próprio usuário conectado.",
            "erro"
        )
        return redirect("/usuarios")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    cargo_logado = session.get("usuario_cargo")
    escola_id = session.get("escola_id")

    if cargo_logado == "Administrador Geral":
        cursor.execute("""
            SELECT id
            FROM usuarios
            WHERE id = ?
            LIMIT 1
        """, (id,))
    else:
        cursor.execute("""
            SELECT id
            FROM usuarios
            WHERE id = ?
              AND escola_id = ?
            LIMIT 1
        """, (
            id,
            escola_id
        ))

    usuario = cursor.fetchone()

    if usuario is None:
        banco.close()

        flash(
            "Usuário não encontrado ou sem permissão para excluir.",
            "erro"
        )

        return redirect("/usuarios")

    try:
        cursor.execute("""
            DELETE FROM coordenador_turmas
            WHERE usuario_id = ?
        """, (id,))

        cursor.execute("""
            DELETE FROM usuarios
            WHERE id = ?
        """, (id,))

        banco.commit()

        flash(
            "Usuário excluído com sucesso.",
            "success"
        )

    except Exception as erro:
        banco.rollback()

        print(f"Erro ao excluir usuário: {erro}")

        flash(
            "Não foi possível excluir o usuário.",
            "erro"
        )

    finally:
        banco.close()

    return redirect("/usuarios")

# =========================================================
# API - COMPONENTES CURRICULARES DA TURMA
# =========================================================

@app.route(
    "/api/turmas/<int:turma_id>/componentes",
    methods=["GET"]
)
def api_componentes_da_turma(turma_id):

    if not cargo_permitido([
        "Administrador Geral",
        "Administrador da Instituição"
    ]):
        return jsonify({
            "sucesso": False,
            "mensagem": "Você não possui permissão para consultar componentes.",
            "componentes": []
        }), 403

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        cargo_logado = session.get(
            "usuario_cargo",
            ""
        ).strip()

        escola_logada_id = session.get("escola_id")

        if escola_logada_id not in (None, ""):
            try:
                escola_logada_id = int(escola_logada_id)
            except (TypeError, ValueError):
                escola_logada_id = None

        cursor.execute("""
            SELECT
                id,
                escola_id,
                etapa,
                ano,
                nome,
                turno
            FROM turmas
            WHERE id = ?
            LIMIT 1
        """, (turma_id,))

        turma = cursor.fetchone()

        if turma is None:
            return jsonify({
                "sucesso": False,
                "mensagem": "Turma não encontrada.",
                "componentes": []
            }), 404

        if (
            cargo_logado == "Administrador da Instituição"
            and turma["escola_id"] != escola_logada_id
        ):
            return jsonify({
                "sucesso": False,
                "mensagem": "Você não possui acesso a esta turma.",
                "componentes": []
            }), 403

        etapa_turma = (turma["etapa"] or "").strip()

        if not etapa_turma:
            return jsonify({
                "sucesso": False,
                "mensagem": "A turma não possui uma etapa de ensino definida.",
                "componentes": []
            }), 400

        cursor.execute("""
            SELECT
                MIN(id) AS id,
                TRIM(nome) AS nome,
                TRIM(etapa_ensino) AS etapa_ensino
            FROM componentes_curriculares
            WHERE escola_id = ?
              AND ativo = 1
              AND TRIM(COALESCE(nome, '')) <> ''
              AND LOWER(TRIM(etapa_ensino)) = LOWER(TRIM(?))
            GROUP BY
                LOWER(TRIM(nome)),
                LOWER(TRIM(etapa_ensino))
            ORDER BY
                TRIM(nome) COLLATE NOCASE ASC
        """, (
            turma["escola_id"],
            etapa_turma
        ))

        componentes = [
            {
                "id": componente["id"],
                "nome": componente["nome"],
                "etapa_ensino": componente["etapa_ensino"]
            }
            for componente in cursor.fetchall()
        ]

        return jsonify({
            "sucesso": True,
            "turma": {
                "id": turma["id"],
                "etapa": turma["etapa"],
                "ano": turma["ano"],
                "nome": turma["nome"],
                "turno": turma["turno"]
            },
            "componentes": componentes
        })

    except sqlite3.Error as erro:
        print("ERRO AO BUSCAR COMPONENTES DA TURMA:", erro)

        return jsonify({
            "sucesso": False,
            "mensagem": "Não foi possível carregar os componentes curriculares.",
            "componentes": []
        }), 500

    finally:
        banco.close()


# =====================================================
# VÍNCULOS DO PROFESSOR
# =====================================================

@app.route(
    "/professor/<int:professor_id>/vinculos",
    methods=["GET", "POST"]
)
def professor_vinculos(professor_id):

    # Apenas administradores podem gerenciar vínculos
    if not cargo_permitido([
        "Administrador Geral",
        "Administrador da Instituição"
    ]):
        flash(
            "Você não possui permissão para acessar os vínculos.",
            "erro"
        )
        return redirect("/acesso_negado")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:

        cursor.execute("PRAGMA foreign_keys = ON")

        cargo_logado = session.get(
            "usuario_cargo",
            ""
        ).strip()

        escola_logada_id = session.get("escola_id")

        # =================================================
        # BUSCA O PROFESSOR
        # =================================================

        cursor.execute("""
            SELECT
                usuarios.id,
                usuarios.nome,
                usuarios.email,
                usuarios.escola_id,
                usuarios.ativo,
                cargos.nome AS cargo,
                escolas.nome_instituicao
            FROM usuarios

            INNER JOIN cargos
                ON cargos.id = usuarios.cargo_id

            LEFT JOIN escolas
                ON escolas.id = usuarios.escola_id

            WHERE usuarios.id = ?
            LIMIT 1
        """, (
            professor_id,
        ))

        professor = cursor.fetchone()

        if professor is None:
            flash(
                "Professor não encontrado.",
                "erro"
            )
            return redirect("/usuarios")

        # Confirma que o usuário selecionado é professor
        if professor["cargo"] != "Professor":
            flash(
                "O usuário selecionado não possui o cargo Professor.",
                "erro"
            )
            return redirect("/usuarios")

        # Confirma que o professor está ativo
        if professor["ativo"] != 1:
            flash(
                "O professor selecionado está inativo.",
                "erro"
            )
            return redirect("/usuarios")

        escola_professor_id = professor["escola_id"]

        if not escola_professor_id:
            flash(
                "O professor não está vinculado a uma instituição.",
                "erro"
            )
            return redirect("/usuarios")

        # Administrador da instituição só acessa professores
        # da própria instituição
        if (
            cargo_logado == "Administrador da Instituição"
            and escola_professor_id != escola_logada_id
        ):
            flash(
                "Você não possui permissão para acessar esse professor.",
                "erro"
            )
            return redirect("/usuarios")

        # =================================================
        # SALVA VÍNCULOS EM LOTE
        # =================================================

        if request.method == "POST":
            selecionados = request.form.getlist("vinculos")
            pares = set()
            for valor in selecionados:
                try:
                    turma_txt, componente_txt = valor.split(":", 1)
                    pares.add((int(turma_txt), int(componente_txt)))
                except (ValueError, TypeError):
                    continue

            # Carrega combinações válidas da instituição. O componente só pode
            # ser vinculado a turmas da mesma etapa de ensino.
            cursor.execute("""
                SELECT t.id AS turma_id, cc.id AS componente_id
                FROM turmas AS t
                INNER JOIN componentes_curriculares AS cc
                    ON cc.escola_id = t.escola_id
                   AND cc.ativo = 1
                   AND LOWER(TRIM(cc.etapa_ensino)) = LOWER(TRIM(t.etapa))
                WHERE t.escola_id = ?
            """, (escola_professor_id,))
            validos = {(int(r["turma_id"]), int(r["componente_id"])) for r in cursor.fetchall()}
            pares = pares.intersection(validos)

            # Sincroniza apenas os vínculos deste professor. Não altera turmas,
            # componentes, usuários ou qualquer outro cadastro.
            cursor.execute("SELECT turma_id, componente_id FROM professor_vinculos WHERE professor_id = ?", (professor_id,))
            atuais = {(int(r["turma_id"]), int(r["componente_id"])) for r in cursor.fetchall()}

            adicionar = pares - atuais
            remover = atuais - pares

            for turma_id, componente_id in adicionar:
                cursor.execute("""
                    INSERT OR IGNORE INTO professor_vinculos
                        (professor_id, turma_id, componente_id)
                    VALUES (?, ?, ?)
                """, (professor_id, turma_id, componente_id))

            for turma_id, componente_id in remover:
                cursor.execute("""
                    DELETE FROM professor_vinculos
                    WHERE professor_id = ? AND turma_id = ? AND componente_id = ?
                """, (professor_id, turma_id, componente_id))

            banco.commit()
            flash(
                f"Vínculos atualizados com sucesso: {len(adicionar)} adicionado(s) e {len(remover)} removido(s).",
                "success"
            )
            return redirect(f"/professor/{professor_id}/vinculos")

        # =================================================
        # LISTA AS TURMAS DA INSTITUIÇÃO
        # =================================================

        cursor.execute("""
            SELECT
                id,
                etapa,
                ano,
                nome,
                turno
            FROM turmas
            WHERE escola_id = ?
            ORDER BY
                etapa COLLATE NOCASE ASC,
                CAST(ano AS INTEGER) ASC,
                nome COLLATE NOCASE ASC,
                turno COLLATE NOCASE ASC
        """, (
            escola_professor_id,
        ))

        turmas = cursor.fetchall()

        # =================================================
        # COMPONENTES DISPONÍVEIS POR ETAPA
        # =================================================
        cursor.execute("""
            SELECT id, TRIM(nome) AS nome, TRIM(etapa_ensino) AS etapa_ensino
            FROM componentes_curriculares
            WHERE escola_id = ? AND ativo = 1
            ORDER BY etapa_ensino COLLATE NOCASE, nome COLLATE NOCASE
        """, (escola_professor_id,))
        componentes = cursor.fetchall()

        # =================================================
        # LISTA OS VÍNCULOS DO PROFESSOR
        # =================================================

        cursor.execute("""
            SELECT
                professor_vinculos.id,

                turmas.id AS turma_id,
                turmas.etapa,
                turmas.ano,
                turmas.nome AS turma_nome,
                turmas.turno,

                componentes_curriculares.id AS componente_id,
                componentes_curriculares.nome AS componente_nome

            FROM professor_vinculos

            INNER JOIN turmas
                ON turmas.id = professor_vinculos.turma_id

            INNER JOIN componentes_curriculares
                ON componentes_curriculares.id =
                   professor_vinculos.componente_id

            WHERE professor_vinculos.professor_id = ?

            ORDER BY
                turmas.etapa COLLATE NOCASE ASC,
                CAST(turmas.ano AS INTEGER) ASC,
                turmas.nome COLLATE NOCASE ASC,
                componentes_curriculares.nome COLLATE NOCASE ASC
        """, (
            professor_id,
        ))

        vinculos = cursor.fetchall()

        return render_template(
            "gestao/professor_vinculos.html",
            professor=professor,
            turmas=turmas,
            componentes=componentes,
            vinculos=vinculos
        )

    except sqlite3.Error as erro:

        banco.rollback()

        print(
            "ERRO NOS VÍNCULOS DO PROFESSOR:",
            erro
        )

        flash(
            f"Erro ao carregar os vínculos: {erro}",
            "erro"
        )

        return redirect("/usuarios")

    finally:
        banco.close()


# =====================================================
# EXCLUIR VÍNCULO DO PROFESSOR
# =====================================================

@app.route(
    "/professor/<int:professor_id>/vinculos/<int:vinculo_id>/excluir",
    methods=["POST"]
)
def excluir_professor_vinculo(
    professor_id,
    vinculo_id
):

    if not cargo_permitido([
        "Administrador Geral",
        "Administrador da Instituição"
    ]):
        flash(
            "Você não possui permissão para excluir vínculos.",
            "erro"
        )
        return redirect("/acesso_negado")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:

        cargo_logado = session.get(
            "usuario_cargo",
            ""
        ).strip()

        escola_logada_id = session.get("escola_id")

        # Busca o vínculo e a instituição do professor
        cursor.execute("""
            SELECT
                professor_vinculos.id,
                usuarios.escola_id

            FROM professor_vinculos

            INNER JOIN usuarios
                ON usuarios.id =
                   professor_vinculos.professor_id

            WHERE professor_vinculos.id = ?
              AND professor_vinculos.professor_id = ?

            LIMIT 1
        """, (
            vinculo_id,
            professor_id
        ))

        vinculo = cursor.fetchone()

        if vinculo is None:
            flash(
                "Vínculo não encontrado.",
                "erro"
            )
            return redirect(
                f"/professor/{professor_id}/vinculos"
            )

        if (
            cargo_logado == "Administrador da Instituição"
            and vinculo["escola_id"] != escola_logada_id
        ):
            flash(
                "Você não possui permissão para excluir esse vínculo.",
                "erro"
            )
            return redirect("/usuarios")

        cursor.execute("""
            DELETE FROM professor_vinculos
            WHERE id = ?
              AND professor_id = ?
        """, (
            vinculo_id,
            professor_id
        ))

        banco.commit()

        flash(
            "Vínculo excluído com sucesso.",
            "success"
        )

    except sqlite3.Error as erro:

        banco.rollback()

        print(
            "ERRO AO EXCLUIR VÍNCULO:",
            erro
        )

        flash(
            f"Erro ao excluir o vínculo: {erro}",
            "erro"
        )

    finally:
        banco.close()

    return redirect(
        f"/professor/{professor_id}/vinculos"
    )

# =====================================================
# COMPONENTES CURRICULARES
# =====================================================

@app.route("/componentes_curriculares", methods=["GET", "POST"])
def componentes_curriculares():

    if not cargo_permitido([
        "Administrador Geral",
        "Administrador da Instituição",
        "Coordenador"
    ]):
        flash(
            "Você não possui permissão para acessar os componentes curriculares.",
            "erro"
        )
        return redirect("/acesso_negado")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:

        usuario_cargo = session.get(
            "usuario_cargo",
            ""
        ).strip()

        usuario_id = session.get("usuario_id")
        escola_id = session.get("escola_id")

        # =================================================
        # RECUPERA A INSTITUIÇÃO DO USUÁRIO
        # =================================================

        if (
            usuario_cargo != "Administrador Geral"
            and not escola_id
            and usuario_id
        ):

            cursor.execute("""
                SELECT escola_id
                FROM usuarios
                WHERE id = ?
                LIMIT 1
            """, (
                usuario_id,
            ))

            usuario = cursor.fetchone()

            if usuario and usuario["escola_id"]:
                escola_id = usuario["escola_id"]
                session["escola_id"] = escola_id

        # =================================================
        # CADASTRA NOVO COMPONENTE
        # =================================================

        if request.method == "POST":

            nome = request.form.get(
                "nome",
                ""
            ).strip()

            escola_formulario = request.form.get(
                "escola_id",
                ""
            ).strip()

            if usuario_cargo == "Administrador Geral":
                escola_cadastro_id = escola_formulario
            else:
                escola_cadastro_id = escola_id

            if not escola_cadastro_id:
                flash(
                    "Selecione uma instituição.",
                    "erro"
                )
                return redirect("/componentes_curriculares")

            if not nome:
                flash(
                    "Informe o nome do componente curricular.",
                    "erro"
                )
                return redirect("/componentes_curriculares")

            cursor.execute("""
                SELECT id
                FROM escolas
                WHERE id = ?
                  AND COALESCE(status, 1) = 1
                LIMIT 1
            """, (
                escola_cadastro_id,
            ))

            escola = cursor.fetchone()

            if escola is None:
                flash(
                    "A instituição selecionada é inválida ou está inativa.",
                    "erro"
                )
                return redirect("/componentes_curriculares")

            cursor.execute("""
                SELECT id
                FROM componentes_curriculares
                WHERE escola_id = ?
                  AND LOWER(TRIM(nome)) = LOWER(TRIM(?))
                LIMIT 1
            """, (
                escola_cadastro_id,
                nome
            ))

            componente_existente = cursor.fetchone()

            if componente_existente:
                flash(
                    "Esse componente curricular já está cadastrado.",
                    "erro"
                )
                return redirect("/componentes_curriculares")

            cursor.execute("""
                INSERT INTO componentes_curriculares (
                    escola_id,
                    nome,
                    ativo,
                    criado_em
                )
                VALUES (?, ?, 1, ?)
            """, (
                escola_cadastro_id,
                nome,
                datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            ))

            banco.commit()

            flash(
                "Componente curricular cadastrado com sucesso.",
                "success"
            )

            return redirect("/componentes_curriculares")

        # =================================================
        # LISTA INSTITUIÇÕES PARA ADMINISTRADOR GERAL
        # =================================================

        escolas = []

        if usuario_cargo == "Administrador Geral":

            cursor.execute("""
                SELECT
                    id,
                    nome_instituicao
                FROM escolas
                WHERE COALESCE(status, 1) = 1
                ORDER BY nome_instituicao COLLATE NOCASE ASC
            """)

            escolas = cursor.fetchall()

        # =================================================
        # LISTA COMPONENTES
        # =================================================

        if usuario_cargo == "Administrador Geral":

            cursor.execute("""
                SELECT
                    componentes_curriculares.id,
                    componentes_curriculares.nome,
                    componentes_curriculares.ativo,
                    componentes_curriculares.escola_id,
                    escolas.nome_instituicao

                FROM componentes_curriculares

                INNER JOIN escolas
                    ON escolas.id =
                       componentes_curriculares.escola_id

                ORDER BY
                    escolas.nome_instituicao COLLATE NOCASE ASC,
                    componentes_curriculares.nome COLLATE NOCASE ASC
            """)

        else:

            if not escola_id:
                flash(
                    "Seu usuário não está vinculado a uma instituição.",
                    "erro"
                )
                return redirect("/")

            cursor.execute("""
                SELECT
                    componentes_curriculares.id,
                    componentes_curriculares.nome,
                    componentes_curriculares.ativo,
                    componentes_curriculares.escola_id,
                    escolas.nome_instituicao

                FROM componentes_curriculares

                INNER JOIN escolas
                    ON escolas.id =
                       componentes_curriculares.escola_id

                WHERE componentes_curriculares.escola_id = ?

                ORDER BY
                    componentes_curriculares.nome COLLATE NOCASE ASC
            """, (
                escola_id,
            ))

        componentes = cursor.fetchall()

        return render_template(
            "gestao/componentes_curriculares.html",
            componentes=componentes,
            escolas=escolas,
            usuario_cargo=usuario_cargo
        )

    except sqlite3.Error as erro:

        banco.rollback()

        print(
            "ERRO NOS COMPONENTES CURRICULARES:",
            erro
        )

        flash(
            f"Erro ao carregar os componentes curriculares: {erro}",
            "erro"
        )

        return redirect("/")

    finally:
        banco.close()


# =====================================================
# ALTERAR STATUS DO COMPONENTE
# =====================================================

@app.route(
    "/componentes_curriculares/<int:componente_id>/status",
    methods=["POST"]
)
def alterar_status_componente(componente_id):

    if not cargo_permitido([
        "Administrador Geral",
        "Administrador da Instituição",
        "Coordenador"
    ]):
        flash(
            "Você não possui permissão para alterar componentes curriculares.",
            "erro"
        )
        return redirect("/acesso_negado")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:

        usuario_cargo = session.get(
            "usuario_cargo",
            ""
        ).strip()

        escola_id = session.get("escola_id")

        cursor.execute("""
            SELECT
                id,
                escola_id,
                ativo
            FROM componentes_curriculares
            WHERE id = ?
            LIMIT 1
        """, (
            componente_id,
        ))

        componente = cursor.fetchone()

        if componente is None:
            flash(
                "Componente curricular não encontrado.",
                "erro"
            )
            return redirect("/componentes_curriculares")

        if (
            usuario_cargo != "Administrador Geral"
            and componente["escola_id"] != escola_id
        ):
            flash(
                "Você não possui permissão para alterar esse componente.",
                "erro"
            )
            return redirect("/componentes_curriculares")

        novo_status = 0 if componente["ativo"] == 1 else 1

        cursor.execute("""
            UPDATE componentes_curriculares
            SET ativo = ?
            WHERE id = ?
        """, (
            novo_status,
            componente_id
        ))

        banco.commit()

        if novo_status == 1:
            flash(
                "Componente curricular ativado com sucesso.",
                "success"
            )
        else:
            flash(
                "Componente curricular desativado com sucesso.",
                "success"
            )

    except sqlite3.Error as erro:

        banco.rollback()

        print(
            "ERRO AO ALTERAR COMPONENTE:",
            erro
        )

        flash(
            f"Erro ao alterar o componente curricular: {erro}",
            "erro"
        )

    finally:
        banco.close()

    return redirect("/componentes_curriculares")


# =====================================================
# EXCLUIR COMPONENTE
# =====================================================

# =========================================================
# LISTAR ANOS LETIVOS
# =========================================================

@app.route("/anos-letivos")
def anos_letivos():

    if not cargo_permitido([
        "Administrador Geral",
        "Administrador da Instituição"
    ]):
        return redirect("/acesso_negado")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    cargo = session.get("usuario_cargo", "").strip()
    escola_id = obter_escola_usuario()

    lista_anos = []
    escolas = []
    configuracao_global = None
    configuracao_instituicao = None
    proximos_anos = {}
    ano_sugerido = datetime.now().year + 1

    try:
        cursor.execute("""
            SELECT *
            FROM configuracao_ano_letivo_global
            WHERE id = 1
            LIMIT 1
        """)
        configuracao_global = cursor.fetchone()

        if cargo == "Administrador Geral":
            cursor.execute("""
                SELECT id, nome_instituicao
                FROM escolas
                WHERE COALESCE(status, 1) = 1
                ORDER BY nome_instituicao COLLATE NOCASE ASC
            """)
            escolas = cursor.fetchall()

            for escola in escolas:
                cursor.execute("""
                    SELECT COALESCE(MAX(ano), ?) + 1 AS proximo
                    FROM anos_letivos
                    WHERE escola_id = ?
                """, (datetime.now().year - 1, escola["id"]))
                proximo = cursor.fetchone()
                proximos_anos[str(escola["id"])] = (
                    proximo["proximo"] if proximo else datetime.now().year
                )

            cursor.execute("""
                SELECT
                    al.id,
                    al.escola_id,
                    al.ano,
                    al.data_inicio,
                    al.data_fim,
                    al.ativo,
                    al.encerrado,
                    al.criado_em,
                    e.nome_instituicao,

                    (SELECT COUNT(*) FROM turmas t
                     WHERE t.ano_letivo_id = al.id) AS total_turmas,

                    (
                        SELECT COUNT(*)
                        FROM (
                            SELECT am.aluno_id
                            FROM aluno_matriculas am
                            WHERE am.ano_letivo_id = al.id
                            UNION
                            SELECT a.id
                            FROM alunos a
                            WHERE a.ano_letivo_id = al.id
                        )
                    ) AS total_alunos,

                    (SELECT COUNT(*) FROM provas p
                     WHERE p.ano_letivo_id = al.id) AS total_provas,

                    (
                        SELECT COUNT(DISTINCT pv.professor_id)
                        FROM professor_vinculos pv
                        INNER JOIN turmas t_prof ON t_prof.id = pv.turma_id
                        WHERE t_prof.ano_letivo_id = al.id
                    ) AS total_professores

                FROM anos_letivos al
                INNER JOIN escolas e ON e.id = al.escola_id
                ORDER BY e.nome_instituicao COLLATE NOCASE ASC, al.ano DESC
            """)
            lista_anos = cursor.fetchall()

        else:
            if not escola_id:
                flash("Não foi possível identificar sua instituição.", "erro")
                return redirect("/")

            atualizar_ano_letivo_na_sessao(escola_id)

            cursor.execute("""
                SELECT *
                FROM configuracao_ano_letivo_instituicao
                WHERE escola_id = ?
                LIMIT 1
            """, (escola_id,))
            configuracao_instituicao = cursor.fetchone()

            cursor.execute("""
                SELECT COALESCE(MAX(ano), ?) + 1 AS proximo
                FROM anos_letivos
                WHERE escola_id = ?
            """, (datetime.now().year - 1, escola_id))
            proximo = cursor.fetchone()
            ano_sugerido = proximo["proximo"] if proximo else datetime.now().year

            cursor.execute("""
                SELECT
                    al.id,
                    al.escola_id,
                    al.ano,
                    al.data_inicio,
                    al.data_fim,
                    al.ativo,
                    al.encerrado,
                    al.criado_em,
                    e.nome_instituicao,

                    (SELECT COUNT(*) FROM turmas t
                     WHERE t.ano_letivo_id = al.id) AS total_turmas,

                    (
                        SELECT COUNT(*)
                        FROM (
                            SELECT am.aluno_id
                            FROM aluno_matriculas am
                            WHERE am.ano_letivo_id = al.id
                            UNION
                            SELECT a.id
                            FROM alunos a
                            WHERE a.ano_letivo_id = al.id
                        )
                    ) AS total_alunos,

                    (SELECT COUNT(*) FROM provas p
                     WHERE p.ano_letivo_id = al.id) AS total_provas,

                    (
                        SELECT COUNT(DISTINCT pv.professor_id)
                        FROM professor_vinculos pv
                        INNER JOIN turmas t_prof ON t_prof.id = pv.turma_id
                        WHERE t_prof.ano_letivo_id = al.id
                    ) AS total_professores

                FROM anos_letivos al
                INNER JOIN escolas e ON e.id = al.escola_id
                WHERE al.escola_id = ?
                ORDER BY al.ano DESC
            """, (escola_id,))
            lista_anos = cursor.fetchall()

        return render_template(
            "gestao/anos_letivos.html",
            anos_letivos=lista_anos,
            escolas=escolas,
            cargo=cargo,
            configuracao_global=configuracao_global,
            configuracao_instituicao=configuracao_instituicao,
            proximos_anos=proximos_anos,
            ano_sugerido=ano_sugerido
        )

    except sqlite3.Error as erro:
        import traceback
        traceback.print_exc()
        flash(f"Erro ao carregar os anos letivos: {erro}", "erro")

        return render_template(
            "gestao/anos_letivos.html",
            anos_letivos=[],
            escolas=escolas,
            cargo=cargo,
            configuracao_global=configuracao_global,
            configuracao_instituicao=configuracao_instituicao,
            proximos_anos=proximos_anos,
            ano_sugerido=ano_sugerido
        )

    finally:
        banco.close()


# =========================================================
# ABRIR NOVO ANO LETIVO
# =========================================================

@app.route("/anos-letivos/abrir", methods=["POST"])
def abrir_ano_letivo():

    if not cargo_permitido([
        "Administrador Geral",
        "Administrador da Instituição"
    ]):
        return redirect("/acesso_negado")

    cargo = session.get("usuario_cargo", "").strip()

    if cargo == "Administrador Geral":
        escola_id = request.form.get("escola_id", "").strip()
    else:
        escola_id = obter_escola_usuario()

    ano = request.form.get("ano", "").strip()
    data_inicio = request.form.get("data_inicio", "").strip()
    data_fim = request.form.get("data_fim", "").strip()

    copiar_turmas = request.form.get("copiar_turmas") == "1"
    copiar_vinculos = (
        copiar_turmas
        and request.form.get("copiar_vinculos") == "1"
    )

    if not escola_id:
        flash("Selecione uma instituição.", "erro")
        return redirect("/anos-letivos")

    try:
        escola_id = int(escola_id)
    except (TypeError, ValueError):
        flash("A instituição selecionada é inválida.", "erro")
        return redirect("/anos-letivos")

    try:
        ano = int(ano)
    except (TypeError, ValueError):
        flash("Informe um ano letivo válido.", "erro")
        return redirect("/anos-letivos")

    if ano < 2000 or ano > 2100:
        flash("Informe um ano entre 2000 e 2100.", "erro")
        return redirect("/anos-letivos")

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:

        cursor.execute("""
            SELECT id, nome_instituicao
            FROM escolas
            WHERE id = ?
              AND COALESCE(status, 1) = 1
            LIMIT 1
        """, (
            escola_id,
        ))

        escola = cursor.fetchone()

        if escola is None:
            flash(
                "A instituição não existe ou está inativa.",
                "erro"
            )
            return redirect("/anos-letivos")

        cursor.execute("""
            SELECT id
            FROM anos_letivos
            WHERE escola_id = ?
              AND ano = ?
            LIMIT 1
        """, (
            escola_id,
            ano
        ))

        if cursor.fetchone():
            flash(
                f"O ano letivo {ano} já está cadastrado.",
                "erro"
            )
            return redirect("/anos-letivos")

        # Guarda o ano anterior para copiar as turmas.
        cursor.execute("""
            SELECT id, ano
            FROM anos_letivos
            WHERE escola_id = ?
              AND ativo = 1
              AND encerrado = 0
            ORDER BY ano DESC
            LIMIT 1
        """, (
            escola_id,
        ))

        ano_anterior = cursor.fetchone()

        # Encerra o ano ativo atual.
        cursor.execute("""
            UPDATE anos_letivos
            SET
                ativo = 0,
                encerrado = 1
            WHERE escola_id = ?
              AND ativo = 1
        """, (
            escola_id,
        ))

        # Cria o novo ano.
        cursor.execute("""
            INSERT INTO anos_letivos (
                escola_id,
                ano,
                data_inicio,
                data_fim,
                ativo,
                encerrado
            )
            VALUES (?, ?, ?, ?, 1, 0)
        """, (
            escola_id,
            ano,
            data_inicio or None,
            data_fim or None
        ))

        novo_ano_id = cursor.lastrowid

        # Copia as turmas do ano anterior, quando solicitado.
        total_turmas_copiadas = 0
        total_vinculos_copiados = 0
        mapa_turmas = {}

        if copiar_turmas and ano_anterior:
            cursor.execute("""
                SELECT id, nome, etapa, ano, turno
                FROM turmas
                WHERE escola_id = ?
                  AND ano_letivo_id = ?
                ORDER BY id
            """, (escola_id, ano_anterior["id"]))

            turmas_anteriores = cursor.fetchall()

            for turma in turmas_anteriores:
                cursor.execute("""
                    INSERT INTO turmas (
                        nome,
                        etapa,
                        ano,
                        turno,
                        escola_id,
                        ano_letivo,
                        ano_letivo_id
                    )
                    VALUES (?, ?, ?, ?, ?, ?, ?)
                """, (
                    turma["nome"],
                    turma["etapa"],
                    turma["ano"],
                    turma["turno"],
                    escola_id,
                    ano,
                    novo_ano_id
                ))

                nova_turma_id = cursor.lastrowid
                mapa_turmas[turma["id"]] = nova_turma_id
                total_turmas_copiadas += 1

            if copiar_vinculos and mapa_turmas:
                for turma_antiga_id, turma_nova_id in mapa_turmas.items():
                    cursor.execute("""
                        SELECT professor_id, componente_id
                        FROM professor_vinculos
                        WHERE turma_id = ?
                    """, (turma_antiga_id,))

                    for vinculo in cursor.fetchall():
                        cursor.execute("""
                            INSERT OR IGNORE INTO professor_vinculos (
                                professor_id,
                                turma_id,
                                componente_id
                            )
                            VALUES (?, ?, ?)
                        """, (
                            vinculo["professor_id"],
                            turma_nova_id,
                            vinculo["componente_id"]
                        ))
                        if cursor.rowcount:
                            total_vinculos_copiados += 1

        # Mantém o campo antigo sincronizado.
        cursor.execute("""
            UPDATE escolas
            SET ano_letivo = ?
            WHERE id = ?
        """, (
            ano,
            escola_id
        ))

        banco.commit()

        # Remove eventual consulta a um ano antigo.
        session.pop("ano_letivo_selecionado_id", None)

        if cargo != "Administrador Geral":
            session["ano_letivo_id"] = novo_ano_id
            session["ano_letivo"] = ano
            session["ano_letivo_visualizado"] = ano
        else:
            session["ano_letivo_visualizado"] = ano
            session["ano_letivo"] = ano

        mensagem = f"Ano letivo {ano} aberto com sucesso."

        if copiar_turmas:
            mensagem += f" {total_turmas_copiadas} turma(s) foram copiadas."

        if copiar_vinculos:
            mensagem += (
                f" {total_vinculos_copiados} vínculo(s) de professor "
                "foram copiados."
            )

        cursor.execute("""
            INSERT INTO ano_letivo_auditoria (
                escola_id,
                ano_letivo_id,
                usuario_id,
                acao,
                detalhes
            )
            VALUES (?, ?, ?, ?, ?)
        """, (
            escola_id,
            novo_ano_id,
            session.get("usuario_id"),
            "ABRIR_ANO",
            f"Ano {ano} aberto manualmente."
        ))
        banco.commit()

        flash(mensagem, "success")

        return redirect("/anos-letivos")

    except sqlite3.Error as erro:

        banco.rollback()

        import traceback
        traceback.print_exc()

        print("ERRO AO ABRIR ANO LETIVO:", erro)

        flash(
            f"Erro ao abrir o ano letivo: {erro}",
            "erro"
        )

        return redirect("/anos-letivos")

    finally:
        banco.close()


# =========================================================
# CONSULTAR UM ANO LETIVO
# =========================================================

@app.route(
    "/anos-letivos/<int:ano_letivo_id>/selecionar",
    methods=["POST"]
)
def selecionar_ano_letivo(ano_letivo_id):

    if not cargo_permitido([
        "Administrador Geral",
        "Administrador da Instituição"
    ]):
        return redirect("/acesso_negado")

    cargo = session.get("usuario_cargo", "").strip()
    escola_id = obter_escola_usuario()

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:

        if cargo == "Administrador Geral":

            cursor.execute("""
                SELECT id, escola_id, ano
                FROM anos_letivos
                WHERE id = ?
                LIMIT 1
            """, (
                ano_letivo_id,
            ))

        else:

            cursor.execute("""
                SELECT id, escola_id, ano
                FROM anos_letivos
                WHERE id = ?
                  AND escola_id = ?
                LIMIT 1
            """, (
                ano_letivo_id,
                escola_id
            ))

        ano_letivo = cursor.fetchone()

        if ano_letivo is None:
            flash(
                "Ano letivo não encontrado ou sem permissão.",
                "erro"
            )
            return redirect("/anos-letivos")

        if cargo == "Administrador Geral":
            session["ano_letivo_visualizado"] = ano_letivo["ano"]
            session["ano_letivo"] = ano_letivo["ano"]
            session.pop("ano_letivo_id", None)
            session.pop("ano_letivo_selecionado_id", None)
        else:
            session["ano_letivo_selecionado_id"] = ano_letivo["id"]
            session["ano_letivo_id"] = ano_letivo["id"]
            session["ano_letivo"] = ano_letivo["ano"]
            session["ano_letivo_visualizado"] = ano_letivo["ano"]

        flash(
            f"A plataforma está consultando o ano "
            f"{ano_letivo['ano']}.",
            "success"
        )

        return redirect("/anos-letivos")

    except sqlite3.Error as erro:

        print("ERRO AO SELECIONAR ANO LETIVO:", erro)

        flash(
            f"Erro ao selecionar o ano letivo: {erro}",
            "erro"
        )

        return redirect("/anos-letivos")

    finally:
        banco.close()


# =========================================================
# VOLTAR AO ANO ATIVO
# =========================================================

@app.route(
    "/anos-letivos/usar-ativo",
    methods=["POST"]
)
def usar_ano_letivo_ativo():

    if not cargo_permitido([
        "Administrador Geral",
        "Administrador da Instituição"
    ]):
        return redirect("/acesso_negado")

    session.pop("ano_letivo_selecionado_id", None)
    session.pop("ano_letivo_visualizado", None)

    escola_id = obter_escola_usuario()

    if escola_id:
        atualizar_ano_letivo_na_sessao(escola_id)
    else:
        session.pop("ano_letivo_id", None)
        session.pop("ano_letivo", None)

    flash(
        "A plataforma voltou a utilizar o ano letivo ativo.",
        "success"
    )

    return redirect("/anos-letivos")


# =========================================================
# SALVAR AGENDAMENTO GLOBAL
# =========================================================

@app.route("/anos-letivos/agendamento-global", methods=["POST"])
def salvar_agendamento_global():

    if not cargo_permitido(["Administrador Geral"]):
        return redirect("/acesso_negado")

    banco = conectar_banco()
    cursor = banco.cursor()

    try:
        ano = int(request.form.get("ano", "").strip())
        data_execucao = request.form.get("data_execucao", "").strip()
        data_inicio = request.form.get("data_inicio", "").strip() or None
        data_fim = request.form.get("data_fim", "").strip() or None

        if ano < 2000 or ano > 2100:
            raise ValueError("Informe um ano entre 2000 e 2100.")

        if not data_execucao:
            raise ValueError("Informe a data de execução.")

        if data_inicio and data_fim and data_fim < data_inicio:
            raise ValueError("A data final não pode ser anterior à inicial.")

        cursor.execute("""
            INSERT INTO configuracao_ano_letivo_global (
                id, ativo, ano, data_execucao, data_inicio, data_fim,
                copiar_turmas, copiar_vinculos, encerrar_anterior,
                executado, atualizado_em
            )
            VALUES (1, ?, ?, ?, ?, ?, ?, ?, ?, 0, CURRENT_TIMESTAMP)
            ON CONFLICT(id) DO UPDATE SET
                ativo = excluded.ativo,
                ano = excluded.ano,
                data_execucao = excluded.data_execucao,
                data_inicio = excluded.data_inicio,
                data_fim = excluded.data_fim,
                copiar_turmas = excluded.copiar_turmas,
                copiar_vinculos = excluded.copiar_vinculos,
                encerrar_anterior = excluded.encerrar_anterior,
                executado = 0,
                atualizado_em = CURRENT_TIMESTAMP
        """, (
            1 if request.form.get("ativo") == "1" else 0,
            ano,
            data_execucao,
            data_inicio,
            data_fim,
            1 if request.form.get("copiar_turmas") == "1" else 0,
            1 if request.form.get("copiar_vinculos") == "1" else 0,
            1 if request.form.get("encerrar_anterior") == "1" else 0
        ))

        banco.commit()
        flash("Agendamento global salvo com sucesso.", "success")

    except (ValueError, sqlite3.Error) as erro:
        banco.rollback()
        flash(f"Erro ao salvar o agendamento global: {erro}", "erro")

    finally:
        banco.close()

    return redirect("/anos-letivos")


# =========================================================
# SALVAR CONFIGURAÇÃO DA INSTITUIÇÃO
# =========================================================

@app.route("/anos-letivos/agendamento-instituicao", methods=["POST"])
def salvar_agendamento_instituicao():

    if not cargo_permitido([
        "Administrador Geral",
        "Administrador da Instituição"
    ]):
        return redirect("/acesso_negado")

    cargo = session.get("usuario_cargo", "").strip()

    if cargo == "Administrador Geral":
        escola_id = request.form.get("escola_id", "").strip()
    else:
        escola_id = obter_escola_usuario()

    try:
        escola_id = int(escola_id)
    except (TypeError, ValueError):
        flash("Selecione uma instituição válida.", "erro")
        return redirect("/anos-letivos")

    modo = request.form.get("modo", "global").strip().lower()

    if modo not in ("global", "proprio", "manual"):
        modo = "global"

    ano_texto = request.form.get("ano", "").strip()
    data_execucao = request.form.get("data_execucao", "").strip() or None
    data_inicio = request.form.get("data_inicio", "").strip() or None
    data_fim = request.form.get("data_fim", "").strip() or None

    try:
        ano = int(ano_texto) if ano_texto else None

        if ano is not None and (ano < 2000 or ano > 2100):
            raise ValueError("Informe um ano entre 2000 e 2100.")

        if modo == "proprio" and (ano is None or not data_execucao):
            raise ValueError(
                "No agendamento próprio, informe o ano e a data de execução."
            )

        if data_inicio and data_fim and data_fim < data_inicio:
            raise ValueError("A data final não pode ser anterior à inicial.")

        banco = conectar_banco()
        cursor = banco.cursor()

        cursor.execute("""
            INSERT INTO configuracao_ano_letivo_instituicao (
                escola_id, modo, ativo, ano, data_execucao,
                data_inicio, data_fim, copiar_turmas,
                copiar_vinculos, encerrar_anterior,
                executado, atualizado_em
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 0, CURRENT_TIMESTAMP)
            ON CONFLICT(escola_id) DO UPDATE SET
                modo = excluded.modo,
                ativo = excluded.ativo,
                ano = excluded.ano,
                data_execucao = excluded.data_execucao,
                data_inicio = excluded.data_inicio,
                data_fim = excluded.data_fim,
                copiar_turmas = excluded.copiar_turmas,
                copiar_vinculos = excluded.copiar_vinculos,
                encerrar_anterior = excluded.encerrar_anterior,
                executado = 0,
                atualizado_em = CURRENT_TIMESTAMP
        """, (
            escola_id,
            modo,
            1 if request.form.get("ativo") == "1" else 0,
            ano,
            data_execucao,
            data_inicio,
            data_fim,
            1 if request.form.get("copiar_turmas") == "1" else 0,
            1 if request.form.get("copiar_vinculos") == "1" else 0,
            1 if request.form.get("encerrar_anterior") == "1" else 0
        ))

        banco.commit()
        banco.close()

        flash("Configuração da instituição salva com sucesso.", "success")

    except (ValueError, sqlite3.Error) as erro:
        try:
            banco.rollback()
            banco.close()
        except Exception:
            pass
        flash(f"Erro ao salvar a configuração: {erro}", "erro")

    return redirect("/anos-letivos")


# =========================================================
# REABRIR ANO LETIVO
# =========================================================

@app.route(
    "/anos-letivos/<int:ano_letivo_id>/reabrir",
    methods=["POST"]
)
def reabrir_ano_letivo(ano_letivo_id):

    if not cargo_permitido([
        "Administrador Geral",
        "Administrador da Instituição"
    ]):
        return redirect("/acesso_negado")

    cargo = session.get("usuario_cargo", "").strip()
    escola_usuario = obter_escola_usuario()
    modo = request.form.get("modo", "edicao").strip().lower()

    if modo not in ("edicao", "ativar"):
        modo = "edicao"

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        if cargo == "Administrador Geral":
            cursor.execute("""
                SELECT id, escola_id, ano
                FROM anos_letivos
                WHERE id = ?
                LIMIT 1
            """, (ano_letivo_id,))
        else:
            cursor.execute("""
                SELECT id, escola_id, ano
                FROM anos_letivos
                WHERE id = ?
                  AND escola_id = ?
                LIMIT 1
            """, (ano_letivo_id, escola_usuario))

        ano_letivo = cursor.fetchone()

        if not ano_letivo:
            flash("Ano letivo não encontrado ou sem permissão.", "erro")
            return redirect("/anos-letivos")

        if modo == "ativar":
            cursor.execute("""
                UPDATE anos_letivos
                SET ativo = 0,
                    encerrado = 1
                WHERE escola_id = ?
                  AND id <> ?
                  AND ativo = 1
            """, (ano_letivo["escola_id"], ano_letivo_id))

            cursor.execute("""
                UPDATE anos_letivos
                SET ativo = 1,
                    encerrado = 0
                WHERE id = ?
            """, (ano_letivo_id,))

            cursor.execute("""
                UPDATE escolas
                SET ano_letivo = ?
                WHERE id = ?
            """, (ano_letivo["ano"], ano_letivo["escola_id"]))

            acao = "REABRIR_E_ATIVAR"
            mensagem = (
                f"O ano letivo {ano_letivo['ano']} foi reaberto e ativado."
            )

            if cargo != "Administrador Geral":
                session["ano_letivo_selecionado_id"] = ano_letivo_id
                session["ano_letivo_id"] = ano_letivo_id
            else:
                session.pop("ano_letivo_id", None)
                session.pop("ano_letivo_selecionado_id", None)

            session["ano_letivo"] = ano_letivo["ano"]
            session["ano_letivo_visualizado"] = ano_letivo["ano"]

        else:
            cursor.execute("""
                UPDATE anos_letivos
                SET ativo = 0,
                    encerrado = 0
                WHERE id = ?
            """, (ano_letivo_id,))

            acao = "REABRIR_EDICAO"
            mensagem = (
                f"O ano letivo {ano_letivo['ano']} foi reaberto para edição."
            )

        cursor.execute("""
            INSERT INTO ano_letivo_auditoria (
                escola_id, ano_letivo_id, usuario_id, acao, detalhes
            )
            VALUES (?, ?, ?, ?, ?)
        """, (
            ano_letivo["escola_id"],
            ano_letivo_id,
            session.get("usuario_id"),
            acao,
            mensagem
        ))

        banco.commit()
        flash(mensagem, "success")

    except sqlite3.Error as erro:
        banco.rollback()
        flash(f"Erro ao reabrir o ano letivo: {erro}", "erro")

    finally:
        banco.close()

    return redirect("/anos-letivos")

# =========================================================
# RELATÓRIOS — VISÃO GERAL
# Cole este bloco no app.py antes do:
# 
# =========================================================


def _recalcular_notas_aplicacoes_por_peso(
    cursor,
    prova_id=None,
    aplicacao_id=None,
    aluno_id=None
):
    """
    Recalcula notas objetivas usando prova_questoes.peso.

    Também atualiza a nota final quando não há discursivas ou quando todas
    as discursivas do aluno já foram corrigidas. Isso corrige registros
    antigos que foram gravados com a fórmula proporcional por quantidade.
    """
    condicoes = ["aa.objetiva_corrigida = 1"]
    parametros = []

    if prova_id is not None:
        condicoes.append("ap.prova_id = ?")
        parametros.append(prova_id)

    if aplicacao_id is not None:
        condicoes.append("aa.aplicacao_id = ?")
        parametros.append(aplicacao_id)

    if aluno_id is not None:
        condicoes.append("aa.aluno_id = ?")
        parametros.append(aluno_id)

    cursor.execute(f"""
        SELECT
            aa.aplicacao_id,
            aa.aluno_id,
            ap.prova_id,
            ROUND(
                COALESCE(
                    SUM(
                        CASE
                            WHEN COALESCE(aro.acertou, 0) = 1
                                THEN COALESCE(pq.peso, 0)
                            ELSE 0
                        END
                    ),
                    0
                ),
                2
            ) AS nota_objetiva_calculada,
            COUNT(aro.id) AS total_respostas_objetivas,
            COALESCE(aa.nota_discursiva, 0) AS nota_discursiva
        FROM aplicacao_alunos AS aa
        INNER JOIN aplicacoes AS ap
            ON ap.id = aa.aplicacao_id
        INNER JOIN aplicacao_respostas_objetivas AS aro
            ON aro.aplicacao_id = aa.aplicacao_id
           AND aro.aluno_id = aa.aluno_id
        LEFT JOIN prova_questoes AS pq
            ON pq.prova_id = ap.prova_id
           AND pq.questao_id = aro.questao_id
        WHERE {" AND ".join(condicoes)}
        GROUP BY
            aa.aplicacao_id,
            aa.aluno_id,
            ap.prova_id,
            aa.nota_discursiva
        HAVING COUNT(aro.id) > 0
    """, parametros)

    registros = cursor.fetchall()

    for registro in registros:
        aplicacao_atual = registro["aplicacao_id"]
        aluno_atual = registro["aluno_id"]
        nota_objetiva = round(
            float(registro["nota_objetiva_calculada"] or 0),
            2
        )
        nota_discursiva = round(
            float(registro["nota_discursiva"] or 0),
            2
        )

        cursor.execute("""
            SELECT
                COUNT(*) AS total,
                SUM(
                    CASE
                        WHEN corrigida = 1 THEN 1
                        ELSE 0
                    END
                ) AS corrigidas
            FROM respostas_discursivas_aplicacao
            WHERE aplicacao_id = ?
              AND aluno_id = ?
              AND TRIM(COALESCE(imagem_resposta, '')) <> ''
        """, (aplicacao_atual, aluno_atual))

        discursivas = cursor.fetchone()
        total_discursivas = int(discursivas["total"] or 0)
        corrigidas = int(discursivas["corrigidas"] or 0)

        if total_discursivas == 0:
            nota_final = nota_objetiva
        elif corrigidas == total_discursivas:
            nota_final = round(
                nota_objetiva + nota_discursiva,
                2
            )
        else:
            nota_final = None

        cursor.execute("""
            UPDATE aplicacao_alunos
            SET nota_objetiva = ?,
                nota_final = ?
            WHERE aplicacao_id = ?
              AND aluno_id = ?
        """, (
            nota_objetiva,
            nota_final,
            aplicacao_atual,
            aluno_atual
        ))


@app.route("/relatorios")
def relatorios():
    """
    Visão geral dos resultados.

    Consolida:
    1. resultados das novas aplicações, armazenados em aplicacao_alunos;
    2. resultados legados da tabela resultados, quando ainda não existir
       um resultado equivalente nas novas aplicações.
    """
    if not permissao_modulo("Relatórios"):
        return redirect("/acesso_negado")

    # Garante que as tabelas e colunas das aplicações existam antes
    # de montar as consultas consolidadas.
    _garantir_tabelas_aplicacoes()

    contexto = obter_contexto_plataforma()
    cargo = contexto["cargo"]
    usuario_id = contexto["usuario_id"]
    escola_usuario_id = contexto["escola_id"]
    ano_contexto = contexto["ano"]
    ano_letivo_id = contexto["ano_letivo_id"]

    escola_filtro = request.args.get("escola_id", type=int)
    turma_filtro = request.args.get("turma_id", type=int)
    disciplina_filtro = request.args.get("disciplina", "").strip()
    prova_filtro = request.args.get("prova_id", type=int)

    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        _recalcular_notas_aplicacoes_por_peso(cursor)
        banco.commit()

        # ---------------------------------------------------------
        # Restrições de acesso e filtros
        # ---------------------------------------------------------
        condicoes = ["1 = 1"]
        parametros = []

        if cargo == "Administrador Geral":
            if escola_filtro:
                condicoes.append("COALESCE(p.escola_id, t.escola_id) = ?")
                parametros.append(escola_filtro)

            # Para o Administrador Geral, o ano é comparado pelo número,
            # pois cada instituição possui um ano_letivo_id diferente.
            if ano_contexto is not None:
                condicoes.append("al.ano = ?")
                parametros.append(ano_contexto)
        else:
            if not escola_usuario_id:
                flash(
                    "Não foi possível identificar sua instituição.",
                    "erro"
                )
                return redirect("/")

            condicoes.append("COALESCE(p.escola_id, t.escola_id) = ?")
            parametros.append(escola_usuario_id)

            if ano_letivo_id:
                condicoes.append("p.ano_letivo_id = ?")
                parametros.append(ano_letivo_id)

            if cargo == "Professor":
                condicoes.append("p.professor_id = ?")
                parametros.append(usuario_id)

        if turma_filtro:
            condicoes.append("p.turma_id = ?")
            parametros.append(turma_filtro)

        if disciplina_filtro:
            condicoes.append(
                "LOWER(TRIM(p.disciplina)) = LOWER(TRIM(?))"
            )
            parametros.append(disciplina_filtro)

        if prova_filtro:
            condicoes.append("p.id = ?")
            parametros.append(prova_filtro)

        where_sql = " AND ".join(condicoes)

        base_joins = """
            FROM provas AS p
            INNER JOIN turmas AS t
                ON t.id = p.turma_id
            LEFT JOIN escolas AS e
                ON e.id = COALESCE(p.escola_id, t.escola_id)
            LEFT JOIN anos_letivos AS al
                ON al.id = p.ano_letivo_id
        """

        # Um aluno entra no relatório quando possui nota final ou quando
        # concluiu uma aplicação somente objetiva.
        resultados_cte = """
            WITH resultados_consolidados AS (
                SELECT
                    ap.prova_id,
                    aa.aluno_id,
                    ap.id AS aplicacao_id,
                    CASE
                        WHEN aa.nota_final IS NOT NULL
                            THEN aa.nota_final
                        WHEN aa.objetiva_corrigida = 1
                         AND COALESCE(aa.discursiva_pendente, 0) = 0
                            THEN aa.nota_objetiva
                        ELSE NULL
                    END AS nota
                FROM aplicacoes AS ap
                INNER JOIN aplicacao_alunos AS aa
                    ON aa.aplicacao_id = ap.id
                WHERE
                    CASE
                        WHEN aa.nota_final IS NOT NULL
                            THEN aa.nota_final
                        WHEN aa.objetiva_corrigida = 1
                         AND COALESCE(aa.discursiva_pendente, 0) = 0
                            THEN aa.nota_objetiva
                        ELSE NULL
                    END IS NOT NULL
                  AND LOWER(TRIM(COALESCE(aa.status, ''))) NOT IN (
                      'ausente',
                      'faltou'
                  )

                UNION ALL

                SELECT
                    rl.prova_id,
                    rl.aluno_id,
                    NULL AS aplicacao_id,
                    rl.nota
                FROM resultados AS rl
                WHERE rl.nota IS NOT NULL
                  AND NOT EXISTS (
                      SELECT 1
                      FROM aplicacoes AS ap2
                      INNER JOIN aplicacao_alunos AS aa2
                          ON aa2.aplicacao_id = ap2.id
                      WHERE ap2.prova_id = rl.prova_id
                        AND aa2.aluno_id = rl.aluno_id
                        AND (
                            aa2.nota_final IS NOT NULL
                            OR (
                                aa2.objetiva_corrigida = 1
                                AND COALESCE(
                                    aa2.discursiva_pendente,
                                    0
                                ) = 0
                                AND aa2.nota_objetiva IS NOT NULL
                            )
                        )
                  )
            )
        """

        respostas_cte = """
            WITH respostas_consolidadas AS (
                SELECT
                    ap.prova_id,
                    aro.aluno_id,
                    aro.numero_questao,
                    aro.acertou
                FROM aplicacoes AS ap
                INNER JOIN aplicacao_respostas_objetivas AS aro
                    ON aro.aplicacao_id = ap.id

                UNION ALL

                SELECT
                    ra.prova_id,
                    ra.aluno_id,
                    ra.numero_questao,
                    ra.acertou
                FROM respostas_alunos AS ra
                WHERE NOT EXISTS (
                    SELECT 1
                    FROM aplicacoes AS ap2
                    INNER JOIN aplicacao_respostas_objetivas AS aro2
                        ON aro2.aplicacao_id = ap2.id
                    WHERE ap2.prova_id = ra.prova_id
                      AND aro2.aluno_id = ra.aluno_id
                      AND aro2.numero_questao = ra.numero_questao
                )
            )
        """

        # ---------------------------------------------------------
        # Opções dos filtros
        # ---------------------------------------------------------
        if cargo == "Administrador Geral":
            cursor.execute("""
                SELECT id, nome_instituicao
                FROM escolas
                WHERE COALESCE(status, 1) = 1
                ORDER BY nome_instituicao COLLATE NOCASE
            """)
            escolas = cursor.fetchall()
        else:
            cursor.execute("""
                SELECT id, nome_instituicao
                FROM escolas
                WHERE id = ?
            """, (escola_usuario_id,))
            escolas = cursor.fetchall()

        opcoes_condicoes = ["1 = 1"]
        opcoes_parametros = []

        if cargo == "Administrador Geral":
            if escola_filtro:
                opcoes_condicoes.append("t.escola_id = ?")
                opcoes_parametros.append(escola_filtro)

            if ano_contexto is not None:
                opcoes_condicoes.append("al.ano = ?")
                opcoes_parametros.append(ano_contexto)
        else:
            opcoes_condicoes.append("t.escola_id = ?")
            opcoes_parametros.append(escola_usuario_id)

            if ano_letivo_id:
                opcoes_condicoes.append("t.ano_letivo_id = ?")
                opcoes_parametros.append(ano_letivo_id)

        if cargo == "Professor":
            opcoes_condicoes.append("""
                EXISTS (
                    SELECT 1
                    FROM professor_vinculos AS pv
                    WHERE pv.professor_id = ?
                      AND pv.turma_id = t.id
                )
            """)
            opcoes_parametros.append(usuario_id)

        cursor.execute(f"""
            SELECT DISTINCT
                t.id,
                t.nome,
                t.ano,
                t.turno
            FROM turmas AS t
            LEFT JOIN anos_letivos AS al
                ON al.id = t.ano_letivo_id
            WHERE {" AND ".join(opcoes_condicoes)}
            ORDER BY
                t.ano COLLATE NOCASE,
                t.nome COLLATE NOCASE
        """, opcoes_parametros)
        turmas = cursor.fetchall()

        cursor.execute(f"""
            SELECT DISTINCT TRIM(p.disciplina) AS disciplina
            {base_joins}
            WHERE {where_sql}
              AND TRIM(COALESCE(p.disciplina, '')) <> ''
            ORDER BY disciplina COLLATE NOCASE
        """, parametros)
        disciplinas = [
            linha["disciplina"]
            for linha in cursor.fetchall()
        ]

        # Lista de provas sem aplicar o próprio filtro de prova.
        condicoes_provas = [
            condicao
            for condicao in condicoes
            if condicao != "p.id = ?"
        ]
        parametros_provas = list(parametros)

        if prova_filtro:
            parametros_provas.pop()

        cursor.execute(f"""
            SELECT
                p.id,
                p.nome,
                t.nome AS turma_nome,
                p.disciplina
            {base_joins}
            WHERE {" AND ".join(condicoes_provas)}
            ORDER BY
                COALESCE(
                    p.data_aplicacao,
                    p.data_geracao,
                    p.id
                ) DESC,
                p.id DESC
        """, parametros_provas)
        provas_filtro = cursor.fetchall()

        # ---------------------------------------------------------
        # Indicadores gerais
        # ---------------------------------------------------------
        cursor.execute(
            resultados_cte + f"""
            SELECT
                COUNT(DISTINCT p.id) AS total_avaliacoes,
                COUNT(DISTINCT rc.aluno_id) AS alunos_avaliados,
                ROUND(AVG(rc.nota), 1) AS media_geral,
                ROUND(MAX(rc.nota), 1) AS maior_nota,
                ROUND(MIN(rc.nota), 1) AS menor_nota,
                SUM(
                    CASE
                        WHEN p.media_ativa = 1
                         AND p.media_aprovacao IS NOT NULL
                         AND rc.nota >= p.media_aprovacao
                        THEN 1
                        ELSE 0
                    END
                ) AS aprovados_com_media,
                SUM(
                    CASE
                        WHEN p.media_ativa = 1
                         AND p.media_aprovacao IS NOT NULL
                         AND rc.nota IS NOT NULL
                        THEN 1
                        ELSE 0
                    END
                ) AS resultados_com_media
            {base_joins}
            LEFT JOIN resultados_consolidados AS rc
                ON rc.prova_id = p.id
            WHERE {where_sql}
            """,
            parametros
        )
        indicadores = cursor.fetchone()

        total_com_media = indicadores["resultados_com_media"] or 0
        aprovados_com_media = indicadores["aprovados_com_media"] or 0

        taxa_aprovacao = (
            round(
                (aprovados_com_media / total_com_media) * 100,
                1
            )
            if total_com_media > 0
            else 0
        )

        # ---------------------------------------------------------
        # Desempenho por turma
        # ---------------------------------------------------------
        cursor.execute(
            resultados_cte + f"""
            SELECT
                t.id AS turma_id,
                t.nome AS turma_nome,
                t.ano AS turma_ano,
                COUNT(DISTINCT p.id) AS total_avaliacoes,
                COUNT(DISTINCT rc.aluno_id) AS alunos_avaliados,
                ROUND(AVG(rc.nota), 1) AS media,
                ROUND(MAX(rc.nota), 1) AS maior_nota,
                ROUND(MIN(rc.nota), 1) AS menor_nota
            {base_joins}
            LEFT JOIN resultados_consolidados AS rc
                ON rc.prova_id = p.id
            WHERE {where_sql}
            GROUP BY
                t.id,
                t.nome,
                t.ano
            HAVING COUNT(rc.aluno_id) > 0
            ORDER BY
                media DESC,
                t.nome COLLATE NOCASE
            LIMIT 12
            """,
            parametros
        )
        desempenho_turmas = cursor.fetchall()

        # ---------------------------------------------------------
        # Avaliações recentes
        # ---------------------------------------------------------
        cursor.execute(
            resultados_cte + f"""
            SELECT
                p.id,
                p.nome,
                p.disciplina,
                COALESCE(
                    MAX(ap_recente.data_aplicacao),
                    p.data_aplicacao
                ) AS data_aplicacao,
                p.status,
                p.media_ativa,
                p.media_aprovacao,
                t.nome AS turma_nome,
                e.nome_instituicao,
                COUNT(rc.aluno_id) AS total_resultados,
                ROUND(AVG(rc.nota), 1) AS media,
                ROUND(MAX(rc.nota), 1) AS maior_nota,
                ROUND(MIN(rc.nota), 1) AS menor_nota
            {base_joins}
            LEFT JOIN resultados_consolidados AS rc
                ON rc.prova_id = p.id
            LEFT JOIN aplicacoes AS ap_recente
                ON ap_recente.prova_id = p.id
            WHERE {where_sql}
            GROUP BY
                p.id,
                p.nome,
                p.disciplina,
                p.data_aplicacao,
                p.status,
                p.media_ativa,
                p.media_aprovacao,
                t.nome,
                e.nome_instituicao
            ORDER BY
                COALESCE(
                    MAX(ap_recente.data_aplicacao),
                    p.data_aplicacao,
                    p.data_geracao,
                    p.id
                ) DESC,
                p.id DESC
            LIMIT 10
            """,
            parametros
        )
        avaliacoes = cursor.fetchall()

        # ---------------------------------------------------------
        # Questões objetivas com maior índice de erro
        # ---------------------------------------------------------
        cursor.execute(
            respostas_cte + f"""
            SELECT
                rq.prova_id,
                rq.numero_questao,
                p.nome AS prova_nome,
                p.disciplina,
                COUNT(*) AS respondentes,
                SUM(
                    CASE WHEN rq.acertou = 1 THEN 1 ELSE 0 END
                ) AS acertos,
                SUM(
                    CASE WHEN rq.acertou = 0 THEN 1 ELSE 0 END
                ) AS erros,
                ROUND(
                    100.0
                    * SUM(
                        CASE WHEN rq.acertou = 0 THEN 1 ELSE 0 END
                    )
                    / NULLIF(COUNT(*), 0),
                    1
                ) AS percentual_erros
            {base_joins}
            INNER JOIN respostas_consolidadas AS rq
                ON rq.prova_id = p.id
            WHERE {where_sql}
            GROUP BY
                rq.prova_id,
                rq.numero_questao,
                p.nome,
                p.disciplina
            HAVING COUNT(*) > 0
            ORDER BY
                percentual_erros DESC,
                respondentes DESC
            LIMIT 8
            """,
            parametros
        )
        questoes_criticas = cursor.fetchall()

        # ---------------------------------------------------------
        # Ranking de alunos
        # ---------------------------------------------------------
        cursor.execute(
            resultados_cte + f"""
            SELECT
                alu.id AS aluno_id,
                alu.nome AS aluno_nome,
                t.nome AS turma_nome,
                COUNT(rc.nota) AS avaliacoes_realizadas,
                ROUND(AVG(rc.nota), 1) AS media
            {base_joins}
            INNER JOIN resultados_consolidados AS rc
                ON rc.prova_id = p.id
            INNER JOIN alunos AS alu
                ON alu.id = rc.aluno_id
            WHERE {where_sql}
            GROUP BY
                alu.id,
                alu.nome,
                t.nome
            HAVING COUNT(rc.nota) > 0
            ORDER BY
                media DESC,
                avaliacoes_realizadas DESC,
                alu.nome COLLATE NOCASE
            LIMIT 5
            """,
            parametros
        )
        ranking_alunos = cursor.fetchall()

        return render_template(
            "relatorios/index.html",
            escolas=escolas,
            turmas=turmas,
            disciplinas=disciplinas,
            provas_filtro=provas_filtro,
            indicadores=indicadores,
            taxa_aprovacao=taxa_aprovacao,
            desempenho_turmas=desempenho_turmas,
            avaliacoes=avaliacoes,
            questoes_criticas=questoes_criticas,
            ranking_alunos=ranking_alunos,
            filtros={
                "escola_id": escola_filtro,
                "turma_id": turma_filtro,
                "disciplina": disciplina_filtro,
                "prova_id": prova_filtro
            }
        )

    except sqlite3.Error as erro:
        import traceback
        traceback.print_exc()
        flash(
            f"Não foi possível carregar os relatórios: {erro}",
            "erro"
        )
        return redirect("/")

    finally:
        banco.close()


# =========================================================
# MÓDULO DE APLICAÇÕES DE AVALIAÇÕES — ARK EDUS
# =========================================================

def _garantir_tabelas_aplicacoes():
    banco = conectar_banco()
    cursor = banco.cursor()
    cursor.execute("PRAGMA foreign_keys = ON")
    cursor.executescript("""
        CREATE TABLE IF NOT EXISTS aplicacoes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            prova_id INTEGER NOT NULL,
            turma_id INTEGER NOT NULL,
            escola_id INTEGER,
            ano_letivo_id INTEGER,
            nome TEXT,
            data_aplicacao TEXT,
            quantidade_modelos INTEGER NOT NULL DEFAULT 1,
            status TEXT NOT NULL DEFAULT 'Agendada',
            observacoes TEXT,
            criado_por INTEGER,
            criado_em TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (prova_id) REFERENCES provas(id) ON DELETE CASCADE,
            FOREIGN KEY (turma_id) REFERENCES turmas(id) ON DELETE CASCADE,
            FOREIGN KEY (escola_id) REFERENCES escolas(id) ON DELETE CASCADE,
            FOREIGN KEY (ano_letivo_id) REFERENCES anos_letivos(id) ON DELETE SET NULL,
            FOREIGN KEY (criado_por) REFERENCES usuarios(id) ON DELETE SET NULL
        );

        CREATE TABLE IF NOT EXISTS aplicacao_alunos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            aplicacao_id INTEGER NOT NULL,
            aluno_id INTEGER NOT NULL,
            modelo INTEGER NOT NULL DEFAULT 1,
            status TEXT NOT NULL DEFAULT 'Aguardando aplicação',
            objetiva_corrigida INTEGER NOT NULL DEFAULT 0,
            discursiva_pendente INTEGER NOT NULL DEFAULT 0,
            acertos_objetivos INTEGER NOT NULL DEFAULT 0,
            total_objetivas INTEGER NOT NULL DEFAULT 0,
            nota_objetiva REAL,
            criado_em TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (aplicacao_id) REFERENCES aplicacoes(id) ON DELETE CASCADE,
            FOREIGN KEY (aluno_id) REFERENCES alunos(id) ON DELETE CASCADE,
            UNIQUE (aplicacao_id, aluno_id)
        );

        CREATE TABLE IF NOT EXISTS aplicacao_importacoes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            aplicacao_id INTEGER NOT NULL,
            aluno_id INTEGER,
            modelo INTEGER,
            nome_arquivo TEXT NOT NULL,
            caminho_arquivo TEXT NOT NULL,
            status TEXT NOT NULL DEFAULT 'Processado',
            mensagem TEXT,
            importado_em TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (aplicacao_id) REFERENCES aplicacoes(id) ON DELETE CASCADE,
            FOREIGN KEY (aluno_id) REFERENCES alunos(id) ON DELETE SET NULL
        );

        CREATE TABLE IF NOT EXISTS respostas_discursivas_aplicacao (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            aplicacao_id INTEGER NOT NULL,
            aluno_id INTEGER NOT NULL,
            questao_id INTEGER NOT NULL,
            numero_exibicao INTEGER NOT NULL,
            imagem_resposta TEXT,
            nota REAL,
            comentario TEXT,
            corrigida INTEGER NOT NULL DEFAULT 0,
            corrigido_por INTEGER,
            corrigido_em TEXT,
            FOREIGN KEY (aplicacao_id) REFERENCES aplicacoes(id) ON DELETE CASCADE,
            FOREIGN KEY (aluno_id) REFERENCES alunos(id) ON DELETE CASCADE,
            FOREIGN KEY (questao_id) REFERENCES questoes(id) ON DELETE CASCADE,
            FOREIGN KEY (corrigido_por) REFERENCES usuarios(id) ON DELETE SET NULL,
            UNIQUE (aplicacao_id, aluno_id, questao_id)
        );

        CREATE TABLE IF NOT EXISTS aplicacao_respostas_objetivas (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            aplicacao_id INTEGER NOT NULL,
            aluno_id INTEGER NOT NULL,
            numero_questao INTEGER NOT NULL,
            questao_id INTEGER NOT NULL,
            modelo INTEGER NOT NULL,
            resposta TEXT,
            situacao TEXT NOT NULL DEFAULT 'em_branco',
            resposta_correta TEXT,
            acertou INTEGER NOT NULL DEFAULT 0,
            origem TEXT NOT NULL DEFAULT 'automatica',
            atualizado_em TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (aplicacao_id) REFERENCES aplicacoes(id) ON DELETE CASCADE,
            FOREIGN KEY (aluno_id) REFERENCES alunos(id) ON DELETE CASCADE,
            FOREIGN KEY (questao_id) REFERENCES questoes(id) ON DELETE CASCADE,
            UNIQUE (aplicacao_id, aluno_id, numero_questao)
        );

        CREATE INDEX IF NOT EXISTS idx_aplicacoes_prova ON aplicacoes(prova_id);
        CREATE INDEX IF NOT EXISTS idx_aplicacao_alunos_app ON aplicacao_alunos(aplicacao_id);
        CREATE INDEX IF NOT EXISTS idx_app_resp_objetivas
            ON aplicacao_respostas_objetivas(aplicacao_id, aluno_id);
    """)

    def garantir_coluna_importacao(nome, definicao):
        cursor.execute("PRAGMA table_info(aplicacao_importacoes)")
        existentes = {linha[1] for linha in cursor.fetchall()}
        if nome not in existentes:
            cursor.execute(
                f"ALTER TABLE aplicacao_importacoes ADD COLUMN {nome} {definicao}"
            )

    garantir_coluna_importacao("qr_texto", "TEXT")
    garantir_coluna_importacao("tipo_folha", "TEXT")
    garantir_coluna_importacao("revisado", "INTEGER NOT NULL DEFAULT 0")
    garantir_coluna_importacao("revisado_por", "INTEGER")
    garantir_coluna_importacao("revisado_em", "TEXT")

    def garantir_coluna_resposta_objetiva(nome, definicao):
        cursor.execute("PRAGMA table_info(aplicacao_respostas_objetivas)")
        existentes = {linha[1] for linha in cursor.fetchall()}
        if nome not in existentes:
            cursor.execute(
                f"ALTER TABLE aplicacao_respostas_objetivas ADD COLUMN {nome} {definicao}"
            )

    garantir_coluna_resposta_objetiva(
        "anulada", "INTEGER NOT NULL DEFAULT 0"
    )

    def garantir_coluna_aplicacao_aluno(nome, definicao):
        cursor.execute("PRAGMA table_info(aplicacao_alunos)")
        existentes = {linha[1] for linha in cursor.fetchall()}
        if nome not in existentes:
            cursor.execute(
                f"ALTER TABLE aplicacao_alunos ADD COLUMN {nome} {definicao}"
            )

    garantir_coluna_aplicacao_aluno("nota_discursiva", "REAL")
    garantir_coluna_aplicacao_aluno("nota_final", "REAL")

    # Garante compatibilidade com provas antigas. Uma questão anulada recebe
    # crédito automático para todos os alunos e nunca exige revisão do cartão.
    cursor.execute("PRAGMA table_info(prova_questoes)")
    colunas_pq = {linha[1] for linha in cursor.fetchall()}
    if "anulada" not in colunas_pq:
        cursor.execute(
            "ALTER TABLE prova_questoes "
            "ADD COLUMN anulada INTEGER NOT NULL DEFAULT 0"
        )

    banco.commit()
    banco.close()


def _sincronizar_alunos_aplicacao(cursor, aplicacao_id):
    """Garante que os alunos matriculados na turma existam em aplicacao_alunos.

    A rotina é idempotente: preserva vínculos já existentes e cria apenas os
    que estiverem ausentes. Isso corrige aplicações antigas/migradas sem
    apagar resultados já registrados.
    """
    cursor.execute("""
        SELECT a.turma_id, a.ano_letivo_id, a.quantidade_modelos
        FROM aplicacoes a
        WHERE a.id = ?
    """, (aplicacao_id,))
    aplicacao = cursor.fetchone()
    if not aplicacao:
        return 0

    cursor.execute("""
        SELECT DISTINCT al.id, al.nome
        FROM alunos al
        LEFT JOIN aluno_matriculas am
          ON am.aluno_id = al.id
         AND am.turma_id = ?
         AND (? IS NULL OR am.ano_letivo_id = ?)
         AND COALESCE(am.situacao, 'Cursando') = 'Cursando'
        WHERE am.id IS NOT NULL OR al.turma_id = ?
        ORDER BY al.nome COLLATE NOCASE, al.id
    """, (
        aplicacao["turma_id"],
        aplicacao["ano_letivo_id"],
        aplicacao["ano_letivo_id"],
        aplicacao["turma_id"],
    ))
    alunos_turma = cursor.fetchall()
    quantidade_modelos = max(1, int(aplicacao["quantidade_modelos"] or 1))

    criados = 0
    for indice, aluno in enumerate(alunos_turma):
        modelo_padrao = (indice % quantidade_modelos) + 1
        cursor.execute("""
            INSERT OR IGNORE INTO aplicacao_alunos
                (aplicacao_id, aluno_id, modelo)
            VALUES (?, ?, ?)
        """, (aplicacao_id, aluno["id"], modelo_padrao))
        criados += max(0, cursor.rowcount)

    return criados


def _garantir_aluno_qr_na_aplicacao(cursor, aplicacao_id, aluno_id, modelo=1):
    """Garante o vínculo do aluno identificado pelo QR sem aceitar aluno de outra turma.

    É um segundo nível de proteção para aplicações antigas: se a sincronização em
    lote não encontrou o vínculo por diferenças de migração, valida diretamente
    se o aluno pertence à turma da aplicação e cria somente o vínculo ausente.
    """
    cursor.execute("""
        SELECT a.turma_id, a.ano_letivo_id
        FROM aplicacoes a
        WHERE a.id = ?
    """, (aplicacao_id,))
    aplicacao = cursor.fetchone()
    if not aplicacao:
        return False

    cursor.execute("""
        SELECT 1
        FROM aplicacao_alunos
        WHERE aplicacao_id = ? AND aluno_id = ?
    """, (aplicacao_id, aluno_id))
    if cursor.fetchone():
        return True

    pertence = False
    # Compatibilidade com o vínculo legado alunos.turma_id.
    try:
        cursor.execute("""
            SELECT 1 FROM alunos
            WHERE id = ? AND turma_id = ?
        """, (aluno_id, aplicacao["turma_id"]))
        pertence = cursor.fetchone() is not None
    except sqlite3.Error:
        pertence = False

    # Compatibilidade com a estrutura atual aluno_matriculas.
    if not pertence:
        try:
            cursor.execute("""
                SELECT 1
                FROM aluno_matriculas am
                WHERE am.aluno_id = ?
                  AND am.turma_id = ?
                  AND (? IS NULL OR am.ano_letivo_id = ?)
                  AND LOWER(TRIM(COALESCE(am.situacao, 'Cursando'))) = 'cursando'
                LIMIT 1
            """, (
                aluno_id,
                aplicacao["turma_id"],
                aplicacao["ano_letivo_id"],
                aplicacao["ano_letivo_id"],
            ))
            pertence = cursor.fetchone() is not None
        except sqlite3.Error:
            pertence = False

    if not pertence:
        return False

    cursor.execute("""
        INSERT OR IGNORE INTO aplicacao_alunos
            (aplicacao_id, aluno_id, modelo)
        VALUES (?, ?, ?)
    """, (aplicacao_id, aluno_id, max(1, int(modelo or 1))))
    return True


def _questoes_modelo_aplicacao(cursor, aplicacao_id, modelo):
    """Monta um modelo estável preservando texto, HTML e imagens das questões."""
    import random

    cursor.execute("""
        SELECT a.prova_id
        FROM aplicacoes a
        WHERE a.id = ?
    """, (aplicacao_id,))
    app_reg = cursor.fetchone()
    if not app_reg:
        return []

    cursor.execute("""
        SELECT
            q.*,
            COALESCE(NULLIF(pq.ordem, 0), pq.id) AS ordem_original,
            COALESCE(pq.peso, 0) AS peso,
            COALESCE(pq.anulada, 0) AS anulada
        FROM prova_questoes pq
        INNER JOIN questoes q ON q.id = pq.questao_id
        WHERE pq.prova_id = ?
        ORDER BY COALESCE(NULLIF(pq.ordem, 0), pq.id), pq.id
    """, (app_reg["prova_id"],))

    questoes = [dict(registro) for registro in cursor.fetchall()]
    modelo = int(modelo)

    questoes_objetivas = [
        questao for questao in questoes
        if not _tipo_discursivo_aplicacao(questao.get("tipo_questao"))
    ]
    questoes_discursivas = [
        questao for questao in questoes
        if _tipo_discursivo_aplicacao(questao.get("tipo_questao"))
    ]

    gerador_questoes = random.Random(
        f"ARKEDUS:{aplicacao_id}:MODELO:{modelo}:QUESTOES"
    )
    gerador_questoes.shuffle(questoes_objetivas)
    questoes = questoes_objetivas + questoes_discursivas

    for questao in questoes:
        if _tipo_discursivo_aplicacao(questao.get("tipo_questao")):
            questao["alternativas"] = []
            continue

        # Fonte principal: JSON moderno, que preserva imagens das alternativas.
        try:
            alternativas_origem = json.loads(questao.get("alternativas_json") or "[]")
            if not isinstance(alternativas_origem, list):
                alternativas_origem = []
        except (TypeError, ValueError, json.JSONDecodeError):
            alternativas_origem = []

        # Compatibilidade com questões antigas que só possuem A/B/C/D legados.
        if not alternativas_origem:
            alternativas_origem = []
            for indice, letra in enumerate(["A", "B", "C", "D"]):
                texto_alt = questao.get(f"alternativa_{letra.lower()}") or ""
                if str(texto_alt).strip():
                    alternativas_origem.append({
                        "letra": letra,
                        "texto": texto_alt,
                        "imagem": "",
                    })

        corretas_originais = set()
        try:
            lista_corretas = json.loads(questao.get("respostas_corretas") or "[]")
            if isinstance(lista_corretas, list):
                corretas_originais = {str(x).strip().upper() for x in lista_corretas if str(x).strip()}
        except (TypeError, ValueError, json.JSONDecodeError):
            pass
        if not corretas_originais and (questao.get("correta") or "").strip():
            corretas_originais.add((questao.get("correta") or "").strip().upper())

        alternativas = []
        for indice, alt in enumerate(alternativas_origem):
            if not isinstance(alt, dict):
                continue
            letra_original = str(alt.get("letra") or chr(65 + indice)).strip().upper()
            texto_alt = str(alt.get("texto") or "").strip()
            imagem_alt = str(alt.get("imagem") or "").strip()
            if texto_alt or imagem_alt:
                alternativas.append({
                    "letra_original": letra_original,
                    "texto": texto_alt,
                    "imagem": imagem_alt,
                })

        gerador_alternativas = random.Random(
            f"ARKEDUS:{aplicacao_id}:MODELO:{modelo}:QUESTAO:{questao['id']}:ALTERNATIVAS"
        )
        gerador_alternativas.shuffle(alternativas)

        alternativas_exibicao = []
        novas_corretas = []
        for indice, alternativa in enumerate(alternativas):
            nova_letra = chr(65 + indice)
            alternativas_exibicao.append({
                "letra": nova_letra,
                "texto": alternativa["texto"],
                "imagem": alternativa.get("imagem") or "",
            })
            if alternativa["letra_original"] in corretas_originais:
                novas_corretas.append(nova_letra)

        questao["alternativas"] = alternativas_exibicao
        questao["correta_original"] = next(iter(corretas_originais), "")
        questao["correta"] = novas_corretas[0] if novas_corretas else ""
        questao["respostas_corretas_modelo"] = novas_corretas

        # Mantém compatibilidade com trechos antigos que ainda leem A/B/C/D.
        for indice, letra in enumerate(["A", "B", "C", "D"]):
            questao[f"alternativa_{letra.lower()}"] = (
                alternativas_exibicao[indice]["texto"]
                if indice < len(alternativas_exibicao)
                else ""
            )

    return questoes


def _tipo_discursivo_aplicacao(tipo):
    valor = (tipo or "").strip().lower().replace("-", "_")
    return valor in {"discursiva", "dissertativa", "resposta_aberta", "resposta aberta"}



def _sincronizar_status_prova(cursor, prova_id):
    """Recalcula o status da prova sem alterar avaliações em rascunho."""
    cursor.execute("""
        SELECT LOWER(TRIM(COALESCE(status, 'rascunho'))) AS status
        FROM provas
        WHERE id = ?
    """, (prova_id,))
    prova = cursor.fetchone()

    if not prova:
        return None

    status_atual = prova["status"] or "rascunho"
    if status_atual == "rascunho":
        return "rascunho"

    cursor.execute("""
        SELECT
            COUNT(*) AS total_aplicacoes,
            SUM(CASE WHEN LOWER(TRIM(COALESCE(status, ''))) = 'corrigida'
                     THEN 1 ELSE 0 END) AS corrigidas,
            SUM(CASE WHEN LOWER(TRIM(COALESCE(status, ''))) = 'em correção'
                          OR LOWER(TRIM(COALESCE(status, ''))) = 'aguardando correção'
                     THEN 1 ELSE 0 END) AS em_correcao
        FROM aplicacoes
        WHERE prova_id = ?
    """, (prova_id,))
    resumo = cursor.fetchone()

    total = int(resumo["total_aplicacoes"] or 0)
    corrigidas = int(resumo["corrigidas"] or 0)
    em_correcao = int(resumo["em_correcao"] or 0)

    if total > 0 and corrigidas == total:
        novo_status = "finalizada"
    elif em_correcao > 0:
        novo_status = "em_correcao"
    else:
        novo_status = "agendada"

    cursor.execute("""
        UPDATE provas
        SET status = ?, atualizado_em = ?
        WHERE id = ?
    """, (
        novo_status,
        datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        prova_id
    ))
    return novo_status


def _sincronizar_status_aplicacao(cursor, aplicacao_id):
    """
    Atualiza a aplicação e, em seguida, o status geral da prova.

    Agendada: nenhum cartão foi importado e nenhuma ausência foi marcada.
    Em correção: o processamento começou, mas há aluno presente pendente.
    Corrigida: todos os presentes foram resolvidos e os demais estão ausentes.
    """
    cursor.execute("""
        SELECT a.prova_id,
               SUM(CASE WHEN LOWER(TRIM(COALESCE(q.tipo_questao, ''))) IN (
                    'discursiva', 'dissertativa', 'resposta_aberta',
                    'resposta aberta'
               ) THEN 1 ELSE 0 END) AS total_discursivas,
               SUM(CASE WHEN LOWER(TRIM(COALESCE(q.tipo_questao, ''))) NOT IN (
                    'discursiva', 'dissertativa', 'resposta_aberta',
                    'resposta aberta'
               ) THEN 1 ELSE 0 END) AS total_objetivas
        FROM aplicacoes a
        LEFT JOIN prova_questoes pq ON pq.prova_id = a.prova_id
        LEFT JOIN questoes q ON q.id = pq.questao_id
        WHERE a.id = ?
        GROUP BY a.id, a.prova_id
    """, (aplicacao_id,))
    dados = cursor.fetchone()

    if not dados:
        return None

    prova_id = dados["prova_id"]
    possui_objetivas = int(dados["total_objetivas"] or 0) > 0
    possui_discursivas = int(dados["total_discursivas"] or 0) > 0

    cursor.execute("""
        SELECT
            COUNT(*) AS importacoes,
            EXISTS (
                SELECT 1 FROM aplicacao_alunos aa
                WHERE aa.aplicacao_id = ?
                  AND LOWER(TRIM(COALESCE(aa.status, ''))) = 'ausente'
            ) AS possui_ausencia
        FROM aplicacao_importacoes ai
        WHERE ai.aplicacao_id = ?
          AND LOWER(TRIM(COALESCE(ai.status, ''))) <> 'erro'
    """, (aplicacao_id, aplicacao_id))
    atividade = cursor.fetchone()
    iniciou = (
        int(atividade["importacoes"] or 0) > 0
        or int(atividade["possui_ausencia"] or 0) == 1
    )

    if not iniciou:
        novo_status = "Agendada"
    else:
        cursor.execute("""
            SELECT COUNT(*) AS pendentes
            FROM aplicacao_alunos aa
            WHERE aa.aplicacao_id = ?
              AND LOWER(TRIM(COALESCE(aa.status, ''))) <> 'ausente'
              AND (
                    (? = 1 AND COALESCE(aa.objetiva_corrigida, 0) = 0)
                 OR (? = 1 AND COALESCE(aa.discursiva_pendente, 1) = 1)
              )
        """, (
            aplicacao_id,
            1 if possui_objetivas else 0,
            1 if possui_discursivas else 0
        ))
        pendentes = int(cursor.fetchone()["pendentes"] or 0)
        novo_status = "Aguardando correção" if pendentes else "Corrigida"

    cursor.execute("""
        UPDATE aplicacoes
        SET status = ?
        WHERE id = ?
    """, (novo_status, aplicacao_id))

    _sincronizar_status_prova(cursor, prova_id)
    return novo_status


@app.route("/provas/<int:prova_id>/aplicacoes")
def aplicacoes_prova(prova_id):
    if not permissao_modulo("Provas"):
        return redirect("/acesso_negado")

    _garantir_tabelas_aplicacoes()
    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()
    try:
        if not _pode_gerenciar_prova(cursor, prova_id, exigir_edicao=False):
            return _redirecionar_acesso_negado_prova()

        cursor.execute("""
            SELECT p.*, t.nome AS turma_nome, t.escola_id AS turma_escola_id,
                   t.ano_letivo_id, e.nome_instituicao,
                   COALESCE(pr.nome, 'Não informado') AS professor_nome
            FROM provas p
            INNER JOIN turmas t ON t.id = p.turma_id
            LEFT JOIN escolas e ON e.id = COALESCE(p.escola_id, t.escola_id)
            LEFT JOIN professores pr ON pr.id = p.professor_id
            WHERE p.id = ?
        """, (prova_id,))
        prova = cursor.fetchone()
        if not prova:
            flash("Avaliação não encontrada.", "erro")
            return redirect("/provas")

        cursor.execute("""
            SELECT id
            FROM aplicacoes
            WHERE prova_id = ?
        """, (prova_id,))
        for item_aplicacao in cursor.fetchall():
            _sincronizar_status_aplicacao(
                cursor,
                item_aplicacao["id"]
            )
        banco.commit()

        cursor.execute("""
            SELECT a.*,
                   COUNT(aa.id) AS total_alunos,
                   SUM(
                       CASE
                           WHEN LOWER(TRIM(COALESCE(aa.status, ''))) <> 'ausente'
                            AND aa.objetiva_corrigida = 1
                           THEN 1 ELSE 0
                       END
                   ) AS objetivas_corrigidas,
                   SUM(
                       CASE
                           WHEN LOWER(TRIM(COALESCE(aa.status, ''))) <> 'ausente'
                            AND aa.discursiva_pendente = 1
                           THEN 1 ELSE 0
                       END
                   ) AS discursivas_pendentes,
                   SUM(
                       CASE
                           WHEN LOWER(TRIM(COALESCE(aa.status, ''))) = 'ausente'
                           THEN 1 ELSE 0
                       END
                   ) AS total_ausentes,
                   EXISTS (
                       SELECT 1
                       FROM prova_questoes pq
                       INNER JOIN questoes q ON q.id = pq.questao_id
                       WHERE pq.prova_id = a.prova_id
                         AND LOWER(TRIM(COALESCE(q.tipo_questao, ''))) IN (
                             'discursiva', 'dissertativa',
                             'resposta_aberta', 'resposta aberta'
                         )
                   ) AS possui_discursivas
            FROM aplicacoes a
            LEFT JOIN aplicacao_alunos aa ON aa.aplicacao_id = a.id
            WHERE a.prova_id = ?
            GROUP BY a.id
            ORDER BY a.criado_em DESC, a.id DESC
        """, (prova_id,))
        aplicacoes = cursor.fetchall()

        return render_template("aplicacoes/index.html", prova=prova, aplicacoes=aplicacoes)
    finally:
        banco.close()


@app.route("/provas/<int:prova_id>/aplicacoes/nova", methods=["GET", "POST"])
def nova_aplicacao(prova_id):
    if not permissao_modulo("Provas"):
        return redirect("/acesso_negado")

    _garantir_tabelas_aplicacoes()
    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()
    try:
        if not _pode_gerenciar_prova(
            cursor,
            prova_id,
            exigir_edicao=True,
            permitir_finalizada=True
        ):
            return _redirecionar_acesso_negado_prova()

        cursor.execute("""
            SELECT p.*, t.nome AS turma_nome, t.escola_id AS turma_escola_id,
                   t.ano_letivo_id, e.nome_instituicao
            FROM provas p
            INNER JOIN turmas t ON t.id = p.turma_id
            LEFT JOIN escolas e ON e.id = COALESCE(p.escola_id, t.escola_id)
            WHERE p.id = ?
        """, (prova_id,))
        prova = cursor.fetchone()
        if not prova:
            flash("Avaliação não encontrada.", "erro")
            return redirect("/provas")

        cursor.execute("""
            SELECT DISTINCT a.id, a.nome, COALESCE(a.matricula, '') AS matricula
            FROM alunos a
            LEFT JOIN aluno_matriculas am
              ON am.aluno_id = a.id
             AND am.turma_id = ?
             AND (? IS NULL OR am.ano_letivo_id = ?)
             AND COALESCE(am.situacao, 'Cursando') = 'Cursando'
            WHERE am.id IS NOT NULL OR a.turma_id = ?
            ORDER BY a.nome COLLATE NOCASE
        """, (prova["turma_id"], prova["ano_letivo_id"], prova["ano_letivo_id"], prova["turma_id"]))
        alunos = cursor.fetchall()

        if request.method == "POST":
            quantidade_modelos = request.form.get("quantidade_modelos", "1")
            try:
                quantidade_modelos = max(1, min(4, int(quantidade_modelos)))
            except ValueError:
                quantidade_modelos = 1

            modo_alunos = request.form.get("modo_alunos", "turma")
            ids_validos = {int(a["id"]) for a in alunos}
            if modo_alunos == "selecionados":
                selecionados = []
                for valor in request.form.getlist("alunos"):
                    try:
                        aluno_id = int(valor)
                        if aluno_id in ids_validos:
                            selecionados.append(aluno_id)
                    except ValueError:
                        pass
            else:
                selecionados = sorted(ids_validos)

            if not selecionados:
                flash("Selecione pelo menos um aluno para criar a aplicação.", "erro")
                return render_template("aplicacoes/nova.html", prova=prova, alunos=alunos)

            cursor.execute("""
                INSERT INTO aplicacoes (
                    prova_id, turma_id, escola_id, ano_letivo_id, nome,
                    data_aplicacao, quantidade_modelos, status,
                    observacoes, criado_por
                ) VALUES (?, ?, ?, ?, ?, ?, ?, 'Agendada', ?, ?)
            """, (
                prova_id, prova["turma_id"], prova["escola_id"] or prova["turma_escola_id"],
                prova["ano_letivo_id"], request.form.get("nome") or f"Aplicação - {prova['nome']}",
                request.form.get("data_aplicacao") or prova["data_aplicacao"], quantidade_modelos,
                request.form.get("observacoes", "").strip(), session.get("usuario_id")
            ))
            aplicacao_id = cursor.lastrowid

            for indice, aluno_id in enumerate(selecionados):
                modelo = (indice % quantidade_modelos) + 1
                cursor.execute("""
                    INSERT INTO aplicacao_alunos (aplicacao_id, aluno_id, modelo)
                    VALUES (?, ?, ?)
                """, (aplicacao_id, aluno_id, modelo))

            banco.commit()
            flash("Aplicação criada com sucesso.", "sucesso")
            return redirect(f"/provas/{prova_id}/aplicacoes")

        return render_template("aplicacoes/nova.html", prova=prova, alunos=alunos)
    except sqlite3.Error as erro:
        banco.rollback()
        flash(f"Não foi possível criar a aplicação: {erro}", "erro")
        return redirect(f"/provas/{prova_id}/aplicacoes")
    finally:
        banco.close()


@app.route("/aplicacoes/<int:aplicacao_id>/modelo/<int:modelo>")
def imprimir_modelo_aplicacao(aplicacao_id, modelo):
    _garantir_tabelas_aplicacoes()
    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()
    try:
        cursor.execute("""
            SELECT a.*, p.nome AS prova_nome, p.disciplina, p.professor_id,
                   CASE WHEN CAST(COALESCE(p.tem_nota, 0) AS INTEGER) = 1 THEN 1 ELSE 0 END AS tem_nota,
                   t.nome AS turma_nome, e.nome_instituicao, e.logo,
                   e.cidade, e.estado,
                   COALESCE(pr.nome, 'Não informado') AS professor_nome
            FROM aplicacoes a
            INNER JOIN provas p ON p.id = a.prova_id
            INNER JOIN turmas t ON t.id = a.turma_id
            LEFT JOIN escolas e ON e.id = a.escola_id
            LEFT JOIN professores pr ON pr.id = p.professor_id
            WHERE a.id = ?
        """, (aplicacao_id,))
        aplicacao = cursor.fetchone()
        if not aplicacao or modelo < 1 or modelo > int(aplicacao["quantidade_modelos"]):
            flash("Modelo de prova não encontrado.", "erro")
            return redirect("/provas")

        if not _pode_gerenciar_prova(
            cursor, aplicacao["prova_id"], exigir_edicao=False
        ):
            return _redirecionar_acesso_negado_prova()

        questoes = _questoes_modelo_aplicacao(cursor, aplicacao_id, modelo)
        return render_template("aplicacoes/modelo_prova.html", aplicacao=aplicacao, questoes=questoes, modelo=modelo)
    finally:
        banco.close()


@app.route("/aplicacoes/<int:aplicacao_id>/cartoes")
def cartoes_aplicacao(aplicacao_id):
    _garantir_tabelas_aplicacoes()
    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()
    try:
        cursor.execute("""
            SELECT a.*, p.nome AS prova_nome, p.disciplina, p.data_aplicacao,
                   t.nome AS turma_nome, e.nome_instituicao, e.logo,
                   COALESCE(al.ano, e.ano_letivo) AS ano_letivo,
                   COALESCE(pr.nome, 'Não informado') AS professor_nome
            FROM aplicacoes a
            INNER JOIN provas p ON p.id = a.prova_id
            INNER JOIN turmas t ON t.id = a.turma_id
            LEFT JOIN escolas e ON e.id = a.escola_id
            LEFT JOIN anos_letivos al ON al.id = a.ano_letivo_id
            LEFT JOIN professores pr ON pr.id = p.professor_id
            WHERE a.id = ?
        """, (aplicacao_id,))
        aplicacao = cursor.fetchone()
        if not aplicacao:
            flash("Aplicação não encontrada.", "erro")
            return redirect("/provas")

        if not _pode_gerenciar_prova(
            cursor, aplicacao["prova_id"], exigir_edicao=False
        ):
            return _redirecionar_acesso_negado_prova()

        # Corrige automaticamente aplicações antigas que ficaram sem
        # aplicacao_alunos. A operação só cria vínculos ausentes.
        _sincronizar_alunos_aplicacao(cursor, aplicacao_id)
        banco.commit()

        cursor.execute("""
            SELECT aa.modelo, al.id, al.nome, COALESCE(al.matricula, '') AS matricula
            FROM aplicacao_alunos aa
            INNER JOIN alunos al ON al.id = aa.aluno_id
            WHERE aa.aplicacao_id = ?
            ORDER BY al.nome COLLATE NOCASE
        """, (aplicacao_id,))
        alunos = [dict(linha) for linha in cursor.fetchall()]

        # Compatibilidade/recuperação para aplicações antigas: algumas aplicações
        # podem existir sem linhas em aplicacao_alunos (por exemplo, aplicações
        # criadas antes da estrutura atual ou após migração). Nessa situação a
        # página de cartões não deve ficar completamente em branco. Usamos, apenas
        # para a geração desta tela, os alunos atualmente matriculados na turma e
        # distribuímos os modelos de forma determinística. Não gravamos nada no
        # banco aqui, portanto não alteramos a aplicação nem os dados existentes.
        if not alunos:
            cursor.execute("""
                SELECT DISTINCT al.id, al.nome,
                       COALESCE(al.matricula, '') AS matricula
                FROM alunos al
                LEFT JOIN aluno_matriculas am
                  ON am.aluno_id = al.id
                 AND am.turma_id = ?
                 AND (? IS NULL OR am.ano_letivo_id = ?)
                 AND COALESCE(am.situacao, 'Cursando') = 'Cursando'
                WHERE am.id IS NOT NULL OR al.turma_id = ?
                ORDER BY al.nome COLLATE NOCASE
            """, (
                aplicacao["turma_id"],
                aplicacao["ano_letivo_id"],
                aplicacao["ano_letivo_id"],
                aplicacao["turma_id"],
            ))
            quantidade_modelos = max(1, int(aplicacao["quantidade_modelos"] or 1))
            alunos = []
            for indice, linha in enumerate(cursor.fetchall()):
                item = dict(linha)
                item["modelo"] = (indice % quantidade_modelos) + 1
                alunos.append(item)

        cartoes = []
        for aluno in alunos:
            questoes = _questoes_modelo_aplicacao(cursor, aplicacao_id, aluno["modelo"])
            itens = []
            discursivas = []
            for numero, q in enumerate(questoes, 1):
                item = {"numero": numero, "questao_id": q["id"], "discursiva": _tipo_discursivo_aplicacao(q["tipo_questao"]), "linhas": int(q["linhas_resposta"] or 5) if "linhas_resposta" in q.keys() else 5}
                itens.append(item)
                if item["discursiva"]:
                    discursivas.append(item)

            # Código visual da prova sempre com seis algarismos.
            codigo_prova = f"{int(aplicacao['prova_id']):06d}"

            # Mapa para vincular cada área discursiva à questão correta.
            mapa_discursivas = ",".join(
                f"{item['numero']}:{item['questao_id']}"
                for item in discursivas
            )

            def gerar_qr_cartao(tipo_folha):
                conteudo_qr = (
                    f"CODIGO:{codigo_prova}|PROVA:{aplicacao['prova_id']}|"
                    f"ALUNO:{aluno['id']}|APLICACAO:{aplicacao_id}|"
                    f"MODELO:{aluno['modelo']}|FOLHA:{tipo_folha}"
                )

                if tipo_folha == "DISCURSIVAS" and mapa_discursivas:
                    conteudo_qr += f"|QUESTOES:{mapa_discursivas}"

                qr = qrcode.QRCode(
                    version=None,
                    error_correction=qrcode.constants.ERROR_CORRECT_M,
                    box_size=5,
                    border=2
                )
                qr.add_data(conteudo_qr)
                qr.make(fit=True)
                imagem_qr = qr.make_image(
                    fill_color="black",
                    back_color="white"
                )
                buffer = BytesIO()
                imagem_qr.save(buffer, format="PNG")
                return base64.b64encode(buffer.getvalue()).decode("utf-8")

            cartoes.append({
                "aluno": aluno,
                "questoes": itens,
                "discursivas": discursivas,
                "mapa_discursivas": mapa_discursivas,
                "codigo_prova": codigo_prova,
                "qr_base64": gerar_qr_cartao("OBJETIVAS"),
                "qr_discursiva_base64": gerar_qr_cartao("DISCURSIVAS")
            })

        return render_template(
            "aplicacoes/cartoes.html",
            aplicacao=aplicacao,
            cartoes=cartoes,
            codigo_prova=f"{int(aplicacao['prova_id']):06d}"
        )
    finally:
        banco.close()



def _decodificar_qr_cartao(caminho_imagem):
    """Lê o QR Code usando somente OpenCV.

    Faz tentativas na imagem inteira e na região superior direita, aplicando
    escalas, rotações e tratamentos de contraste para fotos e PDFs digitalizados.
    """
    imagem = cv2.imread(caminho_imagem)
    if imagem is None:
        return None

    detector = cv2.QRCodeDetector()

    def tentar_decodificar(candidata):
        if candidata is None or candidata.size == 0:
            return None

        # Primeiro tenta detectar vários QR Codes (quando suportado pelo OpenCV).
        try:
            resultado = detector.detectAndDecodeMulti(candidata)
            if len(resultado) == 4:
                sucesso, textos, _, _ = resultado
                if sucesso and textos:
                    for texto in textos:
                        if texto and texto.strip():
                            return texto.strip()
        except (cv2.error, AttributeError, TypeError, ValueError):
            pass

        # Depois tenta o detector tradicional para um único QR Code.
        try:
            texto, _, _ = detector.detectAndDecode(candidata)
            if texto and texto.strip():
                return texto.strip()
        except cv2.error:
            pass

        return None

    altura, largura = imagem.shape[:2]
    regioes = [imagem]

    # Nos cartões do ARK EDUS, o QR costuma ficar no topo direito.
    recorte_superior_direito = imagem[
        0:max(1, int(altura * 0.32)),
        max(0, int(largura * 0.58)):largura
    ]
    if recorte_superior_direito.size:
        regioes.append(recorte_superior_direito)

    for regiao in regioes:
        for angulo in (0, 90, 180, 270):
            if angulo == 0:
                rotacionada = regiao
            elif angulo == 90:
                rotacionada = cv2.rotate(regiao, cv2.ROTATE_90_CLOCKWISE)
            elif angulo == 180:
                rotacionada = cv2.rotate(regiao, cv2.ROTATE_180)
            else:
                rotacionada = cv2.rotate(regiao, cv2.ROTATE_90_COUNTERCLOCKWISE)

            for escala in (1.0, 1.5, 2.0, 3.0, 4.0):
                candidata = rotacionada if escala == 1.0 else cv2.resize(
                    rotacionada,
                    None,
                    fx=escala,
                    fy=escala,
                    interpolation=cv2.INTER_CUBIC
                )

                texto = tentar_decodificar(candidata)
                if texto:
                    return texto

                if len(candidata.shape) == 3:
                    cinza = cv2.cvtColor(candidata, cv2.COLOR_BGR2GRAY)
                else:
                    cinza = candidata

                versoes = [cinza]

                # Melhora contraste local em digitalizações com iluminação irregular.
                clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
                versoes.append(clahe.apply(cinza))

                _, otsu = cv2.threshold(
                    cinza,
                    0,
                    255,
                    cv2.THRESH_BINARY + cv2.THRESH_OTSU
                )
                versoes.append(otsu)

                versoes.append(cv2.adaptiveThreshold(
                    cinza,
                    255,
                    cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                    cv2.THRESH_BINARY,
                    31,
                    5
                ))

                for versao in versoes:
                    texto = tentar_decodificar(versao)
                    if texto:
                        return texto

    return None


def _converter_pdf_em_imagens(caminho_pdf, pasta_destino):
    """Converte cada página de um PDF em PNG usando PyMuPDF.

    Instale a dependência com: pip install PyMuPDF
    """
    try:
        import fitz
    except ImportError as erro:
        raise RuntimeError(
            "Para importar PDF, instale a dependência PyMuPDF: "
            "pip install PyMuPDF"
        ) from erro

    caminhos = []
    documento = fitz.open(caminho_pdf)
    try:
        if documento.page_count < 1:
            raise ValueError("O PDF não possui páginas.")

        for indice in range(documento.page_count):
            pagina = documento.load_page(indice)
            matriz = fitz.Matrix(2.2, 2.2)
            pixmap = pagina.get_pixmap(matrix=matriz, alpha=False)
            nome = f"{uuid.uuid4().hex}_pagina_{indice + 1:03d}.png"
            caminho = os.path.join(pasta_destino, nome)
            pixmap.save(caminho)
            caminhos.append((caminho, indice + 1))
    finally:
        documento.close()

    return caminhos


def _remover_arquivo_silenciosamente(caminho):
    try:
        if caminho and os.path.isfile(caminho):
            os.remove(caminho)
    except OSError as erro:
        print(f"Não foi possível remover o arquivo {caminho}: {erro}")


@app.route("/aplicacoes/<int:aplicacao_id>/importar", methods=["GET", "POST"])
def importar_cartoes_aplicacao(aplicacao_id):
    """Importa cartões objetivos.

    Ausência de modelo, questão em branco ou dupla marcação nunca descarta o
    cartão. Esses casos são registrados como revisão necessária.
    """
    _garantir_tabelas_aplicacoes()
    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        cursor.execute("""
            SELECT a.*, p.nome AS prova_nome, p.disciplina,
                   t.nome AS turma_nome,
                   EXISTS (
                       SELECT 1
                       FROM prova_questoes pq
                       INNER JOIN questoes q ON q.id = pq.questao_id
                       WHERE pq.prova_id = a.prova_id
                         AND LOWER(TRIM(COALESCE(q.tipo_questao, ''))) IN (
                             'discursiva', 'dissertativa',
                             'resposta_aberta', 'resposta aberta'
                         )
                   ) AS possui_discursivas
            FROM aplicacoes a
            INNER JOIN provas p ON p.id = a.prova_id
            INNER JOIN turmas t ON t.id = a.turma_id
            WHERE a.id = ?
        """, (aplicacao_id,))
        aplicacao = cursor.fetchone()

        if not aplicacao:
            flash("Aplicação não encontrada.", "erro")
            return redirect("/provas")

        if not _pode_gerenciar_prova(
            cursor,
            aplicacao["prova_id"],
            exigir_edicao=True,
            permitir_finalizada=True
        ):
            return _redirecionar_acesso_negado_prova()

        # Importação e cartões precisam usar a mesma fonte de alunos.
        # Para aplicações antigas, recompõe apenas os vínculos faltantes.
        _sincronizar_alunos_aplicacao(cursor, aplicacao_id)
        banco.commit()

        if request.method == "POST":
            arquivos = [
                arquivo for arquivo in request.files.getlist("arquivos")
                if arquivo and arquivo.filename
            ]

            if not arquivos:
                flash("Selecione pelo menos um PDF ou uma imagem.", "erro")
                return redirect(url_for(
                    "importar_cartoes_aplicacao",
                    aplicacao_id=aplicacao_id
                ))

            extensoes_permitidas = {".pdf", ".jpg", ".jpeg", ".png"}
            pasta = os.path.join(
                app.config["UPLOAD_FOLDER"], "aplicacoes", str(aplicacao_id)
            )
            os.makedirs(pasta, exist_ok=True)

            importados = 0
            revisoes = 0
            descartados = 0
            erros_importacao = []

            for arquivo in arquivos:
                nome_original = secure_filename(arquivo.filename)
                extensao = os.path.splitext(nome_original)[1].lower()

                if extensao not in extensoes_permitidas:
                    descartados += 1
                    continue

                nome_temporario = f"{uuid.uuid4().hex}_{nome_original}"
                caminho_temporario = os.path.join(pasta, nome_temporario)
                arquivo.save(caminho_temporario)

                try:
                    if extensao == ".pdf":
                        paginas = _converter_pdf_em_imagens(
                            caminho_temporario, pasta
                        )
                        _remover_arquivo_silenciosamente(caminho_temporario)
                    else:
                        paginas = [(caminho_temporario, None)]
                except Exception as erro:
                    print(f"ERRO AO ABRIR {arquivo.filename}: {erro}")
                    _remover_arquivo_silenciosamente(caminho_temporario)
                    descartados += 1
                    continue

                for caminho, numero_pagina in paginas:
                    cursor.execute("SAVEPOINT importar_pagina")

                    try:
                        qr_texto = _decodificar_qr_cartao(caminho)
                        if not qr_texto:
                            raise ValueError("QR Code não identificado.")

                        dados_qr = {}
                        for parte in qr_texto.split("|"):
                            if ":" in parte:
                                chave, valor = parte.split(":", 1)
                                dados_qr[chave.strip().upper()] = valor.strip()

                        if int(dados_qr.get("APLICACAO", 0)) != aplicacao_id:
                            raise ValueError(
                                "O cartão pertence a outra aplicação."
                            )

                        aluno_id = int(dados_qr["ALUNO"])
                        modelo_qr = int(dados_qr.get("MODELO", 1))

                        # Aplicações antigas podem ter cartões válidos gerados
                        # antes da criação dos vínculos em aplicacao_alunos. O QR
                        # identifica o aluno de forma inequívoca; valida a turma e
                        # recompõe somente esse vínculo antes da correção.
                        if not _garantir_aluno_qr_na_aplicacao(
                            cursor, aplicacao_id, aluno_id, modelo_qr
                        ):
                            raise ValueError(
                                "O aluno identificado pelo QR não pertence à turma desta aplicação."
                            )

                        modelo_visual = ler_modelo_cartao(
                            caminho,
                            aplicacao["quantidade_modelos"] or 1
                        )

                        # O QR é usado somente para conseguir montar a leitura
                        # provisória. Se a bolha do modelo estiver vazia, o
                        # professor obrigatoriamente escolherá o modelo na revisão.
                        modelo_calculo = modelo_visual or modelo_qr
                        tipo_folha = dados_qr.get(
                            "FOLHA", "OBJETIVAS"
                        ).upper()

                        if tipo_folha == "DISCURSIVAS":
                            cursor.execute("""
                                SELECT 1
                                FROM aplicacao_alunos
                                WHERE aplicacao_id = ? AND aluno_id = ?
                            """, (aplicacao_id, aluno_id))
                            if not cursor.fetchone():
                                raise ValueError(
                                    "O aluno não pertence a esta aplicação."
                                )

                            questoes_mapeadas = []
                            mapa_qr = dados_qr.get("QUESTOES", "")
                            for item_mapa in mapa_qr.split(","):
                                item_mapa = item_mapa.strip()
                                if not item_mapa or ":" not in item_mapa:
                                    continue
                                numero_texto, questao_texto = item_mapa.split(":", 1)
                                try:
                                    questoes_mapeadas.append(
                                        (int(numero_texto), int(questao_texto))
                                    )
                                except (TypeError, ValueError):
                                    continue

                            if not questoes_mapeadas:
                                questoes_modelo = _questoes_modelo_aplicacao(
                                    cursor, aplicacao_id, modelo_calculo
                                )
                                questoes_mapeadas = [
                                    (numero, questao["id"])
                                    for numero, questao in enumerate(
                                        questoes_modelo, start=1
                                    )
                                    if _tipo_discursivo_aplicacao(
                                        questao["tipo_questao"]
                                    )
                                ]

                            if not questoes_mapeadas:
                                raise ValueError(
                                    "Esta aplicação não possui questões discursivas."
                                )

                            nome_exibicao = nome_original
                            if numero_pagina is not None:
                                nome_exibicao = (
                                    f"{nome_original} · página {numero_pagina}"
                                )

                            for numero_exibicao, questao_id in questoes_mapeadas:
                                cursor.execute("""
                                    INSERT INTO respostas_discursivas_aplicacao
                                    (
                                        aplicacao_id, aluno_id, questao_id,
                                        numero_exibicao, imagem_resposta,
                                        corrigida
                                    )
                                    VALUES (?, ?, ?, ?, ?, 0)
                                    ON CONFLICT(
                                        aplicacao_id, aluno_id, questao_id
                                    ) DO UPDATE SET
                                        numero_exibicao = excluded.numero_exibicao,
                                        imagem_resposta = excluded.imagem_resposta,
                                        nota = NULL,
                                        comentario = NULL,
                                        corrigida = 0,
                                        corrigido_por = NULL,
                                        corrigido_em = NULL
                                """, (
                                    aplicacao_id,
                                    aluno_id,
                                    questao_id,
                                    numero_exibicao,
                                    caminho
                                ))

                            cursor.execute("""
                                INSERT INTO aplicacao_importacoes
                                (
                                    aplicacao_id, aluno_id, modelo,
                                    nome_arquivo, caminho_arquivo,
                                    status, mensagem, qr_texto,
                                    tipo_folha, revisado, revisado_em
                                )
                                VALUES (
                                    ?, ?, ?, ?, ?, 'Processado', ?, ?,
                                    'DISCURSIVAS', 1, ?
                                )
                            """, (
                                aplicacao_id,
                                aluno_id,
                                modelo_calculo,
                                nome_exibicao,
                                caminho,
                                (
                                    f"Folha discursiva reconhecida com "
                                    f"{len(questoes_mapeadas)} questão(ões)."
                                ),
                                qr_texto,
                                datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                            ))

                            cursor.execute("""
                                UPDATE aplicacao_alunos
                                SET modelo = ?,
                                    discursiva_pendente = 1,
                                    status = CASE
                                        WHEN objetiva_corrigida = 1
                                            THEN 'Aguardando correção discursiva'
                                        ELSE status
                                    END
                                WHERE aplicacao_id = ? AND aluno_id = ?
                            """, (
                                modelo_calculo,
                                aplicacao_id,
                                aluno_id
                            ))

                            importados += 1
                            cursor.execute(
                                "RELEASE SAVEPOINT importar_pagina"
                            )
                            continue

                        if tipo_folha != "OBJETIVAS":
                            raise ValueError(
                                "Tipo de folha não reconhecido."
                            )

                        quantidade_modelos = int(
                            aplicacao["quantidade_modelos"] or 1
                        )
                        if not 1 <= modelo_calculo <= quantidade_modelos:
                            raise ValueError("Modelo de prova inválido.")

                        cursor.execute("""
                            SELECT 1
                            FROM aplicacao_alunos
                            WHERE aplicacao_id = ? AND aluno_id = ?
                        """, (aplicacao_id, aluno_id))
                        if not cursor.fetchone():
                            raise ValueError(
                                "O aluno não pertence a esta aplicação."
                            )

                        questoes = _questoes_modelo_aplicacao(
                            cursor, aplicacao_id, modelo_calculo
                        )
                        objetivas = [
                            (numero, questao)
                            for numero, questao in enumerate(questoes, 1)
                            if not _tipo_discursivo_aplicacao(
                                questao["tipo_questao"]
                            )
                        ]
                        if not objetivas:
                            raise ValueError(
                                "Esta aplicação não possui questões objetivas."
                            )

                        leitura = ler_respostas_cartao_detalhado(
                            caminho, len(objetivas)
                        )

                        cursor.execute("""
                            DELETE FROM aplicacao_respostas_objetivas
                            WHERE aplicacao_id = ? AND aluno_id = ?
                        """, (aplicacao_id, aluno_id))
                        cursor.execute("""
                            DELETE FROM respostas_alunos
                            WHERE prova_id = ? AND aluno_id = ?
                        """, (aplicacao["prova_id"], aluno_id))

                        acertos = 0
                        nota_objetiva = 0.0
                        valor_total_objetivas = round(
                            sum(
                                float(questao.get("peso") or 0)
                                for _, questao in objetivas
                            ),
                            2
                        )
                        total_objetivas = len(objetivas)
                        total_discursivas = 0
                        ambiguas = 0
                        brancas = 0
                        indice_leitura = 0

                        for numero, questao in enumerate(questoes, 1):
                            if _tipo_discursivo_aplicacao(
                                questao["tipo_questao"]
                            ):
                                total_discursivas += 1
                                cursor.execute("""
                                    INSERT OR IGNORE INTO
                                    respostas_discursivas_aplicacao
                                    (aplicacao_id, aluno_id, questao_id,
                                     numero_exibicao)
                                    VALUES (?, ?, ?, ?)
                                """, (
                                    aplicacao_id,
                                    aluno_id,
                                    questao["id"],
                                    numero
                                ))
                                continue

                            indice_leitura += 1
                            dado = leitura.get(indice_leitura, {
                                "resposta": "",
                                "situacao": "em_branco"
                            })
                            resposta = dado.get("resposta", "")
                            situacao = dado.get(
                                "situacao", "em_branco"
                            )
                            correta = (
                                questao["correta"] or ""
                            ).strip().upper()
                            anulada = int(questao["anulada"] or 0) if "anulada" in questao.keys() else 0

                            if anulada:
                                situacao = "anulada"
                                acertou = 1
                            else:
                                acertou = int(
                                    situacao == "respondida"
                                    and resposta == correta
                                )
                                if situacao == "dupla_marcacao":
                                    ambiguas += 1
                                elif situacao == "em_branco":
                                    brancas += 1

                            acertos += acertou

                            peso_questao = round(
                                float(questao.get("peso") or 0),
                                2
                            )

                            # Cada questão objetiva vale exatamente o peso
                            # cadastrado na prova. Questão errada, em branco
                            # ou com dupla marcação soma zero. Questão anulada
                            # recebe o peso integral.
                            if acertou:
                                nota_objetiva = round(
                                    nota_objetiva + peso_questao,
                                    2
                                )

                            cursor.execute("""
                                INSERT INTO aplicacao_respostas_objetivas
                                (aplicacao_id, aluno_id, numero_questao,
                                 questao_id, modelo, resposta, situacao,
                                 resposta_correta, acertou, origem,
                                 atualizado_em, anulada)
                                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                            """, (
                                aplicacao_id,
                                aluno_id,
                                numero,
                                questao["id"],
                                modelo_calculo,
                                resposta,
                                situacao,
                                correta,
                                acertou,
                                "automatica",
                                datetime.now().strftime(
                                    "%Y-%m-%d %H:%M:%S"
                                ),
                                anulada
                            ))

                            resposta_legada = (
                                "__ANULADA__" if anulada
                                else resposta
                                if situacao == "respondida"
                                else "__DUPLA__"
                                if situacao == "dupla_marcacao"
                                else "__BRANCO__"
                            )
                            cursor.execute("""
                                INSERT INTO respostas_alunos
                                (prova_id, aluno_id, numero_questao,
                                 resposta_aluno, resposta_correta, acertou)
                                VALUES (?, ?, ?, ?, ?, ?)
                            """, (
                                aplicacao["prova_id"],
                                aluno_id,
                                numero,
                                resposta_legada,
                                "ANULADA" if anulada else correta,
                                acertou
                            ))

                        modelo_pendente = modelo_visual is None
                        precisa_revisao = bool(
                            modelo_pendente or ambiguas or brancas
                        )
                        status_importacao = (
                            "Revisão necessária"
                            if precisa_revisao else "Processado"
                        )

                        partes_mensagem = [
                            f"{acertos}/{total_objetivas} corretas",
                            f"{brancas} em branco",
                            f"{ambiguas} com dupla marcação"
                        ]
                        if modelo_pendente:
                            partes_mensagem.append("modelo não marcado")
                        partes_mensagem.append(
                            f"nota objetiva {nota_objetiva:.2f}/"
                            f"{valor_total_objetivas:.2f}"
                        )
                        mensagem = "; ".join(partes_mensagem) + "."

                        nota_objetiva = round(
                            nota_objetiva,
                            2
                        )

                        cursor.execute("""
                            UPDATE aplicacao_alunos
                            SET modelo = ?,
                                objetiva_corrigida = ?,
                                discursiva_pendente = ?,
                                acertos_objetivos = ?,
                                total_objetivas = ?,
                                nota_objetiva = ?,
                                status = ?
                            WHERE aplicacao_id = ? AND aluno_id = ?
                        """, (
                            modelo_calculo,
                            0 if precisa_revisao else 1,
                            1 if total_discursivas else 0,
                            acertos,
                            total_objetivas,
                            nota_objetiva,
                            "Revisão necessária"
                            if precisa_revisao
                            else (
                                "Aguardando correção discursiva"
                                if total_discursivas else "Corrigido"
                            ),
                            aplicacao_id,
                            aluno_id
                        ))

                        nome_exibicao = nome_original
                        if numero_pagina is not None:
                            nome_exibicao = (
                                f"{nome_original} · página {numero_pagina}"
                            )

                        cursor.execute("""
                            INSERT INTO aplicacao_importacoes
                            (aplicacao_id, aluno_id, modelo, nome_arquivo,
                             caminho_arquivo, status, mensagem, qr_texto,
                             tipo_folha, revisado, revisado_em)
                            VALUES (?, ?, ?, ?, ?, ?, ?, ?, 'OBJETIVAS', ?, ?)
                        """, (
                            aplicacao_id,
                            aluno_id,
                            modelo_visual,
                            nome_exibicao,
                            caminho,
                            status_importacao,
                            mensagem,
                            qr_texto,
                            0 if precisa_revisao else 1,
                            None if precisa_revisao else
                            datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                        ))

                        if precisa_revisao:
                            revisoes += 1
                        else:
                            importados += 1

                        cursor.execute(
                            "RELEASE SAVEPOINT importar_pagina"
                        )

                    except Exception as erro:
                        cursor.execute(
                            "ROLLBACK TO SAVEPOINT importar_pagina"
                        )
                        cursor.execute(
                            "RELEASE SAVEPOINT importar_pagina"
                        )
                        print(
                            f"ERRO AO IMPORTAR {arquivo.filename}"
                            f"{f' PÁGINA {numero_pagina}' if numero_pagina else ''}: "
                            f"{erro}"
                        )
                        erros_importacao.append(
                            f"{nome_original}{f' · página {numero_pagina}' if numero_pagina else ''}: {erro}"
                        )
                        _remover_arquivo_silenciosamente(caminho)
                        descartados += 1

            _sincronizar_status_aplicacao(cursor, aplicacao_id)
            banco.commit()

            mensagem_flash = (
                f"Importação concluída: {importados} corrigido(s), "
                f"{revisoes} para revisão e {descartados} descartado(s)."
            )
            if erros_importacao:
                mensagem_flash += " Motivo: " + " | ".join(erros_importacao[:3])
            flash(
                mensagem_flash,
                "sucesso" if not descartados else "aviso"
            )
            return redirect(url_for(
                "importar_cartoes_aplicacao",
                aplicacao_id=aplicacao_id
            ))

        # Exibe todos os alunos da aplicação, inclusive aqueles que ainda
        # não tiveram cartão importado.
        cursor.execute("""
            SELECT
                aa.id AS aplicacao_aluno_id,
                aa.aluno_id,
                aa.modelo AS modelo_aluno,
                aa.status AS aluno_status,
                aa.objetiva_corrigida,
                aa.nota_objetiva,
                aa.acertos_objetivos,
                aa.total_objetivas,
                al.nome AS aluno_nome,
                al.matricula AS aluno_matricula,
                ai.id AS importacao_id,
                ai.modelo AS modelo_importacao,
                ai.nome_arquivo,
                ai.importado_em,
                ai.status AS importacao_status,
                ai.mensagem,
                CASE
                    WHEN aa.status = 'Ausente' THEN 'Ausente'
                    WHEN ai.id IS NULL THEN 'Pendente'
                    WHEN ai.status = 'Revisão necessária'
                        THEN 'Revisão necessária'
                    ELSE 'Processado'
                END AS status_lista
            FROM aplicacao_alunos aa
            INNER JOIN alunos al ON al.id = aa.aluno_id
            LEFT JOIN aplicacao_importacoes ai
              ON ai.id = (
                    SELECT MAX(ai2.id)
                    FROM aplicacao_importacoes ai2
                    WHERE ai2.aplicacao_id = aa.aplicacao_id
                      AND ai2.aluno_id = aa.aluno_id
                      AND ai2.tipo_folha = 'OBJETIVAS'
                      AND ai2.status <> 'Erro'
                )
            WHERE aa.aplicacao_id = ?
            ORDER BY al.nome COLLATE NOCASE
        """, (aplicacao_id,))
        alunos_cartoes = cursor.fetchall()

        resumo = {
            "alunos": len(alunos_cartoes),
            "total": sum(
                1 for item in alunos_cartoes
                if item["importacao_id"] is not None
            ),
            "processados": sum(
                1 for item in alunos_cartoes
                if item["status_lista"] == "Processado"
            ),
            "revisao": sum(
                1 for item in alunos_cartoes
                if item["status_lista"] == "Revisão necessária"
            ),
            "pendentes": sum(
                1 for item in alunos_cartoes
                if item["status_lista"] == "Pendente"
            ),
            "ausentes": sum(
                1 for item in alunos_cartoes
                if item["status_lista"] == "Ausente"
            )
        }

        return render_template(
            "aplicacoes/importar.html",
            aplicacao=aplicacao,
            alunos_cartoes=alunos_cartoes,
            resumo=resumo
        )

    finally:
        banco.close()


@app.route(
    "/aplicacoes/<int:aplicacao_id>/importacoes/<int:importacao_id>/dados"
)
def dados_importacao_cartao(aplicacao_id, importacao_id):
    _garantir_tabelas_aplicacoes()
    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        cursor.execute("""
            SELECT ai.*, al.nome AS aluno_nome,
                   al.matricula AS aluno_matricula,
                   a.quantidade_modelos, a.prova_id,
                   t.nome AS turma_nome
            FROM aplicacao_importacoes ai
            INNER JOIN aplicacoes a ON a.id = ai.aplicacao_id
            INNER JOIN turmas t ON t.id = a.turma_id
            LEFT JOIN alunos al ON al.id = ai.aluno_id
            WHERE ai.id = ? AND ai.aplicacao_id = ?
        """, (importacao_id, aplicacao_id))
        importacao = cursor.fetchone()

        if not importacao:
            return jsonify({"erro": "Importação não encontrada."}), 404

        respostas = []
        if importacao["aluno_id"]:
            cursor.execute("""
                SELECT numero_questao, questao_id, modelo, resposta,
                       situacao, resposta_correta, acertou, origem, anulada
                FROM aplicacao_respostas_objetivas
                WHERE aplicacao_id = ? AND aluno_id = ?
                ORDER BY numero_questao
            """, (aplicacao_id, importacao["aluno_id"]))
            respostas = [dict(item) for item in cursor.fetchall()]

        caminho = importacao["caminho_arquivo"] or ""
        relativo = ""
        try:
            relativo = os.path.relpath(
                caminho, app.config["UPLOAD_FOLDER"]
            ).replace(os.sep, "/")
        except ValueError:
            relativo = ""

        return jsonify({
            "importacao": {
                **dict(importacao),
                "imagem_url": (
                    f"/static/uploads/{relativo}" if relativo else ""
                )
            },
            "respostas": respostas
        })

    finally:
        banco.close()


@app.route(
    "/aplicacoes/<int:aplicacao_id>/importacoes/<int:importacao_id>/salvar",
    methods=["POST"]
)
def salvar_revisao_cartao(aplicacao_id, importacao_id):
    """Salva qualquer ajuste feito pelo professor e recalcula o resultado."""
    _garantir_tabelas_aplicacoes()
    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        cursor.execute("""
            SELECT ai.*, a.prova_id, a.quantidade_modelos
            FROM aplicacao_importacoes ai
            INNER JOIN aplicacoes a ON a.id = ai.aplicacao_id
            WHERE ai.id = ? AND ai.aplicacao_id = ?
        """, (importacao_id, aplicacao_id))
        importacao = cursor.fetchone()

        if not importacao:
            return jsonify({"erro": "Importação não encontrada."}), 404

        if not _pode_gerenciar_prova(
            cursor,
            importacao["prova_id"],
            exigir_edicao=True,
            permitir_finalizada=True
        ):
            return jsonify({
                "erro": "Você não possui permissão para editar este cartão."
            }), 403

        dados = request.get_json(silent=True) or {}
        aluno_id = dados.get("aluno_id") or importacao["aluno_id"]

        try:
            modelo = int(dados.get("modelo") or 0)
        except (TypeError, ValueError):
            modelo = 0

        if not aluno_id:
            return jsonify({
                "erro": "Não foi possível identificar o aluno."
            }), 400

        quantidade_modelos = int(importacao["quantidade_modelos"] or 1)
        if modelo < 1 or modelo > quantidade_modelos:
            return jsonify({
                "erro": "Selecione o modelo correto da prova."
            }), 400

        respostas_enviadas = {
            int(item["numero"]): item
            for item in dados.get("respostas", [])
            if item.get("numero") is not None
        }

        questoes = _questoes_modelo_aplicacao(
            cursor, aplicacao_id, modelo
        )

        cursor.execute("""
            DELETE FROM aplicacao_respostas_objetivas
            WHERE aplicacao_id = ? AND aluno_id = ?
        """, (aplicacao_id, aluno_id))
        cursor.execute("""
            DELETE FROM respostas_alunos
            WHERE prova_id = ? AND aluno_id = ?
        """, (importacao["prova_id"], aluno_id))

        acertos = 0
        nota_objetiva = 0.0
        valor_total_objetivas = 0.0
        total_objetivas = 0
        total_discursivas = 0

        for numero, questao in enumerate(questoes, 1):
            if _tipo_discursivo_aplicacao(questao["tipo_questao"]):
                total_discursivas += 1
                cursor.execute("""
                    INSERT OR IGNORE INTO respostas_discursivas_aplicacao
                    (aplicacao_id, aluno_id, questao_id, numero_exibicao)
                    VALUES (?, ?, ?, ?)
                """, (
                    aplicacao_id,
                    aluno_id,
                    questao["id"],
                    numero
                ))
                continue

            total_objetivas += 1

            peso_questao = round(
                float(questao.get("peso") or 0),
                2
            )
            valor_total_objetivas = round(
                valor_total_objetivas + peso_questao,
                2
            )

            anulada = int(questao["anulada"] or 0) if "anulada" in questao.keys() else 0
            recebido = respostas_enviadas.get(numero, {})
            situacao = recebido.get("situacao", "em_branco")
            resposta = (
                recebido.get("resposta", "").strip().upper()
                if situacao == "respondida" else ""
            )

            if situacao not in {
                "respondida", "em_branco", "dupla_marcacao"
            }:
                situacao = "em_branco"
                resposta = ""

            if situacao == "respondida":
                if resposta not in {"A", "B", "C", "D"}:
                    situacao = "em_branco"
                    resposta = ""
            elif situacao not in {"em_branco", "dupla_marcacao"}:
                situacao = "em_branco"
                resposta = ""

            correta = (questao["correta"] or "").strip().upper()

            if anulada:
                situacao = "anulada"
                acertou = 1
            elif situacao == "respondida":
                acertou = int(resposta == correta)
            else:
                # Questão em branco ou com dupla marcação vale zero.
                acertou = 0

            acertos += acertou

            # Soma o peso cadastrado somente quando a questão estiver
            # correta ou anulada.
            if acertou:
                nota_objetiva = round(
                    nota_objetiva + peso_questao,
                    2
                )

            cursor.execute("""
                INSERT INTO aplicacao_respostas_objetivas
                (aplicacao_id, aluno_id, numero_questao, questao_id,
                 modelo, resposta, situacao, resposta_correta,
                 acertou, origem, atualizado_em, anulada)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, 'manual', ?, ?)
            """, (
                aplicacao_id,
                aluno_id,
                numero,
                questao["id"],
                modelo,
                resposta,
                situacao,
                correta,
                acertou,
                datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                anulada
            ))

            resposta_legada = (
                "__ANULADA__" if anulada
                else resposta if situacao == "respondida"
                else "__DUPLA__" if situacao == "dupla_marcacao"
                else "__BRANCO__"
            )

            cursor.execute("""
                INSERT INTO respostas_alunos
                (prova_id, aluno_id, numero_questao, resposta_aluno,
                 resposta_correta, acertou)
                VALUES (?, ?, ?, ?, ?, ?)
            """, (
                importacao["prova_id"],
                aluno_id,
                numero,
                resposta_legada,
                "ANULADA" if anulada else correta,
                acertou
            ))

        nota = round(nota_objetiva, 2)

        cursor.execute("""
            UPDATE aplicacao_alunos
            SET modelo = ?,
                objetiva_corrigida = 1,
                discursiva_pendente = ?,
                acertos_objetivos = ?,
                total_objetivas = ?,
                nota_objetiva = ?,
                status = ?
            WHERE aplicacao_id = ? AND aluno_id = ?
        """, (
            modelo,
            1 if total_discursivas else 0,
            acertos,
            total_objetivas,
            nota,
            "Aguardando correção discursiva"
            if total_discursivas else "Corrigido",
            aplicacao_id,
            aluno_id
        ))

        cursor.execute("""
            UPDATE aplicacao_importacoes
            SET aluno_id = ?,
                modelo = ?,
                status = 'Processado',
                mensagem = ?,
                revisado = 1,
                revisado_por = ?,
                revisado_em = ?
            WHERE id = ? AND aplicacao_id = ?
        """, (
            aluno_id,
            modelo,
            (
                f"Conferência concluída pelo professor: "
                f"{acertos}/{total_objetivas} corretas; "
                f"nota objetiva {nota:.2f}/"
                f"{valor_total_objetivas:.2f}."
            ),
            session.get("usuario_id"),
            datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            importacao_id,
            aplicacao_id
        ))

        banco.commit()
        return jsonify({
            "sucesso": True,
            "acertos": acertos,
            "total": total_objetivas,
            "nota": nota
        })

    except Exception as erro:
        banco.rollback()
        return jsonify({"erro": str(erro)}), 400

    finally:
        banco.close()



@app.route(
    "/aplicacoes/<int:aplicacao_id>/alunos/<int:aluno_id>/ausencia",
    methods=["POST"]
)
def alterar_ausencia_aplicacao(aplicacao_id, aluno_id):
    """Marca ou desfaz a ausência de um aluno na aplicação."""
    _garantir_tabelas_aplicacoes()
    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        cursor.execute("""
            SELECT a.prova_id, aa.status
            FROM aplicacao_alunos aa
            INNER JOIN aplicacoes a ON a.id = aa.aplicacao_id
            WHERE aa.aplicacao_id = ? AND aa.aluno_id = ?
        """, (aplicacao_id, aluno_id))
        registro = cursor.fetchone()

        if not registro:
            return jsonify({
                "erro": "Aluno não encontrado nesta aplicação."
            }), 404

        if not _pode_gerenciar_prova(
            cursor,
            registro["prova_id"],
            exigir_edicao=True,
            permitir_finalizada=True
        ):
            return jsonify({
                "erro": "Você não possui permissão para alterar a presença."
            }), 403

        dados = request.get_json(silent=True) or {}
        ausente = bool(dados.get("ausente"))

        if ausente:
            cursor.execute("""
                SELECT COUNT(*) AS total
                FROM aplicacao_importacoes
                WHERE aplicacao_id = ?
                  AND aluno_id = ?
                  AND tipo_folha = 'OBJETIVAS'
                  AND status <> 'Erro'
            """, (aplicacao_id, aluno_id))

            if int(cursor.fetchone()["total"] or 0) > 0:
                return jsonify({
                    "erro": (
                        "Este aluno já possui cartão importado. "
                        "Exclua a importação antes de marcar falta."
                    )
                }), 409

            cursor.execute("""
                UPDATE aplicacao_alunos
                SET status = 'Ausente',
                    objetiva_corrigida = 0,
                    acertos_objetivos = 0,
                    total_objetivas = 0,
                    nota_objetiva = NULL
                WHERE aplicacao_id = ? AND aluno_id = ?
            """, (aplicacao_id, aluno_id))
            mensagem = "Aluno marcado como ausente."
        else:
            cursor.execute("""
                UPDATE aplicacao_alunos
                SET status = 'Aguardando aplicação',
                    objetiva_corrigida = 0,
                    acertos_objetivos = 0,
                    total_objetivas = 0,
                    nota_objetiva = NULL
                WHERE aplicacao_id = ? AND aluno_id = ?
            """, (aplicacao_id, aluno_id))
            mensagem = "Ausência desfeita. O aluno voltou para pendentes."

        _sincronizar_status_aplicacao(cursor, aplicacao_id)
        banco.commit()
        return jsonify({
            "sucesso": True,
            "ausente": ausente,
            "mensagem": mensagem
        })

    except Exception as erro:
        banco.rollback()
        return jsonify({"erro": str(erro)}), 400

    finally:
        banco.close()


@app.route(
    "/aplicacoes/<int:aplicacao_id>/importacoes/<int:importacao_id>/excluir",
    methods=["POST"]
)
def excluir_importacao_cartao(aplicacao_id, importacao_id):
    """Exclui a importação, a imagem e o resultado ligado a ela."""
    _garantir_tabelas_aplicacoes()
    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()
    caminho_arquivo = None

    try:
        cursor.execute("""
            SELECT ai.*, a.prova_id
            FROM aplicacao_importacoes ai
            INNER JOIN aplicacoes a ON a.id = ai.aplicacao_id
            WHERE ai.id = ? AND ai.aplicacao_id = ?
        """, (importacao_id, aplicacao_id))
        importacao = cursor.fetchone()

        if not importacao:
            return jsonify({"erro": "Importação não encontrada."}), 404

        if not _pode_gerenciar_prova(
            cursor,
            importacao["prova_id"],
            exigir_edicao=True,
            permitir_finalizada=True
        ):
            return jsonify({
                "erro": "Você não possui permissão para excluir este cartão."
            }), 403

        aluno_id = importacao["aluno_id"]
        caminho_arquivo = importacao["caminho_arquivo"]

        cursor.execute("""
            DELETE FROM aplicacao_importacoes
            WHERE id = ? AND aplicacao_id = ?
        """, (importacao_id, aplicacao_id))

        if aluno_id:
            cursor.execute("""
                SELECT COUNT(*) AS total
                FROM aplicacao_importacoes
                WHERE aplicacao_id = ?
                  AND aluno_id = ?
                  AND tipo_folha = 'OBJETIVAS'
            """, (aplicacao_id, aluno_id))
            possui_outra = int(cursor.fetchone()["total"] or 0) > 0

            if not possui_outra:
                cursor.execute("""
                    DELETE FROM aplicacao_respostas_objetivas
                    WHERE aplicacao_id = ? AND aluno_id = ?
                """, (aplicacao_id, aluno_id))
                cursor.execute("""
                    DELETE FROM respostas_alunos
                    WHERE prova_id = ? AND aluno_id = ?
                """, (importacao["prova_id"], aluno_id))
                cursor.execute("""
                    UPDATE aplicacao_alunos
                    SET objetiva_corrigida = 0,
                        acertos_objetivos = 0,
                        total_objetivas = 0,
                        nota_objetiva = NULL,
                        status = 'Aguardando aplicação'
                    WHERE aplicacao_id = ? AND aluno_id = ?
                """, (aplicacao_id, aluno_id))

        banco.commit()
        _remover_arquivo_silenciosamente(caminho_arquivo)

        return jsonify({
            "sucesso": True,
            "mensagem": "Importação excluída com sucesso."
        })

    except Exception as erro:
        banco.rollback()
        return jsonify({"erro": str(erro)}), 400

    finally:
        banco.close()

@app.route("/aplicacoes/<int:aplicacao_id>/excluir", methods=["POST"])
def excluir_aplicacao(aplicacao_id):
    """
    Exclui a aplicação e todos os registros vinculados.

    A remoção é feita explicitamente para também funcionar em bancos antigos
    nos quais as chaves estrangeiras com ON DELETE CASCADE podem não ter sido
    criadas corretamente.
    """
    _garantir_tabelas_aplicacoes()
    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()
    arquivos_para_excluir = []

    try:
        cursor.execute("PRAGMA foreign_keys = ON")

        cursor.execute("""
            SELECT prova_id
            FROM aplicacoes
            WHERE id = ?
        """, (aplicacao_id,))
        aplicacao = cursor.fetchone()

        if not aplicacao:
            flash("Aplicação não encontrada.", "erro")
            return redirect("/provas")

        if not _pode_gerenciar_prova(
            cursor,
            aplicacao["prova_id"],
            exigir_edicao=True,
            permitir_finalizada=True
        ):
            return _redirecionar_acesso_negado_prova()

        prova_id = aplicacao["prova_id"]

        # Guarda os caminhos antes de remover os registros do banco.
        cursor.execute("""
            SELECT caminho_arquivo
            FROM aplicacao_importacoes
            WHERE aplicacao_id = ?
              AND TRIM(COALESCE(caminho_arquivo, '')) <> ''
        """, (aplicacao_id,))
        arquivos_para_excluir.extend(
            linha["caminho_arquivo"] for linha in cursor.fetchall()
        )

        cursor.execute("""
            SELECT imagem_resposta
            FROM respostas_discursivas_aplicacao
            WHERE aplicacao_id = ?
              AND TRIM(COALESCE(imagem_resposta, '')) <> ''
        """, (aplicacao_id,))
        arquivos_para_excluir.extend(
            linha["imagem_resposta"] for linha in cursor.fetchall()
        )

        # Remove primeiro as tabelas filhas, garantindo compatibilidade
        # com estruturas antigas sem cascata ativa.
        tabelas_filhas = (
            "aplicacao_respostas_objetivas",
            "respostas_discursivas_aplicacao",
            "aplicacao_importacoes",
            "aplicacao_alunos",
        )

        for tabela in tabelas_filhas:
            cursor.execute(
                f"DELETE FROM {tabela} WHERE aplicacao_id = ?",
                (aplicacao_id,)
            )

        cursor.execute(
            "DELETE FROM aplicacoes WHERE id = ?",
            (aplicacao_id,)
        )

        if cursor.rowcount == 0:
            raise RuntimeError("A aplicação não foi excluída.")

        banco.commit()

        # Exclui os arquivos somente depois de confirmar a transação.
        for caminho in set(arquivos_para_excluir):
            _remover_arquivo_silenciosamente(caminho)

        flash("Aplicação excluída com sucesso.", "sucesso")
        return redirect(f"/provas/{prova_id}/aplicacoes")

    except Exception as erro:
        banco.rollback()
        print("ERRO AO EXCLUIR APLICAÇÃO:", erro)
        flash(
            f"Não foi possível excluir a aplicação: {erro}",
            "erro"
        )
        destino = request.referrer
        if not destino:
            if 'aplicacao' in locals() and aplicacao:
                destino = f"/provas/{aplicacao['prova_id']}/aplicacoes"
            else:
                destino = "/provas"
        return redirect(destino)

    finally:
        banco.close()

@app.route("/aplicacoes/<int:aplicacao_id>/discursivas")
def corrigir_discursivas_aplicacao(aplicacao_id):
    """Exibe as respostas discursivas digitalizadas para correção."""
    _garantir_tabelas_aplicacoes()
    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        cursor.execute("""
            SELECT
                a.*,
                p.nome AS prova_nome,
                p.disciplina,
                COALESCE(p.peso_total, 10) AS peso_total,
                t.nome AS turma_nome
            FROM aplicacoes a
            INNER JOIN provas p ON p.id = a.prova_id
            INNER JOIN turmas t ON t.id = a.turma_id
            WHERE a.id = ?
        """, (aplicacao_id,))
        aplicacao = cursor.fetchone()

        if not aplicacao:
            flash("Aplicação não encontrada.", "erro")
            return redirect("/provas")

        if not _pode_gerenciar_prova(
            cursor,
            aplicacao["prova_id"],
            exigir_edicao=True,
            permitir_finalizada=True
        ):
            return _redirecionar_acesso_negado_prova()

        cursor.execute("""
            SELECT
                rd.id,
                rd.aluno_id,
                rd.questao_id,
                rd.numero_exibicao,
                rd.imagem_resposta,
                rd.nota,
                rd.comentario,
                rd.corrigida,
                al.nome AS aluno_nome,
                COALESCE(al.matricula, '') AS aluno_matricula,
                q.enunciado,
                COALESCE(q.resposta_esperada, '') AS resposta_esperada,
                COALESCE(q.criterios_correcao, '') AS criterios_correcao,
                ROUND(COALESCE(pq.peso, 0), 2) AS peso
            FROM respostas_discursivas_aplicacao rd
            INNER JOIN alunos al ON al.id = rd.aluno_id
            INNER JOIN questoes q ON q.id = rd.questao_id
            INNER JOIN prova_questoes pq
              ON pq.prova_id = ?
             AND pq.questao_id = rd.questao_id
            WHERE rd.aplicacao_id = ?
              AND TRIM(COALESCE(rd.imagem_resposta, '')) <> ''
            ORDER BY
                al.nome COLLATE NOCASE,
                rd.numero_exibicao
        """, (aplicacao["prova_id"], aplicacao_id))
        respostas = [dict(item) for item in cursor.fetchall()]

        # Totais atuais por aluno, exibidos durante a correção.
        cursor.execute("""
            SELECT
                aa.aluno_id,
                aa.nota_objetiva,
                aa.nota_discursiva,
                aa.nota_final
            FROM aplicacao_alunos aa
            WHERE aa.aplicacao_id = ?
        """, (aplicacao_id,))
        totais_alunos = {
            item["aluno_id"]: dict(item)
            for item in cursor.fetchall()
        }

        for resposta in respostas:
            totais = totais_alunos.get(resposta["aluno_id"], {})
            resposta["nota_objetiva"] = totais.get("nota_objetiva")
            resposta["nota_discursiva_aluno"] = totais.get("nota_discursiva")
            resposta["nota_final"] = totais.get("nota_final")

            caminho = resposta.get("imagem_resposta") or ""
            caminho_normalizado = caminho.replace("\\", "/")
            marcador = "/static/"

            if marcador in caminho_normalizado:
                resposta["imagem_url"] = (
                    "/static/" + caminho_normalizado.split(marcador, 1)[1]
                )
            elif caminho_normalizado.startswith("static/"):
                resposta["imagem_url"] = "/" + caminho_normalizado
            else:
                resposta["imagem_url"] = ""

            resposta["peso"] = round(float(resposta.get("peso") or 0), 2)
            if resposta.get("nota") is not None:
                resposta["nota"] = round(float(resposta["nota"]), 2)

        total = len(respostas)
        corrigidas = sum(1 for item in respostas if item["corrigida"])
        pendentes = total - corrigidas
        valor_total_discursivas = round(
            sum(float(item.get("peso") or 0) for item in respostas), 2
        )

        return render_template(
            "aplicacoes/corrigir_discursivas.html",
            aplicacao=aplicacao,
            respostas=respostas,
            resumo={
                "total": total,
                "corrigidas": corrigidas,
                "pendentes": pendentes,
                "valor_total": valor_total_discursivas
            }
        )

    finally:
        banco.close()


@app.route(
    "/aplicacoes/<int:aplicacao_id>/discursivas/<int:resposta_id>/salvar",
    methods=["POST"]
)
def salvar_correcao_discursiva(aplicacao_id, resposta_id):
    """Salva a pontuação da discursiva e recalcula a nota final do aluno."""
    _garantir_tabelas_aplicacoes()
    banco = conectar_banco()
    banco.row_factory = sqlite3.Row
    cursor = banco.cursor()

    try:
        dados = request.get_json(silent=True) or request.form
        nota_texto = str(dados.get("nota", "")).replace(",", ".").strip()
        comentario = str(dados.get("comentario", "")).strip()

        try:
            nota = round(float(nota_texto), 2)
        except (TypeError, ValueError):
            return jsonify({"erro": "Informe uma pontuação válida."}), 400

        cursor.execute("""
            SELECT
                rd.aluno_id,
                rd.questao_id,
                a.prova_id,
                ROUND(COALESCE(pq.peso, 0), 2) AS peso
            FROM respostas_discursivas_aplicacao rd
            INNER JOIN aplicacoes a ON a.id = rd.aplicacao_id
            INNER JOIN prova_questoes pq
              ON pq.prova_id = a.prova_id
             AND pq.questao_id = rd.questao_id
            WHERE rd.id = ?
              AND rd.aplicacao_id = ?
        """, (resposta_id, aplicacao_id))
        resposta = cursor.fetchone()

        if not resposta:
            return jsonify({"erro": "Resposta não encontrada."}), 404

        if not _pode_gerenciar_prova(
            cursor,
            resposta["prova_id"],
            exigir_edicao=True,
            permitir_finalizada=True
        ):
            return jsonify({
                "erro": "Você não possui permissão para corrigir."
            }), 403

        peso_questao = round(float(resposta["peso"] or 0), 2)

        if peso_questao <= 0:
            return jsonify({
                "erro": (
                    "Esta questão não possui peso válido na prova. "
                    "Configure os pesos antes de corrigir."
                )
            }), 400

        if nota < 0 or nota > peso_questao:
            return jsonify({
                "erro": (
                    f"A pontuação deve estar entre 0 e "
                    f"{peso_questao:.2f} pontos."
                )
            }), 400

        agora = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        cursor.execute("""
            UPDATE respostas_discursivas_aplicacao
            SET nota = ?,
                comentario = ?,
                corrigida = 1,
                corrigido_por = ?,
                corrigido_em = ?
            WHERE id = ?
              AND aplicacao_id = ?
        """, (
            nota,
            comentario,
            session.get("usuario_id"),
            agora,
            resposta_id,
            aplicacao_id
        ))

        aluno_id = resposta["aluno_id"]

        cursor.execute("""
            SELECT
                COUNT(*) AS total,
                SUM(CASE WHEN rd.corrigida = 1 THEN 1 ELSE 0 END) AS corrigidas,
                ROUND(SUM(
                    CASE
                        WHEN rd.corrigida = 1 THEN COALESCE(rd.nota, 0)
                        ELSE 0
                    END
                ), 2) AS nota_discursiva,
                ROUND(SUM(COALESCE(pq.peso, 0)), 2) AS valor_discursivas
            FROM respostas_discursivas_aplicacao rd
            INNER JOIN aplicacoes a ON a.id = rd.aplicacao_id
            INNER JOIN prova_questoes pq
              ON pq.prova_id = a.prova_id
             AND pq.questao_id = rd.questao_id
            WHERE rd.aplicacao_id = ?
              AND rd.aluno_id = ?
              AND TRIM(COALESCE(rd.imagem_resposta, '')) <> ''
        """, (aplicacao_id, aluno_id))
        consolidado = cursor.fetchone()

        total = int(consolidado["total"] or 0)
        corrigidas = int(consolidado["corrigidas"] or 0)
        nota_discursiva = round(
            float(consolidado["nota_discursiva"] or 0), 2
        )
        valor_discursivas = round(
            float(consolidado["valor_discursivas"] or 0), 2
        )
        concluiu = total > 0 and corrigidas == total

        cursor.execute("""
            SELECT
                nota_objetiva,
                objetiva_corrigida
            FROM aplicacao_alunos
            WHERE aplicacao_id = ?
              AND aluno_id = ?
        """, (aplicacao_id, aluno_id))
        aluno_aplicacao = cursor.fetchone()

        nota_objetiva = (
            round(float(aluno_aplicacao["nota_objetiva"]), 2)
            if aluno_aplicacao
            and aluno_aplicacao["nota_objetiva"] is not None
            else None
        )

        # As questões já possuem pesos próprios. Portanto, a nota final é
        # a soma da parte objetiva com a soma das pontuações discursivas.
        if nota_objetiva is not None:
            nota_final = round(nota_objetiva + nota_discursiva, 2)
        elif corrigidas > 0:
            nota_final = nota_discursiva
        else:
            nota_final = None

        cursor.execute("""
            UPDATE aplicacao_alunos
            SET nota_discursiva = ?,
                nota_final = ?,
                discursiva_pendente = ?,
                status = CASE
                    WHEN ? = 1 AND objetiva_corrigida = 1 THEN 'Corrigido'
                    WHEN ? = 1 THEN 'Discursivas corrigidas'
                    ELSE 'Aguardando correção discursiva'
                END
            WHERE aplicacao_id = ?
              AND aluno_id = ?
        """, (
            nota_discursiva,
            nota_final,
            0 if concluiu else 1,
            1 if concluiu else 0,
            1 if concluiu else 0,
            aplicacao_id,
            aluno_id
        ))

        _sincronizar_status_aplicacao(cursor, aplicacao_id)

        banco.commit()

        return jsonify({
            "sucesso": True,
            "nota": nota,
            "peso": peso_questao,
            "nota_objetiva": nota_objetiva,
            "nota_discursiva": nota_discursiva,
            "valor_discursivas": valor_discursivas,
            "nota_final": nota_final,
            "aluno_concluido": concluiu,
            "mensagem": "Correção salva com sucesso."
        })

    except Exception as erro:
        banco.rollback()
        return jsonify({"erro": str(erro)}), 400

    finally:
        banco.close()



# =========================================================
# BACKUPS DO SISTEMA
# =========================================================
@app.route("/gestao/backups")
def gestao_backups():
    if session.get("usuario_cargo") != "Administrador Geral":
        return redirect(url_for("index"))
    registros = backup_manager.listar()
    return render_template(
        "gestao/backups.html",
        backups=registros,
        backup_dir=str(backup_manager.backup_dir),
        retencao_dias=backup_manager.retention_days,
        intervalo_horas=backup_manager.interval_hours,
        automatico_ativo=backup_manager.enabled,
        inclui_uploads=backup_manager.include_uploads,
        criptografia_ativa=bool(backup_manager.encryption_key),
    )


@app.route("/gestao/backups/criar", methods=["POST"])
def criar_backup_manual():
    if session.get("usuario_cargo") != "Administrador Geral":
        return redirect(url_for("index"))
    try:
        resultado = backup_manager.criar_backup(
            tipo="manual", criado_por=session.get("usuario_id")
        )
        registrar_auditoria(
            "BACKUP_MANUAL", recurso="backup", recurso_id=resultado["id"],
            detalhes={"arquivo": resultado["nome_arquivo"]},
        )
        flash("Backup criado e validado com sucesso.", "success")
    except BackupError as exc:
        flash(str(exc), "erro")
    except Exception:
        app.logger.exception("Falha ao criar backup manual")
        flash("Não foi possível criar o backup. Consulte o registro de falha.", "erro")
    return redirect(url_for("gestao_backups"))


@app.route("/gestao/backups/<int:backup_id>/baixar")
def baixar_backup(backup_id):
    if session.get("usuario_cargo") != "Administrador Geral":
        return redirect(url_for("index"))
    banco = conectar_banco(); banco.row_factory = sqlite3.Row
    registro = banco.execute(
        "SELECT * FROM backups_sistema WHERE id=? AND status='concluido'", (backup_id,)
    ).fetchone()
    banco.close()
    if not registro or not registro["caminho"]:
        abort(404)
    caminho = os.path.realpath(registro["caminho"])
    raiz = os.path.realpath(str(backup_manager.backup_dir)) + os.sep
    if not caminho.startswith(raiz) or not os.path.isfile(caminho):
        abort(404)
    registrar_auditoria("DOWNLOAD_BACKUP", recurso="backup", recurso_id=backup_id)
    return send_file(caminho, as_attachment=True, download_name=registro["nome_arquivo"])


# =========================================================
# PRIVACIDADE, SEGURANÇA E LGPD
# =========================================================
@app.route("/aceite-legal", methods=["GET", "POST"])
def aceite_legal():
    usuario_id = session.get("usuario_id")
    if not usuario_id:
        return redirect(url_for("login"))
    if request.method == "POST":
        if request.form.get("aceite_termos") != "1" or request.form.get("aceite_privacidade") != "1":
            flash("É necessário aceitar os dois documentos para continuar.", "aviso")
            return render_template("lgpd/aceite.html", termos_versao=TERMOS_VERSAO, privacidade_versao=PRIVACIDADE_VERSAO, canal_privacidade=CANAL_PRIVACIDADE)
        ip = request.headers.get("X-Forwarded-For", request.remote_addr)
        agente = request.headers.get("User-Agent", "")[:500]
        banco = conectar_banco()
        banco.execute("""UPDATE usuarios SET termos_aceitos_em=CURRENT_TIMESTAMP, termos_versao=?, privacidade_aceita_em=CURRENT_TIMESTAMP, privacidade_versao=? WHERE id=?""", (TERMOS_VERSAO, PRIVACIDADE_VERSAO, usuario_id))
        banco.executemany("INSERT INTO aceites_legais(usuario_id,documento,versao,ip,user_agent) VALUES(?,?,?,?,?)", [
            (usuario_id,"Termos de Uso",TERMOS_VERSAO,ip,agente),
            (usuario_id,"Política de Privacidade",PRIVACIDADE_VERSAO,ip,agente),
        ])
        banco.commit(); banco.close()
        registrar_auditoria("ACEITE_LEGAL", recurso="documentos_legais", detalhes={"termos":TERMOS_VERSAO,"privacidade":PRIVACIDADE_VERSAO})
        flash("Documentos aceitos. Bem-vindo à ARK EDUS!", "success")
        return redirect(url_for("index"))
    return render_template("lgpd/aceite.html", termos_versao=TERMOS_VERSAO, privacidade_versao=PRIVACIDADE_VERSAO, canal_privacidade=CANAL_PRIVACIDADE)


@app.route("/seguranca-privacidade", methods=["GET", "POST"])
def seguranca_privacidade():
    usuario_id = session.get("usuario_id")
    if not usuario_id:
        return redirect(url_for("login"))
    banco = conectar_banco(); banco.row_factory = sqlite3.Row
    if request.method == "POST":
        tipo = request.form.get("tipo", "").strip()
        descricao = request.form.get("descricao", "").strip()
        tipos_validos = {"acesso", "correcao", "exportacao", "exclusao", "revogacao", "outro"}
        if tipo not in tipos_validos:
            flash("Selecione um tipo de solicitação válido.", "erro")
        else:
            banco.execute("INSERT INTO solicitacoes_lgpd(usuario_id,escola_id,tipo,descricao) VALUES(?,?,?,?)", (usuario_id,session.get("escola_id"),tipo,descricao))
            banco.commit(); registrar_auditoria("SOLICITACAO_LGPD", recurso="solicitacoes_lgpd", detalhes={"tipo":tipo})
            flash("Solicitação registrada com sucesso.", "success")
    usuario = banco.execute("SELECT id,nome,email,cpf,termos_aceitos_em,termos_versao,privacidade_aceita_em,privacidade_versao,ultimo_login_em,ultimo_login_ip FROM usuarios WHERE id=?", (usuario_id,)).fetchone()
    sessoes = banco.execute("SELECT * FROM sessoes_usuario WHERE usuario_id=? ORDER BY criado_em DESC LIMIT 20", (usuario_id,)).fetchall()
    acessos = banco.execute("SELECT * FROM logs_auditoria WHERE usuario_id=? ORDER BY criado_em DESC LIMIT 30", (usuario_id,)).fetchall()
    solicitacoes = banco.execute("SELECT * FROM solicitacoes_lgpd WHERE usuario_id=? ORDER BY criado_em DESC", (usuario_id,)).fetchall()
    banco.close()
    return render_template("lgpd/seguranca.html", usuario=usuario, sessoes=sessoes, acessos=acessos, solicitacoes=solicitacoes, canal_privacidade=CANAL_PRIVACIDADE)


@app.route("/seguranca-privacidade/encerrar-sessao/<int:sessao_id>", methods=["POST"])
def encerrar_sessao_usuario(sessao_id):
    usuario_id = session.get("usuario_id")
    if not usuario_id: return redirect(url_for("login"))
    banco = conectar_banco()
    registro = banco.execute("SELECT token FROM sessoes_usuario WHERE id=? AND usuario_id=?", (sessao_id,usuario_id)).fetchone()
    if registro:
        banco.execute("UPDATE sessoes_usuario SET ativa=0,encerrado_em=CURRENT_TIMESTAMP WHERE id=?", (sessao_id,)); banco.commit()
        if registro[0] == session.get("token_sessao"):
            banco.close(); return redirect(url_for("logout"))
        flash("Sessão encerrada.", "success")
    banco.close(); return redirect(url_for("seguranca_privacidade"))


@app.route("/meus-dados.json")
def exportar_meus_dados():
    usuario_id = session.get("usuario_id")
    if not usuario_id: return redirect(url_for("login"))
    banco = conectar_banco(); banco.row_factory = sqlite3.Row
    u = banco.execute("SELECT id,nome,email,cpf,telefone,data_nascimento,tipo_conta,criado_em,termos_aceitos_em,termos_versao,privacidade_aceita_em,privacidade_versao FROM usuarios WHERE id=?", (usuario_id,)).fetchone()
    sols = banco.execute("SELECT tipo,descricao,status,criado_em,concluido_em FROM solicitacoes_lgpd WHERE usuario_id=?", (usuario_id,)).fetchall(); banco.close()
    dados = dict(u) if u else {}; dados["cpf"] = mascarar_cpf(dados.get("cpf")); dados["solicitacoes_lgpd"] = [dict(x) for x in sols]
    registrar_auditoria("EXPORTACAO_DADOS", recurso="usuario", recurso_id=usuario_id)
    resposta = jsonify(dados); resposta.headers["Content-Disposition"] = 'attachment; filename="meus_dados_ark_edus.json"'; return resposta


@app.route("/gestao/lgpd", methods=["GET", "POST"])
def painel_lgpd():
    if session.get("usuario_cargo") not in {"Administrador Geral", "Administrador da Instituição"}:
        return redirect(url_for("index"))
    banco = conectar_banco(); banco.row_factory = sqlite3.Row
    escola_id = session.get("escola_id")
    filtro, params = ("", []) if session.get("usuario_cargo") == "Administrador Geral" else (" WHERE (s.escola_id=? OR s.escola_id IS NULL)", [escola_id])
    if request.method == "POST":
        sid = request.form.get("solicitacao_id"); status=request.form.get("status")
        if status in {"aberta","em_analise","concluida","negada"}:
            banco.execute("UPDATE solicitacoes_lgpd SET status=?, concluido_em=CASE WHEN ?='concluida' THEN CURRENT_TIMESTAMP ELSE concluido_em END WHERE id=?", (status,status,sid)); banco.commit()
            flash("Solicitação atualizada.", "success")
    solicitacoes = banco.execute("SELECT s.*,u.nome AS usuario_nome,u.email FROM solicitacoes_lgpd s LEFT JOIN usuarios u ON u.id=s.usuario_id"+filtro+" ORDER BY s.criado_em DESC", params).fetchall()
    logs = banco.execute("SELECT l.*,u.nome AS usuario_nome FROM logs_auditoria l LEFT JOIN usuarios u ON u.id=l.usuario_id ORDER BY l.criado_em DESC LIMIT 150").fetchall()
    aceites = banco.execute("SELECT a.*,u.nome AS usuario_nome,u.email FROM aceites_legais a JOIN usuarios u ON u.id=a.usuario_id ORDER BY a.aceito_em DESC LIMIT 150").fetchall()
    incidentes = banco.execute("SELECT i.*,u.nome AS criado_por_nome FROM incidentes_seguranca i LEFT JOIN usuarios u ON u.id=i.criado_por ORDER BY i.criado_em DESC").fetchall()
    banco.close(); return render_template("lgpd/painel.html", solicitacoes=solicitacoes,logs=logs,aceites=aceites,incidentes=incidentes)


@app.route("/gestao/lgpd/incidente", methods=["POST"])
def registrar_incidente_seguranca():
    if session.get("usuario_cargo") not in {"Administrador Geral", "Administrador da Instituição"}: return redirect(url_for("index"))
    titulo=request.form.get("titulo","").strip(); descricao=request.form.get("descricao","").strip(); gravidade=request.form.get("gravidade","baixa")
    if titulo and descricao and gravidade in {"baixa","media","alta","critica"}:
        banco=conectar_banco(); banco.execute("INSERT INTO incidentes_seguranca(titulo,descricao,gravidade,escola_id,criado_por) VALUES(?,?,?,?,?)",(titulo,descricao,gravidade,session.get("escola_id"),session.get("usuario_id"))); banco.commit(); banco.close(); flash("Incidente registrado.","success")
    return redirect(url_for("painel_lgpd"))

backup_manager.iniciar_agendador()

# =========================================================
# MÓDULO SAAS: professores autônomos, planos e assinaturas
# =========================================================
from saas_arkedu import init_saas
init_saas(app, conectar_banco, sincronizar_ano_letivo_instituicao)

if __name__ == "__main__":
    app.run(
        host="0.0.0.0",
        port=int(os.environ.get("PORT", "5000")),
        debug=os.environ.get("FLASK_DEBUG", "false").lower() == "true"
    )
